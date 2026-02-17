# main.py
from fastapi import FastAPI, HTTPException, Depends, Response, Query, Path, File, UploadFile, Body 
from passlib.context import CryptContext
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr, Field
import docx
from datetime import date
 
from io import BytesIO 
from typing import List
from sendgrid import SendGridAPIClient
from datetime import datetime, timedelta,date, timezone
from sqlalchemy.dialects.postgresql import JSONB
import mammoth
from collections import defaultdict
from sqlalchemy.exc import IntegrityError
from docx.oxml.ns import qn
import copy

import pandas as pd
import os
import random 
import time
from sendgrid.helpers.mail import Mail
from typing import List, Dict, Any, Optional
import re 
import traceback
from openai import AsyncOpenAI
from google.oauth2.service_account import Credentials

import uuid
from sqlalchemy.dialects.postgresql import UUID
from google.cloud import storage
import ast




 
import json

from sqlalchemy import create_engine, Column, Integer, String, JSON, DateTime, ForeignKey, select, func, text, Boolean, Date, Text, desc, Float, Index, case, Numeric, distinct, cast, UniqueConstraint
from fastapi.responses import JSONResponse



from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session, relationship
from sqlalchemy.ext.mutable import MutableDict



from openai import OpenAI
from apscheduler.schedulers.background import BackgroundScheduler
print("🔥 FILE EXECUTION:", __name__)
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
GLOBAL_IMAGE_MAP = {}

# ---------------------------
# Database Setup PGUSER,PGPASSWORD,PGHOST,PGPORT,PGDATABASE rewrite "" using PGPASSWORD=lgZmFsBTApVPJIyTegBttTLfdWnvccHj psql -h metro.proxy.rlwy.net -U postgres -p 31631 -d railway
# ---------------------------
#DATABASE_URL = os.getenv("DATABASE_URL")
# Confirm connection details
print("PGUSER:", os.environ.get("PGUSER"))
print("PGHOST:", os.environ.get("PGHOST"))
print("PGPORT:", os.environ.get("PGPORT"))
print("PGDATABASE:", os.environ.get("PGDATABASE"))

DATABASE_URL = os.environ.get("DATABASE_URL") or (
    f"postgresql://{os.environ['PGUSER']}:{os.environ['PGPASSWORD']}"
    f"@{os.environ['PGHOST']}:{os.environ['PGPORT']}/{os.environ['PGDATABASE']}"
)
print("Connecting to DATABASE_URL:", DATABASE_URL)
engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

QuestionSchema = {
    "type": "object",
    "properties": {
        "class_name": {"type": "string"},
        "subject": {"type": "string"},
        "topic": {"type": "string"},
        "difficulty": {"type": "string"},

        # Full textual content only (concatenated text blocks)
        "question_text": {"type": "string"},

        # Ordered visual structure (text + images)
        "question_blocks": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "type": {
                        "type": "string",
                        "enum": ["text", "image"]
                    },
                    "content": {
                        "type": "string"
                    },
                    "src": {
                        "type": "string"
                    }
                },
                "required": ["type"],
                "additionalProperties": False
            }
        },

        # Flat image list (legacy / convenience)
        "images": {
            "type": "array",
            "items": {"type": "string"}
        },

        # MCQ options
        "options": {
            "type": "object",
            "additionalProperties": {"type": "string"}
        },

        "correct_answer": {"type": "string"},
        "partial": {"type": "boolean"}
    },

    # Required fields for a valid question
    "required": [
        "question_text",
        "question_blocks",
        "options",
        "correct_answer"
    ],

    "additionalProperties": False
}

# -----------------------------
# Google Cloud Storage Setup
# -----------------------------
# Path to your service account JSON for GCS
service_account_json = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS_JSON")
if not service_account_json:
    raise ValueError("Environment variable 'GOOGLE_APPLICATION_CREDENTIALS_JSON' is missing.")

# 2️⃣ Replace literal "\n" with actual newlines for the private key
service_account_info = json.loads(service_account_json)
service_account_info["private_key"] = service_account_info["private_key"].replace("\\n", "\n")

# 3️⃣ Initialize GCS client with credentials
gcs_client = storage.Client(
    credentials=Credentials.from_service_account_info(service_account_info),
    project=service_account_info["project_id"]
)

# 4️⃣ Access your bucket (GLOBAL)
BUCKET_NAME = "exammoduleimages"
gcs_bucket = gcs_client.bucket(BUCKET_NAME)

print(f"✅ Initialized GCS client for bucket: {gcs_bucket}")

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ---------------------------
# OpenAI Client
# ---------------------------
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY", "your_openai_api_key"))
client_save_questions = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))
SENDGRID_API_KEY = os.environ.get("SENDGRID_API_KEY")
otp_store = {}
# ---------------------------
# Models
# ---------------------------
class ParsingCursor:
    def __init__(self, lines):
        self.lines = lines
        self.index = 0

    def peek(self):
        if self.index < len(self.lines):
            return self.lines[self.index]
        return None

    def next(self):
        if self.index < len(self.lines):
            current = self.lines[self.index]
            self.index += 1
            return current
        return None

class UploadSummary:
    def __init__(self):
        self.saved = 0
        self.skipped_partial = 0
        self.blocks = []

    def block_success(self, block_idx, questions):
        self.blocks.append({
            "block": block_idx,
            "status": "success",
            "details": f"Saved {len(questions)} question(s)"
        })

    def block_failure(self, block_idx, reason):
        self.blocks.append({
            "block": block_idx,
            "status": "failed",
            "details": reason
        })

    def response(self):
        return {
            "status": "success" if self.saved > 0 else "failed",
            "summary": {
                "saved": self.saved,
                "skipped_partial": self.skipped_partial,
            },
            "blocks": self.blocks,
        }


class NaplanTopicCreate(BaseModel):
    name: str
    ai: Optional[int] = Field(default=0, ge=0)
    db: int = Field(..., ge=0)
    total: int = Field(..., ge=0)
 
class NaplanQuizCreate(BaseModel):
    class_name: str
    subject: str
    year: int
    difficulty: str
    num_topics: int
    total_questions: int
    topics: List[NaplanTopicCreate]
class ReviewQuestion(BaseModel):
    q_id: int
    blocks: List[Dict[str, Any]]
    options: Dict[str, Dict[str, Any]]
    student_answer: Optional[str]
    correct_answer: str


class ExamReviewResponse(BaseModel):
    exam_attempt_id: int
    questions: List[ReviewQuestion]
class StudentExamResponseMathematicalReasoning(Base):
    __tablename__ = "student_exam_response_mathematical_reasoning"

    id = Column(Integer, primary_key=True, index=True)

    # 🔑 Internal student FK (students.id)
    student_id = Column(
        Integer,
        ForeignKey("students.id", ondelete="CASCADE"),
        nullable=False,
        index=True
    )
    exam_id = Column(
        Integer,
        ForeignKey("exams.id", ondelete="CASCADE"),
        nullable=False,
        index=True
    ) 

    # 🔑 Attempt FK (student_exam_mathematical_reasoning.id)
    exam_attempt_id = Column(
        Integer,
        ForeignKey(
            "student_exam_mathematical_reasoning.id",
            ondelete="CASCADE"
        ),
        nullable=False,
        index=True
    )

    # 🧮 Question metadata
    q_id = Column(Integer, nullable=False, index=True)
    topic = Column(Text, nullable=True)

    # 📝 Answer evaluation
    selected_option = Column(Text, nullable=True)
    correct_option = Column(Text, nullable=True)
    is_correct = Column(Boolean, nullable=True)

    # --------------------------------------------------
    # SAFE relationships (optional)
    # --------------------------------------------------
    student = relationship("Student", lazy="joined")
    attempt = relationship("StudentExamMathematicalReasoning", lazy="joined")


# Helpful indexes
Index(
    "idx_math_response_attempt_q",
    StudentExamResponseMathematicalReasoning.exam_attempt_id,
    StudentExamResponseMathematicalReasoning.q_id
)
class StudentExamMathematicalReasoning(Base):
    __tablename__ = "student_exam_mathematical_reasoning"

    id = Column(Integer, primary_key=True, index=True)

    # 🔑 INTERNAL student FK (students.id)"
    student_id = Column(
        Integer,
        ForeignKey("students.id", ondelete="CASCADE"),
        nullable=False,
        index=True
    )

    # 🔑 Exam definition FK (exams.id)
    exam_id = Column(
        Integer,
        ForeignKey("exams.id", ondelete="RESTRICT"),
        nullable=False
    )

    # 🕒 Attempt lifecycle
    started_at = Column(
        DateTime(timezone=True),
        nullable=False,
        default=lambda: datetime.now(timezone.utc)
    )

    completed_at = Column(
        DateTime(timezone=True),
        nullable=True
    )

    # --------------------------------------------------
    # OPTIONAL (SAFE) relationships
    # --------------------------------------------------
    student = relationship(
        "Student",
        backref="math_exam_attempts",
        lazy="joined"
    )

    exam = relationship(
        "Exam",
        lazy="joined"
    )

    # ⚠️ DO NOT add relationship to results/responses yet
    # Query those tables explicitly to avoid mapper issues


# Helpful indexes (Postgres will love you for this)
Index(
    "idx_math_exam_active_attempt",
    StudentExamMathematicalReasoning.student_id,
    StudentExamMathematicalReasoning.completed_at
)
class UploadedImage(Base):
    __tablename__ = "uploaded_images"

    id = Column(Integer, primary_key=True)
    original_name = Column(String, unique=True, nullable=False)
    gcs_url = Column(Text, nullable=False)
    created_at = Column(DateTime(timezone=True), server_default=func.now())
 
class AdminOverallSelectiveReport(Base):
    __tablename__ = "admin_overall_selective_reports"

    id = Column(Integer, primary_key=True)
    student_id = Column(String, index=True)
    exam_date = Column(Date, index=True)

    overall_percent = Column(Float)
    readiness_band = Column(String)
    school_recommendation = Column(JSON)

    override_flag = Column(Boolean, default=False)
    override_message = Column(Text, nullable=True)

    components = Column(JSON)  # {reading: 78, maths: 65, ...}

    created_at = Column(DateTime, default=datetime.utcnow)

class AdminExamRawScore(Base):
    __tablename__ = "admin_exam_raw_scores"

    id = Column(Integer, primary_key=True, index=True)

    # internal student identifier (NO foreign key)
    student_id = Column(Integer, nullable=False)

    # exam attempt identifier (one row per attempt)
    exam_attempt_id = Column(Integer, nullable=False, unique=True)

    subject = Column(String, nullable=False)

    total_questions = Column(Integer, nullable=False)
    correct_answers = Column(Integer, nullable=False)
    wrong_answers = Column(Integer, nullable=False)

    accuracy_percent = Column(Float, nullable=False)

    created_at = Column(DateTime(timezone=True), server_default=func.now())

class AdminExamReport(Base):
    __tablename__ = "admin_exam_reports"

    id = Column(Integer, primary_key=True, index=True)

    # external student identifier (e.g. "Gem002")
    student_id = Column(String, nullable=False, index=True)

    # references exam attempt table (thinking / reading / etc)
    exam_attempt_id = Column(Integer, nullable=False, index=True)

    # subject / exam type
    exam_type = Column(String, nullable=False, index=True)

    overall_score = Column(Float, nullable=True)

    readiness_band = Column(String, nullable=False)

    school_guidance_level = Column(String, nullable=False)

    summary_notes = Column(String, nullable=True)

    created_at = Column(
        DateTime(timezone=True),
        server_default=func.now(),
        nullable=False
    )

    __table_args__ = (
        UniqueConstraint(
            "exam_attempt_id",
            "exam_type",
            name="uq_admin_exam_reports_attempt_subject"
        ),
    )

class AdminExamSectionResult(Base):
    __tablename__ = "admin_exam_section_results"

    id = Column(Integer, primary_key=True, index=True)

    # references AdminExamReport.id
    admin_report_id = Column(Integer, nullable=False, index=True)

    section_name = Column(String, nullable=False)
    # "reading", "maths", "thinking", "writing"

    raw_score = Column(Float, nullable=True)
    # percent or marks

    performance_band = Column(String, nullable=False)
    # A / B / C / etc

    strengths_summary = Column(String, nullable=True)
    improvement_summary = Column(String, nullable=True)

    created_at = Column(
        DateTime(timezone=True),
        server_default=func.now(),
        nullable=False
    )
class AdminReadinessRuleApplied(Base):
    __tablename__ = "admin_readiness_rules_applied"

    id = Column(Integer, primary_key=True, index=True)

    # references AdminExamReport.id
    admin_report_id = Column(Integer, nullable=False, index=True)

    rule_code = Column(String, nullable=False)
    # e.g. "WRITING_MINIMUM"

    rule_description = Column(String, nullable=False)
    # human-readable explanation

    rule_result = Column(String, nullable=False)
    # "passed" or "failed"

    created_at = Column(
        DateTime(timezone=True),
        server_default=func.now(),
        nullable=False
    )


class StudentExamResultsMathematicalReasoning(Base):
    __tablename__ = "student_exam_results_mathematical_reasoning"

    id = Column(Integer, primary_key=True, index=True)

    student_id = Column(String, ForeignKey("students.id"), nullable=False)
    exam_attempt_id = Column(
        Integer,
        ForeignKey("student_exams.id", ondelete="CASCADE"),
        nullable=False,
        unique=True
    )

    total_questions = Column(Integer, nullable=False)
    correct_answers = Column(Integer, nullable=False)
    wrong_answers = Column(Integer, nullable=False)
    accuracy_percent = Column(Numeric(5, 2), nullable=False)

    created_at = Column(
        DateTime(timezone=True),
        server_default=func.now(),
        nullable=False
    )

    # relationships (optional but recommended)
    attempt = relationship("StudentExam")
    student = relationship("Student")
class StudentExamThinkingSkills(Base):
    __tablename__ = "student_exam_thinking_skills"

    id = Column(Integer, primary_key=True, index=True)

    student_id = Column(Integer, ForeignKey("students.id"), nullable=False)
    exam_id = Column(Integer, ForeignKey("exams.id"), nullable=False)

    started_at = Column(
        DateTime(timezone=True),
        default=lambda: datetime.now(timezone.utc),
        nullable=False
    )

    completed_at = Column(
        DateTime(timezone=True),
        nullable=True
    )

    duration_minutes = Column(Integer, nullable=False)

    # -----------------------------
    # Relationships (optional but recommended)
    # -----------------------------
    student = relationship("Student")
    exam = relationship("Exam")
    responses = relationship(
        "StudentExamResponseThinkingSkills",
        back_populates="attempt",
        cascade="all, delete-orphan"
    )
class StudentExamNaplanLanguageConventions(Base):
    __tablename__ = "student_exam_naplan_language_conventions"

    id = Column(Integer, primary_key=True, index=True)

    # -----------------------------
    # Foreign keys
    # -----------------------------
    student_id = Column(
        String,
        ForeignKey("students.id"),
        nullable=False
    )

    exam_id = Column(
        Integer,
        ForeignKey("exam_naplan_language_conventions.id"),
        nullable=False
    )

    # -----------------------------
    # Timing
    # -----------------------------
    started_at = Column(
        DateTime(timezone=True),
        default=lambda: datetime.now(timezone.utc),
        nullable=False
    )

    completed_at = Column(
        DateTime(timezone=True),
        nullable=True
    )

    duration_minutes = Column(
        Integer,
        nullable=False
    )

    # -----------------------------
    # Relationships
    # -----------------------------
    student = relationship("Student")

    exam = relationship("ExamNaplanLanguageConventions")

    responses = relationship(
        "StudentExamResponseNaplanLanguageConventions",
        back_populates="attempt",
        cascade="all, delete-orphan"
    )

class StudentExamNaplanNumeracy(Base):
    __tablename__ = "student_exam_naplan_numeracy"

    id = Column(Integer, primary_key=True, index=True)

    # -----------------------------
    # Foreign keys
    # -----------------------------
    student_id = Column(
        String,                    
        ForeignKey("students.id"),
        nullable=False
    )

    exam_id = Column(
        Integer,
        ForeignKey("exam_naplan_numeracy.id"),
        nullable=False
    )

    # -----------------------------
    # Timing
    # -----------------------------
    started_at = Column(
        DateTime(timezone=True),
        default=lambda: datetime.now(timezone.utc),
        nullable=False
    )

    completed_at = Column(
        DateTime(timezone=True),
        nullable=True
    )

    duration_minutes = Column(
        Integer,
        nullable=False
    )

    # -----------------------------
    # Relationships
    # -----------------------------
    student = relationship("Student")

    exam = relationship("ExamNaplanNumeracy")

    responses = relationship(
        "StudentExamResponseNaplanNumeracy",
        back_populates="attempt",
        cascade="all, delete-orphan"
    )



class StudentExamResponseThinkingSkills(Base):
    __tablename__ = "student_exam_response_thinking_skills"

    id = Column(Integer, primary_key=True, index=True)

    student_id = Column(Integer, ForeignKey("students.id"), nullable=False)
    exam_id = Column(Integer, ForeignKey("exams.id"), nullable=False)
    exam_attempt_id = Column(
        Integer,
        ForeignKey("student_exam_thinking_skills.id"),
        nullable=False
    )

    q_id = Column(Integer, nullable=False)      # stable per-attempt index
    topic = Column(String, nullable=True)

    selected_option = Column(String, nullable=True)
    correct_option = Column(String, nullable=True)

    is_correct = Column(Boolean, nullable=True) # NULL = not attempted

    # -----------------------------
    # Relationships
    # -----------------------------
    attempt = relationship(
        "StudentExamThinkingSkills",
        back_populates="responses"
    )
    student = relationship("Student")
    exam = relationship("Exam")
class StudentExamResponseNaplanLanguageConventions(Base):
    __tablename__ = "student_exam_response_naplan_language_conventions"

    id = Column(Integer, primary_key=True, index=True)

    student_id = Column(
        String,
        ForeignKey("students.id"),
        nullable=False
    )

    exam_id = Column(
        Integer,
        ForeignKey("exam_naplan_language_conventions.id"),
        nullable=False
    )

    exam_attempt_id = Column(
        Integer,
        ForeignKey("student_exam_naplan_language_conventions.id"),
        nullable=False
    )

    q_id = Column(Integer, nullable=False)
    topic = Column(String, nullable=True)

    selected_option = Column(String, nullable=True)
    correct_option = Column(String, nullable=True)

    is_correct = Column(Boolean, nullable=True)

    # -----------------------------
    # Relationships
    # -----------------------------
    attempt = relationship(
        "StudentExamNaplanLanguageConventions",
        back_populates="responses"
    )

    student = relationship("Student")

    exam = relationship("ExamNaplanLanguageConventions")

class StudentExamResponseNaplanNumeracy(Base):
    __tablename__ = "student_exam_response_naplan_numeracy"

    id = Column(Integer, primary_key=True, index=True)

    student_id = Column(
        String,
        ForeignKey("students.id"),
        nullable=False
    )


    exam_id = Column(
        Integer,
        ForeignKey("exam_naplan_numeracy.id"),
        nullable=False
    )

    exam_attempt_id = Column(
        Integer,
        ForeignKey("student_exam_naplan_numeracy.id"),
        nullable=False
    )

    q_id = Column(Integer, nullable=False)
    topic = Column(String, nullable=True)

    selected_option = Column(String, nullable=True)
    correct_option = Column(String, nullable=True)

    is_correct = Column(Boolean, nullable=True)

    # -----------------------------
    # Relationships
    # -----------------------------
    attempt = relationship(
        "StudentExamNaplanNumeracy",
        back_populates="responses"
    )
    student = relationship("Student")
    exam = relationship("ExamNaplanNumeracy")


 

class TopicConfigMathematicalReasoning(BaseModel):
    name: str
    ai: int
    db: int
    total: int

class TopicConfigMathematicalReasoningCreate(BaseModel):
    class_name: str
    subject: str
    difficulty: str
    num_topics: int
    topics: List[TopicConfigMathematicalReasoning]

class QuizMathematicalReasoningCreate(BaseModel):
    class_name: str
    subject: str
    difficulty: str
    num_topics: int
    topics: List[TopicConfigMathematicalReasoning]
class QuizMathematicalReasoning(Base):
    __tablename__ = "quiz_mathematical_reasoning"

    id = Column(Integer, primary_key=True, index=True)

    class_name = Column(String, nullable=False)
    subject = Column(String, nullable=False)  # always mathematical_reasoning
    difficulty = Column(String, nullable=False)

    num_topics = Column(Integer, nullable=False)
    

    # Stores topic configs as JSON
    topics = Column(JSON, nullable=False)
 
class QuizNaplanNumeracy(Base):
    __tablename__ = "quiz_naplan_numeracy"

    id = Column(Integer, primary_key=True, index=True)

    class_name = Column(String, nullable=False)
    subject = Column(String, nullable=False)  
    year = Column(Integer, nullable=False)
    difficulty = Column(String, nullable=False)

    num_topics = Column(Integer, nullable=False)
    total_questions = Column(Integer, nullable=False)

    # Stores topic configs as JSON
    topics = Column(JSON, nullable=False)

class QuizNaplanLanguageConventions(Base):
    __tablename__ = "quiz_naplan_Language_Conventions"

    id = Column(Integer, primary_key=True, index=True)

    class_name = Column(String, nullable=False)
    subject = Column(String, nullable=False)  
    year = Column(Integer, nullable=False)
    difficulty = Column(String, nullable=False)

    num_topics = Column(Integer, nullable=False)
    total_questions = Column(Integer, nullable=False)

    # Stores topic configs as JSON
    topics = Column(JSON, nullable=False)


class UpdateStudentRequest(BaseModel):
    student_id: str  # identifier (cannot be changed)

    name: Optional[str] = None
    parent_email: Optional[str] = None
    class_name: Optional[str] = None
    class_day: Optional[str] = None
    password: Optional[str] = None  # plain text (as per your system)
 
class StudentExamReportReading(Base):
    __tablename__ = "student_exam_report_reading"

    id = Column(Integer, primary_key=True, index=True)

    # 🔑 Identity
    student_id = Column(String, index=True, nullable=False)
    exam_id = Column(Integer, index=True, nullable=False)
    session_id = Column(Integer, index=True, nullable=False, unique=True)

    # 🔎 Question-level data
    topic = Column(String, index=True, nullable=False)
    question_id = Column(String, index=True, nullable=False)

    selected_answer = Column(String, nullable=True)
    correct_answer = Column(String, nullable=False)
    is_correct = Column(Boolean, nullable=False)

    # 🧾 Snapshot (optional but powerful)
    question_snapshot = Column(JSON, nullable=True)

    # 🕒 Metadata
    created_at = Column(DateTime(timezone=True), server_default=func.now())
 
class StudentExamResponseFoundational(Base):
     __tablename__ = "student_exam_response_foundational"
 
     id = Column(Integer, primary_key=True, index=True)
 
     student_id = Column(String, index=True, nullable=False)
     exam_id = Column(Integer, index=True, nullable=False)
     attempt_id = Column(Integer, index=True, nullable=False)
     topic = Column(String)

 
     section_name = Column(String, index=True, nullable=False)  # topic
     question_id = Column(String, index=True, nullable=False)
 
     selected_answer = Column(String, nullable=True)
     correct_answer = Column(String, nullable=False)
     is_correct = Column(Boolean, nullable=False)
 
     created_at = Column(DateTime(timezone=True), server_default=func.now())


class StartExamRequestFoundational(BaseModel):
    student_id: str
class StudentExamResultsThinkingSkills(Base):
    __tablename__ = "student_exam_results_thinking_skills"

    __table_args__ = (
        Index("idx_results_student", "student_id"),
    )

    id = Column(Integer, primary_key=True, index=True)

    student_id = Column(
        Text,
        ForeignKey("students.id"),
        nullable=False
    )

    exam_attempt_id = Column(
        Integer,
        ForeignKey("student_exams.id"),
        nullable=False,
        unique=True
    )

    total_questions = Column(Integer, nullable=False)
    correct_answers = Column(Integer, nullable=False)
    wrong_answers = Column(Integer, nullable=False)
    accuracy_percent = Column(Float, nullable=False)

    created_at = Column(
        DateTime(timezone=True),
        server_default=func.now()
    )
 
class ExamSubmissionRequest(BaseModel):
    session_id: int
    answers: Dict[int, str]

class StudentExamResponse(Base):
    __tablename__ = "student_exam_responses"

    id = Column(Integer, primary_key=True)

    student_id = Column(Integer, nullable=False)
    exam_id = Column(Integer, nullable=False)
    exam_attempt_id = Column(Integer, nullable=False)

    q_id = Column(Integer, nullable=False)
    topic = Column(Text, nullable=False)

    selected_option = Column(Text, nullable=True)
    correct_option = Column(Text, nullable=True)
    is_correct = Column(Boolean, nullable=False)

    created_at = Column(
        DateTime(timezone=True),
        server_default=func.now(),
        nullable=False
    )


class StudentExamAttemptThinkingSkills(Base):
    __tablename__ = "student_exam_attempts_thinkingskills"

    id = Column(Integer, primary_key=True)
    session_id = Column(Text, unique=True, nullable=False)
    student_id = Column(Text, nullable=False)
    exam_id = Column(Integer, nullable=False)

    started_at = Column(DateTime(timezone=True), nullable=False)
    submitted_at = Column(DateTime(timezone=True))
    is_completed = Column(Boolean, default=False)
    score = Column(Integer)

    created_at = Column(DateTime(timezone=True), server_default=func.now())


class StudentExamAnswerThinkingSkills(Base):
    __tablename__ = "student_exam_answers_thinkingskills"

    id = Column(Integer, primary_key=True)
    attempt_id = Column(
        Integer,
        ForeignKey(
            "student_exam_attempts_thinkingskills.id",
            ondelete="CASCADE"
        ),
        nullable=False
    )
    question_id = Column(Integer, nullable=False)
    selected_option = Column(Text, nullable=False)
    is_correct = Column(Boolean)

    created_at = Column(DateTime(timezone=True), server_default=func.now())



class StudentExamReading(Base):
    __tablename__ = "student_exams_reading"

    id = Column(Integer, primary_key=True, index=True)

    # Keep consistent with writing exam (supports IDs like "Gem002")
    student_id = Column(Text, nullable=False)

    exam_id = Column(Integer, nullable=False)

    started_at = Column(
        DateTime(timezone=True),
        server_default=func.now(),
        nullable=False
    )

    completed_at = Column(
        DateTime(timezone=True),
        nullable=True
    )

    finished = Column(Boolean, default=False, nullable=False)

    created_at = Column(
        DateTime(timezone=True),
        server_default=func.now()
    )
    report_json = Column(JSON, nullable=True)


class WritingGenerateSchema(BaseModel):
    class_name: str
    difficulty: str
 
class StudentExamWriting(Base):
    __tablename__ = "student_exam_writing"

    id = Column(Integer, primary_key=True, index=True)

    # External student identifier (e.g. "Gem002")
    student_id = Column(Text, nullable=False)

    # Generated writing exam ID
    exam_id = Column(Integer, nullable=False)

    exam_attempt_id = Column(
        Integer,
        ForeignKey("student_exam_thinking_skills.id"),
        nullable=False
    )

    # --------------------------------------------------
    # Exam lifecycle timestamps
    # --------------------------------------------------
    started_at = Column(
        DateTime(timezone=True),
        server_default=func.now(),
        nullable=False
    )

    completed_at = Column(
        DateTime(timezone=True),
        nullable=True
    )

    # --------------------------------------------------
    # Exam configuration
    # --------------------------------------------------
    duration_minutes = Column(
        Integer,
        default=40,
        nullable=False
    )

    # --------------------------------------------------
    # Student response
    # --------------------------------------------------
    answer_text = Column(
        Text,
        nullable=True
    )

    # --------------------------------------------------
    # AI Writing Evaluation (Summary fields)
    # --------------------------------------------------
    ai_score = Column(
        Integer,
        nullable=True,
        comment="AI-evaluated writing score (0–20)"
    )

    ai_strengths = Column(
        Text,
        nullable=True,
        comment="AI-identified writing strengths (summary)"
    )

    ai_improvements = Column(
        Text,
        nullable=True,
        comment="AI-identified areas for improvement (summary)"
    )

    # --------------------------------------------------
    # AI Writing Evaluation (Structured JSON)
    # --------------------------------------------------
    ai_evaluation_json = Column(
        JSON,
        nullable=True,
        comment="Full structured AI evaluation for frontend rendering"
    )

    # --------------------------------------------------
    # Record creation timestamp
    # --------------------------------------------------
    created_at = Column(
        DateTime(timezone=True),
        server_default=func.now(),
        nullable=False
    )

class WritingSubmitSchema(BaseModel):
    answer_text: str

class GeneratedExamWriting(Base):
    __tablename__ = "generated_exam_writing"

    id = Column(Integer, primary_key=True, index=True)

    class_name = Column(String, nullable=False)
    subject = Column(String, nullable=False)
    topic = Column(String, nullable=False)
    difficulty = Column(String, nullable=False)

    question_text = Column(Text, nullable=False)

    duration_minutes = Column(Integer, default=40)

    is_current = Column(Boolean, default=True)

    created_at = Column(DateTime(timezone=True), server_default=func.now())

class WritingQuestionBank(Base):
    __tablename__ = "writing_question_bank"

    id = Column(Integer, primary_key=True, index=True)

    # -----------------------------
    # FILTER / METADATA FIELDS
    # -----------------------------
    class_name = Column(String, nullable=False)
    subject = Column(String, nullable=False)
    topic = Column(String, nullable=False)
    difficulty = Column(String, nullable=False)

    # -----------------------------
    # QUESTION CONTENT (MATCH DB)
    # -----------------------------
    title = Column(String, nullable=True)

    question_text = Column(Text, nullable=False)
    question_prompt = Column(Text, nullable=False)

    statement = Column(Text, nullable=True)
    opening_sentence = Column(Text, nullable=True)

    guidelines = Column(Text, nullable=True)

    # -----------------------------
    # SOURCE / AUDIT
    # -----------------------------
    source_file = Column(String, nullable=True)
    created_at = Column(
        DateTime(timezone=True),
        server_default=func.now(),
        nullable=False
    ) 

class WritingQuizSchema(BaseModel):
    class_name: str
    subject: str
    topic: str
    difficulty: str
 
class GeneratedExamFoundational(Base):
    __tablename__ = "generated_exam_foundational"

    id = Column(Integer, primary_key=True, index=True)

    class_name = Column(String, nullable=False)
    subject = Column(String, nullable=False)

    total_questions = Column(Integer, nullable=False)

    # Total exam duration in minutes
    duration_minutes = Column(Integer, nullable=False)

    exam_json = Column(JSON, nullable=False)

    is_current = Column(Boolean, default=True)

    created_at = Column(DateTime(timezone=True), server_default=func.now())

class StudentsExamFoundational(Base):
    __tablename__ = "students_exam_foundational"

    id = Column(Integer, primary_key=True, index=True)

    # External student identifier (e.g. "Gem002")
    student_id = Column(String, nullable=False, index=True)

    # Reference to generated_exam_foundational.id (NOT a foreign key)
    exam_id = Column(Integer, nullable=False, index=True)

    started_at = Column(DateTime(timezone=True), nullable=False)
    completed_at = Column(DateTime(timezone=True), nullable=True)

    current_section_index = Column(Integer, nullable=False, default=0)

    answers_json = Column(JSONB, nullable=True)
    completion_reason = Column(String, nullable=True)

    created_at = Column(DateTime(timezone=True), server_default=func.now())
 
class EmptyRequest(BaseModel):
    pass

class QuizSetupWriting(Base):
    __tablename__ = "quiz_setup_writing"

    id = Column(Integer, primary_key=True, index=True)

    class_name = Column(String, nullable=False)
    subject = Column(String, nullable=False)

    topic = Column(String, nullable=False)
    difficulty = Column(String, nullable=False)

    created_at = Column(DateTime(timezone=True), server_default=func.now())

class QuizSetupFoundational(Base):
    __tablename__ = "quiz_setup_foundational"

    id = Column(Integer, primary_key=True, index=True)

    class_name = Column(String(50), nullable=False)
    subject = Column(String(50), nullable=False)

    # -------- Section 1 --------
    section1_name = Column(String(20), nullable=False)
    section1_topic = Column(String(100), nullable=False)   # ✅ ADDED
    section1_ai = Column(Integer, default=0)
    section1_db = Column(Integer, default=0)
    section1_total = Column(Integer, default=0)
    section1_time = Column(Integer, default=0)
    section1_intro = Column(Text, default="")

    # -------- Section 2 --------
    section2_name = Column(String(20), nullable=False)
    section2_topic = Column(String(100), nullable=False)   # ✅ ADDED
    section2_ai = Column(Integer, default=0)
    section2_db = Column(Integer, default=0)
    section2_total = Column(Integer, default=0)
    section2_time = Column(Integer, default=0)
    section2_intro = Column(Text, default="")

    # -------- Section 3 (optional) --------
    section3_name = Column(String(20), nullable=True)
    section3_topic = Column(String(100), nullable=True)    # ✅ ADDED
    section3_ai = Column(Integer, nullable=True)
    section3_db = Column(Integer, nullable=True)
    section3_total = Column(Integer, nullable=True)
    section3_time = Column(Integer, nullable=True)
    section3_intro = Column(Text, nullable=True)

class SectionSchema(BaseModel):
    name: str                # Easy / Medium / Hard
    topic: str               # NEW: selected topic for this section
    ai: int
    db: int
    total: int
    time: int
    intro: str

class QuizSetupFoundationalSchema(BaseModel):
    class_name: str
    subject: str
    sections: List[SectionSchema]
 
class ReadingExamRequest(BaseModel):
    class_name: str
    difficulty: str

class AddStudentExamModuleRequest(BaseModel):
    id: str                 # Backend-suggested ID (e.g. "1" or UUID)
    student_id: str         # "Gem001"
    name: str
    parent_email: EmailStr
    class_name: str
    class_day: str

class GeneratedExamReading(Base):
    __tablename__ = "generated_exams_reading"

    id = Column(Integer, primary_key=True, index=True)
    config_id = Column(Integer, ForeignKey("reading_exam_config.id"), nullable=True)
    class_name = Column(String, nullable=False)
    subject = Column(String, nullable=False)
    difficulty = Column(String, nullable=False)
    total_questions = Column(Integer, nullable=False)
    exam_json = Column(JSONB, nullable=False)  # reading_material + questions[]
    created_at = Column(DateTime(timezone=True), server_default=func.now())

class FinishExamRequestFoundational(BaseModel):
    student_id: str
    answers: Dict[str, str]
    reason: Optional[str] = "submitted"
 
class FinishExamRequest(BaseModel):
    student_id: str
    answers: Dict[str, str]
 
class StartExamRequest(BaseModel):
    student_id: str   # <-- MUST BE STRING
 
#when generate exam is pressed we create a row here
class StudentExam(Base):
    __tablename__ = "student_exams"

    id = Column(Integer, primary_key=True, index=True)

    # External student identifier (e.g. "Gem002")
    student_id = Column(String, nullable=False, index=True)

    # Reference to exams.id (NOT a foreign key)
    exam_id = Column(Integer, nullable=False, index=True)

    # Timezone-aware timestamps (UTC)
    started_at = Column(DateTime(timezone=True), nullable=False)
    completed_at = Column(DateTime(timezone=True), nullable=True)

    duration_minutes = Column(Integer, default=40, nullable=False)


class Exam_reading(Base):
    __tablename__ = "exams_reading"   # UPDATED TABLE NAME

    id = Column(Integer, primary_key=True, index=True)
    class_name = Column(String)
    subject = Column(String)
    difficulty = Column(String)
    topic = Column(String)
    total_questions = Column(Integer)
    reading_material = Column(JSON)  # dictionary of extracts/paragraphs



class QuestionReading(Base):
    __tablename__ = "questions_reading"

    id = Column(Integer, primary_key=True, index=True)

    # --------------------------------------------------
    # FILTERABLE / SEARCH FIELDS
    # --------------------------------------------------
    class_name = Column(String, index=True, nullable=False)
    subject = Column(String, index=True, nullable=False)
    difficulty = Column(String, index=True, nullable=False)
    topic = Column(String, index=True, nullable=False)

    # Explicit for admin sanity & validation
    total_questions = Column(Integer, nullable=False)

    # --------------------------------------------------
    # OPTIONAL EXAM LINK (ADMIN-CREATED EXAMS)
    # --------------------------------------------------
    exam_id = Column(
        Integer,
        ForeignKey("exams_reading.id", ondelete="SET NULL"),
        nullable=True
    )

    # --------------------------------------------------
    # COMPLETE SELF-DESCRIBING QUIZ PAYLOAD
    # --------------------------------------------------
    exam_bundle = Column(JSON, nullable=False)



class ReadingQuestionPreview(BaseModel):
    id: int
    question_text: str
    topic: str
    difficulty: str


class QuestionBankResponse(BaseModel):
    questions: List[ReadingQuestionPreview]



class QuestionReadingCreate(BaseModel):
    question_number: int
    question_text: str
    correct_answer: str

class ExamReadingCreate(BaseModel):
    class_name: str
    subject: str
    difficulty: str
    topic: str
    total_questions: int
    reading_material: Dict[str, str]
    answer_options: Dict[str, str] 
    questions: List[QuestionReadingCreate] 


class ReadingExamConfig(Base):
    __tablename__ = "reading_exam_config"

    id = Column(Integer, primary_key=True, index=True)
    class_name = Column(String, nullable=False)
    subject = Column(String, nullable=False)
    difficulty = Column(String, nullable=False)
    num_topics = Column(Integer, nullable=False)
    topics = Column(JSONB, nullable=False)
    created_at = Column(DateTime(timezone=True), server_default=func.now())

class ReadingTopicItem(BaseModel):
    name: str
    num_questions: int


class ReadingExamConfigCreate(BaseModel):
    class_name: str
    subject: str
    difficulty: str
    topics: List[ReadingTopicItem]

class Student(Base):
    __tablename__ = "students"

    id = Column(String, primary_key=True, index=True)   # internal PK
    student_id = Column(String, unique=True, nullable=False)  # e.g. "Gem002"

    password = Column(String, nullable=False)
    name = Column(String, nullable=False)
    parent_email = Column(String, nullable=False)

    class_name = Column(String, nullable=False)
    class_day = Column(String, nullable=True)
class TopicInput(BaseModel):
    name: str
    ai: int
    db: int
    total: int

class QuizCreate(BaseModel):
    class_name: str
    subject: str
    difficulty: str
    num_topics: int
    topics: List[TopicInput]

class Quiz(Base):
    __tablename__ = "quizzes"

    id = Column(Integer, primary_key=True, index=True)

    class_name = Column(String, nullable=False)
    subject = Column(String, nullable=False)
    difficulty = Column(String, nullable=False)

    num_topics = Column(Integer, nullable=False)
    topics = Column(JSON, nullable=False)

    created_at = Column(DateTime(timezone=True), server_default=func.now())

class Quiz_Naplan_Numeracy(Base):
    __tablename__ = "quizzes_naplan_numeracy"

    id = Column(Integer, primary_key=True, index=True)

    class_name = Column(String, nullable=False)
    subject = Column(String, nullable=False)
    difficulty = Column(String, nullable=False)

    num_topics = Column(Integer, nullable=False)
    topics = Column(JSON, nullable=False)

    created_at = Column(DateTime(timezone=True), server_default=func.now())

class QuestionNumeracyLC(Base):
    __tablename__ = "questions_numeracy_lc"

    id = Column(Integer, primary_key=True, index=True)
    question_type = Column(
        Integer,
        nullable=False,
        comment="Enum: 1=MCQ(single), 2=MCQ(multi), 3=Numeric, 4=Text"

    )

    class_name = Column(String(50), nullable=False)
 
    year = Column(Integer, nullable=False)


    # Must be either 'Numeracy' or 'Language Conventions'
    subject = Column(String(50), nullable=False)

    topic = Column(String(100), nullable=True)

    difficulty = Column(String(20), nullable=False)

    
    question_text = Column(Text, nullable=True)

    # Ordered visual blocks (text + resolved images)
    question_blocks = Column(JSON, nullable=True)

    # MCQ options: {"A": "...", "B": "...", "C": "...", "D": "..."}
    options = Column(JSON, nullable=True)

    correct_answer = Column(JSON, nullable=False)


    created_at = Column(
        DateTime(timezone=True),
        server_default=func.now()
    )
class Question(Base):
    __tablename__ = "questions"

    id = Column(Integer, primary_key=True, index=True)

    class_name = Column(String(50), nullable=False)
    subject = Column(String(50), nullable=False)

    # NEW — topic extracted from Word/GPT parser
    topic = Column(String(100), nullable=True)

    difficulty = Column(String(20), nullable=False)

    question_type = Column(String(50), nullable=False)
    # e.g. "domino_multi_image", "chart_mcq", "text_mcq", "two_diagram_mcq"

    question_text = Column(Text, nullable=True)
    # ✅ NEW — ordered visual blocks (text + images)
    question_blocks = Column(JSON, nullable=True)

    # Structured images (list of dicts)
    images = Column(JSON, nullable=True)

    # Structured MCQ options
    # e.g. {"A": "...", "B": "...", "C": "...", "D": "..."}
    options = Column(JSON, nullable=True)

    correct_answer = Column(String(5), nullable=False)

    created_at = Column(DateTime(timezone=True), server_default=func.now())

class QuestionNaplanReading(Base):
    __tablename__ = "questions_naplan_reading"

    id = Column(Integer, primary_key=True, index=True)
    # -----------------------------
    # Passage linkage (CRITICAL)
    # -----------------------------
    passage_id = Column(
        String(36),                 # UUID as string
        index=True,
        nullable=False,
        comment="Groups all questions belonging to the same reading passage"
    )

    # -----------------------------
    # Core metadata
    # -----------------------------
    class_name = Column(String(50), nullable=False)   # e.g. "year 3"
    subject = Column(String(50), nullable=False)      # "reading"
    year = Column(Integer, nullable=False)            # 3, 5
    topic = Column(String(255), nullable=False)
    difficulty = Column(String(20), nullable=False)   # easy | medium | hard

    # -----------------------------
    # Question type (ENUM INT)
    # -----------------------------
    question_type = Column(Integer, nullable=False)
    # e.g.
    # 1 = single_text_comprehension
    # 2 = comparative_analysis
    # 3 = multiple_texts
    # 4 = visual_literacy

    
    # -----------------------------
    # Render-safe bundle
    # -----------------------------
    exam_bundle = Column(JSON, nullable=False)
    """
    {
      "question_type": 1,
      "topic": "...",
      "text_title": "...",
      "reading_material": "...",
      "questions": [
        {
          "question_id": "Q1",
          "question_text": "...",
          "answer_options": { "A": "...", "B": "...", "C": "...", "D": "..." },
          "correct_answer": "B"
        }
      ]
    }
    """

    # -----------------------------
    # Audit
    # -----------------------------
    created_at = Column(
        DateTime(timezone=True),
        server_default=func.now()
    )
class StudentLogin(BaseModel):
    student_id: str  # e.g., Gem001, Gem002
    password: str

class LeaderboardEntry(BaseModel):
    student_name: str
    class_name: str
    class_day: str
    total_score: int
    total_questions: int
    submitted_at: datetime   # <-- change from str to datetime
    week_number: int

class FranchiseLocation(Base):
    __tablename__ = "franchiselocation"
                     
    id = Column(Integer, primary_key=True, autoincrement=True)
    country = Column(String(100), nullable=False)
    state = Column(String(100), nullable=False)

class AdminTerm(Base):
    __tablename__ = "admin_term"

    id = Column(Integer, primary_key=True, index=True)
    year = Column(String, nullable=False)
    term_start_date = Column(Date, nullable=False)
    term_end_date = Column(Date, nullable=False)

class TermDataSchema(BaseModel):
    year: str
    termStartDate: str  # will parse as date
    termEndDate: str

    
class WeekRequest(BaseModel):
    date: str  #

class LoginRequest(BaseModel):
    phone_number: str
    password: str

class ActivityPayload(BaseModel):
    instructions: str
    questions: List[Dict[str, Any]]                 # expects a JSON array of question objects
    score_logic: Optional[str] = None              # can be used as activity type
    admin_prompt: Optional[str] = None             # optional prompt from front end
    class_name: str
    class_day: Optional[str] = None
    week_number: Optional[int] = None
    
class AnswerPayload(BaseModel):
    question_index: int
    selected_option: str
    student_id: int        # ID of the student submitting the answer
    student_name: str      # Name of the student (needed for leaderboard)
    class_name: str
    class_day: str  

class AdminDate(Base):
    __tablename__ = "admin_date"

    date = Column(String, primary_key=True) 

class TermStartDateSchema(BaseModel):
    date: date

class AdminDateSchema(BaseModel):
    date: str  # we can use str since your column is String


    

class TermStartDateResponse(BaseModel):
    date: str  # YYYY-MM-DD

class OTPVerify(BaseModel):
    email: EmailStr  # required email field
    otp: str    
    
class OTP(Base):
    __tablename__ = "otp_store"
    id = Column(Integer, primary_key=True, index=True)
    phone_number = Column(String, unique=True, index=True)
    otp_code = Column(String)
    created_at = Column(DateTime, default=datetime.utcnow)

class QuizResult(Base):
    __tablename__ = "quiz_results"
    
    id = Column(Integer, primary_key=True)
    quiz_id = Column(Integer, nullable=False)
    student_id = Column(Integer, nullable=False)
    student_name = Column(String, nullable=False)
    class_name = Column(String, nullable=False)
    class_day = Column(String, nullable=True)
    week_number = Column(Integer, nullable=True)   # <-- added
    total_score = Column(Integer, nullable=False)
    total_questions = Column(Integer, nullable=False)
    submitted_at = Column(DateTime, default=datetime.utcnow)
    
class Activity(Base):
    __tablename__ = "activities"

    activity_id = Column(Integer, primary_key=True, index=True)
    instructions = Column(String)
    admin_prompt = Column(String)       # <-- new column for admin prompt
    
    questions = Column(JSON)
    score_logic = Column(String)
    class_name = Column(String)         # new column for class name
    class_day = Column(String)
    week_number = Column(Integer)  
    
class ActivityAttempt(Base):
    __tablename__ = "activity_attempts"
    attempt_id = Column(Integer, primary_key=True, index=True)
    student_id = Column(Integer, ForeignKey("users_temp.id"))
    quiz_id = Column(Integer, ForeignKey("student_quizzes.quiz_id"))
    score = Column(Integer)
    time_taken = Column(Integer)
    week_number = Column(Integer)
    term_number = Column(Integer)
    timestamp = Column(DateTime, default=datetime.utcnow)

class OTPRequest(BaseModel):
    email: str
    
class StudentQuiz(Base):
    __tablename__ = "student_quizzes"

    quiz_id = Column(Integer, primary_key=True)
    
    # Use MutableDict so in-place updates to quiz_json are tracked
    quiz_json = Column(MutableDict.as_mutable(JSON), nullable=False, default=dict)
    
    status = Column(String, default="pending")
    created_at = Column(DateTime, default=datetime.utcnow)
    started_at = Column(DateTime, nullable=True)
    completed_at = Column(DateTime, nullable=True)

    # Class metadata
    class_name = Column(String, nullable=False)
    class_day = Column(String, nullable=True)

"""    
class User(Base):
    __tablename__ = "users_temp"
    id = Column(Integer, primary_key=True, index=True)
    name = Column(String, unique=True, nullable=False)
    email = Column(String, unique=True)
    phone_number = Column(String, unique=True)
    password = Column(String, nullable=False)
    class_name = Column(String)
    class_day = Column(String)
    status = Column(String, default="active")
    created_at = Column(DateTime, default=datetime.utcnow)
"""
class SessionModel(Base):  # Handles authentication sessions
    __tablename__ = "sessions"

    id = Column(Integer, primary_key=True, index=True)
    session_token = Column(UUID(as_uuid=True), unique=True, index=True, default=uuid.uuid4)
    user_id = Column(Integer, ForeignKey("users.id", ondelete="CASCADE"))  # link to User
    is_authenticated = Column(Boolean, default=False)
    public_token = Column(UUID(as_uuid=True), unique=True, index=True, default=uuid.uuid4)

    # Establish relationship back to User
    user = relationship("User", back_populates="sessions")


class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, index=True)
    name = Column(String, nullable=False)
    email = Column(String, unique=True, nullable=False)
    phone_number = Column(String, nullable=False)
    class_name = Column(String, nullable=False)
    class_day = Column(String, nullable=False)  # <-- added
    password = Column(String, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)
    # Establish relationship with sessions
    sessions = relationship("SessionModel", back_populates="user", cascade="all, delete-orphan")

class Exam(Base):
    __tablename__ = "exams"

    id = Column(Integer, primary_key=True, index=True)

    # just a number, NO foreign key
    quiz_id = Column(Integer, nullable=True)

    class_name = Column(String, nullable=False)
    subject = Column(String, nullable=False)
    difficulty = Column(String, nullable=False)

    questions = Column(JSON, nullable=False)

    created_at = Column(DateTime(timezone=True), server_default=func.now())

class ExamNaplanNumeracy(Base):
    __tablename__ = "exam_naplan_numeracy"

    id = Column(Integer, primary_key=True, index=True)

    # just a number, NO foreign key
    quiz_id = Column(Integer, nullable=True)

    class_name = Column(String, nullable=False)
    subject = Column(String, nullable=False)
    difficulty = Column(String, nullable=False)

    questions = Column(JSON, nullable=False)

    created_at = Column(DateTime(timezone=True), server_default=func.now())

class ExamNaplanLanguageConventions(Base):
    __tablename__ = "exam_naplan_language_conventions"

    id = Column(Integer, primary_key=True, index=True)

    # just a number, NO foreign key
    quiz_id = Column(Integer, nullable=True)

    class_name = Column(String, nullable=False)
    subject = Column(String, nullable=False)
    difficulty = Column(String, nullable=False)

    questions = Column(JSON, nullable=False)

    created_at = Column(DateTime(timezone=True), server_default=func.now())




class StudentExamAnswer(Base):
    __tablename__ = "student_exam_answers"

    id = Column(Integer, primary_key=True, index=True)

    # Reference to student_exams.id (NOT a foreign key)
    student_exam_id = Column(Integer, nullable=False, index=True)

    question_id = Column(Integer, nullable=False)
    student_answer = Column(String, nullable=False)
    correct_answer = Column(String, nullable=False)
    is_correct = Column(Boolean, default=False)

    created_at = Column(DateTime(timezone=True), server_default=func.now())





try:
    Base.metadata.create_all(bind=engine)
    print("Tables created successfully")
except Exception as e:
    print("Error creating tables:", e)



# ---------------------------
# FastAPI App
# ---------------------------
app = FastAPI(title="Gem Kids Gamified Quiz API")
 
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://gamified-quiz-delta.vercel.app",
        "https://leader-board-viewer-gamified-quiz.vercel.app",
        "https://leaderboard.gemkidsacademy.com.au",
        "https://gamifiedquiz.gemkidsacademy.com.au",
        "https://exam.gemkidsacademy.com.au",
        "https://exam-module-pink.vercel.app"  # added origin
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
# ---------------------------
# In-memory OTP storage
# ---------------------------
otp_dict = {}
# ---------------------------
# Quiz Generation
# ---------------------------
 

 
def generate_quizzes():
    print("\n==============================")
    print("[SCHEDULER] Task Started: generate_quizzes()")
    print("==============================")

    db = SessionLocal()
    try:
        activities = db.query(Activity).all()
        print(f"[DEBUG] Number of activities fetched: {len(activities)}")
        if not activities:
            print("[DEBUG] No activities found in DB. Task ending.")
            return

        for idx, activity in enumerate(activities, 1):
            print(f"\n[DEBUG] Processing activity {idx}/{len(activities)}: {activity}")
            print("[DEBUG] Activity object dict:", activity.__dict__)

            # --- Extract activity values ---
            raw_prompt = getattr(activity, "admin_prompt", "") or ""
            class_name = getattr(activity, "class_name", "") or ""
            topic_name = getattr(activity, "instructions", "") or ""
            activity_type=getattr(activity,"score_logic","") or ""
            print(f"[DEBUG] Extracted -> class_name: '{class_name}', topic_name: '{topic_name}'")
            print(f"[DEBUG] Admin prompt:\n{raw_prompt}")
            row = db.execute(select(FranchiseLocation)).scalar_one_or_none()
            if row:
               country = row.country
               state = row.state
               print(f"Country: {country}, State: {state}")
            else:
               print("No row found in FranchiseLocation table.")
               country = "UnknownCountry"
               state = "UnknownState"

            # --- Compose strict system prompt (single triple quotes) ---
            system_prompt = '''
            You are an expert quiz-generating AI and a creative educator. Create a gamified quiz for {country} {state} {class_name} class students on the topic: {topic_name}. The activity type is: "{activity_type}".
            
            Follow these rules STRICTLY:
            
            1. Return ONLY a single valid JSON object. NO explanations, NO extra text.
            2. Use standard double quotes " only.
            3. JSON MUST follow this exact structure ONLY:
            {{
              "quiz_title": "Sample Quiz",
              "instructions": "Answer all questions carefully",
              "questions": [
                {{"category": "{topic_name}", "prompt": "Question 1 here", "options": ["A","B","C","D"], "answer": "A"}},
                {{"category": "{topic_name}", "prompt": "Question 2 here", "options": ["A","B","C","D"], "answer": "B"}},
                {{"category": "{topic_name}", "prompt": "Question 3 here", "options": ["A","B","C","D"], "answer": "C"}},
                {{"category": "{topic_name}", "prompt": "Question 4 here", "options": ["A","B","C","D"], "answer": "D"}},
                {{"category": "{topic_name}", "prompt": "Question 5 here", "options": ["A","B","C","D"], "answer": "A"}}
              ]
            }}
            
            4. Each 'prompt' MUST contain only ONE clear question. Never include two questions in a single prompt.
            5. Each question MUST have exactly 4 options tied ONLY to that question.
            6. The 'answer' MUST match exactly one of the 4 options for that specific question.
            7. ENFORCED ACTIVITY TYPE LOGIC: Apply the rules below depending on the activity type.
               - If activity type is "Mini quiz":
                   *Create a short multiple-choice quiz with a clear learning goal.*
               - If activity type is "Single challenge question":
                   *Create one challenging question focused on problem-solving or reasoning.*
               - If activity type is "Story + question":
                   *Provide a brief story followed by a related multiple-choice question.*
               - If activity type is "Riddle":
                   *Create a fun riddle with 4 options, the correct answer must be among them.*
               - If activity type is "Image-based puzzle":
                   *Describe a visual puzzle clearly in the prompt. Provide 4 options describing possible solutions. Answer must be one of the options.*
               - If activity type is "Logic puzzle":
                   *Present a reasoning problem requiring logical deduction. Include 4 options, one of which is correct.*
               - If activity type is "Word scramble":
                   *Present 5 scrambled words. Each prompt shows a scrambled word, provide 4 options (the correct word + 3 plausible distractors). The 'answer' must match the correct unscrambled word exactly.*
               - If activity type is "Mini path-choice scenario":
                   *Present a short scenario with a decision to make. Student selects the correct action from 4 options.*
               - If activity type is "Word maze":
                   *Create a vocabulary-based puzzle inspired by a 'maze' theme. 
                    Each question presents a clue, 
                    and the 4 options are possible words the student might 'choose' while navigating the maze. 
                    The answer MUST be exactly one of the 4 options. 
                    NO directional paths, NO grids, NO letters like A/B/C/D, NO maze drawings.*
            
            8. For ANY other activity type passed through {activity_type}:
                   *Use the name to guide theme and creativity ONLY. DO NOT modify structure or rules.*
            
            9. Questions MUST be fun, clear, engaging, and age-appropriate for Year 5.
            10. NEVER blend two questions together. NEVER ask "What is X? Is it Y?" inside the same prompt.
            11. Admin prompt/context for this quiz: {raw_prompt}
            '''.format(
                country=country,
                state=state,
                class_name=class_name,
                topic_name=topic_name,
                activity_type=activity_type,
                raw_prompt=raw_prompt
            )



            parsed_json = None
            try:
                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "system", "content": system_prompt}],
                    temperature=0.5
                )
                quiz_text = response.choices[0].message.content.strip()
                print(f"[DEBUG] Raw GPT response for class '{class_name}':\n{quiz_text}")

                # --- Direct JSON parsing (no regex) ---
                parsed_json = json.loads(quiz_text)
                print(f"[DEBUG] Successfully parsed JSON for class '{class_name}'")

            except json.JSONDecodeError as e_json:
                print(f"[ERROR] JSON parsing failed for class '{class_name}': {e_json}")
            except Exception as e_ai:
                print(f"[ERROR] AI call failed for class '{class_name}': {e_ai}")

            # ---- FALLBACK QUIZ ----
            if not parsed_json:
                print(f"[DEBUG] Using fallback quiz for class '{class_name}'")
                parsed_json = {
                    "quiz_title": f"Fallback Quiz for {class_name}",
                    "instructions": "This is a fallback quiz.",
                    "questions": [
                        {
                            "category": "Math",
                            "prompt": "What is 10 + 5?",
                            "options": ["12", "15", "20", "25"],
                            "answer": "15"
                        },
                        {
                            "category": "Science",
                            "prompt": "Which planet is closest to the Sun?",
                            "options": ["Earth", "Mercury", "Venus", "Mars"],
                            "answer": "Mercury"
                        },
                        {
                            "category": "English",
                            "prompt": "Plural of 'mouse'?","options": ["Mouses", "Mice", "Mouse", "Mices"],
                            "answer": "Mice"
                        }
                    ]
                }

            # ---- SAVE QUIZ TO DB ----
            try:
                class_quiz = StudentQuiz(
                    quiz_json=parsed_json,
                    status="pending",
                    created_at=datetime.now(timezone.utc),
                    class_name=class_name,
                    class_day=getattr(activity, "class_day", "")
                )
                db.add(class_quiz)
                print(f"[DEBUG] Quiz added to DB for class '{class_name}'")
            except Exception as e_db:
                print(f"[ERROR] Failed to save quiz for class '{class_name}': {e_db}")

        db.commit()
        print("[SCHEDULER] Successfully generated quizzes for all classes.")

    except Exception as e:
        print("[FATAL ERROR] Scheduler crashed:", e)
        db.rollback()

    finally:
        db.close()
        print("[DEBUG] Database session closed.")


# ---------------------------
# APScheduler Setup (weekly run)
# ---------------------------
scheduler = BackgroundScheduler()
scheduler.add_job(generate_quizzes, 'interval', weeks=1)

scheduler.start()

# ---------------------------
# Endpoints
# ---------------------------
def get_response_model(exam: str):
    if exam == "thinking_skills":
        return StudentExamResponseThinkingSkills
    if exam == "reading":
        return StudentExamReportReading
    if exam == "mathematical_reasoning":
        return StudentExamResponseMathematicalReasoning

    if exam == "writing":
        return StudentExamWriting
    raise HTTPException(status_code=400, detail="Invalid exam type")

def chunk_into_pages(paragraphs, per_page=18):
    pages = []
    for i in range(0, len(paragraphs), per_page):
        pages.append("\n\n".join(paragraphs[i:i+per_page]))
    return pages

def group_pages(pages, size=5):
    return [ "\n".join(pages[i:i+size]) for i in range(0, len(pages), size) ]
 


def serialize_blocks_for_gpt(blocks: list[dict]) -> str:
    """
    Converts ordered blocks into a deterministic text format
    that GPT can parse reliably.
    """
    lines = []

    for block in blocks:
        if block["type"] == "text":
            lines.append(block["content"])

        elif block["type"] == "image":
            # IMPORTANT: GPT must see image placeholders
            name = block.get("name") or block.get("src") or ""
            lines.append(f"[IMAGE: {name}]")

    return "\n\n".join(lines)

def serialize_blocks_for_gpt_numeracy_LC(blocks: list[dict]) -> str:
    """
    Flattens ordered Numeracy / Language Conventions blocks into a
    deterministic text representation for GPT parsing.

    Responsibilities:
    - Preserve ordering and wording
    - Normalize legacy metadata keys
    - Ensure question_type and year appear inside METADATA if present
    - Emit image placeholders explicitly
    """

    lines = []

    pending_question_type = None
    inside_metadata = False
    metadata_lines = []

    def normalize_key(line: str) -> str:
        # Normalize common legacy variants
        replacements = {
            r"^Question_Type\s*:": "question_type:",
            r"^QUESTION_TYPE\s*:": "question_type:",
            r"^Year\s*:": "year:",
            r"^YEAR\s*:": "year:",
            r"^CLASS\s*:": "class_name:",
            r"^SUBJECT\s*:": "subject:",
            r"^TOPIC\s*:": "topic:",
            r"^DIFFICULTY\s*:": "difficulty:",
        }

        for pattern, replacement in replacements.items():
            if re.match(pattern, line):
                return re.sub(pattern, replacement, line)

        return line

    for block in blocks:
        block_type = block.get("type")

        if block_type == "text":
            raw = block.get("content", "").strip()

            # Capture question_type if it appears outside METADATA
            if re.match(r"^question_type\s*:\s*\d+", raw, re.IGNORECASE):
                pending_question_type = normalize_key(raw)
                continue

            # Detect METADATA start
            if raw.upper() == "METADATA:":
                inside_metadata = True
                metadata_lines = []
                continue

            # Detect end of METADATA block
            if inside_metadata and raw.endswith(":") and raw.upper() not in {
                "METADATA:",
            }:
                # Inject pending question_type if needed
                if pending_question_type:
                    metadata_lines.insert(0, pending_question_type)
                    pending_question_type = None

                lines.append("METADATA:")
                lines.extend(metadata_lines)
                inside_metadata = False

            if inside_metadata:
                metadata_lines.append(normalize_key(raw))
            else:
                lines.append(raw)

        elif block_type == "image":
            image_ref = block.get("name") or block.get("src") or ""
            lines.append(f"[IMAGE: {image_ref}]")

    # If file ends while still inside METADATA
    if inside_metadata:
        if pending_question_type:
            metadata_lines.insert(0, pending_question_type)
        lines.append("METADATA:")
        lines.extend(metadata_lines)

    return "\n\n".join(line for line in lines if line)
def serialize_blocks_for_gpt_cloze(blocks: list[dict]) -> str:
    """
    Serializes CLOZE (question_type 5) blocks into a deterministic,
    lossless text format for GPT or backend parsing.

    This serializer is:
    - CLOZE-specific
    - Order-preserving
    - Non-normalizing
    - Non-destructive

    It intentionally preserves:
    - METADATA
    - QUESTION_TEXT
    - CLOZE
    - OPTIONS
    - CORRECT_ANSWER
    """

    lines = []

    for block in blocks:
        block_type = block.get("type")

        if block_type == "text":
            content = block.get("content", "").strip()
            if content:
                lines.append(content)

        elif block_type == "image":
            image_ref = block.get("name") or block.get("src") or ""
            lines.append(f"[IMAGE: {image_ref}]")

    serialized = "\n\n".join(lines)

    return serialized
 
async def parse_with_gpt(payload: dict, retries: int = 2):

    SYSTEM_PROMPT = (
        "You are a deterministic exam-question parser.\n\n"
    
        "INPUT CONTRACT:\n"
        "- You will receive text representing EXACTLY ONE exam\n"
        "- The exam contains AT MOST ONE question\n"
        "- Metadata (class, subject, topic, difficulty) may appear before the question\n\n"
    
        "CONTENT RULES:\n"
        "- Preserve wording exactly\n"
        "- Do NOT invent missing fields\n"
        "- Do NOT merge multiple questions\n\n"
    
        "EXTRACTION RULES:\n"
        "- Extract the following fields if present:\n"
        "  class_name\n"
        "  subject\n"
        "  topic\n"
        "  difficulty\n"
        "  options (A–D)\n"
        "  correct_answer\n\n"
    
        "OMISSION RULES:\n"
        "- Extract all fields that are explicitly present\n"
        "- If options or correct_answer are not explicitly present, set them to null\n"
        "- Partial extraction is allowed\n\n"
    
        "OUTPUT RULES:\n"
        "- Return ONLY valid JSON\n"
        "- No markdown, no commentary\n"
        "- Always return a single question object\n\n"
    
        "OUTPUT FORMAT:\n"
        '{ "question": { ... } }'
    )


    serialized = serialize_blocks_for_gpt(payload["blocks"])

    # 🔍 DEBUG: input size + preview
    print("[GPT DEBUG] Serialized length:", len(serialized))
    
    for attempt in range(retries + 1):
        print(f"\n[GPT DEBUG] Attempt {attempt}")

        completion = await client_save_questions.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0,
            response_format={
                "type": "json_schema",
                "json_schema": {
                    "name": "SingleQuestion",
                    "schema": {
                        "type": "object",
                        "properties": {
                            "question": {
                                "type": ["object", "null"],
                                "properties": QuestionSchema["properties"]                                
                            }
                        },
                        "required": ["question"]
                    }
                }
            },
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": serialized}
            ]
        )

        raw = completion.choices[0].message.content

        

        try:
            parsed = json.loads(raw)
        except json.JSONDecodeError as e:
            
            if attempt < retries:
                continue
            return {"question": None}

        question = parsed.get("question")        

        if question:
            return parsed

        
        if attempt < retries:
            continue

        return {"question": None}


async def parse_with_gpt_numeracy_lc(payload: dict, retries: int = 2):
    """
    Deterministic GPT parser for NAPLAN Numeracy / Language Conventions.

    GPT is responsible ONLY for semantic extraction:
    - class_name
    - question_type
    - year
    - subject
    - topic
    - difficulty
    - options (if applicable)
    - correct_answer

    Visual order, images, and question_blocks are handled
    entirely by the backend and MUST NOT be inferred by GPT.
    """
    SYSTEM_PROMPT = """
    You are a deterministic exam-question parser.
    
    Your task is to extract structured data from a single exam question.
    
    INPUT RULES:
    - You will receive plain text extracted from a Word document.
    - The text may span multiple paragraphs.
    - The text contains no images and no layout information.
    
    EXTRACTION RULES:
    Extract the following fields if present:
    
    - class_name
    - question_type (integer)
    - year (integer)
    - subject
    - topic
    - difficulty
    - options (ONLY for question_type 1 and 2)
    - correct_answer
    
    ANSWER FORMAT RULES:
    - If question_type = 1 (MCQ single):
      - correct_answer MUST be a single string
      - Example: "B"
    
    - If question_type = 2 (MCQ multi):
      - correct_answer MUST be an array of strings
      - Example: ["B", "D"]
    
    - If question_type = 3 (Numeric input):
      - correct_answer MUST be a number or numeric string
      - Example: 23
    
    - If question_type = 4 (Text input):
      - correct_answer MUST be a single string
      - Example: "stars"
    
    CONTENT RULES:
    - Preserve wording exactly as given.
    - Do NOT summarize, paraphrase, or invent content.
    - Do NOT infer missing information.
    
    OMISSION RULES:
    - If a required field is missing, OMIT the question entirely.
    - Do NOT emit placeholders or partial questions.
    - An empty questions array is valid.
    
    OUTPUT RULES:
    - Return ONLY valid JSON.
    - Output MUST match the provided JSON schema.
    - Do NOT include commentary, markdown, or explanations.
    """




    

    for attempt in range(retries + 1):
        serialized = serialize_blocks_for_gpt_numeracy_LC(payload["blocks"])

        completion = await client_save_questions.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0,
            top_p=1,
            response_format={
                "type": "json_schema",
                "json_schema": {
                    "name": "QuestionList",
                    "schema": {
                        "type": "object",
                        "properties": {
                            "questions": {
                                "type": "array",
                                "items": QuestionSchema,
                            }
                        },
                        "required": ["questions"],
                    },
                },
            },
            messages=[
                {
                    "role": "system",
                    "content": SYSTEM_PROMPT,
                },
                {
                    "role": "user",
                    "content": serialized,  # MUST be a string
                },
            ],
        )

        raw = completion.choices[0].message.content

        try:
            parsed = json.loads(raw)
        except json.JSONDecodeError as e:
            print(
                f"[GPT-NAPLAN] Invalid JSON "
                f"(attempt {attempt + 1}/{retries + 1}) | error={e}"
            )
            if attempt < retries:
                continue
            return {"questions": []}

        questions = parsed.get("questions", [])

        if questions:
            return parsed

        if attempt < retries:
            print(f"[GPT-NAPLAN] Empty result, retry {attempt + 1}")
            continue

        return parsed


# ai_engine.py (for example)
def generate_exam_questions(quiz, db):
    print("\n===================== GENERATE EXAM START =====================\n")
    print(f"Quiz ID       : {quiz.id}")
    print(f"Class         : {quiz.class_name}")
    print(f"Subject       : {quiz.subject}")
    print(f"Difficulty    : {quiz.difficulty}")
    print(f"Topics Config : {quiz.topics}")
    print("===============================================================\n")

    all_questions = []
    q_id = 1

    for topic in quiz.topics:
        print("\n===================== PROCESSING TOPIC =====================")

        topic_name = topic.get("name")
        ai_count = int(topic.get("ai", 0))
        db_count = int(topic.get("db", 0))

        print(f"Topic         : {topic_name}")
        print(f"Expected DB   : {db_count}")
        print(f"Expected AI   : {ai_count}")
        print("============================================================")

        # --------------------------------------------------
        # 1️⃣ DB PRE-FLIGHT CHECK
        # --------------------------------------------------
        available_db = (
            db.query(Question)
              .filter(func.lower(Question.topic) == topic_name.lower())
              .count()
        )

        print(f"[DB CHECK] Available questions: {available_db}")

        if available_db < db_count:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"Not enough DB questions for topic '{topic_name}'. "
                    f"Required: {db_count}, Available: {available_db}"
                )
            )

        # --------------------------------------------------
        # 2️⃣ FETCH DB QUESTIONS
        # --------------------------------------------------
        db_questions = (
            db.query(Question)
              .filter(func.lower(Question.topic) == topic_name.lower())
              .order_by(func.random())
              .limit(db_count)
              .all()
        )

        print(f"[DB FETCH] Retrieved {len(db_questions)} questions")

        for q in db_questions:
            # ---- Normalize blocks (FIX) ----
            blocks = q.question_blocks or []
            
            # Ensure at least text exists
            if not blocks and q.question_text:
                blocks = [
                    {"type": "text", "content": q.question_text}
                ]
            
            # 🔑 Inject missing image blocks
            existing_image_srcs = {
                b.get("src") for b in blocks if b.get("type") == "image"
            }
            
            for img in q.images or []:
                if img not in existing_image_srcs:
                    blocks.append({
                        "type": "image",
                        "src": img
                    })
            
            all_questions.append({
                "q_id": q_id,
                "topic": topic_name,
                "blocks": blocks,
                "options": q.options,
                "correct": q.correct_answer
            })

            q_id += 1


        # --------------------------------------------------
        # 3️⃣ AI QUESTION GENERATION (STRICT)
        # --------------------------------------------------
        if ai_count > 0:
            print(f"\n[AI GEN] Requesting {ai_count} questions for '{topic_name}'")

            location = db.query(FranchiseLocation).first()
            if not location:
                raise HTTPException(500, "No franchise location found")

            system_prompt = (
                "You are an expert exam generator.\n"
                f"Create exactly {ai_count} MCQs.\n\n"
                f"Class: {quiz.class_name}\n"
                f"Subject: {quiz.subject}\n"
                f"Topic: {topic_name}\n"
                f"Country: {location.country}\n"
                f"State: {location.state}\n\n"
                "Return ONLY a valid JSON array.\n"
                "Do NOT include explanations, markdown, or extra text.\n"
                "Do NOT wrap in ```.\n"
                "JSON format:\n"
                "[{\"question\":\"...\",\"options\":[\"A\",\"B\",\"C\",\"D\"],\"correct\":\"A\"}]"

            )

            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "system", "content": system_prompt}],
                temperature=0.4
            )

            raw_output = response.choices[0].message.content.strip()
            # --- DEBUG: inspect AI output ---
            print("\n[AI RAW OUTPUT]")
            print(raw_output)
            print("[END AI RAW OUTPUT]\n")
         
            try:
                generated = json.loads(raw_output)
            except json.JSONDecodeError:
                raise HTTPException(
                    status_code=500,
                    detail="AI returned invalid JSON"
                )


            print(f"[AI CHECK] Generated {len(generated)} questions")

            if len(generated) != ai_count:
                raise HTTPException(
                    status_code=500,
                    detail=(
                        f"AI generation failed for topic '{topic_name}'. "
                        f"Expected {ai_count}, got {len(generated)}."
                    )
                )

            for item in generated:
                all_questions.append({
                    "q_id": q_id,
                    "topic": topic_name,
            
                    # ✅ normalize AI questions to blocks
                    "blocks": [
                        {"type": "text", "content": item["question"]}
                    ],
            
                    "options": item["options"],
                    "correct": item["correct"]
                })
                q_id += 1

    print("\n==================== FINAL EXAM SUMMARY ====================")
    print(f"TOTAL QUESTIONS GENERATED: {len(all_questions)}")
    print("===========================================================\n")

    return all_questions




def parse_question_text_v3(text: str):
    print("\n====== parse_question_text_v3 START ======")
    print("RAW BLOCK:")
    print(text)
    print("-----------------------------------------")

    data = {}

    def extract(label):
        pattern = rf"{label}\s*:\s*(.+)"
        match = re.search(pattern, text, re.IGNORECASE)
        value = match.group(1).strip() if match else None
        print(f"🔎 Extract {label}: {value}")
        return value

    def clean(val):
        if not val:
            return val
        return val.replace('"', "").replace("“", "").replace("”", "").strip()

    # Header fields
    data["class_name"] = clean(extract("CLASS"))
    data["subject"] = clean(extract("SUBJECT"))
    data["topic"] = clean(extract("TOPIC"))
    data["difficulty"] = clean(extract("DIFFICULTY"))
    data["question_type"] = extract("TYPE") or "multi_image_diagram_mcq"

    # Images
    images_block = extract("IMAGES")
    if images_block:
        imgs = [i.strip() for i in images_block.split(",")]
        data["images"] = imgs
    else:
        data["images"] = []

    print("🖼 IMAGES:", data["images"])

    # Question text
    q_text = re.search(
        r"QUESTION_TEXT:\s*(.+?)(OPTIONS:|$)",
        text,
        flags=re.IGNORECASE | re.DOTALL
    )
    if not q_text:
        raise Exception("QUESTION_TEXT missing.")

    data["question_text"] = q_text.group(1).strip()
    print("📝 QUESTION_TEXT:", data["question_text"])

    # Options
    options_match = re.search(
        r"OPTIONS:\s*(.+?)CORRECT_ANSWER:",
        text,
        flags=re.IGNORECASE | re.DOTALL
    )
    if not options_match:
        raise Exception("OPTIONS missing.")

    raw_options = options_match.group(1).strip()
    print("📄 OPTIONS BLOCK:\n", raw_options)

    options = {}
    for line in raw_options.split("\n"):
        line = line.strip()
        m = re.match(r"([A-Z])\.?\s*(.+)", line)
        if m:
            key, value = m.group(1), m.group(2)
            options[key] = value

    data["options"] = options
    print("✅ Parsed OPTIONS:", options)

    # Correct answer
    correct = extract("CORRECT_ANSWER")
    data["correct_answer"] = clean(correct)

    print("✔ CORRECT_ANSWER:", data["correct_answer"])
    print("====== parse_question_text_v3 END ======\n")

    return data
def normalize_doc_text(text: str) -> str:
    return (
        text
        .replace("\xa0", " ")      # non-breaking spaces (Word’s favorite)
        .replace("\u200b", "")     # zero-width spaces
        .replace("\t", " ")        # tabs
        .replace("\r\n", "\n")
        .replace("\r", "\n")
    )

def extract_text_from_docx(file_bytes: bytes) -> str:
    # Wrap bytes in a file-like object
    file_obj = BytesIO(file_bytes)

    result = mammoth.extract_raw_text(file_obj)
    return result.value
 
def extract_comparative_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX exam paper while preserving
    structural boundaries (paragraphs, headings, sections).

    This helper is INTENTIONALLY separate from extract_text_from_docx
    to avoid impacting other endpoints.
    """
    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(file_bytes))
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()

        # Preserve paragraph breaks explicitly
        if text:
            lines.append(text)
        else:
            # Keep blank lines to preserve structure
            lines.append("")

    # Join with newline, then normalize excessive blanks
    full_text = "\n".join(lines)

    # Normalize line endings
    full_text = full_text.replace("\r\n", "\n").replace("\r", "\n")

    # Collapse 3+ blank lines → max 2
    import re
    full_text = re.sub(r"\n{3,}", "\n\n", full_text)

    return full_text

def parse_exam_with_openai(extracted_text: str, question_type: str):

    """
    Extract exam data from a Word document.
    This function ONLY extracts data.
    It does NOT invent structure or fill missing information.
    """

    # -------------------------------
    # Task-specific rules
    # -------------------------------
    if question_type == "gapped_text":
        task_rules = """
GAPPED TEXT RULES:
- This is a GAPPED TEXT exam.
- Questions represent gaps, not full sentences.
- Each question MUST be extracted as:
  {
    "question_number": <gap_number>,
    "question_text": "Gap <gap_number> _____",
    "correct_answer": "<LETTER>"
  }
- DO NOT invent question wording.
- question_number MUST be the gap number (1–N), not inferred from position.
- DO NOT include passage text inside question_text.
- Extract correct answers ONLY from explicit CORRECT_ANSWER fields.
- Do NOT infer or guess missing answers.
"""
    else:
        task_rules = """
STANDARD READING RULES:
- Every question begins with a number followed by a period (e.g., "1. Question text").
- Extract the number as question_number (INT).
- The remaining text MUST be question_text.
- NEVER keep the number inside question_text.
"""

    # -------------------------------
    # Prompt
    # -------------------------------
    prompt = f"""
You will receive a Reading Comprehension exam document.

QUESTION TYPE: {question_type}

You MUST extract exam data into the EXACT JSON structure below.

{task_rules}

OUTPUT FORMAT (STRICT):
{{
  "exams": [
    {{
      "class_name": "",
      "subject": "",
      "difficulty": "",
      "topic": "",
      "total_questions": 0,

      "reading_material": {{}},

      "answer_options": {{}},

      "questions": [
        {{
          "question_number": 1,
          "question_text": "",
          "correct_answer": "A"
        }}
      ]
    }}
  ]
}}

GLOBAL RULES:
- DO NOT add or remove fields.
- DO NOT rename keys.
- DO NOT include explanations or comments.
- Return ONLY valid JSON.
- JSON MUST start with "{{" and end with "}}".
- If something is unclear, leave fields empty — NEVER guess.

INPUT TEXT:
{extracted_text}
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
    )

    content = response.choices[0].message.content.strip()

    # -------------------------------
    # Debug logging
    # -------------------------------
    print("======== RAW GPT OUTPUT ========")
    print(content[:5000])
    print("================================")

    # -------------------------------
    # Strict JSON extraction
    # -------------------------------
    start = content.find("{")
    end = content.rfind("}") + 1

    if start == -1 or end == -1:
        raise ValueError("No JSON object found in model output")

    json_text = content[start:end]

    try:
        parsed = json.loads(json_text)

        if "exams" not in parsed or not isinstance(parsed["exams"], list):
            raise ValueError("JSON missing 'exams' list")

        # IMPORTANT: Do not normalize here.
        # Normalization happens in the endpoint.
        return parsed["exams"]

    except Exception as e:
        raise ValueError(
            "OpenAI returned invalid JSON or schema mismatch:\n"
            + json_text
        ) from e


def save_exam_to_db(db: Session, exam_data: ExamReadingCreate):

    # Still create exam metadata record (optional)
    exam = Exam_reading(
        class_name=exam_data.class_name,
        subject=exam_data.subject,
        difficulty=exam_data.difficulty,
        topic=exam_data.topic,
        total_questions=exam_data.total_questions
    )
    db.add(exam)
    db.flush()

    # Build full exam bundle
    exam_bundle = {
        "reading_material": exam_data.reading_material,
        "answer_options": exam_data.answer_options,  # <-- SAVE THEM
        "questions": [
            {
                "question_number": q.question_number,
                "question_text": q.question_text,
                "correct_answer": q.correct_answer
            }
            for q in exam_data.questions
        ]
    }

    # Save inside questions_reading with FULL METADATA
    qr = Question_reading(
        exam_id=exam.id,
        class_name=exam_data.class_name,
        subject=exam_data.subject,
        difficulty=exam_data.difficulty,
        topic=exam_data.topic,
        total_questions=exam_data.total_questions,
        exam_bundle=exam_bundle
    )

    db.add(qr)
    db.commit()

    return exam 

def sanitize_filename(filename: str) -> str:
    """
    Convert filename to a URL-safe, GCS-safe format.
    """
    filename = filename.split("/")[-1].split("\\")[-1]  # strip folders
    filename = filename.strip()
    filename = filename.replace(" ", "_")
    filename = re.sub(r"[^a-zA-Z0-9._-]", "", filename)
    return filename
 
def upload_to_gcs(file_bytes: bytes, filename: str) -> str:
    """Upload a file to Google Cloud Storage and return the public URL."""

    print("\n================== GCS UPLOAD START ==================")
    print(f"[GCS] Incoming file: {filename}")
    print(f"[GCS] Bucket: {BUCKET_NAME}")

    if not BUCKET_NAME:
        raise Exception("BUCKET_NAME environment variable not set!")

    bucket = gcs_client.bucket(BUCKET_NAME)

    # ✅ SANITIZE filename
    clean_filename = sanitize_filename(filename)
    print(f"[GCS] Sanitized filename: {clean_filename}")

    # ✅ Generate safe unique object name
    unique_name = f"{uuid.uuid4()}_{clean_filename}"
    print(f"[GCS] Unique stored name: {unique_name}")

    blob = bucket.blob(unique_name)

    try:
        blob.upload_from_string(
            file_bytes,
            content_type="image/png"
        )

        # Uniform Bucket Access → direct public URL
        public_url = f"https://storage.googleapis.com/{BUCKET_NAME}/{unique_name}"

        print(f"[GCS] Upload successful → {public_url}")
        print("================== GCS UPLOAD END ==================\n")

        return public_url

    except Exception as e:
        print("[GCS ERROR] Upload failed:", str(e))
        print("================== GCS UPLOAD FAILED ==================\n")
        raise Exception(f"GCS upload failed: {str(e)}")

#api end points
def get_days_for_class(db: Session, class_name: str):
    results = (
        db.query(distinct(Student.class_day))
        .filter(Student.class_name == class_name)
        .order_by(Student.class_day)
        .all()
    )

    # [('Monday',), ('Wednesday',)] → ['Monday', 'Wednesday']
    return [row[0] for row in results if row[0]]
    # [('Monday',), ('Wednesday',)] → ['Monday', 'Wedne

def normalize_questions_exam_review(raw_questions):
    normalized = []

    for q in raw_questions or []:
        fixed = dict(q)
        opts = fixed.get("options")

        # ✅ Preserve options as a DICT (A, B, C, D)
        if isinstance(opts, dict):
            fixed["options"] = opts

        # Defensive fallback (rare, but safe)
        elif isinstance(opts, list):
            fixed["options"] = {
                chr(65 + i): {
                    "type": "text",
                    "content": v
                }
                for i, v in enumerate(opts)
            }

        else:
            fixed["options"] = {}

        normalized.append(fixed)

    return normalized

def normalize_mr_questions_exam_review(raw_questions):
    normalized = []

    print("🧪 normalize_mr_questions_exam_review CALLED")
    print(f"🧾 raw_questions count: {len(raw_questions or [])}")

    for q in raw_questions or []:
        print("--------------------------------------------")
        print(f"🔍 Processing q_id={q.get('q_id')}")
        print("🗂️ Available keys:", list(q.keys()))

        # ✅ FIX BLOCKS
        blocks = q.get("question_blocks", [])
        print(f"🧱 question_blocks count: {len(blocks)}")

        # ✅ FIX OPTIONS
        opts = q.get("options", {})
        normalized_options = {}

        if isinstance(opts, list):
            for i, opt in enumerate(opts):
                key = chr(65 + i)  # A, B, C...
                normalized_options[key] = {
                    "type": "text",
                    "content": opt
                }

        elif isinstance(opts, dict):
            normalized_options = opts

        else:
            normalized_options = {}

        normalized.append({
            "q_id": q.get("q_id"),
            "blocks": blocks,
            "options": normalized_options
        })

    print(f"🧹 Normalized questions count: {len(normalized)}")
    return normalized
 
def normalize_question_blocks(raw_blocks):
    """
    Normalize question_blocks into a list of block dicts.
    Always returns a list.
    """

    # 1️⃣ None → no blocks
    if raw_blocks is None:
        return []

    # 2️⃣ Single string → text block
    if isinstance(raw_blocks, str):
        content = raw_blocks.strip()
        return (
            [{"type": "text", "content": content}]
            if content else []
        )

    # 3️⃣ Single dict → wrap in list
    if isinstance(raw_blocks, dict):
        return [raw_blocks]

    # 4️⃣ List → normalize each element
    if isinstance(raw_blocks, list):
        normalized = []

        for item in raw_blocks:
            if isinstance(item, str):
                content = item.strip()
                if content:
                    normalized.append({
                        "type": "text",
                        "content": content
                    })
            elif isinstance(item, dict):
                if "type" in item:
                    normalized.append(item)
            else:
                # optional debug
                # print(f"⚠️ Ignoring invalid block item: {item}")
                pass

        return normalized

    # 5️⃣ Truly unsupported
    raise ValueError(
        f"Unsupported question_blocks type: {type(raw_blocks)}"
    )


@app.post("/naplan/language-conventions/generate-exam")
def generate_naplan_language_conventions_exam(
    db: Session = Depends(get_db)
):
    print("\n=== START: Generate NAPLAN Language Conventions Exam ===")

    # 1. Load latest quiz config
    quiz = (
        db.query(QuizNaplanLanguageConventions)
        .order_by(QuizNaplanLanguageConventions.id.desc())
        .first()
    )

    if not quiz:
        print("❌ ERROR: No QuizNaplanLanguageConventions found")
        raise HTTPException(
            status_code=404,
            detail="NAPLAN Language Conventions quiz not found"
        )

    normalized_difficulty = quiz.difficulty.strip().lower()
    normalized_subject = "language conventions"

    print(
        f"✅ Quiz loaded | id={quiz.id}, "
        f"year={quiz.year}, difficulty='{quiz.difficulty}' "
        f"(normalized='{normalized_difficulty}')"
    )
    print(f"📘 Topics config: {quiz.topics}")
    print(f"📌 Expected total questions: {quiz.total_questions}")

    assembled_questions = []

    # 2. Iterate topic configs
    for idx, topic_cfg in enumerate(quiz.topics):
        print(f"\n--- Processing topic {idx + 1} ---")
        print(f"Raw topic config: {topic_cfg}")

        topic_name = topic_cfg.get("name")
        db_count = topic_cfg.get("db")

        if not topic_name or not db_count:
            print("⚠️ Skipping topic due to missing name or db count")
            continue

        print(f"➡️ Topic: {topic_name}")
        print(f"➡️ Required DB questions: {db_count}")

        # 3. Diagnostic: strict topic + subject + year existence
        topic_only_count = (
            db.query(QuestionNumeracyLC)
            .filter(
                QuestionNumeracyLC.topic == topic_name,
                func.lower(func.trim(QuestionNumeracyLC.subject)) == normalized_subject,
                QuestionNumeracyLC.year == quiz.year
            )
            .count()
        )

        print(
            f"🔍 Topic+subject+year match count: {topic_only_count}"
        )

        if topic_only_count == 0:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"No Language Conventions questions found for "
                    f"topic '{topic_name}' (Year {quiz.year})"
                )
            )

        # 4. Strict difficulty + subject + year query
        questions = (
            db.query(QuestionNumeracyLC)
            .filter(
                QuestionNumeracyLC.topic == topic_name,
                QuestionNumeracyLC.year == quiz.year,
                func.lower(func.trim(QuestionNumeracyLC.subject)) == normalized_subject,
                func.lower(func.trim(QuestionNumeracyLC.difficulty)) == normalized_difficulty
            )
            .all()
        )

        print(
            f"🔎 Questions after strict difficulty filter "
            f"('{normalized_difficulty}'): {len(questions)}"
        )

        # 5. Safe fallback (subject + year still enforced)
        if len(questions) < db_count:
            print(
                "⚠️ WARNING: Not enough questions after difficulty filter. "
                "Attempting subject+year fallback."
            )

            fallback_questions = (
                db.query(QuestionNumeracyLC)
                .filter(
                    QuestionNumeracyLC.topic == topic_name,
                    QuestionNumeracyLC.year == quiz.year,
                    func.lower(func.trim(QuestionNumeracyLC.subject)) == normalized_subject
                )
                .all()
            )

            print(
                f"🔎 Questions after fallback (subject+year): "
                f"{len(fallback_questions)}"
            )

            if len(fallback_questions) < db_count:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        f"Not enough Language Conventions questions for topic "
                        f"'{topic_name}' (Year {quiz.year}). "
                        f"Required={db_count}, Found={len(fallback_questions)}"
                    )
                )

            questions = fallback_questions

        # 6. Random sampling
        selected = random.sample(questions, db_count)
        print(f"🎯 Selected {len(selected)} questions for topic '{topic_name}'")

        for q in selected:
            assembled_questions.append({
                "id": q.id,
                "question_type": q.question_type,
                "topic": q.topic,
                "difficulty": q.difficulty,
                "question_text": q.question_text,
                "question_blocks": normalize_question_blocks(q.question_blocks),
                "options": q.options,
                "correct_answer": q.correct_answer,
            })

    # 7. Final validation
    print("\n=== FINAL CHECK ===")
    print(f"Total assembled questions: {len(assembled_questions)}")

    if len(assembled_questions) != quiz.total_questions:
        print(
            f"❌ ERROR: Question count mismatch | "
            f"Expected={quiz.total_questions}, "
            f"Got={len(assembled_questions)}"
        )
        raise HTTPException(
            status_code=400,
            detail="Generated question count does not match quiz total"
        )

    print("✅ Question count validated")

    # 8. Delete previous exams
    print("🧹 Resetting NAPLAN Language Conventions data...")

    # 1️⃣ Delete student responses
    deleted_responses = (
        db.query(StudentExamResponseNaplanLanguageConventions)
        .delete()
    )
    print(f"🗑️ Deleted {deleted_responses} student response record(s)")
    
    # 2️⃣ Delete student attempts
    deleted_attempts = (
        db.query(StudentExamNaplanLanguageConventions)
        .delete()
    )
    print(f"🗑️ Deleted {deleted_attempts} student attempt record(s)")
    
    # 3️⃣ Delete existing exams
    deleted_exams = (
        db.query(ExamNaplanLanguageConventions)
        .delete()
    )
    print(f"🗑️ Deleted {deleted_exams} previous exam record(s)")


    exam = ExamNaplanLanguageConventions(
        quiz_id=quiz.id,
        class_name="NAPLAN",
        subject="Language Conventions",
        difficulty=quiz.difficulty,
        questions=assembled_questions,
    )

    db.add(exam)
    db.commit()
    db.refresh(exam)

    print(
        f"🎉 SUCCESS: Language Conventions exam generated | "
        f"exam_id={exam.id}"
    )
    print("=== END: Generate NAPLAN Language Conventions Exam ===\n")

    
    return {
        "message": "NAPLAN Language Conventions exam generated successfully",
        "exam_id": exam.id,
        "total_questions": len(assembled_questions),
    }

def clean_question_blocks(blocks, correct_answer):
    # ✅ ALWAYS normalize first
    blocks = normalize_question_blocks(blocks)

    cleaned = []

    for block in blocks:
        # now block is GUARANTEED to be a dict
        if block.get("type") != "text":
            cleaned.append(block)
            continue

        content = (block.get("content") or "").strip()
        if not content:
            continue

        lower = content.lower()

        if lower.startswith(("question_type", "answer_type", "year:", "question_blocks")):
            continue

        if str(content).strip() == str(correct_answer).strip():
            continue

        cleaned.append({
            "type": "text",
            "content": content
        })

    return cleaned


AUTHORING_MARKERS = {
    "METADATA:",
    "QUESTION_TYPE:",
    "QUESTION_TEXT:",
    "OPTIONS:",
    "CORRECT_ANSWER:",
    "ANSWER_TYPE:",
    "CLASS:",
    "YEAR:",
    "SUBJECT:",
    "TOPIC:",
    "DIFFICULTY:",
}


def is_authoring_blob(text: str) -> bool:
    if not text:
        return False
    upper = text.upper()
    return any(marker in upper for marker in AUTHORING_MARKERS)


def normalize_question_blocks_backend(blocks):
    """
    Normalize question_blocks into render-safe blocks
    and strip authoring-only content.
    """

    if not blocks:
        return []

    normalized = []

    # ------------------------------
    # Case 1: single string
    # ------------------------------
    if isinstance(blocks, str):
        if not is_authoring_blob(blocks):
            normalized.append({
                "type": "text",
                "content": blocks.strip()
            })
        return normalized

    # ------------------------------
    # Case 2: single dict block
    # ------------------------------
    if isinstance(blocks, dict):
        if blocks.get("type") == "text":
            content = blocks.get("content", "")
            if is_authoring_blob(content):
                return []
            normalized.append({
                "type": "text",
                "content": content.strip()
            })
        else:
            normalized.append(blocks)
        return normalized

    # ------------------------------
    # Case 3: list of blocks
    # ------------------------------
    if isinstance(blocks, list):
        for b in blocks:
            if isinstance(b, str):
                if not is_authoring_blob(b):
                    normalized.append({
                        "type": "text",
                        "content": b.strip()
                    })

            elif isinstance(b, dict):
                block_type = b.get("type", "text")

                if block_type == "text":
                    content = b.get("content", "")
                    if is_authoring_blob(content):
                        continue
                    normalized.append({
                        "type": "text",
                        "content": content.strip()
                    })
                else:
                    # interactive / structured blocks pass through
                    normalized.append(b)

        return normalized

    return []

def build_question_blocks(q):
    """
    Build render-safe question_blocks for exam delivery.
    This function is used ONLY at exam generation time.
    """

    blocks = []

    # ==================================================
    # Include question_text for NON-inline types
    # (Type 5 and 7 are handled explicitly below)
    # ==================================================
    if q.question_text and q.question_type not in {2, 5, 7}:
        blocks.append({
            "type": "text",
            "content": q.question_text.strip()
        })

    # ==================================================
    # TYPE 2 — IMAGE SELECTION (MULTI)
    # ==================================================
    if q.question_type == 2:
        # 1️⃣ Instruction text (from raw blocks)
        if isinstance(q.question_blocks, list):
            for block in q.question_blocks:
                content = block.get("content", "").strip()
    
                # Stop once images start
                if content.startswith("IMAGES:"):
                    break
    
                # Skip metadata / noise
                if (
                    not content
                    or content.startswith("question_type")
                    or content.startswith("Year:")
                    or content.startswith("QUESTION_BLOCKS")
                ):
                    continue
    
                blocks.append({
                    "type": "text",
                    "content": content
                })
    
        # 2️⃣ Extract images
        images = []
        if isinstance(q.question_blocks, list):
            for block in q.question_blocks:
                content = block.get("content", "")
                if content.startswith("IMAGES:"):
                    images.append(
                        content.replace("IMAGES:", "").strip()
                    )
    
        # 3️⃣ Image selection block
        if images:
            blocks.append({
                "type": "image-selection",
                "images": images,
                "maxSelections": 2
            })
    
        return blocks

    # ==================================================
    # TYPE 7 — WORD_SELECTION (Grammar – Adverbs)
    # ==================================================
    if q.question_type == 7:
        # 1️⃣ Instruction text (REQUIRED)
        if q.question_text:
            blocks.append({
                "type": "text",
                "content": q.question_text.strip()
            })

        # 2️⃣ Interactive sentence
        if q.question_blocks and isinstance(q.question_blocks, dict):
            sentence = q.question_blocks.get("sentence")
            selectable_words = q.question_blocks.get("selectable_words")

            if sentence and selectable_words:
                blocks.append({
                    "type": "word-selection",
                    "sentence": sentence.strip(),
                    "selectable_words": selectable_words
                })

        return blocks

    # ==================================================
    # TYPE 5 — CLOZE_DROPDOWN (WITH INSTRUCTION)
    # ==================================================
    if q.question_type == 5:
        sentence = None
        instruction = None
    
        if isinstance(q.question_blocks, list):
            for block in q.question_blocks:
                content = block.get("content", "").strip()
    
                if content == "Choose the word to correctly complete this sentence.":
                    instruction = content
    
                if "{{dropdown}}" in content:
                    sentence = content
    
        # 1️⃣ Instruction text (optional but recommended)
        if instruction:
            blocks.append({
                "type": "text",
                "content": instruction
            })
    
        # 2️⃣ Inline dropdown sentence
        if sentence and q.options:
            blocks.append({
                "type": "cloze-dropdown",
                "sentence": sentence.strip(),
                "options": list(q.options.values())
            })
    
        return blocks

    # ==================================================
    # ALL OTHER TYPES (legacy-safe)
    # ==================================================
    if q.question_blocks:
        normalized = normalize_question_blocks_backend(q.question_blocks)
        blocks.extend(normalized)

    return blocks


@app.post("/naplan/numeracy/generate-exam")
def generate_naplan_numeracy_exam(
    db: Session = Depends(get_db)
):
    print("\n=== START: Generate NAPLAN Numeracy Exam ===")

    # 1. Load latest quiz config
    quiz = (
        db.query(QuizNaplanNumeracy)
        .order_by(QuizNaplanNumeracy.id.desc())
        .first()
    )

    if not quiz:
        print("❌ ERROR: No QuizNaplanNumeracy found")
        raise HTTPException(
            status_code=404,
            detail="NAPLAN Numeracy quiz not found"
        )

    normalized_difficulty = quiz.difficulty.strip().lower()
    normalized_subject = "numeracy"

    print(
        f"✅ Quiz loaded | id={quiz.id}, "
        f"year={quiz.year}, difficulty='{quiz.difficulty}' "
        f"(normalized='{normalized_difficulty}')"
    )
    print(f"📘 Topics config: {quiz.topics}")
    print(f"📌 Expected total questions: {quiz.total_questions}")

    assembled_questions = []

    # 2. Iterate topics
    for idx, topic_cfg in enumerate(quiz.topics):
        print(f"\n--- Processing topic {idx + 1} ---")
        print(f"Raw topic config: {topic_cfg}")

        topic_name = topic_cfg.get("name")
        db_count = topic_cfg.get("db")

        if not topic_name or not db_count:
            print("⚠️ Skipping topic due to missing name or db count")
            continue

        print(f"➡️ Topic: {topic_name}")
        print(f"➡️ Required DB questions: {db_count}")

        # 3. Diagnostic: strict topic + subject + year existence
        topic_only_count = (
            db.query(QuestionNumeracyLC)
            .filter(
                QuestionNumeracyLC.topic == topic_name,
                QuestionNumeracyLC.year == quiz.year,
                func.lower(func.trim(QuestionNumeracyLC.subject)) == normalized_subject
            )
            .count()
        )

        print(
            f"🔍 Topic+subject+year match count: {topic_only_count}"
        )

        if topic_only_count == 0:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"No Numeracy questions found for topic "
                    f"'{topic_name}' (Year {quiz.year})"
                )
            )

        # 4. Strict difficulty + subject + year query
        questions = (
            db.query(QuestionNumeracyLC)
            .filter(
                QuestionNumeracyLC.topic == topic_name,
                QuestionNumeracyLC.year == quiz.year,
                func.lower(func.trim(QuestionNumeracyLC.subject)) == normalized_subject,
                func.lower(func.trim(QuestionNumeracyLC.difficulty)) == normalized_difficulty
            )
            .all()
        )

        print(
            f"🔎 Questions after strict difficulty filter "
            f"('{normalized_difficulty}'): {len(questions)}"
        )

        # 5. Safe fallback (subject + year still enforced)
        if len(questions) < db_count:
            print(
                "⚠️ WARNING: Not enough questions after difficulty filter. "
                "Attempting subject+year fallback."
            )

            fallback_questions = (
                db.query(QuestionNumeracyLC)
                .filter(
                    QuestionNumeracyLC.topic == topic_name,
                    QuestionNumeracyLC.year == quiz.year,
                    func.lower(func.trim(QuestionNumeracyLC.subject)) == normalized_subject
                )
                .all()
            )

            print(
                f"🔎 Questions after fallback (subject+year): "
                f"{len(fallback_questions)}"
            )

            if len(fallback_questions) < db_count:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        f"Not enough Numeracy questions for topic "
                        f"'{topic_name}' (Year {quiz.year}). "
                        f"Required={db_count}, Found={len(fallback_questions)}"
                    )
                )

            questions = fallback_questions

        # 6. Random sampling
        selected = random.sample(questions, db_count)
        print(f"🎯 Selected {len(selected)} questions for topic '{topic_name}'")

        for q in selected:        
            assembled_questions.append({
                "id": q.id,
                "question_type": q.question_type,
                "topic": q.topic,
                "difficulty": q.difficulty,
                "question_blocks": build_question_blocks(q),
                "options": q.options,
                "correct_answer": q.correct_answer,
            })



    # 7. Final validation
    print("\n=== FINAL CHECK ===")
    print(f"Total assembled questions: {len(assembled_questions)}")

    if len(assembled_questions) != quiz.total_questions:
        print(
            f"❌ ERROR: Question count mismatch | "
            f"Expected={quiz.total_questions}, "
            f"Got={len(assembled_questions)}"
        )
        raise HTTPException(
            status_code=400,
            detail="Generated question count does not match quiz total"
        )

    print("✅ Question count validated")

    # 8. Delete previous exams
    db.query(StudentExamResponseNaplanNumeracy).delete()
    db.commit()
    db.query(StudentExamNaplanNumeracy).delete()
    db.commit()

    
    print("🧹 Deleting existing NAPLAN Numeracy exams...")

    deleted_count = (
        db.query(ExamNaplanNumeracy)
        .delete()
    )

    print(f"🗑️ Deleted {deleted_count} previous exam(s)")

    # 9. Persist exam
    print("💾 Saving exam to exam_naplan_numeracy table...")

    exam = ExamNaplanNumeracy(
        quiz_id=quiz.id,
        class_name="NAPLAN",
        subject="Numeracy",
        difficulty=quiz.difficulty,
        questions=assembled_questions,
    )

    db.add(exam)
    db.commit()
    db.refresh(exam)

    print(
        f"🎉 SUCCESS: Numeracy exam generated | "
        f"exam_id={exam.id}"
    )
    print("=== END: Generate NAPLAN Numeracy Exam ===\n")

    return {
        "message": "NAPLAN Numeracy exam generated successfully",
        "exam_id": exam.id,
        "total_questions": len(assembled_questions),
    }

@app.get("/api/admin/question-bank/naplan")
def get_naplan_question_bank(
    subject: str = Query(...),  # numeracy | language_conventions
    year: int = Query(...),
    db: Session = Depends(get_db),
):
    print("\n=== QUESTION BANK DEBUG START ===")
    print(f"[INPUT] subject (raw): {subject}")
    print(f"[INPUT] year: {year}")

    # -----------------------------
    # Normalize subject (frontend → DB)
    # -----------------------------
    SUBJECT_MAP = {
        "numeracy": "Numeracy",
        "language_conventions": "Language Conventions",
    }

    subject_key = subject.lower().strip()
    subject_db = SUBJECT_MAP.get(subject_key)

    print(f"[NORMALIZATION] subject_key: {subject_key}")
    print(f"[NORMALIZATION] subject_db: {subject_db}")

    if not subject_db:
        print("[ERROR] Invalid subject received")
        raise HTTPException(
            status_code=400,
            detail=f"Invalid subject: {subject}",
        )

    # -----------------------------
    # DEBUG STEP 1: Raw rows check
    # (CASE-INSENSITIVE class_name)
    # -----------------------------
    raw_rows = (
        db.query(QuestionNumeracyLC)
        .filter(
            func.lower(QuestionNumeracyLC.class_name) == "naplan",
            QuestionNumeracyLC.subject == subject_db,
            QuestionNumeracyLC.year == year,
        )
        .all()
    )

    print(f"[RAW QUERY] rows found: {len(raw_rows)}")

    if not raw_rows:
        print("[RAW QUERY] No rows matched filters:")
        print("  class_name (ci) = naplan")
        print(f"  subject         = {subject_db}")
        print(f"  year            = {year}")
        print("=== QUESTION BANK DEBUG END (NO DATA) ===\n")
        return []

    for i, row in enumerate(raw_rows[:5]):
        print(
            f"[RAW ROW {i}] "
            f"id={row.id}, "
            f"difficulty={row.difficulty}, "
            f"topic={row.topic}"
        )

    # -----------------------------
    # DEBUG STEP 2: Aggregated query
    # (NORMALIZE difficulty)
    # -----------------------------
    results = (
        db.query(
            func.lower(QuestionNumeracyLC.difficulty).label("difficulty"),
            QuestionNumeracyLC.topic,
            func.count(QuestionNumeracyLC.id).label("total_questions"),
        )
        .filter(
            func.lower(QuestionNumeracyLC.class_name) == "naplan",
            QuestionNumeracyLC.subject == subject_db,
            QuestionNumeracyLC.year == year,
        )
        .group_by(
            func.lower(QuestionNumeracyLC.difficulty),
            QuestionNumeracyLC.topic,
        )
        .order_by(
            func.lower(QuestionNumeracyLC.difficulty),
            QuestionNumeracyLC.topic,
        )
        .all()
    )

    print(f"[AGG QUERY] grouped rows returned: {len(results)}")

    for i, r in enumerate(results):
        print(
            f"[AGG ROW {i}] "
            f"difficulty={r.difficulty}, "
            f"topic={r.topic}, "
            f"total={r.total_questions}"
        )

    print("=== QUESTION BANK DEBUG END (SUCCESS) ===\n")

    return [
        {
            "difficulty": r.difficulty,   # now "easy", "medium", etc.
            "topic": r.topic,
            "total_questions": r.total_questions,
        }
        for r in results
    ]




@app.get("/api/exams/review-reading")
def review_reading_exam(
    session_id: int = Query(..., description="Reading exam session ID"),
    db: Session = Depends(get_db)
):
    print("\n================ REVIEW READING EXAM ================")
    print("🆔 session_id:", session_id)

    # --------------------------------------------------
    # 1️⃣ Load finished session
    # --------------------------------------------------
    session = (
        db.query(StudentExamReading)
        .filter(StudentExamReading.id == session_id)
        .first()
    )

    if not session:
        print("❌ Session not found")
        raise HTTPException(status_code=404, detail="Session not found")

    if not session.finished:
        print("❌ Session not finished yet")
        raise HTTPException(status_code=400, detail="Exam not finished yet")

    print("✅ Session loaded:", {
        "session_id": session.id,
        "student_id": session.student_id,
        "exam_id": session.exam_id,
        "completed_at": session.completed_at
    })

    # --------------------------------------------------
    # 2️⃣ Load exam content
    # --------------------------------------------------
    exam = (
        db.query(GeneratedExamReading)
        .filter(GeneratedExamReading.id == session.exam_id)
        .first()
    )

    if not exam or not exam.exam_json:
        print("❌ Exam content missing")
        raise HTTPException(status_code=500, detail="Exam content missing")

    # --------------------------------------------------
    # 3️⃣ Normalize exam_json (CRITICAL FIX)
    # --------------------------------------------------
    exam_json = exam.exam_json

    if isinstance(exam_json, str):
        try:
            exam_json = json.loads(exam_json)
        except json.JSONDecodeError as e:
            print("❌ exam_json is invalid JSON:", e)
            raise HTTPException(
                status_code=500,
                detail="Invalid exam JSON structure"
            )

    sections = exam_json.get("sections", [])
    print(f"📘 Sections loaded: {len(sections)}")

    # --------------------------------------------------
    # 4️⃣ Load student answers for this session
    # --------------------------------------------------
    reports = (
        db.query(StudentExamReportReading)
        .filter(StudentExamReportReading.session_id == session.id)
        .all()
    )

    print(f"📝 Answer rows found: {len(reports)}")
    report_map = {r.question_id: r for r in reports}

    # --------------------------------------------------
    # 5️⃣ Build review payload
    # --------------------------------------------------
    review_questions = []

    for section_idx, section in enumerate(sections):
        topic = section.get("topic", "Other")
        passage_style = section.get("passage_style", "informational")
        reading_material = section.get("reading_material")

        questions = section.get("questions", [])
        print(f"🧱 Section {section_idx + 1}: {topic} | questions: {len(questions)}")

        for q in questions:
            qid = q.get("question_id")
            r = report_map.get(qid)

            # Resolve answer options correctly
            options_scope = section.get("options_scope", "per_question")
            if options_scope == "shared":
                answer_options = section.get("answer_options", {})
            else:
                answer_options = q.get("answer_options", {})

            answer_options = answer_options or {}

            review_questions.append({
                "question_id": qid,
                "question_number": q.get("question_number"),
                "question_text": q.get("question_text"),
                "answer_options": answer_options,
                "student_answer": r.selected_answer if r else None,
                "correct_answer": r.correct_answer if r else None,
                "is_correct": r.is_correct if r else False,
                "topic": topic,
                "passage_style": passage_style,
                "reading_material": reading_material
            })

    # --------------------------------------------------
    # 6️⃣ Final payload logging
    # --------------------------------------------------
    response_payload = {
        "session_id": session.id,
        "exam_id": session.exam_id,
        "questions": review_questions
    }

    print("📤 REVIEW PAYLOAD SUMMARY")
    print("   session_id:", response_payload["session_id"])
    print("   exam_id:", response_payload["exam_id"])
    print("   total_questions:", len(review_questions))

    if review_questions:
        print("📌 SAMPLE QUESTION SENT TO FRONTEND:")
        print({
            k: response_payload["questions"][0][k]
            for k in [
                "question_id",
                "question_number",
                "student_answer",
                "correct_answer",
                "is_correct",
                "topic"
            ]
        })

    print("================ END REVIEW READING EXAM ================\n")

    return response_payload

@app.get(
    "/api/student/exam-review/mathematical-reasoning",
    response_model=ExamReviewResponse
)
def get_exam_review_mathematical_reasoning(
    student_id: str,  # external/public ID (e.g. Gem_temp3)
    db: Session = Depends(get_db)
):
    print("\n============= EXAM REVIEW (MATHEMATICAL REASONING) =============")
    print(f"➡️ Incoming request (external student_id): {student_id}")

    # ==================================================
    # 0️⃣ Resolve INTERNAL student ID
    # ==================================================
    print("🔐 Resolving internal student ID...")

    student = (
        db.query(Student)
        .filter(Student.student_id == student_id)
        .first()
    )

    if not student:
        print("❌ Student NOT FOUND for external_id =", student_id)
        raise HTTPException(
            status_code=404,
            detail="Student not found"
        )

    internal_student_id = student.id
    print(f"✅ Internal student_id resolved: {internal_student_id}")

    # ==================================================
    # 1️⃣ Resolve LATEST exam_attempt_id
    # ==================================================
    print("🔍 Resolving latest exam_attempt_id from responses...")

    latest_attempt = (
        db.query(StudentExamResponseMathematicalReasoning.exam_attempt_id)
        .filter(
            StudentExamResponseMathematicalReasoning.student_id
            == internal_student_id
        )
        .order_by(
            StudentExamResponseMathematicalReasoning.exam_attempt_id.desc()
        )
        .first()
    )

    if not latest_attempt:
        print(
            "❌ No exam responses found for internal_student_id =",
            internal_student_id
        )
        raise HTTPException(
            status_code=404,
            detail="No completed exam attempt found for this student"
        )

    exam_attempt_id = latest_attempt.exam_attempt_id
    print(f"✅ Latest exam_attempt_id resolved: {exam_attempt_id}")

    # ==================================================
    # 2️⃣ Validate exam attempt ownership
    # ==================================================
    print("🔍 Validating exam attempt ownership...")

    attempt = (
        db.query(StudentExamMathematicalReasoning)
        .filter(
            StudentExamMathematicalReasoning.id == exam_attempt_id,
            StudentExamMathematicalReasoning.student_id
            == internal_student_id
        )
        .first()
    )

    if not attempt:
        print(
            "❌ Exam attempt ownership mismatch:",
            f"attempt_id={exam_attempt_id},",
            f"internal_student_id={internal_student_id}"
        )
        raise HTTPException(
            status_code=404,
            detail="Exam attempt not found for this student"
        )

    print(
        "✅ Exam attempt verified:",
        f"id={attempt.id},",
        f"exam_id={attempt.exam_id},",
        f"student_id={attempt.student_id}"
    )

    # ==================================================
    # 3️⃣ Fetch all responses for this attempt
    # ==================================================
    print("📥 Fetching student responses...")

    responses = (
        db.query(StudentExamResponseMathematicalReasoning)
        .filter(
            StudentExamResponseMathematicalReasoning.exam_attempt_id
            == exam_attempt_id
        )
        .order_by(StudentExamResponseMathematicalReasoning.q_id)
        .all()
    )

    print(f"📊 Response rows fetched: {len(responses)}")

    if not responses:
        print("⚠️ No responses found for exam_attempt_id =", exam_attempt_id)
        return {
            "exam_attempt_id": exam_attempt_id,
            "questions": []
        }

    for r in responses[:3]:
        print(
            f"   ↳ q_id={r.q_id}, "
            f"selected={r.selected_option}, "
            f"correct={r.correct_option}, "
            f"is_correct={r.is_correct}"
        )

    # ==================================================
    # 4️⃣ Load exam definition
    # ==================================================
    print("📘 Loading exam definition...")

    exam = (
        db.query(Exam)
        .filter(Exam.id == attempt.exam_id)
        .first()
    )

    if not exam:
        print("❌ Exam NOT FOUND for exam_id =", attempt.exam_id)
        raise HTTPException(
            status_code=404,
            detail="Exam not found"
        )

    raw_questions = exam.questions or []
    print(f"📚 Raw questions loaded: {len(raw_questions)}")

    normalized = normalize_mr_questions_exam_review(raw_questions)

    print(f"🧹 Normalized questions count: {len(normalized)}")

    question_map = {q["q_id"]: q for q in normalized}
    print(f"🗺️ Question map size: {len(question_map)}")

    # ==================================================
    # 5️⃣ Build review payload
    # ==================================================
    print("🧩 Building review payload...")

    response_map = {r.q_id: r for r in responses}
    review_questions = []

    for q in normalized:
        r = response_map.get(q["q_id"])

        review_questions.append({
            "q_id": q["q_id"],
            "blocks": q.get("blocks", []),
            "options": q.get("options", {}),
            "student_answer": r.selected_option if r else None,
            "correct_answer": r.correct_option if r else None,
        })

    print(f"✅ Review questions prepared: {len(review_questions)}")

    skipped = sum(
        1 for q in review_questions
        if q["student_answer"] is None
    )

    print(
        "📈 Review summary:",
        f"total={len(review_questions)},",
        f"attempted={len(review_questions) - skipped},",
        f"skipped={skipped}"
    )

    print("============= END EXAM REVIEW =================\n")

    return {
        "exam_attempt_id": exam_attempt_id,
        "questions": review_questions
    }
def normalize_options_thinking_skills_review(raw_options: dict) -> dict:
    """
    Ensures every option value is a dictionary so it matches
    the FastAPI / Pydantic response model.

    Supported outputs:
    - text options
    - image options
    """

    normalized = {}

    if not isinstance(raw_options, dict):
        return normalized

    for key, value in raw_options.items():
        # Case 1: already normalized
        if isinstance(value, dict):
            normalized[key] = value
            continue

        # Case 2: string value (text or image filename)
        if isinstance(value, str):
            lower = value.lower()

            if lower.endswith((".png", ".jpg", ".jpeg", ".webp")):
                normalized[key] = {
                    "type": "image",
                    "value": value
                }
            else:
                normalized[key] = {
                    "type": "text",
                    "value": value
                }
            continue

        # Case 3: numbers or unexpected types
        normalized[key] = {
            "type": "text",
            "value": str(value)
        }

    return normalized

@app.get(
    "/api/student/exam-review/thinking-skills",
    response_model=ExamReviewResponse
)
def get_exam_review_thinking_skills(
    student_id: str,  # external/public ID (e.g. Gem_temp3)
    db: Session = Depends(get_db)
):
    print("\n================ EXAM REVIEW (THINKING SKILLS) =================")
    print(f"➡️ Incoming request (external student_id): {student_id}")

    # ==================================================
    # 0️⃣ Resolve INTERNAL student ID (MANDATORY STEP)
    # ==================================================
    print("🔐 Resolving internal student ID...")

    student = (
        db.query(Student)
        .filter(Student.student_id == student_id)
        .first()
    )

    if not student:
        print("❌ Student NOT FOUND for external_id =", student_id)
        raise HTTPException(
            status_code=404,
            detail="Student not found"
        )

    internal_student_id = student.id
    print(f"✅ Internal student_id resolved: {internal_student_id}")

    # ==================================================
    # 1️⃣ Resolve LATEST exam_attempt_id for student
    # ==================================================
    print("🔍 Resolving latest exam_attempt_id from responses...")

    latest_attempt = (
        db.query(StudentExamResponseThinkingSkills.exam_attempt_id)
        .filter(
            StudentExamResponseThinkingSkills.student_id == internal_student_id
        )
        .order_by(StudentExamResponseThinkingSkills.exam_attempt_id.desc())
        .first()
    )

    if not latest_attempt:
        print(
            "❌ No exam responses found for internal_student_id =",
            internal_student_id
        )
        raise HTTPException(
            status_code=404,
            detail="No completed exam attempt found for this student"
        )

    exam_attempt_id = latest_attempt.exam_attempt_id
    print(f"✅ Latest exam_attempt_id resolved: {exam_attempt_id}")

    # ==================================================
    # 2️⃣ Validate exam attempt ownership
    # ==================================================
    print("🔍 Validating exam attempt ownership...")

    attempt = (
        db.query(StudentExamThinkingSkills)
        .filter(
            StudentExamThinkingSkills.id == exam_attempt_id,
            StudentExamThinkingSkills.student_id == internal_student_id
        )
        .first()
    )

    if not attempt:
        print(
            "❌ Exam attempt ownership mismatch:",
            f"attempt_id={exam_attempt_id},",
            f"internal_student_id={internal_student_id}"
        )
        raise HTTPException(
            status_code=404,
            detail="Exam attempt not found for this student"
        )

    print(
        "✅ Exam attempt verified:",
        f"id={attempt.id},",
        f"exam_id={attempt.exam_id},",
        f"student_id={attempt.student_id}"
    )

    # ==================================================
    # 3️⃣ Fetch all responses for this attempt
    # ==================================================
    print("📥 Fetching student responses...")

    responses = (
        db.query(StudentExamResponseThinkingSkills)
        .filter(
            StudentExamResponseThinkingSkills.exam_attempt_id == exam_attempt_id
        )
        .order_by(StudentExamResponseThinkingSkills.q_id)
        .all()
    )

    print(f"📊 Response rows fetched: {len(responses)}")

    if not responses:
        print("⚠️ No responses found for exam_attempt_id =", exam_attempt_id)
        return {
            "exam_attempt_id": exam_attempt_id,
            "questions": []
        }

    for r in responses[:3]:
        print(
            f"   ↳ q_id={r.q_id}, "
            f"selected={r.selected_option}, "
            f"correct={r.correct_option}, "
            f"is_correct={r.is_correct}"
        )

    # ==================================================
    # 4️⃣ Load exam definition
    # ==================================================
    print("📘 Loading exam definition...")

    exam = (
        db.query(Exam)
        .filter(Exam.id == attempt.exam_id)
        .first()
    )

    if not exam:
        print("❌ Exam NOT FOUND for exam_id =", attempt.exam_id)
        raise HTTPException(
            status_code=404,
            detail="Exam not found"
        )

    raw_questions = exam.questions or []
    print(f"📚 Raw questions loaded: {len(raw_questions)}")

    normalized = normalize_questions_exam_review(raw_questions)
    print(f"🧹 Normalized questions count: {len(normalized)}")

    question_map = {q["q_id"]: q for q in normalized}
    print(f"🗺️ Question map size: {len(question_map)}")

    # ==================================================
    # 5️⃣ Build review payload
    # ==================================================
    print("🧩 Building review payload...")

    review_questions = []
    missing_questions = 0

    # Map responses by q_id for quick lookup
    response_map = {r.q_id: r for r in responses}
    
    review_questions = []
    
    for q in normalized:
        r = response_map.get(q["q_id"])
    
        normalized_options = normalize_options_thinking_skills_review(
            q.get("options", {})
        )
    
        review_questions.append({
            "q_id": q["q_id"],
            "blocks": q.get("blocks", []),
            "options": normalized_options,   # ✅ now always Dict[str, Dict]
            "student_answer": r.selected_option if r else None,
            "correct_answer": r.correct_option if r else None,
        })


    print(f"✅ Review questions prepared: {len(review_questions)}")

    if missing_questions:
        print(f"⚠️ Questions missing definitions: {missing_questions}")

    skipped = sum(1 for q in review_questions if q["student_answer"] is None)

    print(
        "📈 Review summary:",
        f"total={len(review_questions)},",
        f"attempted={len(review_questions) - skipped},",
        f"skipped={skipped}"
    )

    print("================ END EXAM REVIEW =================\n")

    return {
        "exam_attempt_id": exam_attempt_id,
        "questions": review_questions
    }

@app.get("/api/classes/{class_name}/exam-dates")
def get_class_exam_dates(
    class_name: str,
    exam: str,
    db: Session = Depends(get_db),
):
    print("\n==============================")
    print("📥 [CLASS EXAM DATES] REQUEST")
    print("   class_name:", class_name)
    print("   exam:", exam)
    print("==============================")

    dates = (
        db.query(func.date(AdminExamReport.created_at))
        .join(
            Student,
            Student.student_id == AdminExamReport.student_id
        )
        .filter(
            Student.class_name == class_name,
            AdminExamReport.exam_type == exam,
        )
        .distinct()
        .order_by(func.date(AdminExamReport.created_at))
        .all()
    )

    date_list = [d[0].isoformat() for d in dates if d[0]]

    print("✅ dates_found:", date_list)

    return {
        "class_name": class_name,
        "exam": exam,
        "dates": date_list,
    }
def get_exam_response_model(exam: str):
    if exam == "thinking_skills":
        return StudentExamResponseThinkingSkills
    elif exam == "reading":
        return StudentExamReportReading
    elif exam == "mathematical_reasoning":
        return StudentExamResponse
    elif exam == "writing":
        return StudentExamWriting

    raise HTTPException(status_code=400, detail="Unsupported exam type")

def get_question_model(exam: str):
    if exam == "thinking_skills":
        return Question
    elif exam == "mathematical_reasoning":
        return Question
    elif exam == "reading":
        return QuestionReading
    elif exam == "writing":
        return WritingQuestionBank
    else:
        raise RuntimeError(f"Unsupported exam type: {exam}")


# ----------------------------------------
# Helpers
# ----------------------------------------

def normalize_topic_reporting(value: str) -> str:
    value = value.lower()

    # Replace symbols with spaces
    value = value.replace("&", " ")

    # Remove punctuation
    value = re.sub(r"[(),-]", " ", value)

    # Normalize whitespace
    value = re.sub(r"\s+", " ", value).strip()

    # Remove standalone word "and"
    tokens = [
        word for word in value.split(" ")
        if word != "and"
    ]

    return "_".join(tokens)
def response_has_own_topic(ResponseModel):
    return hasattr(ResponseModel, "topic")


# ----------------------------------------
# Endpoint
# ----------------------------------------


@app.get("/api/reports/student/cumulative")
def get_student_cumulative_report(
    student_id: str,
    exam: str,
    topic: str,
    attempt_dates: list[str] = Query(...),
    db: Session = Depends(get_db),
):
    print("\n==============================")
    print("📥 [CUMULATIVE] REQUEST RECEIVED")
    print("   student_id:", student_id)
    print("   exam:", exam)
    print("   topic (raw):", repr(topic))
    print("   attempt_dates:", attempt_dates)
    print("==============================")

    try:
        # --------------------------------------------------
        # 0️⃣ Guard: attempt_dates
        # --------------------------------------------------
        if not attempt_dates:
            print("❌ [ABORT @0] attempt_dates missing or empty")
            raise HTTPException(
                status_code=400,
                detail="At least one attempt date is required",
            )

        print("🔥 attempt_dates type:", type(attempt_dates), attempt_dates)

        # --------------------------------------------------
        # 1️⃣ Resolve internal student
        # --------------------------------------------------
        print("\n[1] Resolving internal student...")

        student = (
            db.query(Student)
            .filter(Student.student_id == student_id)
            .first()
        )

        if not student:
            print("❌ [ABORT @1] Student not found:", student_id)
            raise HTTPException(
                status_code=404,
                detail="Student not found",
            )

        print("✅ Student resolved")
        print("   internal_id:", student.id)
        print("   name:", student.name)

        # --------------------------------------------------
        # 2️⃣ Resolve exam attempts
        # --------------------------------------------------
        print("\n[2] Resolving exam attempts (AdminExamReport)...")

        attempts = (
            db.query(AdminExamReport)
            .filter(
                AdminExamReport.student_id == student_id,
                AdminExamReport.exam_type == exam,
                func.date(AdminExamReport.created_at).in_(attempt_dates),
            )
            .order_by(AdminExamReport.created_at)
            .all()
        )

        print("   attempts_found:", len(attempts))

        for a in attempts:
            print(
                "   → exam_attempt_id:", a.exam_attempt_id,
                "| date:", a.created_at.date(),
                "| created_at:", a.created_at,
            )

        if not attempts:
            print("❌ [ABORT @2] No exam attempts matched filters")
            raise HTTPException(
                status_code=404,
                detail="No exam attempts found for given filters",
            )

        # --------------------------------------------------
        # 3️⃣ Resolve response model
        # --------------------------------------------------
        print("\n[3] Resolving response model...")

        ResponseModel = get_response_model(exam)

        if not ResponseModel:
            print("❌ [ABORT @3] Unsupported exam type:", exam)
            raise HTTPException(
                status_code=400,
                detail="Unsupported exam type",
            )

        print(
            "🧪 QUERYING TABLE:",
            ResponseModel.__tablename__,
            "| MODEL:",
            ResponseModel.__name__,
        )

        # --------------------------------------------------
        # 4️⃣ Normalize requested topic
        # --------------------------------------------------
        normalized_request_topic = normalize_topic_reporting(topic)

        print("\n[4] Topic normalization")
        print("   request_topic_raw:", repr(topic))
        print("   request_topic_normalized:", repr(normalized_request_topic))

        results = []

        # --------------------------------------------------
        # 5️⃣ Process each attempt (time series)
        # --------------------------------------------------
        print("\n[5] Processing attempts (time series)...")

        for idx, attempt in enumerate(attempts, start=1):
            print(f"\n   ▶ Attempt {idx}")
            print("     exam_attempt_id:", attempt.exam_attempt_id)
            print("     created_at:", attempt.created_at)

            raw_responses = (
                db.query(ResponseModel)
                .filter(
                    ResponseModel.student_id == student.id,
                    ResponseModel.exam_attempt_id == attempt.exam_attempt_id,
                )
                .all()
            )

            print("     raw_responses_found:", len(raw_responses))

            if not raw_responses:
                print("     ⚠️ No responses found for this attempt")
                continue

            # --- Inspect DB topics ---
            normalized_db_topics = {
                normalize_topic_reporting(r.topic)
                for r in raw_responses
                if r.topic
            }

            print("     🧠 normalized_db_topics:", normalized_db_topics)
            print("     🧠 requested_topic:", normalized_request_topic)

            # --- Topic filtering (DIRECT equality) ---
            responses = [
                r for r in raw_responses
                if r.topic
                and normalize_topic_reporting(r.topic)
                == normalized_request_topic
            ]

            print(
                "     🔍 topic_filtering:",
                f"{len(responses)} / {len(raw_responses)} matched",
            )

            attempted = len(responses)
            correct = sum(1 for r in responses if r.is_correct)

            print("     responses_found:", attempted)
            print("     correct:", correct)

            # Skip empty attempts safely
            if attempted == 0:
                print("     ⚠️ Skipping attempt — no matching topic questions")
                continue

            accuracy = round((correct / attempted) * 100, 2)
            score = accuracy

            print("     accuracy:", accuracy)
            print("     score:", score)

            results.append({
                "date": attempt.created_at.date().isoformat(),
                "questions_attempted": attempted,
                "correct_answers": correct,
                "accuracy": accuracy,
                "score": score,
            })

        # --------------------------------------------------
        # 6️⃣ Final validation
        # --------------------------------------------------
        print("\n[6] Final validation")
        print("   total_valid_attempts:", len(results))
        print("   results_payload:", results)

        if not results:
            print("❌ [ABORT @6] No valid attempts after topic filtering")
            raise HTTPException(
                status_code=400,
                detail="No data found to generate the required report.",
            )

        # --------------------------------------------------
        # 7️⃣ Build cumulative summary
        # --------------------------------------------------
        print("\n[7] Building cumulative summary...")

        first = results[0]
        last = results[-1]

        summary = {
            "first_attempt_score": first["score"],
            "latest_attempt_score": last["score"],
            "score_change": round(last["score"] - first["score"], 2),
            "first_attempt_accuracy": first["accuracy"],
            "latest_attempt_accuracy": last["accuracy"],
            "accuracy_change": round(
                last["accuracy"] - first["accuracy"], 2
            ),
            "trend": (
                "improving"
                if last["accuracy"] > first["accuracy"]
                else "declining"
                if last["accuracy"] < first["accuracy"]
                else "stable"
            ),
        }

        print("   summary:", summary)

        # --------------------------------------------------
        # 8️⃣ Return response
        # --------------------------------------------------
        print("\n[8] Returning cumulative report payload")
        print("==============================\n")

        return {
            "student_id": student_id,
            "student_name": student.name,
            "exam": exam,
            "topic": {
                "key": topic,
                "label": topic.replace("_", " ").title(),
            },
            "attempts": results,
            "summary": summary,
        }

    except HTTPException:
        raise

    except Exception as e:
        print("🔥🔥🔥 UNHANDLED CUMULATIVE ERROR 🔥🔥🔥")
        print("Type:", type(e))
        print("Message:", str(e))
        import traceback
        traceback.print_exc()

        raise HTTPException(
            status_code=500,
            detail="Internal error while generating cumulative report",
        )

@app.get("/api/reports/class")
def class_exam_report(
    class_name: str,
    exam: str,
    date: str,
    db: Session = Depends(get_db)
):
    print("========================================")
    print("[CLASS REPORT] Request received")
    print("class_name:", class_name)
    print("exam:", exam)
    print("date:", date)

    # ===========================
    # STEP 1: Resolve class days
    # ===========================
    print("\n[STEP 1] Resolving distinct class days")

    try:
        class_days_raw = (
            db.query(Student.class_day)
            .filter(Student.class_name == class_name)
            .distinct()
            .all()
        )
    except Exception as e:
        print("[ERROR] Failed resolving class days:", e)
        raise HTTPException(status_code=500, detail="Failed to resolve class days")

    class_days = [d[0] for d in class_days_raw if d[0]]
    print("[STEP 1] Class days found:", class_days)

    if not class_days:
        print("[EARLY EXIT] No class days found")
        return {
            "class_name": class_name,
            "exam": exam,
            "date": date,
            "reports": []
        }

    reports = []

    # ===========================
    # STEP 2: Loop per class day
    # ===========================
    for class_day in class_days:
        print("\n----------------------------------------")
        print(f"[DAY LOOP] Processing day: {class_day}")

        # ---------------------------
        # Resolve students
        # ---------------------------
        print("[STEP 2] Resolving students")

        students = (
            db.query(Student.id, Student.student_id, Student.name)
            .filter(
                Student.class_name == class_name,
                Student.class_day == class_day
            )
            .all()
        )


        students_total = len(students)
        print("[STEP 2] Students found:", students_total)

        if students_total == 0:
            print("[DAY SKIP] No students for this day")
            reports.append({
                "class_day": class_day,
                "summary": {
                    "students_total": 0,
                    "students_attempted": 0,
                    "average_score": 0,
                    "highest_score": 0
                },
                "leaderboard": [],
                "score_distribution": _empty_buckets()
            })
            continue

        student_name_map = {s.id: s.name for s in students}
        student_code_map = {s.id: s.student_id for s in students}  # optional but useful
        
        student_internal_ids = list(student_name_map.keys())


        print("[STEP 2] Student internal IDs:", student_internal_ids)


        # ---------------------------
        # Resolve exam responses
        # ---------------------------
        print("[STEP 3] Resolving exam attempts for date")
        external_ids = list(student_code_map.values())

        attempt_ids = (
            db.query(AdminExamReport.exam_attempt_id)
            .filter(
                AdminExamReport.exam_type == exam,
                func.date(AdminExamReport.created_at) == date,
                AdminExamReport.student_id.in_(external_ids)
            )
            .all()
        )

        
        attempt_ids = [a[0] for a in attempt_ids]
        
        print("[STEP 3] Exam attempt IDs:", attempt_ids)        
        if not attempt_ids:
            print("[STEP 3] No exam attempts found for this date")
            students_attempted = 0
            student_results = []
        else:
            print("[STEP 3] Resolving exam responses")
        
            ResponseModel = get_exam_response_model(exam)
            print("[STEP 3] Using response model:", ResponseModel.__name__)
        
            raw_rows = (
                db.query(
                    ResponseModel.student_id,
                    ResponseModel.exam_attempt_id,
                    ResponseModel.is_correct
                )
                .filter(
                    ResponseModel.exam_attempt_id.in_(attempt_ids)
                )
                .all()
            )
        
            print("[STEP 3] Raw rows fetched:", len(raw_rows))
        
            attempts = {}
            
            for r in raw_rows:
                key = (r.student_id, r.exam_attempt_id)
            
                if key not in attempts:
                    attempts[key] = {"total": 0, "correct": 0}
            
                attempts[key]["total"] += 1
            
                # NULL treated as incorrect
                if r.is_correct is True:
                    attempts[key]["correct"] += 1
            
            student_results = []
            
            for (student_id, _), stats in attempts.items():
                total = stats["total"]
                correct = stats["correct"]
            
                score = round((correct / total) * 100) if total else 0
            
                result = {
                    "student_id": student_id,  # internal ID
                    "student_code": student_code_map.get(student_id),
                    "student_name": student_name_map.get(student_id),
                    "score": score,
                    "accuracy": score
                }

            
                print("[STEP 3] Student result:", result)
                student_results.append(result)
            
            students_attempted = len(student_results)
            print("[STEP 3] Students attempted:", students_attempted)
  

        # ---------------------------
        # STEP 5: Summary metrics
        # ---------------------------
        if students_attempted == 0:
            average_score = 0
            highest_score = 0
        else:
            average_score = round(
                sum(s["score"] for s in student_results) / students_attempted
            )
            highest_score = max(s["score"] for s in student_results)
        
        print("[STEP 5] Average score:", average_score)
        print("[STEP 5] Highest score:", highest_score)

        # ---------------------------
        # STEP 6: Leaderboard
        # ---------------------------
        leaderboard = []
        for idx, s in enumerate(
            sorted(student_results, key=lambda x: x["score"], reverse=True),
            start=1
        ):
            leaderboard.append({
                "rank": idx,
                "student": f"{s['student_name']} ({s['student_code']})",
                "score": s["score"],
                "accuracy": s["accuracy"]
            })

        # ---------------------------
        # STEP 7: Score distribution
        # ---------------------------
        buckets = {
            "0-40": 0,
            "41-60": 0,
            "61-80": 0,
            "81-100": 0
        }

        for s in student_results:
            if s["score"] <= 40:
                buckets["0-40"] += 1
            elif s["score"] <= 60:
                buckets["41-60"] += 1
            elif s["score"] <= 80:
                buckets["61-80"] += 1
            else:
                buckets["81-100"] += 1

        print("[STEP 7] Buckets:", buckets)

        reports.append({
            "class_day": class_day,
            "summary": {
                "students_total": students_total,
                "students_attempted": students_attempted,
                "average_score": average_score,
                "highest_score": highest_score
            },
            "leaderboard": leaderboard,
            "score_distribution": [
                {"range": k, "count": v} for k, v in buckets.items()
            ]
        })

    # ===========================
    # STEP 8: Final response
    # ===========================
    print("\n[STEP 8] Returning multi-day class report")
    print("========================================")

    return {
        "class_name": class_name,
        "exam": exam,
        "date": date,
        "reports": reports
    }


def _empty_buckets():
    return [
        {"range": "0-40", "count": 0},
        {"range": "41-60", "count": 0},
        {"range": "61-80", "count": 0},
        {"range": "81-100", "count": 0}
    ]


@app.get("/api/classes/{class_name}/days")
def fetch_class_days(class_name: str, db: Session = Depends(get_db)):
    days = get_days_for_class(db, class_name)
    return {
        "class": class_name,
        "days": days
    }

@app.get("/api/reading/question-bank")
def get_reading_question_bank_summary(
    subject: str = Query("reading_comprehension"),
    class_name: str = Query("selective"),
    db: Session = Depends(get_db),
):
    """
    Admin overview of reading question bank.

    Each row represents a distinct exam-ready question set,
    grouped by:
    - difficulty
    - topic
    - total_questions (set size)

    This allows admins to clearly see, for example,
    Comparative Analysis sets of size 8 vs 10.
    """

    subject_norm = func.lower(
        func.replace(func.trim(QuestionReading.subject), " ", "_")
    )
    class_norm = func.lower(func.trim(QuestionReading.class_name))
    difficulty_norm = func.lower(func.trim(QuestionReading.difficulty))

    rows = (
        db.query(
            difficulty_norm.label("difficulty"),
            QuestionReading.topic,
            QuestionReading.total_questions.label("set_size"),
            func.count(QuestionReading.id).label("sets_available"),
        )
        .filter(subject_norm == subject.lower())
        .filter(class_norm == class_name.lower())
        .filter(QuestionReading.difficulty.isnot(None))
        .group_by(
            difficulty_norm,
            QuestionReading.topic,
            QuestionReading.total_questions,
        )
        .order_by(
            difficulty_norm,
            QuestionReading.topic,
            QuestionReading.total_questions,
        )
        .all()
    )

    return {
        "rows": [
            {
                "difficulty": r.difficulty.capitalize(),
                "topic": r.topic,
                "set_size": r.set_size,
                "sets_available": r.sets_available,
            }
            for r in rows
        ]
    }


 
@app.get("/api/admin/question-bank-reading")
def get_question_bank_reading(
    db: Session = Depends(get_db)
):
    results = (
        db.query(
            QuestionReading.difficulty,
            QuestionReading.topic,
            func.sum(QuestionReading.total_questions).label("total_questions")
        )
        .filter(
            QuestionReading.subject == "Reading Comprehension"
        )
        .group_by(
            QuestionReading.difficulty,
            QuestionReading.topic
        )
        .order_by(
            QuestionReading.difficulty,
            QuestionReading.topic
        )
        .all()
    )

    return [
        {
            "difficulty": r.difficulty,
            "topic": r.topic,
            "total_questions": r.total_questions
        }
        for r in results
    ]
@app.get("/api/admin/question-bank-mathematical-reasoning")
def get_question_bank_mathematical_reasoning(
    db: Session = Depends(get_db)
):
    results = (
        db.query(
            Question.difficulty,
            Question.topic,
            func.count(Question.id).label("total_questions")
        )
        .filter(
            Question.subject == "Mathematical Reasoning"
        )
        .group_by(
            Question.difficulty,
            Question.topic
        )
        .order_by(
            Question.difficulty,
            Question.topic
        )
        .all()
    )

    return [
        {
            "difficulty": r.difficulty,
            "topic": r.topic,
            "total_questions": r.total_questions
        }
        for r in results
    ]
@app.get("/api/writing/topics")
def get_writing_topics(
    difficulty: str = Query(...),
    db: Session = Depends(get_db),
):
    print("📥 Fetching Writing topics")
    print(f"   difficulty={difficulty}")

    DB_SUBJECT = "Writing"

    topics = (
        db.query(func.distinct(WritingQuestionBank.topic))
        .filter(
            func.lower(func.trim(WritingQuestionBank.subject)) == DB_SUBJECT.lower(),
            func.lower(func.trim(WritingQuestionBank.difficulty)) == difficulty.lower(),
        )
        .order_by(WritingQuestionBank.topic)
        .all()
    )

    topic_list = [{"name": t[0]} for t in topics]

    print(f"✅ Writing topics found: {len(topic_list)}")

    return topic_list

@app.get("/api/reading/topics")
def get_reading_topics(
    difficulty: str = Query(...),
    db: Session = Depends(get_db),
):
    print("📥 Fetching Reading topics")
    print(f"   difficulty={difficulty}")

    # DB VALUE (even if misspelled)
    DB_SUBJECT = "reading_comprehension"

    topics = (
        db.query(func.distinct(QuestionReading.topic))
        .filter(
            func.lower(func.trim(QuestionReading.subject)).like("reading%"),
            func.lower(func.trim(QuestionReading.difficulty)) == difficulty.lower(),
        )

        .order_by(QuestionReading.topic)
        .all()
    )


    topic_list = [{"name": t[0]} for t in topics]

    print(f"✅ Reading topics found: {len(topic_list)}")

    return topic_list



@app.get("/api/topics")
def get_topics(
    class_name: str = Query(...),
    subject: str = Query(...),
    difficulty: str = Query(...),
    db: Session = Depends(get_db),
):
    print("📥 Fetching topics with filters:")
    print(f"   class_name={class_name}")
    print(f"   subject={subject}")
    print(f"   difficulty={difficulty}")

    normalized_subject = subject.lower().replace("_", " ")

    topics = (
        db.query(func.distinct(Question.topic))
        .filter(
            func.lower(func.trim(Question.class_name)) == class_name.lower(),
            func.lower(func.trim(Question.subject)) == normalized_subject,
            func.lower(func.trim(Question.difficulty)) == difficulty.lower(),
            Question.topic.isnot(None),
            Question.topic != "",
        )
        .order_by(Question.topic)
        .all()
    )


    topic_list = [{"name": t[0]} for t in topics]

    print(f"✅ Topics found: {len(topic_list)}")

    return topic_list



def normalize_topic_key(raw: str) -> str:
    """
    Converts:
    'Argument Analysis (Critical Thinking)'
    -> 'argument_analysis'
    """
    raw = raw.lower()
    raw = re.sub(r"\(.*?\)", "", raw)      # remove parentheses
    raw = re.sub(r"[^a-z0-9\s_]", "", raw) # remove symbols
    raw = raw.strip()
    raw = raw.replace(" ", "_")
    raw = re.sub(r"_+", "_", raw)
    return raw


@app.get("/api/exams/{exam}/topics")
def get_exam_topics(
    exam: str,
    db: Session = Depends(get_db)
):
    print(f"📘 [TOPICS] Requested exam = {exam}")

    if exam == "thinking_skills":
        rows = (
            db.query(Question.topic)
            .filter(
                Question.subject == "Thinking Skills",
                Question.topic.isnot(None)
            )
            .distinct()
            .order_by(Question.topic)
            .all()
        )

    elif exam == "mathematical_reasoning":
        rows = (
            db.query(Question.topic)
            .filter(
                Question.subject == "Mathematical Reasoning",
                Question.topic.isnot(None)
            )
            .distinct()
            .order_by(Question.topic)
            .all()
        )

    elif exam == "reading":
        rows = (
            db.query(ReadingQuestion.topic)
            .filter(ReadingQuestion.topic.isnot(None))
            .distinct()
            .order_by(ReadingQuestion.topic)
            .all()
        )

    elif exam == "writing":
        rows = (
            db.query(WritingQuestion.topic)
            .filter(WritingQuestion.topic.isnot(None))
            .distinct()
            .order_by(WritingQuestion.topic)
            .all()
        )

    else:
        raise HTTPException(status_code=400, detail="Unsupported exam type")

    topics = []

    for (raw_topic,) in rows:
        key = normalize_topic_key(raw_topic)
        label = raw_topic.strip()

        print(f"🧩 [TOPIC] raw='{raw_topic}' → key='{key}'")

        topics.append({
            "key": key,
            "label": label
        })

    print(f"✅ [TOPICS] Returning {len(topics)} topics")

    return {
        "exam": exam,
        "topics": topics
    }


@app.get("/api/admin/question-bank-thinking-skills")
def get_question_bank_thinking_skills(
    db: Session = Depends(get_db)
):
    results = (
        db.query(
            Question.difficulty,
            Question.topic,
            func.count(Question.id).label("total_questions")
        )
        .filter(
            Question.subject == "Thinking Skills"
        )
        .group_by(
            Question.difficulty,
            Question.topic
        )
        .order_by(
            Question.difficulty,
            Question.topic
        )
        .all()
    )

    return [
        {
            "difficulty": r.difficulty,
            "topic": r.topic,
            "total_questions": r.total_questions
        }
        for r in results
    ]

def get_attempt_filter(ResponseModel, exam_attempt_id):
    if hasattr(ResponseModel, "exam_attempt_id"):
        return ResponseModel.exam_attempt_id == exam_attempt_id
    elif hasattr(ResponseModel, "session_id"):
        return ResponseModel.session_id == exam_attempt_id
    else:
        raise RuntimeError(
            f"{ResponseModel.__name__} has no attempt identifier column"
        )

@app.get("/api/reports/student")
def get_student_exam_report(
    student_id: str,
    exam: str,
    date: date,
    db: Session = Depends(get_db),
):
    print("\n" + "=" * 80)
    print("📊 STUDENT EXAM REPORT REQUEST")
    print("=" * 80)
    print("➡️ external student_id:", student_id)
    print("➡️ exam:", exam)
    print("➡️ date:", date)

    # --------------------------------------------------
    # 0️⃣ Resolve student (EXTERNAL → INTERNAL ID)
    # --------------------------------------------------
    student = (
        db.query(Student)
        .filter(
            func.lower(Student.student_id) ==
            func.lower(student_id.strip())
        )
        .first()
    )

    if not student:
        print("❌ Student not found for external id:", student_id)
        raise HTTPException(status_code=404, detail="Student not found")

    print("✅ Student resolved → internal id:", student.id)

    # --------------------------------------------------
    # 1️⃣ Resolve exam attempt via AdminExamReport
    # --------------------------------------------------
    admin_report = (
        db.query(AdminExamReport)
        .filter(
            AdminExamReport.student_id == student.student_id,  # external id stored here
            AdminExamReport.exam_type == exam,
            func.date(AdminExamReport.created_at) == date
        )
        .first()
    )

    if not admin_report:
        print("❌ AdminExamReport not found")
        raise HTTPException(status_code=404, detail="Exam attempt not found")

    exam_attempt_id = admin_report.exam_attempt_id
    print("✅ Exam attempt resolved → exam_attempt_id:", exam_attempt_id)

    # --------------------------------------------------
    # 2️⃣ Load correct response table
    # --------------------------------------------------
    ResponseModel = get_response_model(exam)
    print("📘 Response model:", ResponseModel.__name__)

    responses = (
        db.query(ResponseModel)
        .filter(
            ResponseModel.student_id == student.id,
            get_attempt_filter(ResponseModel, exam_attempt_id)
        )
        .all()
    )


    print("📊 Responses fetched:", len(responses))

    if not responses:
        print("❌ No responses found for this attempt")
        raise HTTPException(status_code=404, detail="No responses found")

    # --------------------------------------------------
    # 3️⃣ Overall summary
    # --------------------------------------------------
    total = len(responses)
    attempted = sum(1 for r in responses if r.is_correct is not None)
    correct = sum(1 for r in responses if r.is_correct is True)
    incorrect = sum(1 for r in responses if r.is_correct is False)
    not_attempted = total - attempted

    accuracy = round((correct / attempted) * 100) if attempted else 0
    result = "Pass" if accuracy >= 50 else "Fail"

    summary = {
        "total_questions": total,
        "attempted": attempted,
        "correct": correct,
        "incorrect": incorrect,
        "not_attempted": not_attempted,
        "accuracy": accuracy,
        "score": accuracy,
        "result": result,
    }

    # --------------------------------------------------
    # 4️⃣ Topic-wise aggregation
    # --------------------------------------------------
    topic_rows = (
       db.query(
           ResponseModel.topic,
           func.count().label("total"),
           func.count(case((ResponseModel.is_correct.isnot(None), 1))).label("attempted"),
           func.count(case((ResponseModel.is_correct.is_(True), 1))).label("correct"),
           func.count(case((ResponseModel.is_correct.is_(False), 1))).label("incorrect"),
       )
       .filter(
           ResponseModel.student_id == student.id,
           get_attempt_filter(ResponseModel, exam_attempt_id)
       )
       .group_by(ResponseModel.topic)
       .all()
   )

    topics = []
    improvement_areas = []

    for row in topic_rows:
        accuracy_pct = round((row.correct / row.attempted) * 100) if row.attempted else 0
        weakness = 100 - accuracy_pct

        topics.append({
            "topic": row.topic,
            "total": row.total,
            "attempted": row.attempted,
            "correct": row.correct,
            "incorrect": row.incorrect,
            "accuracy": accuracy_pct,
        })

        improvement_areas.append({
            "topic": row.topic,
            "weakness": weakness,
        })

    # --------------------------------------------------
    # 5️⃣ Final payload
    # --------------------------------------------------
    print("✅ Report generated successfully")
    print("=" * 80 + "\n")

    return {
        "exam": exam,
        "date": date.isoformat(),
        "summary": summary,
        "topics": topics,
        "improvement_areas": improvement_areas,
    }
@app.get("/api/exams/writing/dates")
def get_writing_exam_dates(student_id: str, db: Session = Depends(get_db)):

    student = (
        db.query(Student)
        .filter(func.lower(Student.student_id) == func.lower(student_id))
        .first()
    )

    if not student:
        raise HTTPException(404, "Student not found")

    dates = (
        db.query(func.date(AdminExamReport.created_at))
        .filter(
            AdminExamReport.student_id == student.student_id,
            AdminExamReport.exam_type == "writing"
        )
        .order_by(func.date(AdminExamReport.created_at).desc())
        .distinct()
        .all()
    )

    return {
        "dates": [d[0].isoformat() for d in dates]
    }


@app.get("/api/exams/writing/result")
def get_writing_result(
    student_id: str,
    db: Session = Depends(get_db)
):
    # --------------------------------------------------
    # 1️⃣ Resolve student (EXTERNAL → INTERNAL)
    # --------------------------------------------------
    student = (
        db.query(Student)
        .filter(func.lower(Student.student_id) == func.lower(student_id))
        .first()
    )

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # --------------------------------------------------
    # 2️⃣ Load latest admin report (SOURCE OF TRUTH)
    # --------------------------------------------------
    admin_report = (
        db.query(AdminExamReport)
        .filter(
            AdminExamReport.student_id == student.student_id,  # ✅ external ID
            AdminExamReport.exam_type == "writing"
        )
        .order_by(AdminExamReport.created_at.desc())
        .first()
    )

    if not admin_report:
        raise HTTPException(
            status_code=404,
            detail="Writing result not found"
        )

    # --------------------------------------------------
    # 3️⃣ Load writing attempt (for AI details)
    # --------------------------------------------------
    exam_state = (
        db.query(StudentExamWriting)
        .filter(
            StudentExamWriting.id == admin_report.exam_attempt_id
        )
        .first()
    )

    if not exam_state or not exam_state.ai_evaluation_json:
        raise HTTPException(
            status_code=500,
            detail="AI evaluation missing for writing exam"
        )

    # --------------------------------------------------
    # 4️⃣ Final response (UI SAFE)
    # --------------------------------------------------
    return {
        "exam_type": "Writing",
        "score": admin_report.overall_score,
        "max_score": 25,

        # ✅ canonical readiness band
        "selective_readiness_band": admin_report.readiness_band,

        # ✅ full AI breakdown
        "evaluation": exam_state.ai_evaluation_json,

        "advisory": "This report is advisory only and does not guarantee placement."
    }


@app.get("/api/reports/student/writing")
def get_student_writing_report(
    student_id: str,
    date: date,
    db: Session = Depends(get_db)
):
    # 1️⃣ Resolve student
    student = (
        db.query(Student)
        .filter(func.lower(Student.student_id) == func.lower(student_id))
        .first()
    )

    if not student:
        raise HTTPException(404, "Student not found")

    # 2️⃣ Find admin report for this date
    report = (
        db.query(AdminExamReport)
        .filter(
            AdminExamReport.student_id == student.student_id,  # external ID
            AdminExamReport.exam_type == "writing",
            func.date(AdminExamReport.created_at) == date
        )
        .order_by(AdminExamReport.created_at.desc())
        .first()
    )

    if not report:
        raise HTTPException(404, "Writing report not found")

    # 3️⃣ Load writing attempt
    exam_state = (
        db.query(StudentExamWriting)
        .filter(StudentExamWriting.id == report.exam_attempt_id)
        .first()
    )

    if not exam_state or not exam_state.ai_evaluation_json:
        raise HTTPException(500, "Writing evaluation not available")

    # 4️⃣ Return data in SAME SHAPE as MCQ reports
    return {
        "exam": "writing",
        "date": report.created_at.date().isoformat(),

        "summary": {
            "total_questions": 25,
            "attempted": 25,
            "correct": report.overall_score,
            "incorrect": 25 - report.overall_score,
            "not_attempted": 0,
            "accuracy": round((report.overall_score / 25) * 100),
            "score": round((report.overall_score / 25) * 100),
            "result": report.readiness_band,
        },

        "topics": [],            # writing has no topics
        "improvement_areas": [],

        # 👇 Writing-specific payload
        "evaluation": exam_state.ai_evaluation_json,
    }

def get_all_classes(db: Session):
    results = (
        db.query(distinct(Student.class_name))
        .order_by(Student.class_name)
        .all()
    )

    # results comes as list of tuples → [('Class A',), ('Class B',)]
    return [row[0] for row in results if row[0]]

@app.get("/api/classes")
def fetch_classes(db: Session = Depends(get_db)):
    classes = get_all_classes(db)
    return {
        "classes": classes
    }

@app.get("/api/topics-naplan")
def get_naplan_topics(
    subject: str,
    year: int,
    difficulty: str,
    db: Session = Depends(get_db),
):
    print("\n=== TOPICS-NAPLAN DEBUG START ===")

    print(f"[INPUT] subject={subject}, year={year}, difficulty={difficulty}")

    # Normalize inputs (frontend-friendly → DB-friendly)
    normalized_subject = subject.replace("_", " ").title()
    normalized_difficulty = difficulty.capitalize()

    print(
        f"[NORMALIZED] subject={normalized_subject}, "
        f"difficulty={normalized_difficulty}"
    )

    rows = (
        db.query(distinct(QuestionNumeracyLC.topic))
        # IMPORTANT: case-insensitive match for class_name
        .filter(func.lower(QuestionNumeracyLC.class_name) == "naplan")
        .filter(QuestionNumeracyLC.subject == normalized_subject)
        .filter(QuestionNumeracyLC.year == year)
        .filter(QuestionNumeracyLC.difficulty == normalized_difficulty)
        .filter(QuestionNumeracyLC.topic.isnot(None))
        .order_by(QuestionNumeracyLC.topic.asc())
        .all()
    )

    print(f"[QUERY RESULT] row_count={len(rows)}")

    for idx, (topic,) in enumerate(rows[:5]):
        print(f"[ROW {idx + 1}] topic={topic}")

    print("=== TOPICS-NAPLAN DEBUG END ===\n")

    return [{"name": topic} for (topic,) in rows]


@app.post("/generate-new-mr")
def generate_exam(
    db: Session = Depends(get_db)
):
    # --------------------------------------------------
    # 1️⃣ Fetch latest Mathematical Reasoning quiz
    # --------------------------------------------------
    quiz = (
        db.query(QuizMathematicalReasoning)
        .filter(
            QuizMathematicalReasoning.subject == "mathematical_reasoning"
        )
        .order_by(QuizMathematicalReasoning.id.desc())
        .first()
    )

    if not quiz:
        raise HTTPException(
            status_code=404,
            detail="No Mathematical Reasoning quiz found"
        )

    # ✅ Difficulty now comes from DB
    difficulty = quiz.difficulty

    # --------------------------------------------------
    # 2️⃣ PRE-FLIGHT DB VALIDATION (CRITICAL)
    # --------------------------------------------------
    for topic in quiz.topics:
        topic_name = topic["name"]
        db_required = int(topic.get("db", 0))

        available = (
            db.query(Question)
              .filter(
                  func.lower(Question.topic) == topic_name.lower(),
                  func.lower(Question.difficulty) == difficulty.lower()
              )
              .count()
        )

        if available < db_required:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"Cannot generate exam. "
                    f"Topic '{topic_name}' requires {db_required} DB questions, "
                    f"but only {available} exist. "
                    f"Please add more questions to the database."
                )
            )

    # --------------------------------------------------
    # 3️⃣ Generate questions (DB + AI)
    # --------------------------------------------------
    questions = []
    q_id = 1

    for topic in quiz.topics:
        topic_name = topic["name"]
        ai_count = int(topic.get("ai", 0))
        db_count = int(topic.get("db", 0))

        # ---- DB questions
        db_questions = (
            db.query(Question)
              .filter(
                  func.lower(Question.topic) == topic_name.lower(),
                  func.lower(Question.difficulty) == difficulty.lower()
              )
              .order_by(func.random())
              .limit(db_count)
              .all()
        )

        for q in db_questions:
            if not q.question_blocks:
                raise HTTPException(
                    status_code=500,
                    detail=f"Question {q.id} has no question_blocks"
                )

            sanitized_blocks = sanitize_question_blocks(q.question_blocks)

            if not sanitized_blocks:
                raise HTTPException(
                    status_code=500,
                    detail=f"Question {q.id} has no valid question_blocks after sanitation"
                )

            questions.append({
                "q_id": q_id,
                "topic": topic_name,
                "question_blocks": sanitized_blocks,
                "options": normalize_options(q.options),
                "correct": q.correct_answer
            })
            q_id += 1

        # ---- AI questions
        if ai_count > 0:
            ai_questions = generate_ai_questions_strict_Mathematical_Reasoning(
                quiz=quiz,
                topic_name=topic_name,
                ai_count=ai_count,
                db=db
            )

            for item in ai_questions:
                questions.append({
                    "q_id": q_id,
                    "topic": topic_name,
                    "question_blocks": [
                        {
                            "type": "text",
                            "content": clean_question_text(item["question"])
                        }
                    ],
                    "options": normalize_options(item["options"]),
                    "correct": item["correct"]
                })
                q_id += 1

    # --------------------------------------------------
    # 4️⃣ FINAL TOTAL VALIDATION (NON-NEGOTIABLE)
    # --------------------------------------------------
    expected_total = sum(
        int(t.get("db", 0)) + int(t.get("ai", 0))
        for t in quiz.topics
    )

    actual_total = len(questions)

    if actual_total != expected_total:
        raise HTTPException(
            status_code=500,
            detail=(
                f"Exam generation failed. "
                f"Expected {expected_total} questions, "
                f"but generated {actual_total}. "
                f"This indicates a generation inconsistency."
            )
        )

    # --------------------------------------------------
    # 5️⃣ Clear previous exams (FULL WIPE)
    # --------------------------------------------------
    # 1️⃣ MR-specific responses (DEEPEST CHILD)
    db.query(StudentExamResponseMathematicalReasoning).delete(
        synchronize_session=False
    )
    
    # 2️⃣ MR exam attempts
    db.query(StudentExamMathematicalReasoning).delete(
        synchronize_session=False
    )
    
    # 3️⃣ Aggregated MR results
    db.query(StudentExamResultsMathematicalReasoning).delete(
        synchronize_session=False
    )
    
    # 4️⃣ Generic student exams (if applicable)
    db.query(StudentExam).delete(synchronize_session=False)
    
    # 5️⃣ Parent exams LAST
    db.query(Exam).filter(
        Exam.subject == "mathematical_reasoning"
    ).delete(synchronize_session=False)
    
    db.commit()

    # --------------------------------------------------
    # 6️⃣ Save exam
    # --------------------------------------------------
    new_exam = Exam(
        quiz_id=quiz.id,
        class_name=quiz.class_name,
        subject=quiz.subject,
        difficulty=quiz.difficulty,
        questions=questions,
    )

    db.add(new_exam)
    db.commit()
    db.refresh(new_exam)

    # --------------------------------------------------
    # 7️⃣ Response
    # --------------------------------------------------
    return {
        "message": "Exam generated successfully",
        "exam_id": new_exam.id,
        "quiz_id": quiz.id,
        "class_name": quiz.class_name,
        "subject": quiz.subject,
        "difficulty": quiz.difficulty,
        "total_questions": len(questions),
        "questions": questions,
    }


@app.post("/api/admin/bulk-users-exam-module")
async def bulk_users_exam_module(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    print("📥 Bulk user upload request received")
    print(f"📄 Uploaded filename: {file.filename}")

    if not file.filename.lower().endswith(".csv"):
        raise HTTPException(status_code=400, detail="Only .csv files are supported")

    try:
        contents = await file.read()
        df = pd.read_csv(BytesIO(contents))
    except Exception:
        raise HTTPException(status_code=400, detail="Failed to read CSV file")

    required_columns = {
        "student_id",
        "password",
        "name",
        "parent_email",
        "class_name",
        "class_day",
    }

    if not required_columns.issubset(df.columns):
        missing = required_columns - set(df.columns)
        raise HTTPException(
            status_code=400,
            detail=f"Missing required columns: {', '.join(missing)}",
        )

    # 🔑 Get max numeric ID (ignore non-numeric / legacy UUIDs)
    last_id_int = (
        db.query(func.max(cast(Student.id, Integer)))
        .filter(Student.id.op("~")("^[0-9]+$"))
        .scalar()
    ) or 0

    next_id = last_id_int + 1

    print(f"🔢 Last ID in DB (numeric): {last_id_int}")
    print(f"🆕 Starting next_id: {next_id}")

    success = 0
    failed = 0
    errors = []

    for index, row in df.iterrows():
        print(
            f"➡️ Row {index + 2} | Assigning id={next_id}, "
            f"student_id={row['student_id']}"
        )

        try:
            student = Student(
                id=str(next_id),  # STRING id, sequential
                student_id=str(row["student_id"]).strip(),
                password=str(row["password"]).strip(),
                name=str(row["name"]).strip(),
                parent_email=str(row["parent_email"]).strip(),
                class_name=str(row["class_name"]).strip(),
                class_day=(
                    str(row["class_day"]).strip()
                    if not pd.isna(row["class_day"])
                    else None
                ),
            )

            db.add(student)
            db.commit()
            db.refresh(student)

            success += 1

        except IntegrityError:
            db.rollback()
            failed += 1
            errors.append(
                {
                    "row": index + 2,
                    "student_id": row["student_id"],
                    "error": "Integrity error (likely duplicate id or student_id)",
                }
            )

        except Exception as e:
            db.rollback()
            failed += 1
            errors.append(
                {
                    "row": index + 2,
                    "student_id": row["student_id"],
                    "error": str(e),
                }
            )

        finally:
            next_id += 1   # ✅ ALWAYS increment

    print("📊 Upload summary")
    print(f"   ✅ Success: {success}")
    print(f"   ❌ Failed: {failed}")

    print("❌ BULK UPLOAD ERRORS:")
    for err in errors:
        print(err)

    return {
        "success": success,
        "failed": failed,
        "errors": errors,
    }

@app.get("/api/topics-exam-setup", response_model=List[str])
def get_distinct_topics_exam_setup(
    class_name: str = Query(...),
    subject: str = Query(...),
    difficulty: str = Query(...),
    db: Session = Depends(get_db)
):
    print("\n================ TOPICS EXAM SETUP =================")
    print("class_name:", class_name)
    print("subject:", subject)
    print("difficulty:", difficulty)

    topics = (
        db.query(distinct(Question.topic))
        .filter(
            func.lower(Question.class_name) == class_name.lower(),
            func.lower(Question.subject) == subject.lower(),
            func.lower(Question.difficulty) == difficulty.lower(),
            Question.topic.isnot(None),
            Question.topic != ""
        )
        .order_by(Question.topic)
        .all()
    )

    return [t[0] for t in topics]


@app.get("/api/admin/students/{student_id}/selective-report-dates")
def get_selective_report_dates(
    student_id: str,
    db: Session = Depends(get_db)
):
    rows = (
        db.query(func.date(AdminExamReport.created_at))
        .filter(AdminExamReport.student_id == student_id)
        .distinct()
        .order_by(func.date(AdminExamReport.created_at).desc())
        .all()
    )

    return [r[0].isoformat() for r in rows]


@app.post("/api/admin/students/{student_id}/overall-selective-report")
def generate_overall_selective_report(
    student_id: str,
    exam_date: date,
    db: Session = Depends(get_db)
):
    """
    Generate (once) or fetch Overall Selective Readiness Report
    for a student on a given exam date.
    """

    # --------------------------------------------------
    # 1️⃣ Guard: already exists?
    # --------------------------------------------------
    existing = (
        db.query(AdminOverallSelectiveReport)
        .filter(
            AdminOverallSelectiveReport.student_id == student_id,
            AdminOverallSelectiveReport.exam_date == exam_date
        )
        .first()
    )

    if existing:
        return existing

    # --------------------------------------------------
    # 2️⃣ Fetch per-exam admin reports for that date
    # --------------------------------------------------
    reports = (
        db.query(AdminExamReport)
        .filter(
            AdminExamReport.student_id == student_id,
            func.date(AdminExamReport.created_at) == exam_date
        )
        .all()
    )
    if not reports:
       raise HTTPException(
           status_code=404,
           detail={
               "code": "NO_EXAMS_FOUND",
               "message": "No exam reports exist for this student on the selected exam date.",
               "student_id": student_id,
               "exam_date": str(exam_date)
           }
       )


    
    # --------------------------------------------------
    # 3️⃣ Normalize exam types (🔥 FIX)
    # --------------------------------------------------
    def normalize_exam_type(raw: str) -> str:
        raw = raw.lower().strip()
        if raw == "thinking skills":
            return "thinking_skills"
        return raw

    reports_by_subject = {
        normalize_exam_type(r.exam_type): r for r in reports
    }

    required_subjects = {
        "reading",
        "mathematical_reasoning",
        "thinking_skills",
        "writing"
    }

    missing = required_subjects - reports_by_subject.keys()

    if missing:
        raise HTTPException(
            status_code=409,
            detail={
                "code": "INCOMPLETE_EXAMS",
                "message": "Overall readiness cannot be generated because not all required exams are completed.",
                "missing_subjects": sorted(list(missing)),
                "completed_subjects": sorted(list(reports_by_subject.keys())),
                "student_id": student_id,
                "exam_date": str(exam_date)
            }
        )


    
    # --------------------------------------------------
    # 4️⃣ Compute normalized overall score (FIXED)
    # --------------------------------------------------
    MAX_SCORES = {
        "reading": 100,
        "mathematical_reasoning": 100,
        "thinking_skills": 100,
        "writing": 20,
    }
    
    # Raw scores (keep for overrides + storage)
    components = {
        subject: reports_by_subject[subject].overall_score
        for subject in MAX_SCORES
    }
    
    # Normalize all components to percentage
    normalized_components = {
        subject: round(
            (components[subject] / MAX_SCORES[subject]) * 100,
            2
        )
        for subject in components
    }
    
    # Equal-weight average of normalized scores
    overall_percent = round(
        sum(normalized_components.values()) / len(normalized_components),
        2
    )

    # --------------------------------------------------
    # 5️⃣ Map to readiness band
    # --------------------------------------------------
    if overall_percent >= 80:
        band = "Band 1 – Fully Selective Ready"
    elif overall_percent >= 70:
        band = "Band 2 – Strong Selective Potential"
    elif overall_percent >= 60:
        band = "Band 3 – Borderline Selective"
    else:
        band = "Band 4 – Not Selective Ready Yet"

    # --------------------------------------------------
    # 6️⃣ School recommendation matrix
    # --------------------------------------------------
    school_map = {
        "Band 1 – Fully Selective Ready": [
            "James Ruse",
            "Baulkham Hills",
            "North Sydney Boys/Girls",
            "Sydney Grammar (Academic profile)"
        ],
        "Band 2 – Strong Selective Potential": [
            "Hornsby Girls",
            "Chatswood",
            "Ryde",
            "Sydney Technical",
            "Parramatta",
            "Penrith Selective"
        ],
        "Band 3 – Borderline Selective": [
            "Local / partially selective schools",
            "Lower competition selective schools"
        ],
        "Band 4 – Not Selective Ready Yet": [
            "Focus on foundations",
            "Reassess selective pathway",
            "Consider enrichment before exam prep"
        ]
    }

    school_recommendation = school_map[band]

    # --------------------------------------------------
    # 7️⃣ Override rules (CRITICAL SAFETY)
    # --------------------------------------------------
    override_flag = False
    override_message = None

    if components["writing"] < 60:
        override_flag = True
        override_message = (
            "While the overall score is competitive, improvement in Writing "
            "is required for higher-tier selective schools."
        )

    for subject, score in components.items():
        if score < 55:
            override_flag = True
            override_message = (
                f"While the overall score is competitive, improvement in "
                f"{subject.replace('_', ' ').title()} is required for higher-tier selective schools."
            )
            break

    # --------------------------------------------------
    # 8️⃣ Store snapshot (IMMUTABLE)
    # --------------------------------------------------
    overall_report = AdminOverallSelectiveReport(
        student_id=student_id,
        exam_date=exam_date,
        overall_percent=overall_percent,
        readiness_band=band,
        school_recommendation=school_recommendation,
        override_flag=override_flag,
        override_message=override_message,
        components=components
    )

    db.add(overall_report)
    db.commit()
    db.refresh(overall_report)

    return overall_report


@app.get("/api/admin/students/{student_id}/selective-reports")
def get_student_selective_reports(
    student_id: str,
    exam_date: date,   # ✅ REQUIRED
    db: Session = Depends(get_db)
): 
    """
    Fetch all selective exam reports for a given student.
    Read-only aggregation over admin reporting tables.
    """

    # --------------------------------------------------
    # 1️⃣ Fetch admin exam reports for student
    # --------------------------------------------------
    reports = (
        db.query(AdminExamReport)
        .filter(
            AdminExamReport.student_id == student_id,
            func.date(AdminExamReport.created_at) == exam_date
        )
        .order_by(AdminExamReport.created_at.desc())
        .all()
    )

    if not reports:
        return []

    report_ids = [r.id for r in reports]

    # --------------------------------------------------
    # 2️⃣ Fetch section results
    # --------------------------------------------------
    sections = (
        db.query(AdminExamSectionResult)
        .filter(AdminExamSectionResult.admin_report_id.in_(report_ids))
        .all()
    )

    sections_by_report = defaultdict(list)
    for s in sections:
        sections_by_report[s.admin_report_id].append({
            "section_name": s.section_name,
            "raw_score": s.raw_score,
            "performance_band": s.performance_band,
            "strengths": s.strengths_summary,
            "improvements": s.improvement_summary
        })

    # --------------------------------------------------
    # 3️⃣ Fetch readiness rules (audit trail)
    # --------------------------------------------------
    rules = (
        db.query(AdminReadinessRuleApplied)
        .filter(AdminReadinessRuleApplied.admin_report_id.in_(report_ids))
        .all()
    )

    rules_by_report = defaultdict(list)
    for r in rules:
        rules_by_report[r.admin_report_id].append({
            "rule_code": r.rule_code,
            "result": r.rule_result,
            "description": r.rule_description
        })

    # --------------------------------------------------
    # 4️⃣ Shape final response
    # --------------------------------------------------
    response = []

    for report in reports:
        response.append({
            "id": report.id,                          # 🔑 REQUIRED
            "exam_type": report.exam_type,
            "exam_attempt_id": report.exam_attempt_id,
            "overall_score": report.overall_score,
            "readiness_band": report.readiness_band,
            "school_guidance_level": report.school_guidance_level,
            "summary_notes": report.summary_notes,
            "exam_date": report.created_at.date().isoformat(),  # 🔑 REQUIRED
            "disclaimer": (
                "This report is advisory only and does not guarantee placement."
            ),
    
            # ✅ sections are ALREADY dicts → just pass them through
            "sections": sections_by_report.get(report.id, []),
    
            # ✅ readiness rules are ALREADY dicts
            "readiness_rules": rules_by_report.get(report.id, [])
        })


    return response
@app.get("/api/exams/dates")
def get_exam_dates(
    exam: str,
    student_id: str | None = None,
    db: Session = Depends(get_db),
):
    """
    Returns available exam dates for:
    - Per Student Report (exam + student_id)
    - Per Class Report (exam only)
    """

    # Base query: filter by exam
    query = (
        db.query(AdminExamReport.created_at)
        .filter(AdminExamReport.exam_type == exam)
    )

    # Per Student Report: further restrict by student
    if student_id:
        query = query.filter(AdminExamReport.student_id == student_id)

    # Execute query
    rows = (
        query
        .distinct()
        .order_by(AdminExamReport.created_at.desc())
        .all()
    )

    return {
        "dates": [
            row.created_at.date().isoformat()
            for row in rows
        ]
    }
 

@app.get("/api/admin/students")
def get_admin_students(db: Session = Depends(get_db)):
    students = (
        db.query(Student.student_id, Student.name)
        .order_by(Student.student_id)
        .all()
    )

    return [
        {
            "student_id": student_id,
            "name": name
        }
        for student_id, name in students
    ]


@app.get("/api/admin/students/{student_id}")
def get_student_details(student_id: str, db: Session = Depends(get_db)):
    student = (
        db.query(Student)
        .filter(Student.student_id == student_id)
        .first()
    )

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    return {
        "student_id": student.student_id,
        "name": student.name,
        "class_name": student.class_name,
        "class_day": student.class_day,
        "parent_email": student.parent_email,
    }
@app.get("/students/{student_id}/selective-reports")
def get_student_selective_reports(
    student_id: str,
    db: Session = Depends(get_db)
):
    """
    Return all selective readiness reports for a student.
    Reports are immutable snapshots.
    """

    reports = (
        db.query(AdminExamReport)
        .filter(AdminExamReport.student_id == student_id)
        .order_by(AdminExamReport.created_at.desc())
        .all()
    )

    response = []

    for report in reports:
        sections = (
            db.query(AdminExamSectionResult)
            .filter(
                AdminExamSectionResult.exam_attempt_id == report.exam_attempt_id
            )
            .order_by(AdminExamSectionResult.section_name)
            .all()
        )

        response.append({
            "id": report.id,
            "exam_date": report.created_at.date().isoformat(),
            "readiness_band": report.readiness_band,
            "school_guidance_level": report.school_guidance_level,
            "sections": [
                {
                    "section_name": s.section_name,
                    "performance_band": s.performance_band,
                    "strengths_summary": s.strengths_summary,
                    "improvement_summary": s.improvement_summary,
                }
                for s in sections_by_report.get(report.id, [])
            ],
            "disclaimer": (
                "This report is advisory only and does not guarantee placement."
            )
        })

    return response
 
@app.delete("/delete_student_exam_module/{id}")
def delete_student_exam_module(
    id: str,
    db: Session = Depends(get_db)
):
    print("\n================ DELETE STUDENT (EXAM MODULE) ================")
    print("➡ student id:", id)

    student = db.query(Student).filter(Student.id == id).first()

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    try:
        # -----------------------------
        # 0️⃣ READING MODULE (NO FK)
        # -----------------------------
        db.query(StudentExamReportReading)\
            .filter(StudentExamReportReading.student_id == id)\
            .delete(synchronize_session=False)

        db.query(StudentExamReading)\
            .filter(StudentExamReading.student_id == id)\
            .delete(synchronize_session=False)

        # -----------------------------
        # 1️⃣ RESPONSE tables
        # -----------------------------
        db.query(StudentExamResponseThinkingSkills)\
            .filter(StudentExamResponseThinkingSkills.student_id == id)\
            .delete(synchronize_session=False)

        db.query(StudentExamResponseMathematicalReasoning)\
            .filter(StudentExamResponseMathematicalReasoning.student_id == id)\
            .delete(synchronize_session=False)

        # -----------------------------
        # 2️⃣ RESULT tables
        # -----------------------------
        db.query(StudentExamResultsThinkingSkills)\
            .filter(StudentExamResultsThinkingSkills.student_id == id)\
            .delete(synchronize_session=False)

        db.query(StudentExamResultsMathematicalReasoning)\
            .filter(StudentExamResultsMathematicalReasoning.student_id == id)\
            .delete(synchronize_session=False)

        # -----------------------------
        # 3️⃣ EXAM attempts
        # -----------------------------
        db.query(StudentExamThinkingSkills)\
            .filter(StudentExamThinkingSkills.student_id == id)\
            .delete(synchronize_session=False)

        db.query(StudentExamMathematicalReasoning)\
            .filter(StudentExamMathematicalReasoning.student_id == id)\
            .delete(synchronize_session=False)

        # -----------------------------
        # 4️⃣ STUDENT
        # -----------------------------
        db.delete(student)
        db.commit()

    except Exception as e:
        db.rollback()
        print("❌ Delete failed:", e)
        raise HTTPException(
            status_code=500,
            detail="Failed to delete student and related data"
        )

    return {
        "message": "Student and all related data deleted successfully",
        "deleted_id": id
    }

def generate_ai_questions_strict_Mathematical_Reasoning(
    *,
    quiz,
    topic_name,
    ai_count,
    db,
    max_attempts=6,
    batch_size=3
):
    collected = []
    attempts = 0

    location = db.query(FranchiseLocation).first()
    if not location:
        raise HTTPException(500, "No franchise location found")

    while len(collected) < ai_count and attempts < max_attempts:
        remaining = ai_count - len(collected)
        request_count = min(batch_size, remaining)
        attempts += 1

        print(
            f"[AI GEN] Attempt {attempts} | "
            f"Requesting {request_count} | "
            f"Collected so far: {len(collected)}"
        )

        system_prompt = (
            "You are an expert exam generator.\n"
            f"Create {request_count} MCQs.\n\n"
            f"Class: {quiz.class_name}\n"
            f"Subject: {quiz.subject}\n"
            f"Topic: {topic_name}\n"
            f"Country: {location.country}\n"
            f"State: {location.state}\n\n"
            "Rules:\n"
            "- Exactly one correct option\n"
            "- 4 options per question\n\n"
            "Return STRICT JSON ONLY:\n"
            "[{\"question\":\"...\",\"options\":[\"A\",\"B\",\"C\",\"D\"],\"correct\":\"A\"}]"
        )

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": system_prompt}],
            temperature=0.4
        )

        raw_output = response.choices[0].message.content.strip()

        try:
            batch = json.loads(raw_output)
        except Exception as e:
            print("[AI ERROR] JSON parse failed:", str(e))
            continue

        valid = [
            q for q in batch
            if isinstance(q.get("question"), str)
            and isinstance(q.get("options"), list)
            and len(q["options"]) >= 4
            and q.get("correct") in ["A", "B", "C", "D"]
        ]

        print(f"[AI GEN] Valid questions received: {len(valid)}")
        collected.extend(valid)

    if len(collected) < ai_count:
        raise HTTPException(
            status_code=500,
            detail=(
                f"AI failed to generate required questions for topic "
                f"'{topic_name}'. Expected {ai_count}, got {len(collected)}."
            )
        )

    return collected[:ai_count]



def clean_question_text(raw: str) -> str:
    if not raw:
        return raw

    text = raw.strip()

    if "QUESTION_TEXT:" in text:
        text = text.split("QUESTION_TEXT:", 1)[1]

    # Remove OPTIONS and answers
    text = re.split(r"\n\s*OPTIONS\s*:\s*\n", text, flags=re.IGNORECASE)[0]
    text = re.split(r"\n\s*CORRECT_ANSWER\s*:", text, flags=re.IGNORECASE)[0]

    return text.strip()

def normalize_options(raw_options):
    """
    Convert options into frontend-safe format.
    Keeps text options as strings.
    Skips unsupported image-options for now (safe fallback).
    """
    if not raw_options:
        return []

    normalized = []

    if isinstance(raw_options, list):
        for opt in raw_options:
            # Skip legacy embedded image dict strings
            if isinstance(opt, str) and "{'type': 'image'" in opt:
                continue
            normalized.append(opt)
        return normalized

    if isinstance(raw_options, dict):
        for key, value in raw_options.items():
            if isinstance(value, dict):
                if value.get("type") == "text":
                    normalized.append(f"{key}) {value.get('content', '')}")
                elif value.get("type") == "image":
                    # 🔒 Do not leak broken image-options yet
                    normalized.append(f"{key}) [Image option]")
            else:
                normalized.append(f"{key}) {value}")

        return normalized

    return []

GCS_IMAGE_PREFIX = "https://storage.googleapis.com/exammoduleimages/"

def sanitize_question_blocks(blocks):
    """
    HARD GUARANTEES:
    - Preserves question order
    - Preserves full image URLs (UUID-safe)
    - Rejects legacy IMAGES: filename references
    - Removes leaked answer hints (A. / B) / E.)
    - Never reconstructs or guesses image URLs
    """

    if not isinstance(blocks, list):
        return []

    sanitized = []

    for idx, block in enumerate(blocks):
        if not isinstance(block, dict):
            continue

        block_type = block.get("type")

        # --------------------------------------------------
        # TEXT BLOCK
        # --------------------------------------------------
        if block_type == "text":
            content = (block.get("content") or "").strip()

            if not content:
                continue

            # 🔥 Strip leaked answers (A. ..., B) ..., E. ...)
            if re.match(r"^[A-E][\.\)]\s*", content):
                continue

            # 🚨 HARD FAIL on legacy image references
            legacy_img_match = re.match(
                r"IMAGES?\s*\d*\s*:\s*(.+)",
                content,
                flags=re.IGNORECASE
            )

            if legacy_img_match:
                raise HTTPException(
                    status_code=500,
                    detail=(
                        "Legacy image reference detected in question_blocks. "
                        f"Found text-based image reference: '{content}'. "
                        "All images must be stored as explicit image blocks "
                        "with full GCS URLs. Please migrate this question."
                    )
                )

            sanitized.append({
                "type": "text",
                "content": content
            })

        # --------------------------------------------------
        # IMAGE BLOCK
        # --------------------------------------------------
        elif block_type == "image":
            src = block.get("src")

            if not src or not isinstance(src, str):
                raise HTTPException(
                    status_code=500,
                    detail=(
                        f"Invalid image block detected (missing src). "
                        f"Block index: {idx}"
                    )
                )

            # 🔒 Enforce FULL GCS URL
            if not src.startswith(GCS_IMAGE_PREFIX):
                raise HTTPException(
                    status_code=500,
                    detail=(
                        "Invalid image URL detected. "
                        f"Image src must be a full GCS URL. Found: {src}"
                    )
                )

            sanitized.append({
                "type": "image",
                "src": src
            })

        # --------------------------------------------------
        # UNKNOWN BLOCK TYPE
        # --------------------------------------------------
        else:
            raise HTTPException(
                status_code=500,
                detail=(
                    f"Unknown question block type '{block_type}' "
                    f"at index {idx}. Allowed types are 'text' and 'image'."
                )
            )

    if not sanitized:
        raise HTTPException(
            status_code=500,
            detail="Question has no valid question_blocks after sanitization."
        )

    return sanitized




@app.post("/api/exams/generate-thinking-skills")
def generate_thinking_skills_exam(
    payload: Optional[dict] = Body(default=None),
    db: Session = Depends(get_db)
):
    
    """
    Generate a Thinking Skills exam based on difficulty only.
    """
    print("DATABASE URL:", db.bind.url)
    print("QUIZ COUNT:", db.query(Quiz).count())
    print("QUIZZES:", db.query(Quiz.id, Quiz.subject, Quiz.difficulty).all())
    result = db.execute(text("select current_schema()")).scalar()
    print("SCHEMA:", result)
    
    result = db.execute(text("select count(*) from quizzes")).scalar()
    print("RAW QUIZ COUNT:", result)


    payload = payload or {}
    # --------------------------------------------------
    # 0️⃣ Clear previous Thinking Skills exams
    # --------------------------------------------------
    # 1️⃣ Delete student exams for thinking skills
    # ✅ STEP 1: delete dependent StudentExam rows (SAFE)
    exam_ids_subq = select(Exam.id).where(
        Exam.subject == "thinking_skills"
    )

    
    db.query(StudentExam).filter(
        StudentExam.exam_id.in_(exam_ids_subq)
    ).delete(synchronize_session=False)
    
    # ✅ STEP 2: delete Exam rows
    db.query(Exam).filter(
        Exam.subject == "thinking_skills"
    ).delete(synchronize_session=False)
    
    db.commit()


    # --------------------------------------------------
    # 1️⃣ Fetch latest Thinking Skills quiz for difficulty
    # --------------------------------------------------
    # --------------------------------------------------
    # 1️⃣ Fetch latest Thinking Skills quiz (difficulty optional)
    # --------------------------------------------------
    
    query = db.query(Quiz).filter(
        Quiz.subject == "thinking_skills"
    )
    
    
    
    quiz = query.order_by(Quiz.id.desc()).first()
    
    if not quiz:
        raise HTTPException(
            status_code=404,
            detail="No Thinking Skills quiz found in database"
        )

    # --------------------------------------------------
    # 2️⃣ Generate exam questions
    # --------------------------------------------------
    try:
        questions = generate_exam_questions(quiz, db)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate Thinking Skills exam: {str(e)}"
        )

    if not questions:
        raise HTTPException(
            status_code=500,
            detail="No questions generated for Thinking Skills exam"
        )

    # --------------------------------------------------
    # 3️⃣ Save generated exam
    # --------------------------------------------------
    new_exam = Exam(
        quiz_id=quiz.id,
        class_name=quiz.class_name,
        subject=quiz.subject,
        difficulty=quiz.difficulty,
        questions=questions
    )

    db.add(new_exam)
    db.commit()
    db.refresh(new_exam)

    # --------------------------------------------------
    # 4️⃣ Return response
    # --------------------------------------------------
    return {
        "message": "Thinking Skills exam generated successfully",
        "exam_id": new_exam.id,
        "quiz_id": quiz.id,
        "class_name": quiz.class_name,
        "subject": quiz.subject,
        "difficulty": quiz.difficulty,
        "total_questions": len(questions),
        "questions": questions
    }

@app.post("/api/generate-exam-mathematical-reasoning")
def generate_exam_mathematical_reasoning(
    db: Session = Depends(get_db)
):
    """
    Generate Mathematical Reasoning exam based on latest quiz configuration
    """

    # --------------------------------------------------
    # 1️⃣ Fetch latest Mathematical Reasoning quiz
    # --------------------------------------------------
    quiz = (
        db.query(Quiz)
        .filter(Quiz.subject == "mathematical_reasoning")
        .order_by(Quiz.id.desc())
        .first()
    )

    if not quiz:
        raise HTTPException(
            status_code=404,
            detail="No Mathematical Reasoning quiz configuration found"
        )

    # --------------------------------------------------
    # 2️⃣ Clear previous Mathematical Reasoning exams
    # --------------------------------------------------
    exam_ids_subq = (
        db.query(Exam.id)
        .filter(Exam.subject == "mathematical_reasoning")
        .subquery()
    )

    # Delete dependent student exam attempts first
    db.query(StudentExam).filter(
        StudentExam.exam_id.in_(exam_ids_subq)
    ).delete(synchronize_session=False)

    # Delete previous exams
    db.query(Exam).filter(
        Exam.subject == "mathematical_reasoning"
    ).delete(synchronize_session=False)

    db.commit()

    # --------------------------------------------------
    # 3️⃣ Generate exam questions
    # --------------------------------------------------
    try:
        questions = (quiz, db)
    except Exception as e:
        print("❌ Question generation failed:", e)
        raise HTTPException(
            status_code=500,
            detail="Failed to generate exam questions"
        )

    if not questions:
        raise HTTPException(
            status_code=500,
            detail="No questions generated"
        )

    # --------------------------------------------------
    # 4️⃣ Save generated exam
    # --------------------------------------------------
    new_exam = Exam(
        quiz_id=quiz.id,
        class_name=quiz.class_name,
        subject=quiz.subject,
        difficulty=quiz.difficulty,
        questions=questions
    )

    db.add(new_exam)
    db.commit()
    db.refresh(new_exam)

    # --------------------------------------------------
    # 5️⃣ Frontend-ready response
    # --------------------------------------------------
    return {
        "message": "Exam generated successfully",
        "exam_id": new_exam.id,
        "quiz_id": quiz.id,
        "class_name": quiz.class_name,
        "subject": quiz.subject,
        "difficulty": quiz.difficulty,
        "total_questions": len(questions),
        "questions": questions
    }


@app.post("/api/quizzes-naplan-language-conventions")
def create_naplan_language_conventions_quiz(
    quiz: NaplanQuizCreate,
    db: Session = Depends(get_db),
):
    print("\n========== NAPLAN LANGUAGE CONVENTIONS QUIZ CREATION START ==========")

    print("🔍 Incoming payload:", quiz)

    try:
        quiz_dict = quiz.dict()
        print("📦 Parsed quiz payload:", quiz_dict)
    except Exception as e:
        print("❌ Failed to parse payload:", e)
        traceback.print_exc()
        raise HTTPException(status_code=400, detail="Invalid quiz payload")

    print("➡️ class_name:", quiz.class_name)
    print("➡️ subject:", quiz.subject)
    print("➡️ year:", quiz.year)
    print("➡️ difficulty:", quiz.difficulty)
    print("➡️ num_topics:", quiz.num_topics)
    print("➡️ total_questions:", quiz.total_questions)

    if not isinstance(quiz.topics, list):
        raise HTTPException(status_code=400, detail="topics must be a list")

    print("📝 Topics:")
    for i, t in enumerate(quiz.topics):
        print(f"   └─ Topic {i + 1}: {t}")

    try:
        # 🔥 DELETE PREVIOUS QUIZ CONFIG(S)
        print("\n🧹 Deleting existing NAPLAN Language Conventions quiz configs...")

        deleted_count = (
            db.query(QuizNaplanLanguageConventions)
            .delete()
        )

        print(f"🗑️ Deleted {deleted_count} existing quiz config(s)")
        print("\n--- Creating quiz_naplan_language_conventions row ---")

        new_quiz = QuizNaplanLanguageConventions(
            class_name=quiz.class_name,
            subject=quiz.subject,  # expected: "Language Conventions"
            year=quiz.year,
            difficulty=quiz.difficulty,
            num_topics=quiz.num_topics,
            total_questions=quiz.total_questions,
            topics=[t.dict() for t in quiz.topics],
        )

        db.add(new_quiz)
        db.commit()
        db.refresh(new_quiz)

        print("✅ NAPLAN LANGUAGE CONVENTIONS QUIZ CREATED:", new_quiz)
        print("========== NAPLAN LANGUAGE CONVENTIONS QUIZ CREATION COMPLETE ==========\n")

        return {
            "message": "NAPLAN Language Conventions quiz created successfully",
            "quiz_id": new_quiz.id,
        }

    except Exception as e:
        print("\n❌ EXCEPTION DURING QUIZ CREATION ❌")
        print("Error:", str(e))
        traceback.print_exc()
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))



@app.post("/api/quizzes-naplan-numeracy")
def create_naplan_numeracy_quiz(
    quiz: NaplanQuizCreate,
    db: Session = Depends(get_db),
):
    print("\n========== NAPLAN QUIZ CREATION START ==========")

    print("🔍 Incoming payload:", quiz)

    try:
        quiz_dict = quiz.dict()
        print("📦 Parsed quiz payload:", quiz_dict)
    except Exception as e:
        print("❌ Failed to parse payload:", e)
        traceback.print_exc()
        raise HTTPException(status_code=400, detail="Invalid quiz payload")

    print("➡️ class_name:", quiz.class_name)
    print("➡️ subject:", quiz.subject)
    print("➡️ year:", quiz.year)
    print("➡️ difficulty:", quiz.difficulty)
    print("➡️ num_topics:", quiz.num_topics)
    print("➡️ total_questions:", quiz.total_questions)

    if not isinstance(quiz.topics, list):
        raise HTTPException(status_code=400, detail="topics must be a list")

    print("📝 Topics:")
    for i, t in enumerate(quiz.topics):
        print(f"   └─ Topic {i + 1}: {t}")

    try:
         # 🔥 DELETE PREVIOUS QUIZ CONFIG(S)
        print("\n🧹 Deleting existing NAPLAN Numeracy quiz configs...")

        deleted_count = (
            db.query(QuizNaplanNumeracy)
            .delete()
        )

        print(f"🗑️ Deleted {deleted_count} existing quiz config(s)")
        print("\n--- Creating quizzes_naplan_numeracy row ---")

        new_quiz = QuizNaplanNumeracy(
            class_name=quiz.class_name,
            subject=quiz.subject,
            year=quiz.year,
            difficulty=quiz.difficulty,
            num_topics=quiz.num_topics,
            total_questions=quiz.total_questions,
            topics=[t.dict() for t in quiz.topics],  # JSON column
        )

        db.add(new_quiz)
        db.commit()
        db.refresh(new_quiz)

        print("✅ NAPLAN QUIZ CREATED:", new_quiz)
        print("========== NAPLAN QUIZ CREATION COMPLETE ==========\n")

        return {
            "message": "NAPLAN Numeracy quiz created successfully",
            "quiz_id": new_quiz.id,
        }

    except Exception as e:
        print("\n❌ EXCEPTION DURING QUIZ CREATION ❌")
        print("Error:", str(e))
        traceback.print_exc()
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/quizzes/thinking-skills/difficulty")
def get_thinking_skills_difficulty(db: Session = Depends(get_db)):
    """
    Return the difficulty level for the latest Thinking Skills quiz.
    """

    quiz = (
        db.query(Quiz)
        .filter(Quiz.subject == "thinking_skills")
        .order_by(Quiz.created_at.desc())
        .first()
    )

    if not quiz:
        raise HTTPException(
            status_code=404,
            detail="No Thinking Skills quiz found"
        )

    return quiz.difficulty

@app.post("/api/quizzes/mathematical-reasoning")
def create_quiz_mathematical_reasoning(
    quiz: QuizMathematicalReasoningCreate,
    db: Session = Depends(get_db)
):
    """
    Create a Mathematical Reasoning quiz (admin specification).
    Stored in quiz_mathematical_reasoning table.
    """

    print("\n========== MATHEMATICAL REASONING QUIZ CREATION START ==========")

    print("🔍 Incoming payload:", quiz)

    # Convert to dict for debugging
    try:
        quiz_dict = quiz.dict()
        print("📦 Parsed quiz payload:", quiz_dict)
    except Exception as e:
        print("❌ Invalid payload")
        raise HTTPException(status_code=400, detail=str(e))

    # Enforce subject correctness
    if quiz.subject != "mathematical_reasoning":
        raise HTTPException(
            status_code=400,
            detail="Invalid subject. Expected 'mathematical_reasoning'."
        )

    # Validate topics
    if not isinstance(quiz.topics, list):
        raise HTTPException(
            status_code=400,
            detail="topics must be a list"
        )

    print("📝 Topics count:", len(quiz.topics))
    for i, t in enumerate(quiz.topics):
        print(f"   └─ Topic {i}: {t}")

    try:
        print("\n--- Deleting existing Mathematical Reasoning quizzes ---")

        db.query(QuizMathematicalReasoning).delete()
        db.commit()

        print("🗑️ All previous quizzes deleted")
     
        print("\n--- Creating SQLAlchemy object ---")

        new_quiz = QuizMathematicalReasoning(
            class_name=quiz.class_name.strip(),
            subject="mathematical_reasoning",  # forced
            difficulty=quiz.difficulty.strip(),
            num_topics=quiz.num_topics,
            topics=[t.dict() for t in quiz.topics]
        )

        db.add(new_quiz)
        db.commit()
        db.refresh(new_quiz)

        print("✅ Quiz saved with ID:", new_quiz.id)
        print("========== QUIZ CREATION COMPLETE ==========\n")

        return {
            "message": "Mathematical Reasoning quiz created successfully",
            "quiz_id": new_quiz.id
        }

    except Exception as e:
        print("❌ DB error:", str(e))
        db.rollback()
        raise HTTPException(
            status_code=500,
            detail="Error creating Mathematical Reasoning quiz"
        )

@app.post("/api/quizzes/mathematical-reasoning")
def create_topic_config_mathematical_reasoning(
    payload: TopicConfigMathematicalReasoningCreate,
    db: Session = Depends(get_db)
):
    """
    Create Mathematical Reasoning topic configuration.
    """

    print("\n========== MATHEMATICAL REASONING TOPIC CONFIG ==========")

    try:
        # Debug incoming payload
        payload_dict = payload.dict()
        print("📦 Incoming payload:", payload_dict)

        # Validate subject explicitly
        if payload.subject != "mathematical_reasoning":
            raise HTTPException(
                status_code=400,
                detail="Invalid subject. Expected 'mathematical_reasoning'."
            )

        # Create DB record
        quiz_config = QuizMathematicalReasoning(
            class_name=payload.class_name.strip(),
            subject=payload.subject.strip(),
            difficulty=payload.difficulty.strip(),
            num_topics=payload.num_topics,
            topics=[t.dict() for t in payload.topics]
        )

        db.add(quiz_config)
        db.commit()
        db.refresh(quiz_config)


        print(
            "✅ Saved topicConfig_mathematical_reasoning | ID:",
            quiz_config.id
        )
        print("=======================================================\n")

        return {
            "message": "Mathematical Reasoning topic config saved successfully",
            "topic_config_id": quiz_config.id
        }

    except HTTPException:
        raise

    except Exception as e:
        print("❌ ERROR saving Mathematical Reasoning topic config")
        print("Error:", str(e))
        traceback.print_exc()

        db.rollback()
        raise HTTPException(
            status_code=500,
            detail="Error saving Mathematical Reasoning topic configuration"
        )


@app.get("/get_all_students_exam_module")
def get_all_students_exam_module(
    db: Session = Depends(get_db)
):
    students = db.query(Student).order_by(Student.student_id.asc()).all()

    return [
        {
            "id": s.id,
            "student_id": s.student_id,
            "name": s.name,
            "parent_email": s.parent_email,
            "class_name": s.class_name,
            "class_day": s.class_day
            # ❌ password intentionally excluded
        }
        for s in students
    ]
 
@app.put("/edit_student_exam_module")
def edit_student_exam_module(
    payload: UpdateStudentRequest,
    db: Session = Depends(get_db)
):
    student = (
        db.query(Student)
        .filter(Student.student_id == payload.student_id)
        .first()
    )

    if not student:
        raise HTTPException(
            status_code=404,
            detail="Student not found"
        )

    # --------------------------------------------------
    # Update fields ONLY if provided
    # --------------------------------------------------
    if payload.name is not None:
        student.name = payload.name

    if payload.parent_email is not None:
        student.parent_email = payload.parent_email

    if payload.class_name is not None:
        student.class_name = payload.class_name

    if payload.class_day is not None:
        student.class_day = payload.class_day

    if payload.password is not None:
        student.password = payload.password  # plain text (intentional)

    db.commit()
    db.refresh(student)

    return {
        "message": "Student updated successfully",
        "student_id": student.student_id
    }
 
@app.get("/api/exams/reading-report")
def get_reading_report(
    session_id: int = Query(..., description="Exam session ID"),
    db: Session = Depends(get_db)
):
    print("\n================ GET READING REPORT ================")
    print("🆔 Incoming session_id:", session_id)

    # --------------------------------------------------
    # 1️⃣ Load session
    # --------------------------------------------------
    print("🟡 STEP 1: Loading StudentExamReading session")

    session = (
        db.query(StudentExamReading)
        .filter(StudentExamReading.id == session_id)
        .first()
    )

    if not session:
        print("❌ ERROR: Session not found for session_id =", session_id)
        raise HTTPException(status_code=404, detail="Session not found")

    print("✅ Session loaded:", {
        "session_id": session.id,
        "student_id": session.student_id,
        "exam_id": session.exam_id,
        "started_at": session.started_at,
        "finished": session.finished,
        "has_report_json": bool(session.report_json)
    })

    # --------------------------------------------------
    # 2️⃣ Finished guard
    # --------------------------------------------------
    print("🟡 STEP 2: Checking finished flag")

    if not session.finished:
        print("❌ ERROR: Session exists but exam is NOT finished")
        raise HTTPException(status_code=400, detail="Exam not finished yet")

    print("✅ Session is marked as finished")

    # --------------------------------------------------
    # 3️⃣ Happy path — report snapshot exists
    # --------------------------------------------------
    print("🟡 STEP 3: Checking report_json snapshot")

    if session.report_json:
        print("🟢 report_json FOUND — returning stored report")
        print("📦 Report top-level keys:", list(session.report_json.keys()))
        print("================ END GET READING REPORT ================\n")
        return session.report_json

    print("⚠️ report_json MISSING — entering diagnostic mode")

    # --------------------------------------------------
    # 4️⃣ Inspect per-question report rows
    # --------------------------------------------------
    print("🟡 STEP 4: Counting StudentExamReportReading rows")

    answer_count = (
        db.query(StudentExamReportReading)
        .filter(StudentExamReportReading.session_id == session.id)
        .count()
    )

    print("📝 StudentExamReportReading row count:", answer_count)

    # --------------------------------------------------
    # 5️⃣ Finished but no answers → abandoned / timeout
    # --------------------------------------------------
    if answer_count == 0:
        print("❌ DIAGNOSIS: Finished session with ZERO answers")
        print("➡️ Likely causes: timeout, tab closed, submit not called")
        print("================ END GET READING REPORT ================\n")

        return {
            "status": "no_report",
            "reason": "Exam finished without submission data",
            "attempt_id": session.id,
            "student_id": session.student_id,
            "exam_id": session.exam_id,
            "can_retake": True
        }

    # --------------------------------------------------
    # 6️⃣ Answers exist but snapshot missing → inconsistency
    # --------------------------------------------------
    print("❌ DIAGNOSIS: Answers exist but report_json is missing")
    print("➡️ Likely causes: crash during submit, DB rollback, manual DB edit")

    print("🧨 Diagnostic summary:", {
        "session_id": session.id,
        "answer_count": answer_count,
        "finished": session.finished,
        "has_report_json": False
    })

    print("================ END GET READING REPORT ================\n")

    return {
        "status": "report_missing",
        "reason": "Submission data exists but report snapshot is missing",
        "attempt_id": session.id,
        "student_id": session.student_id,
        "exam_id": session.exam_id,
        "answer_count": answer_count
    }


@app.get("/api/student/exam-report/thinking-skills")
def get_thinking_skills_report(
    student_id: str = Query(..., description="External student id e.g. Gem002"),
    db: Session = Depends(get_db)
):
    # --------------------------------------------------
    # 1️⃣ Resolve student
    # --------------------------------------------------
    student = (
        db.query(Student)
        .filter(Student.student_id == student_id)
        .first()
    )

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # --------------------------------------------------
    # 2️⃣ Get latest completed THINKING SKILLS attempt
    # --------------------------------------------------
    attempt = (
        db.query(StudentExamThinkingSkills)
        .filter(
            StudentExamThinkingSkills.student_id == student.id,
            StudentExamThinkingSkills.completed_at.isnot(None)
        )
        .order_by(StudentExamThinkingSkills.completed_at.desc())
        .first()
    )

    if not attempt:
        raise HTTPException(
            status_code=404,
            detail="No completed Thinking Skills exam found"
        )

    # --------------------------------------------------
    # 3️⃣ Load all responses for this attempt
    # --------------------------------------------------
    responses = (
        db.query(StudentExamResponseThinkingSkills)
        .filter(
            StudentExamResponseThinkingSkills.exam_attempt_id == attempt.id
        )
        .all()
    )

    if not responses:
        raise HTTPException(
            status_code=404,
            detail="No responses found for Thinking Skills exam"
        )

    # --------------------------------------------------
    # 4️⃣ OVERALL SUMMARY (Report B)
    # --------------------------------------------------
    total_questions = len(responses)
    attempted = sum(1 for r in responses if r.is_correct is not None)
    correct = sum(1 for r in responses if r.is_correct is True)
    incorrect = sum(1 for r in responses if r.is_correct is False)
    not_attempted = total_questions - attempted

    accuracy_percent = round((correct / attempted) * 100, 2) if attempted else 0
    score_percent = round((correct / total_questions) * 100, 2) if total_questions else 0

    overall = {
        "total_questions": total_questions,
        "attempted": attempted,
        "correct": correct,
        "incorrect": incorrect,
        "not_attempted": not_attempted,
        "accuracy_percent": accuracy_percent,
        "score_percent": score_percent,
        "pass": None
    }

    # --------------------------------------------------
    # 5️⃣ TOPIC-WISE PERFORMANCE (Report A)
    # --------------------------------------------------
    topic_map = {}

    for r in responses:
        topic = r.topic or "Unknown"

        if topic not in topic_map:
            topic_map[topic] = {
                "topic": topic,
                "total": 0,
                "attempted": 0,
                "correct": 0,
                "incorrect": 0,
                "not_attempted": 0
            }

        topic_map[topic]["total"] += 1

        if r.is_correct is None:
            topic_map[topic]["not_attempted"] += 1
        else:
            topic_map[topic]["attempted"] += 1
            if r.is_correct:
                topic_map[topic]["correct"] += 1
            else:
                topic_map[topic]["incorrect"] += 1

    topic_wise_performance = list(topic_map.values())

    # --------------------------------------------------
    # 6️⃣ TOPIC ACCURACY & RESULT (Report C)
    # --------------------------------------------------
    topic_accuracy = []

    for t in topic_wise_performance:
        attempted_t = t["attempted"]
        accuracy_t = (
            round((t["correct"] / attempted_t) * 100, 2)
            if attempted_t else 0
        )
        score_t = round((t["correct"] / t["total"]) * 100, 2)

        topic_accuracy.append({
            "topic": t["topic"],
            "total_questions": t["total"],
            "attempted": attempted_t,
            "correct": t["correct"],
            "incorrect": t["incorrect"],
            "accuracy_percent": accuracy_t,
            "score_percent": score_t,
            "pass": None
        })

    # --------------------------------------------------
    # 7️⃣ IMPROVEMENT AREAS (Report D)
    # --------------------------------------------------
    improvement_areas = []

    for t in topic_accuracy:
        limited_data = t["total_questions"] < 5

        improvement_areas.append({
            "topic": t["topic"],
            "accuracy_percent": t["accuracy_percent"],
            "score_percent": t["score_percent"],
            "total_questions": t["total_questions"],
            "limited_data": limited_data
        })

    improvement_areas.sort(key=lambda x: x["accuracy_percent"])

    # --------------------------------------------------
    # 8️⃣ Final response
    # --------------------------------------------------
    return {
        "overall": overall,                                 # Report B
        "topic_wise_performance": topic_wise_performance,   # Report A
        "topic_accuracy": topic_accuracy,                   # Report C
        "improvement_areas": improvement_areas              # Report D
    }

# ============================================================
# ❌ OLD / LEGACY MATHEMATICAL REASONING REPORT (DO NOT USE)
# This version reads from GENERIC tables:
#   - StudentExam
#   - StudentExamResponse
# It is incompatible with the current finish-exam implementation
# ============================================================

"""
@app.get("/api/student/exam-report/mathematical-reasoning")
def get_thinking_skills_report(
    student_id: str = Query(..., description="External student id e.g. Gem002"),
    db: Session = Depends(get_db)
):
    # 1️⃣ Resolve student
    student = (
        db.query(Student)
        .filter(Student.student_id == student_id)
        .first()
    )
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # 2️⃣ Get latest completed attempt (GENERIC TABLE ❌)
    attempt = (
        db.query(StudentExam)
        .filter(
            StudentExam.student_id == student.id,
            StudentExam.completed_at.isnot(None)
        )
        .order_by(StudentExam.completed_at.desc())
        .first()
    )

    if not attempt:
        raise HTTPException(status_code=404, detail="No completed exam found")

    # 3️⃣ Load responses (GENERIC TABLE ❌)
    responses = (
        db.query(StudentExamResponse)
        .filter(StudentExamResponse.exam_attempt_id == attempt.id)
        .all()
    )

    if not responses:
        raise HTTPException(status_code=404, detail="No responses found")

    # Remaining logic omitted for brevity
"""



@app.get("/api/student/exam-report/mathematical-reasoning")
def get_mathematical_reasoning_report(
    student_id: str = Query(..., description="External student id e.g. Gem002"),
    db: Session = Depends(get_db)
):
    # --------------------------------------------------
    # 1️⃣ Resolve student
    # --------------------------------------------------
    student = (
        db.query(Student)
        .filter(Student.student_id == student_id)
        .first()
    )

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # --------------------------------------------------
    # 2️⃣ Get latest completed Mathematical Reasoning attempt
    # --------------------------------------------------
    attempt = (
        db.query(StudentExamMathematicalReasoning)
        .filter(
            StudentExamMathematicalReasoning.student_id == student.id,
            StudentExamMathematicalReasoning.completed_at.isnot(None)
        )
        .order_by(StudentExamMathematicalReasoning.completed_at.desc())
        .first()
    )

    if not attempt:
        raise HTTPException(
            status_code=404,
            detail="No completed Mathematical Reasoning exam found"
        )

    # --------------------------------------------------
    # 3️⃣ Load responses for this attempt
    # --------------------------------------------------
    responses = (
        db.query(StudentExamResponseMathematicalReasoning)
        .filter(
            StudentExamResponseMathematicalReasoning.exam_attempt_id == attempt.id
        )
        .all()
    )

    if not responses:
        raise HTTPException(
            status_code=404,
            detail="No responses found for Mathematical Reasoning exam"
        )

    # --------------------------------------------------
    # 4️⃣ OVERALL SUMMARY (Report B)
    # --------------------------------------------------
    total_questions = len(responses)
    attempted = sum(1 for r in responses if r.is_correct is not None)
    correct = sum(1 for r in responses if r.is_correct is True)
    incorrect = sum(1 for r in responses if r.is_correct is False)
    not_attempted = total_questions - attempted

    accuracy_percent = round((correct / attempted) * 100, 2) if attempted else 0
    score_percent = round((correct / total_questions) * 100, 2)

    overall = {
        "total_questions": total_questions,
        "attempted": attempted,
        "correct": correct,
        "incorrect": incorrect,
        "not_attempted": not_attempted,
        "accuracy_percent": accuracy_percent,
        "score_percent": score_percent,
        "pass": None
    }

    # --------------------------------------------------
    # 5️⃣ TOPIC-WISE PERFORMANCE (Report A)
    # --------------------------------------------------
    topic_map = {}

    for r in responses:
        topic = r.topic or "Unknown"

        if topic not in topic_map:
            topic_map[topic] = {
                "topic": topic,
                "total": 0,
                "attempted": 0,
                "correct": 0,
                "incorrect": 0,
                "not_attempted": 0
            }

        topic_map[topic]["total"] += 1

        if r.is_correct is None:
            topic_map[topic]["not_attempted"] += 1
        else:
            topic_map[topic]["attempted"] += 1
            if r.is_correct:
                topic_map[topic]["correct"] += 1
            else:
                topic_map[topic]["incorrect"] += 1

    topic_wise_performance = list(topic_map.values())

    # --------------------------------------------------
    # 6️⃣ TOPIC ACCURACY (Report C)
    # --------------------------------------------------
    topic_accuracy = []

    for t in topic_wise_performance:
        attempted_t = t["attempted"]
        accuracy_t = (
            round((t["correct"] / attempted_t) * 100, 2)
            if attempted_t else 0
        )
        score_t = round((t["correct"] / t["total"]) * 100, 2)

        topic_accuracy.append({
            "topic": t["topic"],
            "total_questions": t["total"],
            "attempted": attempted_t,
            "correct": t["correct"],
            "incorrect": t["incorrect"],
            "accuracy_percent": accuracy_t,
            "score_percent": score_t,
            "pass": None
        })

    # --------------------------------------------------
    # 7️⃣ IMPROVEMENT AREAS (Report D)
    # --------------------------------------------------
    improvement_areas = []

    for t in topic_accuracy:
        limited_data = t["total_questions"] < 5

        improvement_areas.append({
            "topic": t["topic"],
            "accuracy_percent": t["accuracy_percent"],
            "score_percent": t["score_percent"],
            "total_questions": t["total_questions"],
            "limited_data": limited_data
        })

    improvement_areas.sort(key=lambda x: x["accuracy_percent"])

    # --------------------------------------------------
    # 8️⃣ Final response
    # --------------------------------------------------
    return {
        "overall": overall,
        "topic_wise_performance": topic_wise_performance,
        "topic_accuracy": topic_accuracy,
        "improvement_areas": improvement_areas,
        "exam_attempt_id": attempt.id
    }


@app.post("/api/student/save-exam-results-thinkingskills")
def save_exam_results_thinkingskills(
    payload: ExamSubmissionRequest,
    db: Session = Depends(get_db)
):
    # 1️⃣ Validate student
    student = (
        db.query(Student)
        .filter(Student.student_id == payload.student_id)
        .first()
    )

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # 2️⃣ Find ACTIVE exam attempt from student_exams
    attempt = (
        db.query(StudentExam)
        .filter(
            StudentExam.student_id == student.id,
            StudentExam.completed_at.is_(None)
        )
        .order_by(StudentExam.started_at.desc())
        .first()
    )

    if not attempt:
        # No active attempt → exam already completed or never started
        return {"status": "ignored"}

    # 3️⃣ Prevent double submission
    existing_submission = (
        db.query(StudentExamAttemptThinkingSkills)
        .filter(
            StudentExamAttemptThinkingSkills.student_exam_id == attempt.id
        )
        .first()
    )

    if existing_submission and existing_submission.is_completed:
        return {"status": "already_submitted"}

    # 4️⃣ Create attempt record if not exists
    if not existing_submission:
        existing_submission = StudentExamAttemptThinkingSkills(
            student_exam_id=attempt.id,
            submitted_at=None,
            is_completed=False
        )
        db.add(existing_submission)
        db.flush()  # get ID without commit

    # 5️⃣ Save answers (idempotent)
    for qid_str, selected_option in payload.answers.items():
        question_id = int(qid_str)

        exists = (
            db.query(StudentExamAnswerThinkingSkills)
            .filter(
                StudentExamAnswerThinkingSkills.attempt_id == existing_submission.id,
                StudentExamAnswerThinkingSkills.question_id == question_id
            )
            .first()
        )

        if exists:
            continue

        db.add(
            StudentExamAnswerThinkingSkills(
                attempt_id=existing_submission.id,
                question_id=question_id,
                selected_option=selected_option
            )
        )

    # 6️⃣ Mark submission complete
    existing_submission.submitted_at = datetime.now(timezone.utc)
    existing_submission.is_completed = True
    db.commit()

    return {"status": "saved"}

@app.post("/add_student_exam_module")
def add_student_exam_module(
    payload: AddStudentExamModuleRequest,
    db: Session = Depends(get_db)
):
    existing_student = (
        db.query(Student)
        .filter(Student.student_id == payload.student_id)
        .first()
    )

    if existing_student:
        raise HTTPException(
            status_code=400,
            detail="Student with this student_id already exists"
        )

    # ⚠️ Plain text password (as requested)
    plain_password = payload.student_id  # or any value you choose

    student = Student(
        id=payload.id,
        student_id=payload.student_id,
        name=payload.name,
        parent_email=payload.parent_email,
        class_name=payload.class_name,
        class_day=payload.class_day,
        password=plain_password,  # ← STORED AS IS
    )

    db.add(student)
    db.commit()
    db.refresh(student)

    return {
        "message": "Student added successfully",
        "student_id": student.student_id,
        "password": plain_password
    }
#here
def generate_admin_exam_report_reading(
    db: Session,
    student: Student,
    exam_attempt: StudentExamReading,
    score_percent: float,
    accuracy: float,
    coverage: float,
    topic_stats: dict
): 
    """
    Generates an immutable admin report snapshot for a completed Reading exam.
    MUST NOT commit. Caller controls the transaction.
    """

    # -------------------------------
    # 1️⃣ Determine readiness & guidance
    # -------------------------------
    # Coverage guardrail (client concern)
    if coverage < 60:
        readiness_band = "Not Yet Selective Ready"
        school_guidance = "Insufficient Exam Coverage"
    else:
        if score_percent >= 80:
            readiness_band = "Strong Selective Potential"
            school_guidance = "Mid-tier Selective"
        elif score_percent >= 65:
            readiness_band = "Borderline Selective"
            school_guidance = "Selective Preparation Required"
        else:
            readiness_band = "Not Yet Selective Ready"
            school_guidance = "General Stream Recommended"

    # -------------------------------
    # 2️⃣ Create admin report snapshot
    # -------------------------------
    admin_report = AdminExamReport(
        student_id=student.student_id,  # external ID (STRING)
        exam_attempt_id=exam_attempt.id,
        exam_type="reading",
        overall_score=score_percent,
        readiness_band=readiness_band,
        school_guidance_level=school_guidance,
        summary_notes=(
            f"Score: {score_percent}%, "
            f"Accuracy: {accuracy}%, "
            f"Coverage: {coverage}%"
        )


    )

    db.add(admin_report)
    db.flush()  # obtain admin_report.id safely

    # -------------------------------
    # 3️⃣ Create section (topic) results
    # -------------------------------
    for topic, stats in topic_stats.items():
        attempted = stats.get("attempted", 0)
        correct = stats.get("correct", 0)

        accuracy = (
            round((correct / attempted) * 100, 2)
            if attempted > 0 else 0.0
        )

        if accuracy >= 80:
            band = "A"
        elif accuracy >= 60:
            band = "B"
        else:
            band = "C"

        section_result = AdminExamSectionResult(
            admin_report_id=admin_report.id,
            section_name=topic,
            raw_score=accuracy,
            performance_band=band,
            strengths_summary=None,
            improvement_summary=None
        )

        db.add(section_result)

    # -------------------------------
    # 4️⃣ Store applied rule (audit safety)
    # -------------------------------
    rule_applied = AdminReadinessRuleApplied(
        admin_report_id=admin_report.id,
        rule_code="READING_SCORE_WITH_COVERAGE",
        rule_description="Readiness based on total score with minimum coverage requirement",
        rule_result="passed" if readiness_band != "Not Yet Selective Ready" else "failed"
    )

    db.add(rule_applied)


 
@app.post("/api/exams/submit-reading")
def submit_reading_exam(payload: dict, db: Session = Depends(get_db)):

    print("\n================ SUBMIT READING EXAM ================")
    print("📥 Raw payload received:", payload)

    try:
        # --------------------------------------------------
        # 0️⃣ Validate payload
        # --------------------------------------------------
        print("🟡 STEP 0: Validating payload")

        session_id = payload.get("session_id")
        answers = payload.get("answers", {})

        print("🆔 session_id:", session_id)
        print("📝 answers keys:",
              list(answers.keys()) if isinstance(answers, dict) else answers)

        if not session_id:
            print("❌ ERROR: session_id missing")
            raise HTTPException(status_code=400, detail="session_id is required")

        if not isinstance(answers, dict):
            print("❌ ERROR: answers is not a dict")
            raise HTTPException(status_code=400, detail="answers must be a dictionary")

        # --------------------------------------------------
        # 1️⃣ Load session
        # --------------------------------------------------
        print("🟡 STEP 1: Loading StudentExamReading session")

        session = (
            db.query(StudentExamReading)
            .filter(StudentExamReading.id == session_id)
            .first()
        )

        if not session:
            print("❌ ERROR: Session not found:", session_id)
            raise HTTPException(status_code=404, detail="Session not found")

        print("✅ Session loaded:", {
            "id": session.id,
            "student_id": session.student_id,
            "exam_id": session.exam_id,
            "finished": session.finished
        })

        if session.finished:
            print("⚠️ Session already finished — idempotent exit")
            return {
                "status": "ok",
                "message": "Exam already submitted"
            }

        # --------------------------------------------------
        # 2️⃣ Resolve student
        # --------------------------------------------------
        
        print("🟡 Resolving student via INTERNAL ID")

        student = (
            db.query(Student)
            .filter(Student.id == session.student_id)
            .first()
        )
        
        if not student:
            print("❌ ERROR: Student not found for Student.id =", session.student_id)
            raise HTTPException(status_code=404, detail="Student not found")
        
        print("✅ Student resolved:", student.id)


        # --------------------------------------------------
        # 3️⃣ Load exam
        # --------------------------------------------------
        print("🟡 STEP 3: Loading GeneratedExamReading")

        exam = (
            db.query(GeneratedExamReading)
            .filter(GeneratedExamReading.id == session.exam_id)
            .first()
        )

        if not exam or not exam.exam_json:
            print("❌ ERROR: exam or exam_json missing")
            raise HTTPException(status_code=500, detail="Exam data not found")

        sections = exam.exam_json.get("sections", [])
        print("📚 Sections found:", len(sections))

        # --------------------------------------------------
        # 4️⃣ Evaluate questions
        # --------------------------------------------------
        print("🟡 STEP 4: Evaluating answers")
        # Canonical topic keys only (NO labels here)
        CANONICAL_TOPICS = {
            "main_idea": "main_idea",
            "main_idea_and_summary": "main_idea",
            "comparative_analysis": "comparative_analysis",
            "gapped_text": "gapped_text",
        }


        topic_stats = {}
        total_questions = attempted = correct = incorrect = not_attempted = 0
        

        for section in sections:
            raw_topic = section.get("question_type")
            topic = CANONICAL_TOPICS.get(raw_topic, "other")
            

            print("🧪 CANONICAL TOPIC RESOLVED:", {
                "section_id": section.get("section_id"),
                "question_type": raw_topic,
                "stored_topic": topic
            })


            questions = section.get("questions", [])

            topic_stats.setdefault(topic, {
                "total": 0,
                "attempted": 0,
                "correct": 0,
                "incorrect": 0,
                "not_attempted": 0
            })

            for q in questions:
                question_id = q.get("question_id")
                correct_answer = q.get("correct_answer")
                selected_answer = answers.get(question_id)

                is_attempted = selected_answer is not None
                is_correct = is_attempted and selected_answer == correct_answer

                total_questions += 1
                topic_stats[topic]["total"] += 1

                if is_attempted:
                    attempted += 1
                    topic_stats[topic]["attempted"] += 1
                    if is_correct:
                        correct += 1
                        topic_stats[topic]["correct"] += 1
                    else:
                        incorrect += 1
                        topic_stats[topic]["incorrect"] += 1
                else:
                    not_attempted += 1
                    topic_stats[topic]["not_attempted"] += 1

                db.add(StudentExamReportReading(
                    student_id=session.student_id,
                    exam_id=session.exam_id,
                    session_id=session.id,
                    topic=topic,
                    question_id=question_id,
                    selected_answer=selected_answer,
                    correct_answer=correct_answer,
                    is_correct=is_correct
                ))

        print("✅ Evaluation complete")

        # --------------------------------------------------
        # 5️⃣ Build report
        # --------------------------------------------------
        print("🟡 STEP 5: Building report")

        topics_report = []
        for topic, stats in topic_stats.items():
            accuracy = (
                round((stats["correct"] / stats["attempted"]) * 100, 2)
                if stats["attempted"] > 0 else 0.0
            )
            topics_report.append({
                "topic": topic,
                "total": stats["total"],
                "attempted": stats["attempted"],
                "correct": stats["correct"],
                "incorrect": stats["incorrect"],
                "accuracy": accuracy
            })


        accuracy = round((correct / attempted) * 100, 2) if attempted > 0 else 0.0
        coverage = round((attempted / total_questions) * 100, 2) if total_questions > 0 else 0.0
        score_percent = round((correct / total_questions) * 100, 2) if total_questions > 0 else 0.0
        result = "Pass" if score_percent >= 50 else "Fail"

        MIN_ATTEMPTS = max(5, int(total_questions * 0.2))
        has_sufficient_data = attempted >= MIN_ATTEMPTS

        report_json = {
            "overall": {
                "total_questions": total_questions,
                "attempted": attempted,
                "correct": correct,
                "incorrect": incorrect,
                "not_attempted": not_attempted,
            
                # diagnostics
                "accuracy": accuracy,
                "coverage": coverage,
            
                # ranking signal
                "score": score_percent,
                "result": result
            },
            "topics": topics_report,
            "has_sufficient_data": has_sufficient_data,
            "improvement_order": (
                [t["topic"] for t in sorted(topics_report, key=lambda x: x["accuracy"])]
                if has_sufficient_data else []
            )
        }

        print("📈 Report built successfully")

        # --------------------------------------------------
        # 6️⃣ Mark session finished
        # --------------------------------------------------
        print("🟡 STEP 6: Marking session finished")

        session.finished = True
        session.completed_at = datetime.now(timezone.utc)
        session.report_json = report_json
        # --------------------------------------------------
        # 6️⃣b Generate Admin Exam Report (Reading)
        # --------------------------------------------------
        print("🟡 STEP 6b: Generating AdminExamReport (Reading)")
        
        if not admin_report_exists(
            db=db,
            exam_attempt_id=session.id,
            exam_type="reading"
        ):
            generate_admin_exam_report_reading(
               db=db,
               student=student,
               exam_attempt=session,
               score_percent=score_percent,
               accuracy=accuracy,
               coverage=coverage,
               topic_stats=topic_stats
           )
            print("🟢 AdminExamReport created")
        else:
            print("⚠️ AdminExamReport already exists — skipping")
     
  
        print("🟡 STEP 7: About to commit DB")
        db.commit()
        print("🟢 DB commit successful")

        # --------------------------------------------------
        # 7️⃣ Return response
        # --------------------------------------------------
        print("🟢 STEP 8: Returning response")
        print("================ END SUBMIT READING EXAM ================\n")

        return {
            "status": "submitted",
            "message": "Reading exam submitted successfully",
            "report": report_json
        }

    except Exception as e:
        print("❌❌❌ UNHANDLED EXCEPTION IN submit_reading_exam ❌❌❌")
        print("Exception type:", type(e))
        print("Exception message:", str(e))
        import traceback
        traceback.print_exc()
        raise

def auto_submit_reading_exam(session: StudentExamReading, db: Session):
    """
    Auto-submit reading exam on timeout.
    Uses the SAME logic as manual submit, but with empty answers.
    """
    
    print("🟡 AUTO-SUBMIT triggered for session:", session.id)
    if session.finished and session.report_json:
        print("⚠️ Auto-submit skipped — session already finalized")
        return
    existing_rows = (
        db.query(StudentExamReportReading)
        .filter(StudentExamReportReading.session_id == session.id)
        .count()
    )
    
    if existing_rows > 0:
        print("⚠️ Auto-submit skipped — report rows already exist")
        return


    
    exam = (
        db.query(GeneratedExamReading)
        .filter(GeneratedExamReading.id == session.exam_id)
        .first()
    )

    if not exam or not exam.exam_json:
        raise Exception("Exam data missing during auto-submit")

    sections = exam.exam_json.get("sections", [])
    TOPIC_LABELS = {
        "main_idea": "Main Idea and Summary",
        "main_idea_and_summary": "Main Idea and Summary",
        "comparative_analysis": "Comparative Analysis",
        "gapped_text": "Gapped Text",
    }

    topic_stats = {}
    total = attempted = correct = incorrect = not_attempted = 0

    for section in sections:
        raw_topic = section.get("topic") or section.get("question_type")
        topic = TOPIC_LABELS.get(raw_topic, "Other")
        questions = section.get("questions", [])

        topic_stats.setdefault(topic, {
            "total": 0,
            "attempted": 0,
            "correct": 0,
            "incorrect": 0,
            "not_attempted": 0
        })

        for q in questions:
            question_id = q.get("question_id")
            correct_answer = q.get("correct_answer")

            total += 1
            not_attempted += 1

            topic_stats[topic]["total"] += 1
            topic_stats[topic]["not_attempted"] += 1

            db.add(StudentExamReportReading(
                student_id=session.student_id,
                exam_id=session.exam_id,
                session_id=session.id,
                topic=topic,
                question_id=question_id,
                selected_answer=None,
                correct_answer=correct_answer,
                is_correct=False
            ))
    topics_report = []

    for topic, stats in topic_stats.items():
        topics_report.append({
            "topic": topic,
            "total": stats["total"],
            "attempted": 0,
            "correct": 0,
            "incorrect": 0,
            "accuracy": 0.0
        })

    report_json = {
        "overall": {
            "total_questions": total,
            "attempted": attempted,
            "correct": correct,
            "incorrect": incorrect,
            "not_attempted": not_attempted,
            "accuracy": 0.0,
            "score": 0.0,
            "result": "Fail"
        },
        "topics": topics_report,
        "has_sufficient_data": False,
        "improvement_order": []
    }

    session.finished = True
    session.completed_at = datetime.now(timezone.utc)
    session.report_json = report_json

    try:
        db.commit()
    except Exception:
        db.rollback()
        raise


@app.post("/api/exams/start-reading")
def start_exam_reading(
    req: StartExamRequest = Body(...),
    db: Session = Depends(get_db)
):
    print("\n================ START-READING ==================")
    print("📘 Incoming request payload:", req.dict())

    # 1️⃣ Resolve student
    student = (
        db.query(Student)
        .filter(func.lower(Student.student_id) == func.lower(req.student_id.strip()))
        .first()
    )

    if not student:
        print("❌ Student NOT FOUND:", req.student_id)
        raise HTTPException(status_code=404, detail="Student not found")

    print("✅ Student resolved:", {
        "student_pk": student.id,
        "student_id": student.student_id
    })

    # 2️⃣ Fetch latest attempt
    attempt = (
        db.query(StudentExamReading)
        .filter(StudentExamReading.student_id == student.id)
        .order_by(StudentExamReading.started_at.desc())
        .first()
    )

    if attempt:
        print("🧪 Found existing attempt:", {
            "attempt_id": attempt.id,
            "exam_id": attempt.exam_id,
            "started_at": attempt.started_at,
            "finished": attempt.finished
        })
    else:
        print("🆕 No previous attempt found for student")

    # 🚫 Already finished → frontend should load report
    if attempt and attempt.finished:
        payload = {
            "completed": True,
            "attempt_id": attempt.id
        }
        print("🚫 Attempt already finished → returning:", payload)
        print("================================================\n")
        return payload

    # 3️⃣ Latest reading exam
    exam = (
        db.query(GeneratedExamReading)
        .order_by(GeneratedExamReading.id.desc())
        .first()
    )

    if not exam:
        print("❌ No reading exam exists in DB")
        raise HTTPException(status_code=404, detail="Reading exam not found")

    duration_minutes = exam.exam_json.get("duration_minutes", 40)
    print("📘 Active exam resolved:", {
        "exam_id": exam.id,
        "duration_minutes": duration_minutes
    })

    # 🔁 Resume unfinished attempt
    if attempt and not attempt.finished:
        now = datetime.now(timezone.utc)
        started_at = attempt.started_at.replace(tzinfo=timezone.utc)
        elapsed = int((now - started_at).total_seconds())
        remaining = max(0, duration_minutes * 60 - elapsed)

        print("⏱ Resume attempt timing:", {
            "attempt_id": attempt.id,
            "elapsed_seconds": elapsed,
            "remaining_seconds": remaining
        })

        if remaining == 0:
            print("⌛ Time expired — auto-submitting exam")
        
            if not attempt.finished:
                auto_submit_reading_exam(
                    session=attempt,
                    db=db
                )
        
            payload = {
                "completed": True,
                "attempt_id": attempt.id
            }
        
            print("🟢 Auto-submit complete → returning:", payload)
            print("================================================\n")
            return payload

        payload = {
            "completed": False,
            "attempt_id": attempt.id,
            "exam_id": exam.id,
            "remaining_time": remaining
        }

        print("🔁 Resuming attempt → returning:", payload)
        print("================================================\n")
        return payload

    # 🆕 Create new attempt
    new_attempt = StudentExamReading(
        student_id=student.id,
        exam_id=exam.id,
        started_at=datetime.now(timezone.utc),
        finished=False
    )

    db.add(new_attempt)
    db.commit()
    db.refresh(new_attempt)

    payload = {
        "completed": False,
        "attempt_id": new_attempt.id,
        "exam_id": exam.id,
        "remaining_time": duration_minutes * 60
    }

    print("🆕 New attempt created:", {
        "attempt_id": new_attempt.id,
        "exam_id": exam.id
    })
    print("📤 Returning payload:", payload)
    print("================================================\n")

    return payload


@app.get("/api/exams/reading-content/{exam_id}")
def get_reading_content(exam_id: int, db: Session = Depends(get_db)):
    print("\n📘 GET READING CONTENT")
    print("➡ exam_id:", exam_id)

    exam = (
        db.query(GeneratedExamReading)
        .filter(GeneratedExamReading.id == exam_id)
        .first()
    )

    if not exam:
        print("❌ ERROR: GeneratedExamReading not found")
        raise HTTPException(status_code=404, detail="Exam content not found")

    if not exam.exam_json:
        print("❌ ERROR: exam.exam_json is NULL or empty")
        raise HTTPException(status_code=404, detail="Exam content not found")

    # 🔍 DEBUG: structure inspection
    print("✅ exam.exam_json type:", type(exam.exam_json))
    print("🔑 exam.exam_json keys:", exam.exam_json.keys())

    sections = exam.exam_json.get("sections", [])
    print("📚 sections type:", type(sections))
    print("📚 sections length:", len(sections) if isinstance(sections, list) else "N/A")

    if isinstance(sections, list) and len(sections) > 0:
        first_section = sections[0]
        print("🧱 FIRST SECTION KEYS:", first_section.keys())

        for key in ["questions", "items", "question_list"]:
            if key in first_section:
                print(f"🧩 FOUND QUESTIONS UNDER KEY: '{key}'")
                print("🧩 questions count:", len(first_section[key]))
                if len(first_section[key]) > 0:
                    print("🧩 FIRST QUESTION KEYS:", first_section[key][0].keys())
                break
        else:
            print("⚠️ NO QUESTION ARRAY FOUND in first section")

    # 🔧 FIX 1 — normalize reading_material (string → object)
    for section in sections:
        rm = section.get("reading_material")

        if isinstance(rm, str):
            print("🛠 Normalizing reading_material for section:", section.get("section_id"))

            section["reading_material"] = {
                "title": section.get("topic") or "Reading Passage",
                "content": rm
            }

    print("📤 RETURNING exam_json TO FRONTEND\n")

    return {
        "exam_json": exam.exam_json
    }


@app.get("/api/foundational/classes")
def get_foundational_classes(db: Session = Depends(get_db)):
    rows = (
        db.query(QuizSetupFoundational.class_name)
        .distinct()
        .all()
    )

    return [{"class_name": r.class_name} for r in rows]
 
@app.post("/api/student/start-writing-exam")
def start_writing_exam(student_id: str, db: Session = Depends(get_db)):

    student = (
        db.query(Student)
        .filter(func.lower(Student.student_id) == func.lower(student_id))
        .first()
    )
    
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")


    # --------------------------------------------------
    # 1️⃣ Get current writing exam
    # --------------------------------------------------
    exam = (
        db.query(GeneratedExamWriting)
        .filter(GeneratedExamWriting.is_current == True)
        .order_by(GeneratedExamWriting.created_at.desc())
        .first()
    )

    if not exam:
        raise HTTPException(404, "No active writing exam")

    # --------------------------------------------------
    # 2️⃣ Fetch latest attempt for THIS exam
    # --------------------------------------------------
    attempt = (
        db.query(StudentExamWriting)
        .filter(
            StudentExamWriting.student_id == student.id,
            StudentExamWriting.exam_id == exam.id
        )
        .order_by(StudentExamWriting.started_at.desc())
        .first()
    )

    # --------------------------------------------------
    # 🟥 CASE A — Exam already completed
    # --------------------------------------------------
    if attempt and attempt.completed_at:
        return {"completed": True}

    # --------------------------------------------------
    # 🟡 CASE B — Resume active attempt
    # --------------------------------------------------
    if attempt and attempt.completed_at is None:
        started_at = attempt.started_at
        if started_at.tzinfo is None:
            started_at = started_at.replace(tzinfo=timezone.utc)

        now = datetime.now(timezone.utc)
        elapsed = int((now - started_at).total_seconds())
        remaining = max(0, attempt.duration_minutes * 60 - elapsed)

        if remaining == 0:
            attempt.completed_at = now
            db.commit()
            return {"completed": True}

        return {
            "completed": False,
            "remaining_time": remaining
        }

    # --------------------------------------------------
    # 🔵 CASE C — Start new writing attempt
    # --------------------------------------------------
    new_attempt = StudentExamWriting(
        student_id=student.id,
        exam_id=exam.id,
        started_at=datetime.now(timezone.utc),
        duration_minutes=exam.duration_minutes
    )

    db.add(new_attempt)
    db.commit()
    db.refresh(new_attempt)

    return {
        "completed": False,
        "remaining_time": new_attempt.duration_minutes * 60
    }



@app.get("/api/get-quizzes-writing")
def get_quizzes_writing(db: Session = Depends(get_db)):
    rows = (
        db.query(
            QuizSetupWriting.class_name,
            QuizSetupWriting.difficulty
        )
        .distinct()
        .order_by(
            QuizSetupWriting.class_name,
            QuizSetupWriting.difficulty
        )
        .all()
    )

    return [
        {
            "class_name": row.class_name,
            "difficulty": row.difficulty
        }
        for row in rows
    ]



@app.post("/api/exams/writing/submit")
def submit_writing_exam(
    payload: WritingSubmitSchema,
    student_id: str,  # kept for API compatibility
    db: Session = Depends(get_db)
):
    print("\n================ SUBMIT WRITING EXAM =================")
    print("➡️ Incoming student_id:", student_id)
    print("➡️ Payload received:", payload.dict())
    student = (
        db.query(Student)
        .filter(func.lower(Student.student_id) == func.lower(student_id))
        .first()
    )
    
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # --------------------------------------------------
    # 1️⃣ Load active writing attempt
    # --------------------------------------------------
    print("🔍 Looking for active writing attempt...")

    exam_state = (
        db.query(StudentExamWriting)
        .filter(
            StudentExamWriting.student_id == student.id,
            StudentExamWriting.completed_at.is_(None)
        )
        .order_by(StudentExamWriting.started_at.desc())
        .first()
    )

    if not exam_state:
        print("❌ No active writing exam found for student_id:", student_id)
        raise HTTPException(
            status_code=404,
            detail="No active writing exam found"
        )

    print("✅ Writing attempt found:", {
        "attempt_id": exam_state.id,
        "exam_id": exam_state.exam_id,
        "started_at": exam_state.started_at
    })

    # --------------------------------------------------
    # 2️⃣ Persist student answer
    # --------------------------------------------------
    print("💾 Saving student answer...")
    print("✏️ Essay length:", len(payload.answer_text or ""))

    exam_state.answer_text = payload.answer_text
    exam_state.completed_at = datetime.now(timezone.utc)

    # --------------------------------------------------
    # 3️⃣ Evaluate essay using OpenAI
    # --------------------------------------------------
    print("🤖 Preparing OpenAI prompt...")

    prompt = f"""
You are an expert NSW Selective School writing marker.

CRITICAL OUTPUT RULES (MUST FOLLOW EXACTLY):
- Respond with ONLY a valid JSON object
- Do NOT include markdown, headings, emojis, numbering symbols, or formatting characters
- Do NOT include explanations, notes, or extra text outside the JSON
- Do NOT include triple backticks or language tags
- Output must be directly parsable using json.loads()
- Use double quotes for all JSON keys and string values

TASK:
Assess the student's writing response strictly against NSW Selective School exam standards.

INPUTS YOU WILL RECEIVE:
1. Writing prompt
2. Student response

ASSESSMENT INSTRUCTIONS:
- Judge the response holistically, but score using the five selective criteria below
- Prioritise clarity, task fulfilment, and accuracy over creativity
- Be strict but fair; assume this is a competitive selective exam
- Do NOT rewrite or correct the student’s response
- Accuracy issues must meaningfully reduce scores where present

SCORING CRITERIA (TOTAL 25 MARKS):
Each category is scored out of 5.

1. Audience, Purpose and Form
- Matches the required genre
- Appropriate tone for the intended audience
- All parts of the task are addressed

2. Ideas and Content
- Ideas are relevant to the prompt
- Ideas are clearly explained and developed
- No unnecessary, vague, or unrealistic content

3. Structure and Organisation
- Logical sequencing of ideas
- Purposeful paragraphing
- Clear beginning, middle, and end

4. Language and Vocabulary
- Precise and appropriate word choice
- Controlled and varied sentence structures
- Consistent and suitable tone

5. Grammar, Spelling and Punctuation
- Accuracy of spelling and punctuation
- Errors must not limit clarity
- Frequent basic errors must significantly reduce the score

SELECTIVE READINESS BAND (MANDATORY):
Based on the overall score out of 25, assign ONE descriptor only:

- 22–25: Strong selective standard – very competitive
- 18–21: On track for selective with minor improvements
- 14–17: Developing – selective readiness needs strengthening
- 10–13: Below selective standard – significant improvement needed
- Below 10: Well below selective standard at this stage

REQUIRED JSON RESPONSE FORMAT:
Return a JSON object with the following keys ONLY:

- overall_score (integer between 0 and 25)
- selective_readiness_band (string that exactly matches the descriptor wording above)

- categories (object)
- teacher_feedback (3 to 4 sentences, professional selective-exam tone)

The categories object MUST contain the following keys exactly:
- audience_purpose_form
- ideas_content
- structure_organisation
- language_vocabulary
- grammar_spelling_punctuation

Each category MUST contain:
- score (integer between 0 and 5)
- strengths (array of 2–3 concise strings)
- improvements (array of 2–3 concise, actionable strings)

IMPORTANT CONSTRAINTS:
- The readiness band MUST match the overall score
- Do NOT invent information
- Do NOT be lenient
- Clearly identify accuracy issues where present
- Feedback must be suitable for parents and teachers
- Maintain a professional NSW Selective exam tone

Student response:
{payload.answer_text}
"""






    try:
        print("🚀 Calling OpenAI Responses API...")

        response = client.responses.create(
            model="gpt-4o-mini",
            input=prompt,
            temperature=0.4
        )


        print("🧠 Raw OpenAI response object:", response)

        ai_result = response.output_text
        print("🧠 OpenAI output_text:")
        print(ai_result)

    except Exception as e:
        print("❌ OpenAI evaluation failed!")
        print("❌ Exception type:", type(e))
        print("❌ Exception details:", str(e))
        raise HTTPException(
            status_code=500,
            detail="Writing evaluation failed"
        )

    # --------------------------------------------------
    # 4️⃣ Parse AI response
    # --------------------------------------------------
    print("🧪 Parsing AI response as JSON...")

    try:
        
        evaluation = json.loads(ai_result)

        writing_score = int(evaluation.get("overall_score", 0))
        
        categories = evaluation.get("categories", {})
        band = evaluation.get("selective_readiness_band")

        EXPECTED_BANDS = {
            range(22, 26): "Strong selective standard – very competitive",
            range(18, 22): "On track for selective with minor improvements",
            range(14, 18): "Developing – selective readiness needs strengthening",
            range(10, 14): "Below selective standard – significant improvement needed",
            range(0, 10): "Well below selective standard at this stage"
        }
        
        expected_band = None
        for score_range, label in EXPECTED_BANDS.items():
            if writing_score in score_range:
                expected_band = label
                break
        
        if band != expected_band:
            raise HTTPException(
                status_code=500,
                detail=f"Readiness band mismatch: expected '{expected_band}', got '{band}'"
            )

        required_categories = [
            "audience_purpose_form",
            "ideas_content",
            "structure_organisation",
            "language_vocabulary",
            "grammar_spelling_punctuation"
        ]
        
        # ✅ HARD STRUCTURE VALIDATION (CRITICAL)
        for key in required_categories:
            cat = categories.get(key)
        
            if not isinstance(cat, dict):
                raise HTTPException(
                    status_code=500,
                    detail=f"Invalid category format: {key}"
                )
        
            if not all(field in cat for field in ["score", "strengths", "improvements"]):
                raise HTTPException(
                    status_code=500,
                    detail=f"Category missing required fields: {key}"
                )
        
            if not isinstance(cat["strengths"], list) or not isinstance(cat["improvements"], list):
                raise HTTPException(
                    status_code=500,
                    detail=f"Category fields must be arrays: {key}"
                )

        exam_state.ai_evaluation_json = evaluation
        exam_state.ai_score = writing_score       
        
        
        print("✅ Parsed AI evaluation:", {
            "score": writing_score,
            "categories": list(categories.keys())
        })


    except Exception as e:
        print("❌ Failed to parse AI response as JSON")
        print("❌ Raw AI text was:")
        print(ai_result)
        print("❌ Exception:", str(e))
        raise HTTPException(
            status_code=500,
            detail="Invalid AI evaluation format"
        )

    # --------------------------------------------------
    # 5️⃣ Store AI evaluation on writing attempt
    # --------------------------------------------------
    print("💾 Storing AI evaluation on writing attempt...")

    
    # --------------------------------------------------
    # 6️⃣ Admin RAW SCORE snapshot
    # --------------------------------------------------
    print("📊 Creating AdminExamRawScore snapshot...")

    admin_raw_score = AdminExamRawScore(
        student_id=exam_state.student_id,
        exam_attempt_id=exam_state.id,
        subject="writing",
        total_questions=25,
        correct_answers=writing_score,
        wrong_answers=25 - writing_score,
        accuracy_percent=round((writing_score / 25) * 100, 2)

    )

    existing_raw = (
        db.query(AdminExamRawScore)
        .filter(
            AdminExamRawScore.exam_attempt_id == exam_state.id,
            AdminExamRawScore.subject == "writing"
        )
        .first()
    )
    
    if not existing_raw:
        db.add(admin_raw_score)

    print("✅ Admin raw score prepared:", {
        "student_id": admin_raw_score.student_id,
        "attempt_id": admin_raw_score.exam_attempt_id,
        "accuracy": admin_raw_score.accuracy_percent
    })
    # --------------------------------------------------
    # 6️⃣➕ Writing Admin Report Snapshot (FINAL STEP)
    # --------------------------------------------------
    print("📄 Generating Writing admin report snapshot...")
    
    # --------------------------------------------------
    # A) Idempotency guard (CRITICAL)
    # --------------------------------------------------
    existing_report = (
        db.query(AdminExamReport)
        .filter(
            AdminExamReport.exam_attempt_id == exam_state.id,
            AdminExamReport.exam_type == "writing"
        )

        .first()
    )

    
    if existing_report:
        print("⚠️ Writing admin report already exists. Skipping snapshot generation.")
    else:
        # --------------------------------------------------
        # B) Determine performance band & readiness
        # --------------------------------------------------
        if writing_score >= 22:
            performance_band = "Strong"
            readiness_status = "Very Competitive"
            guidance_text = "Student demonstrates strong writing skills and is highly competitive for selective placement."
        elif writing_score >= 18:
            performance_band = "On Track"
            readiness_status = "Ready"
            guidance_text = "Student is on track for selective readiness with minor improvements."
        elif writing_score >= 14:
            performance_band = "Developing"
            readiness_status = "Borderline"
            guidance_text = "Student shows developing writing skills and would benefit from targeted improvement."
        else:
            performance_band = "Weak"
            readiness_status = "Not Ready"
            guidance_text = "Student requires significant improvement in writing fundamentals."


    
        print("🧠 Writing classification:", {
            "score": writing_score,
            "band": performance_band,
            "readiness": readiness_status
        })
    
        # --------------------------------------------------
        # C) Insert admin_exam_reports
        # --------------------------------------------------
        admin_report = AdminExamReport(
            exam_attempt_id=exam_state.id,
            student_id=student.student_id,   # external id string (Gem002)
            exam_type="writing",
            overall_score=writing_score,
            readiness_band=readiness_status,
            school_guidance_level=guidance_text,
            summary_notes=f"Writing score: {writing_score}/25"
        )



    
        db.add(admin_report)
        db.flush()  # 🔑 Needed to get admin_report.id
    
        print("✅ AdminExamReport created:", {
            "report_id": admin_report.id
        })
    
        # --------------------------------------------------
        # D) Insert admin_exam_section_results (single section)
        # --------------------------------------------------
        section_result = AdminExamSectionResult(
            admin_report_id=admin_report.id,
            section_name="Writing",
            raw_score=writing_score,
            performance_band=performance_band,
            strengths_summary="Per-category strengths recorded",
            improvement_summary="Per-category improvements recorded"

        )

    
        db.add(section_result)
    
        print("✅ Writing section result stored")
    
        # --------------------------------------------------
        # E) Insert admin_readiness_rules_applied (audit trail)
        # --------------------------------------------------
        readiness_rule = AdminReadinessRuleApplied(
           admin_report_id=admin_report.id,
           rule_code="writing_score_threshold",
           rule_result=readiness_status,
           rule_description=f"Writing score {writing_score}/25 classified as {performance_band}"
        )

    
        db.add(readiness_rule)
    
        print("✅ Readiness rule audit stored")
    
    
        # --------------------------------------------------
        # 7️⃣ Commit atomically
        # --------------------------------------------------
        print("💾 Committing transaction to database...")
    
    try:
        db.commit()
        print("✅ Database commit successful")
    except Exception as e:
        print("❌ Database commit failed!")
        print("❌ Exception:", str(e))
        db.rollback()
        raise
    
    print("🎉 Writing exam submitted and evaluated successfully")
    print("====================================================\n")

    return {
        "message": "Writing exam submitted successfully",
        "exam_id": exam_state.exam_id,
        "writing_score": writing_score
    }

@app.get("/api/exams/writing/result")
def get_writing_result(
    student_id: str,
    db: Session = Depends(get_db)
):
    # --------------------------------------------------
    # 1️⃣ Resolve student (EXTERNAL → INTERNAL)
    # --------------------------------------------------
    student = (
        db.query(Student)
        .filter(func.lower(Student.student_id) == func.lower(student_id))
        .first()
    )

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # --------------------------------------------------
    # 2️⃣ Load latest admin report (SOURCE OF TRUTH)
    # --------------------------------------------------
    admin_report = (
        db.query(AdminExamReport)
        .filter(
            AdminExamReport.student_id == student.student_id,  # ✅ external ID
            AdminExamReport.exam_type == "writing"
        )
        .order_by(AdminExamReport.created_at.desc())
        .first()
    )

    if not admin_report:
        raise HTTPException(
            status_code=404,
            detail="Writing result not found"
        )

    # --------------------------------------------------
    # 3️⃣ Load writing attempt (for AI details)
    # --------------------------------------------------
    exam_state = (
        db.query(StudentExamWriting)
        .filter(
            StudentExamWriting.id == admin_report.exam_attempt_id
        )
        .first()
    )

    if not exam_state or not exam_state.ai_evaluation_json:
        raise HTTPException(
            status_code=500,
            detail="AI evaluation missing for writing exam"
        )

    # --------------------------------------------------
    # 4️⃣ Final response (UI SAFE)
    # --------------------------------------------------
    return {
        "exam_type": "Writing",
        "score": admin_report.overall_score,
        "max_score": 25,

        # ✅ canonical readiness band
        "selective_readiness_band": admin_report.readiness_band,

        # ✅ full AI breakdown
        "evaluation": exam_state.ai_evaluation_json,

        "advisory": "This report is advisory only and does not guarantee placement."
    }



@app.get("/api/exams/writing/current")
def get_current_writing_exam(student_id: str, db: Session = Depends(get_db)):

    # --------------------------------------------------
    # 0️⃣ Resolve student (external → internal)
    # --------------------------------------------------
    student = (
        db.query(Student)
        .filter(func.lower(Student.student_id) == func.lower(student_id))
        .first()
    )

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # --------------------------------------------------
    # 1️⃣ Fetch ACTIVE writing attempt ONLY
    # --------------------------------------------------
    attempt = (
        db.query(StudentExamWriting)
        .filter(
            StudentExamWriting.student_id == student.id,
            StudentExamWriting.completed_at.is_(None)
        )
        .order_by(StudentExamWriting.started_at.desc())
        .first()
    )

    if not attempt:
        # 🔑 This is intentional.
        # Frontend will call start-writing-exam.
        raise HTTPException(
            status_code=404,
            detail="No active writing exam"
        )

    # --------------------------------------------------
    # 2️⃣ Load exam definition
    # --------------------------------------------------
    exam = (
        db.query(GeneratedExamWriting)
        .filter(GeneratedExamWriting.id == attempt.exam_id)
        .first()
    )

    if not exam:
        raise HTTPException(
            status_code=404,
            detail="Writing exam not found"
        )

    # --------------------------------------------------
    # 3️⃣ Compute remaining time
    # --------------------------------------------------
    started_at = attempt.started_at
    if started_at.tzinfo is None:
        started_at = started_at.replace(tzinfo=timezone.utc)

    now = datetime.now(timezone.utc)
    elapsed = int((now - started_at).total_seconds())
    total_seconds = attempt.duration_minutes * 60
    remaining = max(0, total_seconds - elapsed)

    # --------------------------------------------------
    # 4️⃣ Timeout → auto-complete
    # --------------------------------------------------
    if remaining == 0:
        attempt.completed_at = now
        db.commit()
        return {"completed": True}

    # --------------------------------------------------
    # 5️⃣ Active attempt response
    # --------------------------------------------------
    return {
        "completed": False,
        "remaining_seconds": remaining,
        "exam": {
            "exam_id": exam.id,
            "difficulty": exam.difficulty,
            "question_text": exam.question_text,
            "duration_minutes": attempt.duration_minutes
        }
    }


@app.post("/api/exams/generate-writing")
def generate_exam_writing(
    payload: WritingGenerateSchema,
    db: Session = Depends(get_db)
):
    # ----------------------------------
    # Normalize inputs
    # ----------------------------------
    class_name = payload.class_name.strip().lower()
    difficulty = payload.difficulty.strip().lower()

    # ----------------------------------
    # Clear previous generated exam
    # ----------------------------------
    db.query(GeneratedExamWriting).delete(synchronize_session=False)
    db.commit()
    # ----------------------------------
    # Clear previous student exam attempts
    # ----------------------------------
    db.query(StudentExamWriting).delete(synchronize_session=False)
    db.commit()


    # ----------------------------------
    # Fetch ONE random writing question
    # ----------------------------------
    question = (
        db.query(WritingQuestionBank)
        .filter(
            func.trim(func.lower(WritingQuestionBank.class_name)) == class_name,
            func.trim(func.lower(WritingQuestionBank.difficulty)) == difficulty,
        )
        .order_by(func.random())
        .first()
    )

    if not question:
        raise HTTPException(
            status_code=404,
            detail=f"No writing question found for class '{class_name}' and difficulty '{difficulty}'."
        )

    # ----------------------------------
    # Compose FULL exam text (UI-ready)
    # ----------------------------------
    exam_text_parts = []

    if question.title:
        exam_text_parts.append(f"TITLE:\n{question.title}\n")

    exam_text_parts.append(
        f"TASK:\n{question.question_text}\n"
    )

    if question.statement:
        exam_text_parts.append(
            f"STATEMENT:\n{question.statement}\n"
        )

    exam_text_parts.append(
        f"INSTRUCTIONS:\n{question.question_prompt}\n"
    )

    if question.opening_sentence:
        exam_text_parts.append(
            f"OPENING SENTENCE:\n{question.opening_sentence}\n"
        )

    # Guidelines stored as newline-separated text
    if question.guidelines:
        formatted_guidelines = "\n".join(
            f"- {line.strip()}"
            for line in question.guidelines.splitlines()
            if line.strip()
        )

        exam_text_parts.append(
            f"GUIDELINES:\n{formatted_guidelines}\n"
        )

    full_exam_text = "\n".join(exam_text_parts).strip()

    # ----------------------------------
    # Persist generated exam
    # ----------------------------------
    exam = GeneratedExamWriting(
        class_name=class_name.capitalize(),
        subject="writing",
        difficulty=difficulty.capitalize(),
        question_text=full_exam_text,
        duration_minutes=40,
    )

    db.add(exam)
    db.commit()
    db.refresh(exam)

    # ----------------------------------
    # Return UI-friendly payload
    # ----------------------------------
    return {
        "exam_id": exam.id,
        "class_name": exam.class_name,
        "difficulty": exam.difficulty,
        "duration_minutes": exam.duration_minutes,
        "exam_text": full_exam_text,
    }


def parse_and_normalize_writing_with_openai(text: str) -> list[dict]:
    """
    Normalize loosely formatted writing questions into structured JSON.
    This function is tolerant of missing optional fields and never invents content.
    """

    WRITING_NORMALIZE_PROMPT = f"""
You are a document normalizer for an exam authoring system.

Convert the input text into STRICT, VALID JSON that matches the schema below.

Rules:
- Preserve all instructional meaning
- Do NOT invent content
- If a field is not explicitly present, return an empty string
- Always include ALL fields in the output
- Return JSON only (no markdown, no comments)

Schema (ALL fields required):

{{
  "class_name": string,
  "subject": "Writing",
  "topic": string,
  "difficulty": string,
  "title": string,

  "question_text": string,
  "question_prompt": string,

  "statement": string,
  "opening_sentence": string,
  "guidelines": string
}}

Guidelines formatting rules:
- guidelines = plain text
- newline-separated
- no bullets or numbering
- empty string allowed

Input text:
{text}
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Return valid JSON only."},
            {"role": "user", "content": WRITING_NORMALIZE_PROMPT},
        ],
        temperature=0,
    )

    raw = response.choices[0].message.content.strip()

    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError as e:
        raise ValueError(
            f"OpenAI returned invalid JSON.\nRaw response:\n{raw}"
        ) from e

    # Always normalize to list
    if isinstance(parsed, dict):
        parsed = [parsed]

    if not isinstance(parsed, list):
        raise ValueError("Expected JSON object or array from OpenAI")

    normalized = []

    REQUIRED_STRING_FIELDS = [
        "class_name",
        "subject",
        "topic",
        "difficulty",
        "question_text",
        "question_prompt",
        "statement",
        "opening_sentence",
        "guidelines",
    ]

    for item in parsed:
        if not isinstance(item, dict):
            continue

        # Ensure all required fields exist and are strings
        for field in REQUIRED_STRING_FIELDS:
            value = item.get(field)
            if not isinstance(value, str):
                value = ""
            item[field] = value.replace("\r\n", "\n").replace("\r", "\n").strip()

        # Optional title
        title = item.get("title")
        item["title"] = title.strip() if isinstance(title, str) else ""

        # Subject sanity check (soft)
        if item["subject"] and item["subject"] != "Writing":
            raise ValueError("Invalid subject. Expected 'Writing'.")

        # Normalize guidelines text (strip bullets safely)
        if item["guidelines"]:
            lines = []
            for line in item["guidelines"].splitlines():
                line = line.strip().lstrip("-•* ").strip()
                if line:
                    lines.append(line)
            item["guidelines"] = "\n".join(lines)

        normalized.append(item)

    if not normalized:
        raise ValueError("No valid writing questions detected")

    return normalized

@app.post("/upload-word-writing")
async def upload_word_writing(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    print("\n=== UPLOAD WORD WRITING ===")

    if file.content_type != (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        raise HTTPException(400, "File must be .docx")

    raw = await file.read()
    extracted_text = extract_text_from_docx(raw)

    try:
        parsed_list = parse_and_normalize_writing_with_openai(extracted_text)
    except Exception as e:
        print("❌ Writing normalization error:", e)
        raise HTTPException(400, str(e))

    if not parsed_list:
        raise HTTPException(400, "No writing questions detected")

    saved_ids = []

    for parsed in parsed_list:
        obj = WritingQuestionBank(
            class_name=parsed["class_name"],
            subject=parsed["subject"],
            topic=parsed["topic"],
            difficulty=parsed["difficulty"],
            
            title=parsed.get("title"),
            question_text=parsed["question_text"],
            question_prompt=parsed["question_prompt"],
            statement=parsed.get("statement"),
            opening_sentence=parsed.get("opening_sentence"),
        
            guidelines=parsed["guidelines"],          # ✅ ADD THIS
        
            source_file=file.filename
        )



        db.add(obj)
        db.flush()          # get ID before commit
        saved_ids.append(obj.id)

    db.commit()

    return {
        "message": "Writing questions uploaded successfully",
        "count": len(saved_ids),
        "question_ids": saved_ids
    }


 
@app.post("/api/quizzes-writing")
def save_writing_quiz(payload: WritingQuizSchema, db: Session = Depends(get_db)):
    print("\n--- Deleting previous Writing quiz setups ---")

    db.query(QuizSetupWriting).delete(synchronize_session=False)
    db.commit()

    print("🗑️ Previous writing quiz setups deleted")
    print("\n--- Creating new Writing quiz setup ---")
    quiz = QuizSetupWriting(
        class_name=payload.class_name,
        subject=payload.subject,
        topic=payload.topic,
        difficulty=payload.difficulty
    ) 

    db.add(quiz)
    db.commit()
    db.refresh(quiz)

    return {
        "message": "Writing quiz setup saved",
        "quiz_id": quiz.id
    }


# ------------------------------------------------------------
# 🔧 Helper: Build sections with questions (FOUNDATIONAL)
# ------------------------------------------------------------
def build_sections_with_questions(exam_json):
    print("\n🧠 Normalizing sections from exam_json...")

    sections_meta = exam_json.get("sections", [])
    questions = exam_json.get("questions", [])

    print("   • sections_meta count:", len(sections_meta))
    print("   • questions count:", len(questions))

    section_map = {}

    # ✅ Initialize sections using DIFFICULTY (not name)
    for s in sections_meta:
        difficulty = s.get("difficulty")
        topic = s.get("topic")

        if not difficulty:
            print("⚠️ Skipping section with no difficulty:", s)
            continue

        section_map[difficulty] = {
            "difficulty": difficulty,
            "topic": topic,
            "time": s.get("time", 0),
            "intro": s.get("intro", ""),
            "questions": []
        }

    # ✅ Assign questions to sections
    for q in questions:
        section_key = q.get("section")

        if section_key in section_map:
            section_map[section_key]["questions"].append(q)
        else:
            print("⚠️ Question with unknown section:", q.get("section"))

    sections = list(section_map.values())
    print("📦 Sections normalized:", len(sections))

    for sec in sections:
        print(
            f"   • {sec['difficulty']} | topic={sec['topic']} | "
            f"questions={len(sec['questions'])}"
        )

    return sections

def normalize_naplan_language_conventions_questions_live(raw_questions):
    """
    Build EXAM-SAFE question DTOs.
    No answers. No metadata. No leakage.
    """

    normalized = []

    for q in raw_questions or []:
        blocks = q.get("question_blocks") or []
        display_blocks = []

        correct_answer = str(q.get("correct_answer")).strip()

        for block in blocks:
            content = (block.get("content") or "").strip()
            lowered = content.lower()

            # Skip metadata
            if lowered.startswith("question_type"):
                continue
            if lowered.startswith("year:"):
                continue
            if lowered.startswith("answer_type"):
                continue
            if lowered.startswith("correct_answer"):
                continue

            # Skip raw answer leakage
            if content == correct_answer:
                continue

            display_blocks.append(block)

        normalized.append({
            "id": q["id"],
            "question_type": q["question_type"],
            "topic": q.get("topic"),
            "difficulty": q.get("difficulty"),
            "question_blocks": display_blocks,
            "options": q.get("options")
        })

    return normalized

@app.post("/api/student/start-exam/naplan-language-conventions")
def start_naplan_language_conventions_exam(
    req: StartExamRequest = Body(...),
    db: Session = Depends(get_db)
):
    print("\n================ START NAPLAN LANGUAGE CONVENTIONS EXAM =================")
    print("📥 Incoming payload:", req.dict())

    # --------------------------------------------------
    # 1️⃣ Resolve student
    # --------------------------------------------------
    student = (
        db.query(Student)
        .filter(
            func.lower(Student.student_id) ==
            func.lower(req.student_id.strip())
        )
        .first()
    )

    if not student:
        print("❌ Student not found:", repr(req.student_id))
        raise HTTPException(status_code=404, detail="Student not found")

    print(
        f"✅ Student resolved | "
        f"student_id={student.student_id} | "
        f"internal_id={student.id}"
    )

    # --------------------------------------------------
    # 2️⃣ Fetch latest attempt
    # --------------------------------------------------
    attempt = (
        db.query(StudentExamNaplanLanguageConventions)
        .filter(
            StudentExamNaplanLanguageConventions.student_id == student.id
        )
        .order_by(
            StudentExamNaplanLanguageConventions.started_at.desc()
        )
        .first()
    )

    MAX_DURATION = timedelta(minutes=40)
    now = datetime.now(timezone.utc)

    if attempt:
        print(
            "🔢 Existing attempt found | "
            f"attempt_id={attempt.id} | "
            f"started_at={attempt.started_at} | "
            f"completed_at={attempt.completed_at}"
        )

        started_at = attempt.started_at
        if started_at.tzinfo is None:
            started_at = started_at.replace(tzinfo=timezone.utc)

        expires_at = started_at + MAX_DURATION
        elapsed = int((now - started_at).total_seconds())

        print(
            "⏱️ Timer evaluation | "
            f"now={now} | "
            f"expires_at={expires_at} | "
            f"elapsed_seconds={elapsed}"
        )

        # ⛔ Expired attempt → auto-complete
        if attempt.completed_at is None and now > expires_at:
            attempt.completed_at = expires_at
            db.commit()

            print("➡️ Returning: completed=true (timeout)")
            return {"completed": True}

        # ✅ Already completed
        if attempt.completed_at is not None:
            print("➡️ Returning: completed=true (already completed)")
            return {"completed": True}

        # ▶ Resume active attempt
        remaining = max(0, attempt.duration_minutes * 60 - elapsed)

        exam = (
            db.query(ExamNaplanLanguageConventions)
            .filter(
                func.lower(ExamNaplanLanguageConventions.class_name) ==
                func.lower(student.class_name),
                func.lower(ExamNaplanLanguageConventions.subject) ==
                "language conventions"
            )
            .order_by(
                ExamNaplanLanguageConventions.created_at.desc()
            )
            .first()
        )

        if not exam:
            raise HTTPException(
                status_code=404,
                detail="NAPLAN Language Conventions exam not found"
            )

        # 🔥 SANITIZE BEFORE RETURNING
        normalized_questions = (
            normalize_naplan_language_conventions_questions_live(
                exam.questions or []
            )
        )

        return {
            "completed": False,
            "questions": normalized_questions,
            "remaining_time": remaining
        }

    # --------------------------------------------------
    # 🆕 FIRST ATTEMPT
    # --------------------------------------------------
    print("🆕 No existing attempt → creating first Language Conventions attempt")

    exam = (
        db.query(ExamNaplanLanguageConventions)
        .filter(
            func.lower(ExamNaplanLanguageConventions.class_name) ==
            func.lower(student.class_name),
            func.lower(ExamNaplanLanguageConventions.subject) ==
            "language conventions"
        )
        .order_by(
            ExamNaplanLanguageConventions.created_at.desc()
        )
        .first()
    )

    if not exam:
        raise HTTPException(
            status_code=404,
            detail="NAPLAN Language Conventions exam not found"
        )

    normalized_questions = (
        normalize_naplan_language_conventions_questions_live(
            exam.questions or []
        )
    )

    new_attempt = StudentExamNaplanLanguageConventions(
        student_id=student.id,
        exam_id=exam.id,
        started_at=now,
        duration_minutes=40
    )

    db.add(new_attempt)
    db.commit()
    db.refresh(new_attempt)

    print(
        "🆕 New attempt created | "
        f"attempt_id={new_attempt.id}"
    )

    # --------------------------------------------------
    # Pre-create response rows
    # --------------------------------------------------
    for original_q in exam.questions or []:
        correct_option = original_q["correct_answer"]

        if not isinstance(correct_option, list):
            correct_option = [correct_option]

        db.add(
            StudentExamResponseNaplanLanguageConventions(
                student_id=student.id,
                exam_id=exam.id,
                exam_attempt_id=new_attempt.id,
                q_id=original_q["id"],
                topic=original_q.get("topic"),
                selected_option=None,
                correct_option=correct_option,
                is_correct=None
            )
        )

    db.commit()

    print("================ END START NAPLAN LANGUAGE CONVENTIONS EXAM ================\n")

    return {
        "completed": False,
        "questions": normalized_questions,
        "remaining_time": new_attempt.duration_minutes * 60
    }
 

@app.post("/api/student/start-exam/foundational-skills")
def start_or_resume_foundational_exam(
    payload: StartExamRequestFoundational,
    db: Session = Depends(get_db)
):
    print("\n" + "=" * 70)
    print("🚀 START-EXAM (FOUNDATIONAL) REQUEST RECEIVED")
    print("=" * 70)

    # ------------------------------------------------------------
    # 0️⃣ Payload
    # ------------------------------------------------------------
    print("🧾 Raw payload:", payload)
    student_id = payload.student_id
    print("👤 student_id:", student_id)

    # ------------------------------------------------------------
    # 1️⃣ Fetch latest attempt
    # ------------------------------------------------------------
    print("\n🔍 Fetching latest exam attempt...")
    attempt = (
        db.query(StudentsExamFoundational)
        .filter(StudentsExamFoundational.student_id == student_id)
        .order_by(StudentsExamFoundational.started_at.desc())
        .first()
    )

    print("📌 Attempt found:", bool(attempt))
    if attempt:
        print("   • attempt.id:", attempt.id)
        print("   • completed_at:", attempt.completed_at)
        print("   • current_section_index:", attempt.current_section_index)

    if attempt and attempt.completed_at:
        print("✅ Attempt already completed — returning early")
        return {"completed": True}

    # ------------------------------------------------------------
    # 2️⃣ Load current exam
    # ------------------------------------------------------------
    print("\n📘 Fetching current GeneratedExamFoundational...")
    exam = (
        db.query(GeneratedExamFoundational)
        .filter(GeneratedExamFoundational.is_current == True)
        .order_by(desc(GeneratedExamFoundational.created_at))
        .first()
    )

    if not exam:
        print("❌ ERROR: No active exam found")
        raise HTTPException(404, "No active exam")

    print("✅ Exam loaded")
    print("   • exam.id:", exam.id)
    print("   • class_name:", exam.class_name)
    print("   • subject:", exam.subject)
    print("   • duration_minutes:", exam.duration_minutes)
    print("   • exam_json keys:", exam.exam_json.keys())

    # ------------------------------------------------------------
    # 3️⃣ Normalize sections + questions
    # ------------------------------------------------------------
    print("\n🧠 Normalizing sections from exam_json...")
    sections = build_sections_with_questions(exam.exam_json)

    print("📦 Sections normalized:", len(sections))
    for i, sec in enumerate(sections):
        print(f"   Section {i} keys:", sec.keys())

    if not sections:
        print("❌ ERROR: Sections list is empty after normalization")
        raise HTTPException(500, "No sections after normalization")

    # ------------------------------------------------------------
    # 4️⃣ Create attempt if needed
    # ------------------------------------------------------------
    if not attempt:
        print("\n🆕 Creating new exam attempt...")
        attempt = StudentsExamFoundational(
            student_id=student_id,
            exam_id=exam.id,
            started_at=datetime.now(timezone.utc),
            current_section_index=0
        )
        db.add(attempt)
        db.commit()
        db.refresh(attempt)

        print("✅ Attempt created")
        print("   • attempt.id:", attempt.id)

    # ------------------------------------------------------------
    # 5️⃣ Validate section index
    # ------------------------------------------------------------
    print("\n🧪 Validating section index...")
    print("   • current_section_index:", attempt.current_section_index)
    print("   • total sections:", len(sections))

    if attempt.current_section_index >= len(sections):
        print("❌ ERROR: current_section_index out of bounds")
        raise HTTPException(500, "Invalid section index")

    # ------------------------------------------------------------
    # 6️⃣ Timer calculation
    # ------------------------------------------------------------
    print("\n⏱ Calculating remaining time...")
    elapsed = int((datetime.now(timezone.utc) - attempt.started_at).total_seconds())
    print("   • elapsed seconds:", elapsed)

    if exam.duration_minutes is None:
        print("❌ ERROR: exam.duration_minutes is None")
        raise HTTPException(500, "Invalid exam duration")

    remaining_time = max(exam.duration_minutes * 60 - elapsed, 0)
    print("   • remaining_time (seconds):", remaining_time)

    # ------------------------------------------------------------
    # 7️⃣ Normalize questions for frontend
    # ------------------------------------------------------------
    print("\n🧩 Preparing current section for frontend...")
    current_section = sections[attempt.current_section_index]

    print("📌 Current section object:", current_section)
    print("📌 Current section keys:", current_section.keys())

    questions = current_section.get("questions")
    if questions is None:
        print("❌ ERROR: current_section['questions'] is None")
        raise HTTPException(500, "Section has no questions key")

    print("📊 Question count in section:", len(questions))

    normalized_questions = []
    for idx, q in enumerate(questions):
        print(f"   → Normalizing question {idx + 1}")
        fixed = dict(q)

        opts = fixed.get("options")

        if isinstance(opts, dict) and opts:
           fixed["options"] = [f"{k}) {v}" for k, v in opts.items()]
        else:
           # 🔒 Guarantee at least 1 safe option to avoid frontend crash
           fixed["options"] = ["A) Option unavailable"]


        normalized_questions.append(fixed)

    # ------------------------------------------------------------
    # 8️⃣ Final response
    # ------------------------------------------------------------
    print("\n✅ START-EXAM RESPONSE READY")
    print("   • section name:", current_section.get("difficulty"))
    print("   • questions returned:", len(normalized_questions))
    print("=" * 70 + "\n")

    return {
        "completed": False,
        "current_section_index": attempt.current_section_index,
        "section": {
            "name": current_section["difficulty"],
            "questions": normalized_questions
        },
        "remaining_time": remaining_time
    }



@app.post("/api/exams/foundational/next-section")
def advance_foundational_section(
    student_id: str = Query(...),
    db: Session = Depends(get_db)
):
    print("\n================ NEXT-SECTION (FOUNDATIONAL) ================")
    print("👤 student_id:", student_id)

    # 1️⃣ Get active attempt
    attempt = (
        db.query(StudentsExamFoundational)
        .filter(
            StudentsExamFoundational.student_id == student_id,
            StudentsExamFoundational.completed_at.is_(None)
        )
        .order_by(StudentsExamFoundational.started_at.desc())
        .first()
    )

    if not attempt:
        raise HTTPException(status_code=404, detail="No active attempt")

    print("🧪 attempt.id:", attempt.id)
    print("   ↳ current_section_index:", attempt.current_section_index)

    # 2️⃣ Load exam snapshot
    exam = db.query(GeneratedExamFoundational).get(attempt.exam_id)
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")

    exam_json = exam.exam_json

    # 3️⃣ Build sections DIRECTLY from saved exam_json
    sections_meta = exam_json.get("sections", [])
    all_questions = exam_json.get("questions", [])

    if not sections_meta or not all_questions:
        raise HTTPException(
            status_code=500,
            detail="Exam data is incomplete"
        )

    # Group questions by section name
    sections = []
    for sec in sections_meta:
        sec_name = sec["difficulty"]  # ✅ FIX
    
        sec_questions = [
            q for q in all_questions
            if q.get("section") == sec_name
        ]
    
        sections.append({
            "name": sec_name,
            "questions": sec_questions
        })


    total_sections = len(sections)
    print("📂 Total sections:", total_sections)

    # 4️⃣ Last section check
    if attempt.current_section_index >= total_sections - 1:
        print("🏁 Last section reached")
        return {"completed": True}

    # 5️⃣ Advance section index
    attempt.current_section_index += 1
    db.commit()
    db.refresh(attempt)

    current_section = sections[attempt.current_section_index]

    print("➡️ Advanced to section:", current_section["name"])
    print("   ↳ questions:", len(current_section["questions"]))

    if not current_section["questions"]:
        print("⚠️ WARNING: Section has no questions")

    print("===========================================================\n")

    return {
        "completed": False,
        "current_section_index": attempt.current_section_index,
        "section": {
            "name": current_section["name"],
            "questions": current_section["questions"]
        }
    }


@app.post("/api/student/finish-exam/foundational-skills")
def finish_foundational_exam(
    payload: FinishExamRequestFoundational,
    db: Session = Depends(get_db)
):
    print("\n================ FINISH EXAM (FOUNDATIONAL) ================")
    print("📥 Payload received:", payload.dict())

    # ------------------------------------------------------------
    # 1️⃣ Fetch active attempt
    # ------------------------------------------------------------
    attempt = (
        db.query(StudentsExamFoundational)
        .filter(
            StudentsExamFoundational.student_id == payload.student_id,
            StudentsExamFoundational.completed_at.is_(None)
        )
        .order_by(StudentsExamFoundational.started_at.desc())
        .first()
    )

    print("🔍 Active attempt:", attempt)

    if not attempt:
        print("⚠️ No active attempt found")
        return {
            "success": True,
            "message": "Exam already completed or no active attempt"
        }

    # ------------------------------------------------------------
    # 2️⃣ Load exam definition
    # ------------------------------------------------------------
    exam = db.query(GeneratedExamFoundational).get(attempt.exam_id)

    if not exam:
        print("❌ Exam not found for exam_id:", attempt.exam_id)
        return {"success": False, "message": "Exam not found"}

    questions = exam.exam_json.get("questions", [])
    print(f"📄 Questions in exam_json: {len(questions)}")

    # ------------------------------------------------------------
    # 3️⃣ Build question lookup (q_id → metadata)
    # ------------------------------------------------------------
    question_lookup = {}

    for q in questions:
        qid = q.get("q_id")
        if qid is None:
            continue

        question_lookup[qid] = {
            "section": q.get("section"),
            "topic": q.get("topic"),
            "correct_answer": q.get("correct")
        }

    print("🧠 Question lookup keys:", list(question_lookup.keys()))

    # ------------------------------------------------------------
    # 4️⃣ Persist responses (NO attempt-level idempotency guard)
    # ------------------------------------------------------------
    answers = payload.answers or {}
    print("📝 Answers received:", answers)

    inserted_count = 0

    for q_id_raw, selected_answer in answers.items():
        try:
            q_id = int(q_id_raw)
        except (TypeError, ValueError):
            print("❌ Invalid q_id:", q_id_raw)
            continue

        meta = question_lookup.get(q_id)
        if not meta:
            print("❌ No metadata for q_id:", q_id)
            continue

        correct_answer = meta["correct_answer"]
        is_correct = (
            selected_answer.upper() == correct_answer.upper()
            if selected_answer and correct_answer
            else False
        )

        response = StudentExamResponseFoundational(
            student_id=payload.student_id,
            exam_id=attempt.exam_id,
            attempt_id=attempt.id,
            section_name=meta["section"],
            topic=meta["topic"],                 # ✅ CRITICAL FIX
            question_id=q_id,
            selected_answer=selected_answer,
            correct_answer=correct_answer,
            is_correct=is_correct
        )

        db.add(response)
        inserted_count += 1

    print(f"📊 Responses added: {inserted_count}")

    # ------------------------------------------------------------
    # 5️⃣ Finalize attempt
    # ------------------------------------------------------------
    attempt.answers_json = answers
    attempt.completed_at = datetime.now(timezone.utc)
    attempt.completion_reason = payload.reason

    try:
        db.commit()
        print("✅ DB commit successful")
    except Exception as e:
        print("🔥 DB COMMIT FAILED:", str(e))
        db.rollback()
        raise

    print("🔄 Attempt finalized:", attempt.id)
    print("================ END FINISH EXAM =================\n")

    return {
        "success": True,
        "message": "Exam completed successfully"
    }



@app.get("/api/student/exam-report/foundational-skills")
def get_foundational_exam_report(
    student_id: str = Query(...),
    db: Session = Depends(get_db)
):
    print("\n================ EXAM-REPORT (FOUNDATIONAL) ================")
    print("👤 student_id:", student_id)

    attempt = (
        db.query(StudentsExamFoundational)
        .filter(
            StudentsExamFoundational.student_id == student_id,
            StudentsExamFoundational.completed_at.isnot(None)
        )
        .order_by(StudentsExamFoundational.completed_at.desc())
        .first()
    )

    if not attempt:
        raise HTTPException(
            status_code=404,
            detail="No completed foundational exam found"
        )

    exam = (
        db.query(GeneratedExamFoundational)
        .filter(GeneratedExamFoundational.id == attempt.exam_id)
        .first()
    )

    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")

    sections = build_sections_with_questions(exam.exam_json)

    topic_totals = {}
    total_questions = 0

    for sec in sections:
        for q in sec.get("questions", []):
            topic = q.get("topic", "Unknown")
            topic_totals.setdefault(topic, 0)
            topic_totals[topic] += 1
            total_questions += 1

    responses = (
        db.query(StudentExamResponseFoundational)
        .filter(StudentExamResponseFoundational.attempt_id == attempt.id)
        .all()
    )

    topic_stats = {}

    for topic, total in topic_totals.items():
        topic_stats[topic] = {
            "topic": topic,
            "total": total,
            "attempted": 0,
            "correct": 0,
            "incorrect": 0,
            "not_attempted": total
        }

    for r in responses:
        topic = r.topic
        stats = topic_stats.get(topic)

        if not stats:
            continue

        stats["attempted"] += 1
        stats["not_attempted"] -= 1

        if r.is_correct:
            stats["correct"] += 1
        else:
            stats["incorrect"] += 1

    topic_wise_performance = list(topic_stats.values())

    attempted = sum(t["attempted"] for t in topic_wise_performance)
    correct = sum(t["correct"] for t in topic_wise_performance)
    incorrect = sum(t["incorrect"] for t in topic_wise_performance)
    not_attempted = total_questions - attempted

    accuracy = round((correct / attempted) * 100, 1) if attempted else 0
    score = round((correct / total_questions) * 100, 1) if total_questions else 0

    improvement_areas = []

    for t in topic_wise_performance:
        acc = round((t["correct"] / t["total"]) * 100, 1) if t["total"] else 0
        improvement_areas.append({
            "topic": t["topic"],
            "accuracy": acc,
            "limited_data": t["total"] < 3
        })

    improvement_areas.sort(key=lambda x: x["accuracy"])

    return {
        "overall": {
            "total_questions": total_questions,
            "attempted": attempted,
            "correct": correct,
            "incorrect": incorrect,
            "not_attempted": not_attempted,
            "accuracy_percent": accuracy,
            "score_percent": score
        },
        "topic_wise_performance": topic_wise_performance,
        "improvement_areas": improvement_areas
    }
 
# def generate_ai_questions_foundational(
#     class_name: str,
#     subject: str,
#     difficulty: str,
#     topic: str,
#     count: int
# ):
#     """
#     Generate AI questions for a Foundational exam section.
#     Uses chunking + retries to reduce JSON corruption.
#     Compatible with current OpenAI Python SDK.
#     """
#
#     print("\n🤖 Calling AI to generate questions")
#     print(f"🤖 class={class_name}, subject={subject}, difficulty={difficulty}, count={count}")
#
#     if count <= 0:
#         return []
#
#     CHUNK_SIZE = 5
#     MAX_RETRIES = 2
#
#     remaining = count
#     all_questions = []
#
#     while remaining > 0:
#         batch_size = min(CHUNK_SIZE, remaining)
#
#         prompt = f"""
# You are a STRICT JSON generator for an automated exam system.
#
# TASK:
# Generate EXACTLY {batch_size} multiple-choice questions.
#
# ABSOLUTE RULES:
# - Output MUST be valid JSON
# - Output MUST be a JSON ARRAY
# - NO explanations
# - NO comments
# - NO markdown
# - NO text outside JSON
#
# REQUIRED FORMAT:
# [
#   {{
#     "question_text": "string",
#     "options": {{
#       "A": "string",
#       "B": "string",
#       "C": "string",
#       "D": "string"
#     }},
#     "correct_answer": "A|B|C|D",
#     "topic": "string"
#   }}
# ]
#
# CONSTRAINTS:
# - Class: {class_name}
# - Subject: {subject}
# - Difficulty: {difficulty}
#
# If you cannot comply perfectly, RETURN [] ONLY.
# """
#
#         success = False
#
#         for attempt in range(1, MAX_RETRIES + 1):
#             print(f"🤖 AI batch {batch_size} | attempt {attempt}")
#
#             try:
#                 response = client.responses.create(
#                     model="gpt-4o-mini",
#                     input=prompt,
#                     temperature=0.4
#                 )
#
#                 raw = response.output_text
#                 print("🤖 Raw AI response:", raw)
#
#                 parsed = json.loads(raw)
#
#                 if not isinstance(parsed, list):
#                     raise ValueError("AI output is not a list")
#
#                 if len(parsed) != batch_size:
#                     raise ValueError(
#                         f"Expected {batch_size} questions, got {len(parsed)}"
#                     )
#
#                 for q in parsed:
#                     all_questions.append({
#                         "section": difficulty,
#                         "question_number": None,
#                         "question_text": q["question_text"],
#                         "options": q["options"],
#                         "correct_answer": q["correct_answer"],
#                         "question_type": "mcq",
#                         "images": [],
#                         "topic": q.get("topic", "AI Generated"),
#                     })
#
#                 remaining -= batch_size
#                 success = True
#                 break
#
#             except Exception as e:
#                 print(f"❌ AI batch failed (attempt {attempt}): {e}")
#
#         if not success:
#             print("❌ AI failed after retries. Stopping generation.")
#             break
#
#     print(f"🤖 Successfully generated {len(all_questions)} AI questions")
#     return all_questions

def is_valid_ai_question(q: dict) -> bool:
    try:
        if not isinstance(q, dict):
            return False

        if not isinstance(q.get("question_text"), str):
            return False

        options = q.get("options")
        if not isinstance(options, dict):
            return False

        if set(options.keys()) != {"A", "B", "C", "D"}:
            return False

        if not all(isinstance(v, str) for v in options.values()):
            return False

        if q.get("correct_answer") not in {"A", "B", "C", "D"}:
            return False

        return True

    except Exception:
        return False


def generate_ai_questions_safely(
    class_name: str,
    subject: str,
    difficulty: str,
    topic: str,
    required_count: int,
):
    """
    Safely generate AI questions with strict validation.
    - Accepts partial success
    - Skips invalid questions
    - Bounded retries
    - NEVER raises on AI failure
    """

    print("\n🤖 GENERATE AI QUESTIONS SAFELY")
    print("--------------------------------------------------")
    print(f"Class      : {class_name}")
    print(f"Subject    : {subject}")
    print(f"Difficulty : {difficulty}")
    print(f"Topic      : {topic}")
    print(f"Required   : {required_count}")
    print("--------------------------------------------------")

    if required_count <= 0:
        return []

    CHUNK_SIZE = 5
    MAX_ATTEMPTS = 5

    valid_questions = []
    attempts = 0

    while len(valid_questions) < required_count and attempts < MAX_ATTEMPTS:
        attempts += 1
        remaining = required_count - len(valid_questions)
        batch_size = min(CHUNK_SIZE, remaining)

        print(f"\n🤖 Attempt {attempts}")
        print(f"➡ Requesting {batch_size} questions")

        prompt = f"""
You are a STRICT JSON generator.

RETURN ONLY VALID JSON.
RETURN ONLY A JSON ARRAY.
RETURN EXACTLY {batch_size} OBJECTS.

FORMAT (STRICT):
[
  {{
    "question_text": "string",
    "options": {{
      "A": "string",
      "B": "string",
      "C": "string",
      "D": "string"
    }},
    "correct_answer": "A|B|C|D"
  }}
]

CONSTRAINTS:
- Class: {class_name}
- Subject: {subject}
- Difficulty: {difficulty}
- Topic: {topic}

RULES:
- No extra keys
- No markdown
- No explanations
- If unsure, return []
"""

        try:
            response = client.responses.create(
                model="gpt-4o-mini",
                input=prompt,
                temperature=0.3,
            )

            raw = response.output_text
            print("📥 Raw AI response:", raw)

            try:
                parsed = json.loads(raw)
            except json.JSONDecodeError as e:
                print("❌ JSON parse failed:", e)
                continue

            if not isinstance(parsed, list):
                print("❌ AI output is not a list")
                continue

            print(f"📦 Parsed {len(parsed)} candidate questions")

            for idx, q in enumerate(parsed, start=1):
                if len(valid_questions) >= required_count:
                    break

                if is_valid_ai_question(q):
                    valid_questions.append({
                        "section": difficulty,
                        "question_text": q["question_text"],
                        "options": q["options"],
                        "correct_answer": q["correct_answer"],
                        "question_type": "text_mcq",
                        "images": [],
                        "topic": topic,        # 🔒 FORCE section topic
                        "source": "ai",
                    })
                    print(f"✅ Accepted question {len(valid_questions)}")
                else:
                    print(f"⚠️  Rejected invalid question #{idx}")

        except Exception as e:
            print("❌ AI call failed:", e)

    print("\n--------------------------------------------------")
    print(f"🤖 AI generation finished")
    print(f"✔ Valid questions collected: {len(valid_questions)}")
    print(f"⚠ Attempts used: {attempts}/{MAX_ATTEMPTS}")
    print("--------------------------------------------------\n")

    return valid_questions



@app.post("/api/exams/generate-foundational")
def generate_exam_foundational(
    payload: dict = Body(...),
    db: Session = Depends(get_db)
):
    print("\n" + "=" * 70)
    print("📥 GENERATE FOUNDATIONAL EXAM - REQUEST RECEIVED")
    print("=" * 70)

    # ------------------------------------------------------------
    # 0️⃣ Payload validation
    # ------------------------------------------------------------
    print("🧾 Raw payload:", payload)

    class_name = payload.get("class_name")
    if not class_name:
        print("❌ ERROR: class_name missing from payload")
        raise HTTPException(400, "class_name is required")

    print("✅ class_name:", class_name)

    # ------------------------------------------------------------
    # 1️⃣ Clear previous generated exams
    # ------------------------------------------------------------
    print("🧹 Clearing previous GeneratedExamFoundational records...")
    deleted = (
        db.query(GeneratedExamFoundational)
        .filter(func.lower(GeneratedExamFoundational.class_name) == class_name.lower())
        .delete(synchronize_session=False)
    )
    db.commit()
    print(f"🧹 Deleted {deleted} rows")

    # ------------------------------------------------------------
    # 2️⃣ Load latest quiz setup
    # ------------------------------------------------------------
    print("🔍 Fetching latest QuizSetupFoundational...")

    cfg = (
        db.query(QuizSetupFoundational)
        .filter(func.lower(QuizSetupFoundational.class_name) == class_name.lower())
        .order_by(QuizSetupFoundational.id.desc())
        .first()
    )

    if not cfg:
        print(f"❌ ERROR: No quiz setup found for class '{class_name}'")
        raise HTTPException(404, f"No quiz setup found for class '{class_name}'")

    print("✅ Quiz setup found")
    print("   ▸ Config ID:", cfg.id)
    print("   ▸ Subject:", cfg.subject)

    # ------------------------------------------------------------
    # 3️⃣ Build section metadata
    # ------------------------------------------------------------
    print("\n📦 Building section definitions from config...")

    sections = []

    def add_section(label, name, topic, time, intro):
        print(f"➡ Checking {label}...")
        if name:
            print(f"   ✔ Included")
            print(f"   • Difficulty:", name)
            print(f"   • Topic:", topic)
            print(f"   • Time:", time)
            sections.append({
                "difficulty": name,
                "topic": topic,
                "time": time or 0,
                "intro": intro or "",
                "label": label,
            })
        else:
            print(f"   ⏭ Skipped (no name)")

    add_section(
        "Section 1",
        cfg.section1_name,
        cfg.section1_topic,
        cfg.section1_time,
        cfg.section1_intro,
    )

    add_section(
        "Section 2",
        cfg.section2_name,
        cfg.section2_topic,
        cfg.section2_time,
        cfg.section2_intro,
    )

    add_section(
        "Section 3",
        cfg.section3_name,
        cfg.section3_topic,
        cfg.section3_time,
        cfg.section3_intro,
    )

    if not sections:
        print("❌ ERROR: No sections created from config")
        raise HTTPException(400, "No sections defined in quiz setup")

    print(f"✅ Total sections to generate: {len(sections)}")

    # ------------------------------------------------------------
    # 4️⃣ Generate questions per section
    # ------------------------------------------------------------
    final_questions = []
    warnings = []

    for section in sections:
        print("\n" + "-" * 60)
        print(f"📘 GENERATING SECTION: {section['label']}")
        print("-" * 60)

        difficulty = section["difficulty"]
        topic = section["topic"]

        print("➡ Difficulty:", difficulty)
        print("➡ Topic:", topic)

        # Resolve counts
        if difficulty == cfg.section1_name:
            db_required = cfg.section1_db or 0
            ai_required = cfg.section1_ai or 0
            total_required = cfg.section1_total or 0
        elif difficulty == cfg.section2_name:
            db_required = cfg.section2_db or 0
            ai_required = cfg.section2_ai or 0
            total_required = cfg.section2_total or 0
        elif difficulty == cfg.section3_name:
            db_required = cfg.section3_db or 0
            ai_required = cfg.section3_ai or 0
            total_required = cfg.section3_total or 0
        else:
            print("⚠️ WARNING: Difficulty does not match any config section")
            continue

        print("📊 Required counts:")
        print("   • DB:", db_required)
        print("   • AI:", ai_required)
        print("   • TOTAL:", total_required)

        section_questions = []

        # ---------------- DB QUESTIONS ----------------
        if db_required > 0:
            print("\n🗄 Fetching DB questions...")

            pool = (
                db.query(Question)
                .filter(
                    func.lower(Question.class_name) == class_name.lower(),
                    func.lower(Question.subject) == cfg.subject.lower(),
                    func.lower(Question.difficulty) == difficulty.lower(),
                    func.lower(Question.topic) == topic.lower(),
                )
                .all()
            )

            print(f"📦 DB pool size: {len(pool)}")

            random.shuffle(pool)
            chosen = pool[:db_required]

            print(f"🎯 Selected DB questions: {len(chosen)}")

            if len(chosen) < db_required:
                warning = (
                    f"{difficulty} ({topic}): "
                    f"only {len(chosen)} DB questions available"
                )
                print("⚠️ WARNING:", warning)
                warnings.append(warning)

            for q in chosen:
                section_questions.append({
                    "section": difficulty,
                    "topic": q.topic,
                    "question_text": q.question_text,
                    "options": q.options,
                    "correct_answer": q.correct_answer,
                    "question_type": q.question_type,
                    "images": q.images,
                    "source": "db",
                })

        # ---------------- AI QUESTIONS ----------------
        # ---------------- AI QUESTIONS (SAFE) ----------------
        if ai_required > 0:
            print("\n🤖 Generating AI questions (SAFE MODE)...")
            print("   • Difficulty:", difficulty)
            print("   • Topic:", topic)
            print("   • Required count:", ai_required)
        
            ai_questions = generate_ai_questions_safely(
                class_name=class_name,
                subject=cfg.subject,
                difficulty=difficulty,
                topic=topic,
                required_count=ai_required,
            )

        
            print(f"🤖 AI questions accepted: {len(ai_questions)}")
        
            for q in ai_questions:
                q["section"] = difficulty
                q["topic"] = topic
                q["source"] = "ai"
        
            section_questions.extend(ai_questions)
        
        # ---------------- HARD VALIDATION ----------------
        print("\n🧪 Validating section output...")
        
        if not section_questions:
            print(
                f"❌ ERROR: Section '{difficulty}' ({topic}) produced ZERO questions"
            )
            raise HTTPException(
                status_code=400,
                detail=(
                    f"Section '{difficulty}' ({topic}) produced zero questions. "
                    "Exam generation aborted."
                ),
            )
        
        print(f"✅ Section question count: {len(section_questions)}")
        
        if total_required and len(section_questions) != total_required:
            warning = (
                f"{difficulty} ({topic}): expected {total_required}, "
                f"got {len(section_questions)}"
            )
            print("⚠️ WARNING:", warning)
            warnings.append(warning)
        
        final_questions.extend(section_questions)
        # AFTER final_questions is built
 
    for idx, q in enumerate(final_questions, start=1):
            q["q_id"] = idx
    


    # ------------------------------------------------------------
    # 5️⃣ Final validation
    # ------------------------------------------------------------
    print("\n" + "=" * 70)
    print("📦 FINAL EXAM SUMMARY")
    print("=" * 70)

    print("🧮 Total questions generated:", len(final_questions))

    if warnings:
        print("⚠️ WARNINGS:")
        for w in warnings:
            print("   -", w)
    else:
        print("✅ No warnings")

    if not final_questions:
        print("❌ ERROR: No questions generated at all")
        raise HTTPException(400, "Exam generation failed")

    print("🎉 EXAM GENERATION SUCCESSFUL")
    print("=" * 70 + "\n")
    
    # ------------------------------------------------------------
    # 6️⃣ Build exam JSON (for persistence)
    # ------------------------------------------------------------
    print("📦 Building exam_json payload...")
    normalized_questions = []

    for q in final_questions:
        normalized_questions.append({
            "q_id": q["q_id"],
            "section": q.get("section"),
            "difficulty": q.get("section"),  # same value in your system
            "topic": q.get("topic"),
            "question": q.get("question_text"),
            "options": q.get("options"),
            "correct": q.get("correct_answer"),
            "images": q.get("images") or []
        })
    
    exam_json = {
        "class_name": class_name,
        "subject": cfg.subject,
        "sections": sections,  # keep full section meta here
        "questions": normalized_questions,
        "total_questions": len(normalized_questions)
    }

    
    print("✅ exam_json built successfully")

    # ------------------------------------------------------------
    # 7️⃣ Persist generated exam
    # ------------------------------------------------------------
    print("💾 Saving generated exam to database...")
    
    saved_exam = GeneratedExamFoundational(
        class_name=class_name,
        subject=cfg.subject,
        total_questions=len(final_questions),
        duration_minutes=sum(s["time"] for s in sections),
        exam_json=exam_json,
        is_current=True,
    )
    
    db.add(saved_exam)
    db.commit()
    db.refresh(saved_exam)
    
    print("✅ Exam saved successfully")
    print("   ▸ Exam ID:", saved_exam.id)


    return {
        "generated_exam_id": saved_exam.id,
        "class_name": class_name,
        "subject": cfg.subject,
        "total_questions": len(final_questions),
        "warnings": warnings,
    }


"""
Before topic names were added to quiz setup foundational
@app.post("/api/exams/generate-foundational")
def generate_exam_foundational(
    payload: dict = Body(...),
    db: Session = Depends(get_db)
):
    '''
    Generate a Foundational exam using the LATEST quiz_setup_foundational
    (highest id) for the given class_name.
    Guarantees:
      - subject is taken from config
      - no empty sections
      - section totals respected
      - runtime-safe exam_json
    '''

    print("\n==============================")
    print("📥 Generate Foundational Exam")
    print("==============================")

    # ------------------------------------------------------------
    # 0️⃣ Validate payload
    # ------------------------------------------------------------
    class_name = payload.get("class_name")
    if not class_name:
        raise HTTPException(400, "class_name is required")

    db.query(GeneratedExamFoundational).delete(synchronize_session=False)
    db.commit()

    # ------------------------------------------------------------
    # 1️⃣ Load LATEST config for class (highest ID)
    # ------------------------------------------------------------
    cfg = (
        db.query(QuizSetupFoundational)
        .filter(func.lower(QuizSetupFoundational.class_name) == class_name.lower())
        .order_by(QuizSetupFoundational.id.desc())
        .first()
    )

    if not cfg:
        raise HTTPException(
            404,
            f"No quiz setup found for class '{class_name}'"
        )

    subject = cfg.subject
    print(f"➡ Using config ID={cfg.id}, subject={subject}")

    # ------------------------------------------------------------
    # 2️⃣ Build sections (from config only)
    # ------------------------------------------------------------
    sections = []

    def add_section(name, time, intro):
        if name:
            sections.append({
                "name": name,
                "difficulty": name,
                "time": time or 0,
                "intro": intro
            })

    add_section(cfg.section1_name, cfg.section1_time, cfg.section1_intro)
    add_section(cfg.section2_name, cfg.section2_time, cfg.section2_intro)
    add_section(cfg.section3_name, cfg.section3_time, cfg.section3_intro)

    if not sections:
        raise HTTPException(400, "No sections defined in config")

    final_questions = []
    warnings = []

    # ------------------------------------------------------------
    # 3️⃣ Generate questions PER SECTION
    # ------------------------------------------------------------
    for section in sections:
        difficulty = section["name"]

        if difficulty == cfg.section1_name:
            db_required = cfg.section1_db or 0
            ai_required = cfg.section1_ai or 0
            total_required = cfg.section1_total or 0
        elif difficulty == cfg.section2_name:
            db_required = cfg.section2_db or 0
            ai_required = cfg.section2_ai or 0
            total_required = cfg.section2_total or 0
        elif difficulty == cfg.section3_name:
            db_required = cfg.section3_db or 0
            ai_required = cfg.section3_ai or 0
            total_required = cfg.section3_total or 0
        else:
            continue

        print(f"\n➡ Section {difficulty}: DB={db_required}, AI={ai_required}, TOTAL={total_required}")

        section_questions = []

        if db_required > 0:
            pool = (
                db.query(Question)
                .filter(
                    func.lower(Question.class_name) == class_name.lower(),
                    func.lower(Question.subject) == subject.lower(),
                    func.lower(Question.difficulty) == difficulty.lower(),
                )
                .all()
            )

            random.shuffle(pool)
            chosen = pool[:db_required]

            if len(chosen) < db_required:
                warnings.append(
                    f"{difficulty}: only {len(chosen)} DB questions available"
                )

            for q in chosen:
                section_questions.append({
                    "section": difficulty,
                    "question_text": q.question_text,
                    "options": q.options,
                    "correct_answer": q.correct_answer,
                    "question_type": q.question_type,
                    "images": q.images,
                    "topic": q.topic,
                })

        if ai_required > 0:
            ai_questions = generate_ai_questions_foundational(
                class_name=class_name,
                subject=subject,
                difficulty=difficulty,
                count=ai_required
            )

            for q in ai_questions:
                q["section"] = difficulty

            section_questions.extend(ai_questions)

        if not section_questions:
            raise HTTPException(
                400,
                f"Section '{difficulty}' produced zero questions. Exam rejected."
            )

        if total_required and len(section_questions) != total_required:
            warnings.append(
                f"{difficulty}: expected {total_required}, got {len(section_questions)}"
            )

        final_questions.extend(section_questions)

    # ------------------------------------------------------------
    # 5️⃣ Number questions
    # ------------------------------------------------------------
    for i, q in enumerate(final_questions, start=1):
        q["q_id"] = i
        q["question_number"] = i

    # ------------------------------------------------------------
    # 6️⃣ Compute duration
    # ------------------------------------------------------------
    duration_minutes = sum(s["time"] for s in sections)
    if duration_minutes <= 0:
        raise HTTPException(400, "Invalid exam duration")

    # ------------------------------------------------------------
    # 7️⃣ Build exam JSON
    # ------------------------------------------------------------
    exam_json = {
        "class_name": class_name,
        "subject": subject,
        "total_questions": len(final_questions),
        "sections": sections,
        "questions": final_questions,
        "duration_minutes": duration_minutes,
    }

    # ------------------------------------------------------------
    # 8️⃣ Persist exam
    # ------------------------------------------------------------
    db.query(GeneratedExamFoundational).update({"is_current": False})

    saved = GeneratedExamFoundational(
        class_name=class_name,
        subject=subject,
        total_questions=len(final_questions),
        duration_minutes=duration_minutes,
        exam_json=exam_json,
        is_current=True
    )

    db.add(saved)
    db.commit()
    db.refresh(saved)

    print("✅ Saved Foundational Exam ID:", saved.id)

    return {
        "generated_exam_id": saved.id,
        "total_questions": len(final_questions),
        "duration_minutes": duration_minutes,
        "warnings": warnings
    }
"""




@app.post("/api/quizzes-foundational")
def save_quiz_setup(
    payload: QuizSetupFoundationalSchema,
    db: Session = Depends(get_db)
):
    print("\n================ SAVE QUIZ SETUP =================")
    print("📥 Raw payload received:")
    print(payload)

    print("\n➡️ Basic metadata:")
    print("class_name:", payload.class_name)
    print("subject:", payload.subject)
    print("number of sections:", len(payload.sections))

    if len(payload.sections) < 2:
        print("❌ ERROR: Less than 2 sections provided")
        raise HTTPException(
            status_code=400,
            detail="At least 2 sections required"
        )

    section1 = payload.sections[0]
    section2 = payload.sections[1]
    section3 = payload.sections[2] if len(payload.sections) == 3 else None

    print("\n➡️ Section breakdown:")
    print("SECTION 1:")
    print("  name:", section1.name)
    print("  topic:", section1.topic)
    print("  ai:", section1.ai)
    print("  db:", section1.db)
    print("  total:", section1.total)
    print("  time:", section1.time)

    print("\nSECTION 2:")
    print("  name:", section2.name)
    print("  topic:", section2.topic)
    print("  ai:", section2.ai)
    print("  db:", section2.db)
    print("  total:", section2.total)
    print("  time:", section2.time)

    if section3:
        print("\nSECTION 3:")
        print("  name:", section3.name)
        print("  topic:", section3.topic)
        print("  ai:", section3.ai)
        print("  db:", section3.db)
        print("  total:", section3.total)
        print("  time:", section3.time)
    else:
        print("\nSECTION 3: Not provided")

    total_questions = (
        section1.total +
        section2.total +
        (section3.total if section3 else 0)
    )

    print("\n➡️ Total questions calculated:", total_questions)

    if section3 is None and total_questions > 40:
        print("❌ ERROR: Total questions exceed 40 with 2 sections")
        raise HTTPException(
            status_code=400,
            detail="Total questions cannot exceed 40 with 2 sections"
        )

    if section3 and total_questions > 50:
        print("❌ ERROR: Total questions exceed 50 with 3 sections")
        raise HTTPException(
            status_code=400,
            detail="Total questions cannot exceed 50 with 3 sections"
        )

    print("\n🧹 Clearing existing quiz setup records...")
    db.query(QuizSetupFoundational).delete(synchronize_session=False)
    db.commit()
    print("✅ Existing records cleared")

    print("\n🛠 Creating QuizSetupFoundational object...")

    quiz = QuizSetupFoundational(
        class_name=payload.class_name,
        subject=payload.subject,

        section1_name=section1.name,
        section1_topic=section1.topic,
        section1_ai=section1.ai,
        section1_db=section1.db,
        section1_total=section1.total,
        section1_time=section1.time,
        section1_intro=section1.intro,

        section2_name=section2.name,
        section2_topic=section2.topic,
        section2_ai=section2.ai,
        section2_db=section2.db,
        section2_total=section2.total,
        section2_time=section2.time,
        section2_intro=section2.intro,

        section3_name=section3.name if section3 else None,
        section3_topic=section3.topic if section3 else None,
        section3_ai=section3.ai if section3 else None,
        section3_db=section3.db if section3 else None,
        section3_total=section3.total if section3 else None,
        section3_time=section3.time if section3 else None,
        section3_intro=section3.intro if section3 else None,
    )

    print("\n🧾 Final object values before DB insert:")
    print("section1_topic:", quiz.section1_topic)
    print("section2_topic:", quiz.section2_topic)
    print("section3_topic:", quiz.section3_topic)

    db.add(quiz)
    db.commit()
    db.refresh(quiz)

    print("\n✅ Quiz setup saved successfully")
    print("🆔 quiz_id:", quiz.id)
    print("==================================================\n")

    return {
        "message": "Quiz setup saved successfully",
        "quiz_id": quiz.id,
    }

@app.post("/api/exams/generate-reading")
def generate_exam_reading(
    payload: ReadingExamRequest,
    db: Session = Depends(get_db)
):
    """
    Generate a Reading Comprehension exam.
    Produces a STRICT, frontend-safe exam_json with SECTIONED structure.
    """

    print("\n================ GENERATE READING EXAM ================")
    print("Incoming payload:", payload.dict())

    class_name = payload.class_name.strip()
    difficulty = payload.difficulty.strip()
    print("🧹 Resetting previous reading exam attempts")
    db.query(StudentExamReportReading).delete(synchronize_session=False)
    db.query(StudentExamReading).delete(synchronize_session=False)
    db.query(GeneratedExamReading).delete(synchronize_session=False)
    
    db.commit()
    
    print("🧼 All reading exam attempts cleared")
    # --------------------------------------------------
    # 1️⃣ LOAD CONFIG
    # --------------------------------------------------
    cfg = (
        db.query(ReadingExamConfig)
        .filter(
            func.lower(ReadingExamConfig.class_name) == class_name.lower(),
            func.lower(ReadingExamConfig.difficulty) == difficulty.lower(),
        )
        .first()
    )

    if not cfg:
        raise HTTPException(404, "No reading exam config found")

    subject = cfg.subject
    topics = cfg.topics
    warnings = []

    sections = []

    # --------------------------------------------------
    # 2️⃣ PROCESS EACH TOPIC AS A SECTION
    # --------------------------------------------------
    for section_index, topic_spec in enumerate(topics, start=1):
        topic_name = topic_spec["name"].strip()
        required = int(topic_spec["num_questions"])
        topic_lower = topic_name.lower()

        section_id = f"{topic_lower.replace(' ', '_')}_{section_index}"

        bundles = (
            db.query(QuestionReading)
            .filter(
                func.lower(QuestionReading.class_name) == class_name.lower(),
                func.lower(func.replace(QuestionReading.subject, " ", "_")) == subject.lower(),
                func.lower(QuestionReading.difficulty) == difficulty.lower(),
                func.lower(QuestionReading.topic) == topic_lower,
            )
            .all()
        )

        if not bundles:
            warnings.append(f"No bundles found for topic '{topic_name}'")
            continue

        matched_bundle = next(
            (b for b in bundles if b.total_questions == required),
            None
        )

        if not matched_bundle:
            raise HTTPException(
                400,
                f"Invalid exam config: '{topic_name}' requires {required} questions"
            )

        bundle_json = matched_bundle.exam_bundle or {}

        question_type = bundle_json.get("question_type")
        reading_material = bundle_json.get("reading_material")
        answer_options = bundle_json.get("answer_options")

        # 🔑 NEW: explicit rendering hints
        if question_type in ("main_idea", "literary_analysis"):
            passage_style = "literary"
        else:
            passage_style = "informational"


        render_hint = (
            "gapped_text"
            if question_type == "gapped_text"
            else "standard"
        )

        options_scope = "shared" if answer_options else "per_question"

        collected_questions = [
            copy.deepcopy(q) for q in bundle_json.get("questions", [])
        ]

        for idx, q in enumerate(collected_questions, start=1):
            q["question_number"] = idx
            q["question_id"] = f"{section_id}_Q{idx}"

        section = {
            "section_id": section_id,
            "section_index": section_index,
            "question_type": question_type,
            "topic": topic_name,
            "reading_material": reading_material,
            "passage_style": passage_style,
            "render_hint": render_hint,
            "options_scope": options_scope,
            "questions": collected_questions,
        }

        if answer_options:
            section["answer_options"] = answer_options

        sections.append(section)

    # --------------------------------------------------
    # 3️⃣ FINALIZE EXAM
    # --------------------------------------------------
    if not sections:
        raise HTTPException(400, "No exam sections generated")

    total_questions = sum(len(s["questions"]) for s in sections)

    exam_json = {
        "class_name": class_name,
        "subject": subject,
        "difficulty": difficulty,
        "duration_minutes": 40,
        "total_questions": total_questions,
        "sections": sections,
    }

    saved = GeneratedExamReading(
        config_id=cfg.id,
        class_name=class_name,
        subject=subject,
        difficulty=difficulty,
        total_questions=total_questions,
        exam_json=exam_json,
    )

    db.add(saved)
    db.commit()
    db.refresh(saved)

    print("✅ Exam generated. ID:", saved.id)

    return {
        "generated_exam_id": saved.id,
        "total_questions": total_questions,
        "warnings": warnings,
        "exam_json": exam_json,
    }


@app.get("/api/quizzes-reading")
def get_reading_quiz_dropdown(db: Session = Depends(get_db)):
    """
    Returns unique class_name + difficulty combinations 
    from reading_exam_config for populating the dropdown.
    """

    rows = (
        db.query(
            ReadingExamConfig.class_name,
            ReadingExamConfig.difficulty
        )
        .distinct()
        .all()
    )

    result = []
    seen = set()

    for row in rows:
        key = (row.class_name, row.difficulty)
        if key not in seen:
            seen.add(key)
            result.append({
                "class_name": row.class_name,
                "difficulty": row.difficulty,
                "label": f"{row.class_name} | {row.difficulty}"
            })

    return result

 
CANONICAL_READING_TOPICS = {
    "comparative analysis": "Comparative analysis",
    "gapped text": "Gapped Text",
    "main idea & summary": "Main Idea & Summary",
    "main idea and summary": "Main Idea & Summary",
}
def normalize_topic(value: str) -> str:
    return (
        value.lower()
        .replace("&", "")
        .replace("(", "")
        .replace(")", "")
        .replace("-", " ")
        .replace("/", " ")
        .replace(",", "")
        .strip()
        .replace("  ", " ")
        .replace(" ", "_")
    )

def normalize_topic(raw_topic: str) -> str:
    if not raw_topic:
        raise HTTPException(status_code=400, detail="Missing topic")

    key = raw_topic.strip().lower()

    if key not in CANONICAL_READING_TOPICS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid topic '{raw_topic}'. Must match admin-defined topics."
        )

    return CANONICAL_READING_TOPICS[key]

 
@app.post("/upload-word-reading")
async def upload_word_reading(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    print("\n================ UPLOAD WORD READING =================")

    # --------------------------------------------------
    # 1️⃣ Validate file
    # --------------------------------------------------
    if file.content_type != (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        raise HTTPException(status_code=400, detail="File must be a .docx")

    raw = await file.read()

    extracted = extract_text_from_docx(raw)
    if not extracted or len(extracted.strip()) < 50:
        raise HTTPException(status_code=400, detail="Invalid or empty document")

    # --------------------------------------------------
    # 2️⃣ Detect question type
    # --------------------------------------------------
    lower_text = extracted.lower()
    if "gapped text" in lower_text or "gap 1" in lower_text:
        question_type = "gapped_text"
    else:
        question_type = "standard_reading"

    print("Detected question type:", question_type)

    # --------------------------------------------------
    # 3️⃣ Parse with OpenAI
    # --------------------------------------------------
    try:
        parsed = parse_exam_with_openai(
            extracted_text=extracted,
            question_type=question_type
        )
    except Exception as e:
        print("❌ OpenAI parsing error:", e)
        raise HTTPException(status_code=500, detail="Parsing failed")

    if not parsed:
        raise HTTPException(status_code=400, detail="No exam data parsed")

    exams = parsed if isinstance(parsed, list) else [parsed]
    saved_ids = []

    # --------------------------------------------------
    # 4️⃣ Normalize + Save bundles
    # --------------------------------------------------
    for exam_json in exams:
        raw_topic = exam_json.get("topic")
        if not raw_topic:
            raise HTTPException(
                status_code=400,
                detail="Parsed exam missing topic"
            )

        topic = normalize_topic(raw_topic)

        raw_questions = exam_json.get("questions", [])
        if not raw_questions:
            raise HTTPException(
                status_code=400,
                detail=f"No questions found for topic '{topic}'"
            )

        questions = []
        for q in raw_questions:
            if not q.get("question_text") or not q.get("correct_answer"):
                raise HTTPException(
                    status_code=400,
                    detail="Invalid question structure detected"
                )

            questions.append({
                "question_text": q["question_text"],
                "correct_answer": q["correct_answer"],
            })

        bundle = {
            "topic": topic,
            "reading_material": exam_json.get("reading_material", {}),
            "questions": questions,
            "answer_options": exam_json.get("answer_options", {}),
        }

        # 🔒 SYSTEM FIELDS ARE CONTROLLED HERE (NOT BY OPENAI)
        obj = Question_reading(
            class_name="selective",
            subject="Reading Comprehension",
            difficulty="medium",
            topic=topic,
            exam_bundle=bundle,
        )

        db.add(obj)
        db.commit()
        db.refresh(obj)

        saved_ids.append(obj.id)
        print(f"Saved bundle ID: {obj.id} | Topic: {topic}")

    print("✅ Upload complete")

    return {
        "message": "Word reading exams uploaded successfully",
        "bundle_ids": saved_ids,
    }

@app.post("/upload-word-reading-gapped-multi-ai")
async def upload_word_reading_gapped_multi_ai(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    print("\n" + "=" * 70)
    print("📄 START: MULTI GAPPED TEXT WORD UPLOAD (AI - STRICT MODE)")
    print("=" * 70)

    # --------------------------------------------------
    # 1️⃣ File validation
    # --------------------------------------------------
    print("📥 STEP 1: File validation")
    print("   → Filename:", file.filename)
    print("   → Content-Type:", file.content_type)

    if file.content_type != (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        print("❌ INVALID FILE TYPE")
        raise HTTPException(status_code=400, detail="File must be .docx")

    raw = await file.read()
    print("✅ File read into memory | bytes:", len(raw))

    # --------------------------------------------------
    # 2️⃣ Text extraction
    # --------------------------------------------------
    print("\n📄 STEP 2: Extracting text from DOCX")

    try:
        full_text = extract_text_from_docx(raw)
        full_text = normalize_doc_text(full_text)
    except Exception as e:
        print("❌ DOCX TEXT EXTRACTION FAILED")
        print("   → Exception:", str(e))
        raise HTTPException(status_code=500, detail="Failed to extract document text")

    if not full_text:
        print("❌ Extracted text is EMPTY")
        raise HTTPException(status_code=400, detail="Empty document")

    print("✅ Text extracted successfully")
    print("   → Character count:", len(full_text))

    if len(full_text.strip()) < 300:
        print("❌ Document too short after extraction")
        raise HTTPException(status_code=400, detail="Invalid document content")

    # --------------------------------------------------
    # 3️⃣ Exam block detection
    # --------------------------------------------------
    print("\n🧩 STEP 3: Locating EXAM blocks")

    start_token = "=== EXAM START ==="
    end_token = "=== EXAM END ==="

    parts = full_text.split(start_token)
    print("   → Total START tokens found:", len(parts) - 1)

    blocks = []

    for idx, part in enumerate(parts[1:], start=1):
        print(f"   → Processing potential block {idx}")

        if end_token not in part:
            print("     ⚠️ END token missing, skipping block")
            continue

        block = part.split(end_token)[0].strip()
        print("     → Block length:", len(block))

        if len(block) < 500:
            print("     ⚠️ Block too short, skipped")
            continue

        blocks.append(block)
        print("     ✅ Block accepted")

    if not blocks:
        print("❌ NO VALID EXAM BLOCKS FOUND")
        raise HTTPException(
            status_code=400,
            detail="No exam blocks found. Missing EXAM START / EXAM END markers."
        )

    total_blocks = len(blocks)
    print(f"✅ Total valid exam blocks detected: {total_blocks}")
    
    saved_ids = []


    # --------------------------------------------------
    # 4️⃣ AI extraction prompt
    # --------------------------------------------------
    print("\n🤖 STEP 4: Preparing AI extraction")

    system_prompt = """
You are an exam content extraction engine.

You MUST extract ONE COMPLETE GAPPED TEXT reading exam.

The document contains a METADATA section with the following keys:
- class_name
- subject
- topic
- difficulty

You MUST extract these fields exactly as written.
If any of these metadata fields are missing or empty, RETURN {}.

CRITICAL RULES (FAIL HARD):
- DO NOT generate, infer, or rewrite content
- Extract ONLY what exists in the document
- reading_material.content MUST contain the FULL passage
- answer_options MUST contain ALL options A through G
- questions MUST be EXACTLY 6
- correct_answer MUST be one of: A|B|C|D|E|F|G
- If ANY required field is missing or empty, RETURN {}

OUTPUT:
- VALID JSON ONLY
- No markdown
- No explanations
"""

    # --------------------------------------------------
    # 5️⃣ Process each exam block
    # --------------------------------------------------
    failed_blocks = 0
    failure_reasons = []
    
    for block_idx, block_text in enumerate(blocks, start=1):

        print("\n" + "-" * 70)
        print(f"🔍 STEP 5: Processing block {block_idx}/{len(blocks)}")
        print("   → Block text length:", len(block_text))

        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                temperature=0,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {
                        "role": "user",
                        "content": f"BEGIN_DOCUMENT\n{block_text}\nEND_DOCUMENT"
                    }
                ]
            )
        except Exception as e:
            print("❌ OpenAI API CALL FAILED")
            print("   → Exception:", str(e))
            continue

        raw_output = response.choices[0].message.content.strip()
        print("🧠 AI raw response length:", len(raw_output))

        if not raw_output.startswith("{"):
           print("❌ AI OUTPUT IS NOT JSON")
           print("   → Preview:", raw_output[:200])
       
           failed_blocks += 1
           failure_reasons.append({
               "block": block_idx,
               "reason": "AI output not valid JSON"
           })
       
           continue


        try:
            parsed = json.loads(raw_output)
            print("✅ JSON parsed successfully")
        except Exception as e:
            print("❌ JSON PARSING FAILED")
            print("   → Exception:", str(e))
        
            failed_blocks += 1
            failure_reasons.append({
                "block": block_idx,
                "reason": "JSON parsing failed"
            })
        
            continue


        if not parsed:
           print("⚠️ AI returned EMPTY JSON")
       
           failed_blocks += 1
           failure_reasons.append({
               "block": block_idx,
               "reason": "AI returned empty JSON"
           })
       
           continue


        # --------------------------------------------------
        # 6️⃣ Hard validation
        # --------------------------------------------------
        print("\n🧪 STEP 6: Validating extracted content")

        rm = parsed.get("reading_material", {})
        if not rm.get("content"):
           print("❌ Missing reading material content")
       
           failed_blocks += 1
           failure_reasons.append({
               "block": block_idx,
               "reason": "Missing reading material content"
           })
       
           continue


        if len(rm["content"].strip()) < 300:
           print("❌ Reading material too short")
       
           failed_blocks += 1
           failure_reasons.append({
               "block": block_idx,
               "reason": "Reading material too short"
           })
       
           continue

        opts = parsed.get("answer_options", {})
        required_opts = ["A", "B", "C", "D", "E", "F", "G"]

        missing_opts = [k for k in required_opts if not opts.get(k)]
        if missing_opts:
           print("❌ Missing answer options:", missing_opts)
       
           failed_blocks += 1
           failure_reasons.append({
               "block": block_idx,
               "reason": f"Missing answer options: {missing_opts}"
           })
       
           continue


        questions = parsed.get("questions", [])
        print("   → Questions extracted:", len(questions))

        if len(questions) != 6:
           print("❌ Invalid question count")
       
           failed_blocks += 1
           failure_reasons.append({
               "block": block_idx,
               "reason": "Invalid question count"
           })
       
           continue


        print("✅ Validation passed")

        # --------------------------------------------------
        # 7️⃣ Enrich questions
        # --------------------------------------------------
        enriched_questions = []
        for i, q in enumerate(questions, start=1):
            enriched_questions.append({
                "question_id": f"GT_Q{i}",
                "gap_number": i,
                "question_text": f"Gap {i}",
                "correct_answer": q["correct_answer"]
            })
         
        # --------------------------------------------------
        # 6.5️⃣ Metadata validation (NEW)
        # --------------------------------------------------
        required_meta = ["class_name", "subject", "topic", "difficulty"]
        missing_meta = [k for k in required_meta if not parsed.get(k)]
        
        if missing_meta:
            print("❌ Missing required metadata:", missing_meta)
        
            failed_blocks += 1
            failure_reasons.append({
                "block": block_idx,
                "reason": f"Missing metadata fields: {missing_meta}"
            })
        
            continue

        # --------------------------------------------------
        # 8️⃣ Save to DB
        # --------------------------------------------------
        # --------------------------------------------------
        # 7️⃣ Final bundle (QUIZ-READY)
        # --------------------------------------------------
        bundle = {
            "question_type": "gapped_text",
            "topic": parsed["topic"],
            "reading_material": rm,
            "answer_options": opts,
            "questions": enriched_questions
        }
        
        # --------------------------------------------------
        # 8️⃣ Save ONE row per exam
        # --------------------------------------------------
        try:
            obj = QuestionReading(
                class_name=parsed["class_name"].lower(),
                subject=parsed["subject"],
                difficulty=parsed["difficulty"].lower(),
                topic=parsed["topic"],
                total_questions=len(enriched_questions),
                exam_bundle=bundle
            )
        
            db.add(obj)
            db.commit()
            db.refresh(obj)
        
            saved_ids.append(obj.id)
            print(f"✅ Block {block_idx} saved successfully | ID={obj.id}")
        
        except Exception as e:
               db.rollback()
               print("❌ DATABASE SAVE FAILED")
               print("   → Exception:", str(e))
           
               failed_blocks += 1
               failure_reasons.append({
                   "block": block_idx,
                   "reason": "Database save failed"
               })
           
               continue


    # --------------------------------------------------
    # 9️⃣ Final response
    # --------------------------------------------------
    print("\n" + "=" * 70)
    print("🎉 UPLOAD COMPLETE")
    print("   → Total saved exams:", len(saved_ids))
    print("   → IDs:", saved_ids)
    print("=" * 70)

    return {
        "message": "Gapped Text documents processed",
        "total_detected": total_blocks,
        "successfully_saved": len(saved_ids),
        "failed": failed_blocks,
        "saved_ids": saved_ids,
        "failures": failure_reasons
    }

def split_exam_blocks(full_text: str) -> list[str]:
    import re

    blocks = []

    pattern = re.compile(
        r"===\s*EXAM\s*START\s*===(.*?)===\s*EXAM\s*END\s*===",
        re.IGNORECASE | re.DOTALL
    )

    for match in pattern.finditer(full_text):
        block = match.group(1).strip()
        if len(block) >= 300:
            blocks.append(block)

    return blocks

def normalize_metadata(meta: dict) -> dict:
    return {
        k.lower().strip(): v.strip()
        for k, v in meta.items()
    }

def detect_question_type(block_text: str) -> str | None:
    import re

    match = re.search(
        r"question_type\s*:\s*([a-z_]+)(?=\s*METADATA\s*:|$)",
        block_text,
        re.IGNORECASE
    )

    if match:
        return match.group(1).lower()

    return None


def extract_all_metadata(block_text: str) -> dict[str, str]:
    import re

    metadata = {}

    # This pattern:
    # - Finds KEY:
    # - Captures VALUE
    # - Stops when the NEXT KEY starts (even with no whitespace)
    pattern = re.compile(
        r"([A-Za-z_ ]+)\s*[:：]\s*\"?(.*?)\"?(?=[A-Za-z_ ]+\s*[:：]|$)",
        re.IGNORECASE | re.DOTALL
    )

    for match in pattern.finditer(block_text):
        raw_key = match.group(1)
        value = match.group(2).strip()

        key = raw_key.strip().upper().replace(" ", "_")

        KEY_ALIASES = {
            "CLASS_NAME": "CLASS",
            "TOTALQUESTIONS": "TOTAL_QUESTIONS",
            "TOTAL_QUESTION": "TOTAL_QUESTIONS",
        }

        key = KEY_ALIASES.get(key, key)

        metadata[key] = value

    return metadata


def parse_gapped_block(block_text: str, db: Session) -> list[int]:
    import json
    import re

    print("🧠 [gapped_text] START parsing block")

    saved_ids: list[int] = []

    # --------------------------------------------------
    # 1️⃣ DOCUMENT-AUTHORITATIVE METADATA
    # --------------------------------------------------
    def extract_meta(label: str) -> str | None:
        match = re.search(
            rf"^[^\S\r\n]*{label}[^\S\r\n]*:[^\S\r\n]*\"?(.*?)\"?[^\S\r\n]*$",
            block_text,
            re.MULTILINE | re.IGNORECASE
        )
        return match.group(1).strip() if match else None


    meta = extract_all_metadata(block_text)
    
    class_name = meta.get("CLASS")
    subject = meta.get("SUBJECT")
    topic = meta.get("TOPIC")
    difficulty = meta.get("DIFFICULTY")

    missing = []
    if not class_name: missing.append("CLASS")
    if not subject: missing.append("SUBJECT")
    if not topic: missing.append("TOPIC")
    if not difficulty: missing.append("DIFFICULTY")
    
    if missing:
        raise ValueError(f"Missing METADATA fields: {', '.join(missing)}")

    # --------------------------------------------------
    # 2️⃣ TOTAL QUESTIONS (NORMALIZED)
    # --------------------------------------------------
    tq_match = re.search(
        r"TOTAL[_ ]?QUESTIONS\s*:\s*(\d+)",
        block_text,
        re.IGNORECASE
    )
    
    if not tq_match:
        raise ValueError("TOTAL_QUESTIONS missing")
    
    expected_q_count = int(tq_match.group(1))
    
    if expected_q_count != 6:
        raise ValueError("Gapped text must have exactly 6 gaps")


    # --------------------------------------------------
    # 3️⃣ AI EXTRACTION (STRICT CONTRACT)
    # --------------------------------------------------
    system_prompt = """
You are an exam content extraction engine.

You MUST extract ONE COMPLETE GAPPED TEXT reading exam.

The JSON output MUST contain EXACTLY:

{
  "class_name": string,
  "subject": string,
  "topic": string,
  "difficulty": string,
  "reading_material": {
    "content": string
  },
  "answer_options": {
    "A": string,
    "B": string,
    "C": string,
    "D": string,
    "E": string,
    "F": string,
    "G": string
  },
  "questions": [
    { "correct_answer": "A" }
  ]
}

RULES:
- Extract ONLY what exists in the document
- Preserve wording EXACTLY
- questions MUST be exactly 6
- correct_answer MUST be one of A–G
- If ANY required field is missing or empty, RETURN {}

OUTPUT:
- VALID JSON ONLY
- No markdown
- No explanations
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0,
        messages=[
            {"role": "system", "content": system_prompt},
            {
                "role": "user",
                "content": f"BEGIN_DOCUMENT\n{block_text}\nEND_DOCUMENT"
            }
        ]
    )

    raw_output = response.choices[0].message.content.strip()

    if not raw_output.startswith("{"):
        raise ValueError("AI returned non-JSON output")

    parsed = json.loads(raw_output)

    if not parsed:
        raise ValueError("AI returned empty JSON")

    # --------------------------------------------------
    # 4️⃣ HARD VALIDATION
    # --------------------------------------------------
    rm = parsed.get("reading_material", {})
    content = rm.get("content", "").strip()

    if len(content) < 300:
        raise ValueError("Reading material missing or too short")

    opts = parsed.get("answer_options", {})
    required_opts = {"A", "B", "C", "D", "E", "F", "G"}

    if set(opts.keys()) != required_opts:
        raise ValueError("Answer options must include A–G exactly")

    questions = parsed.get("questions", [])

    if len(questions) != expected_q_count:
        raise ValueError(
            f"Question count mismatch: expected {expected_q_count}, got {len(questions)}"
        )

    for i, q in enumerate(questions, start=1):
        if q.get("correct_answer") not in required_opts:
            raise ValueError(f"Invalid correct_answer in question {i}")

    # --------------------------------------------------
    # 5️⃣ ENRICH QUESTIONS (FRONTEND SAFE)
    # --------------------------------------------------
    enriched_questions = []
    for i, q in enumerate(questions, start=1):
        enriched_questions.append({
            "question_id": f"GT_Q{i}",
            "gap_number": i,
            "question_text": f"Gap {i}",
            "correct_answer": q["correct_answer"]
        })

    # --------------------------------------------------
    # 6️⃣ FINAL BUNDLE
    # --------------------------------------------------
    bundle = {
        "question_type": "gapped_text",
        "topic": topic,
        "reading_material": rm,
        "answer_options": opts,
        "questions": enriched_questions
    }

    # --------------------------------------------------
    # 7️⃣ SAVE
    # --------------------------------------------------
    obj = QuestionReading(
        class_name=class_name.lower(),
        subject=subject,
        difficulty=difficulty.lower(),
        topic=topic,
        total_questions=len(enriched_questions),
        exam_bundle=bundle
    )

    db.add(obj)
    db.commit()
    db.refresh(obj)

    saved_ids.append(obj.id)

    print(f"✅ Gapped text exam saved | ID={obj.id}")
    return saved_ids

def parse_literary_block(block_text: str, db: Session) -> list[int]:
    import json

    saved_ids: list[int] = []

    print("🧠 [literary] START parsing block")

    system_prompt = """
You are an exam content extraction engine.

You MUST extract ONE COMPLETE LITERARY reading exam.

REQUIRED KEYS:
- class_name
- subject
- topic
- difficulty
- total_questions
- reading_material
- questions

Each question MUST contain:
- question_text
- answer_options (A–D)
- correct_answer (A–D)

FAIL HARD if any rule is violated.
RETURN VALID JSON ONLY.
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0,
        messages=[
            {"role": "system", "content": system_prompt},
            {
                "role": "user",
                "content": f"BEGIN_DOCUMENT\n{block_text}\nEND_DOCUMENT"
            }
        ]
    )

    raw_output = response.choices[0].message.content.strip()
    if not raw_output.startswith("{"):
        raise ValueError("AI returned non-JSON output")

    parsed = json.loads(raw_output)
    if not parsed:
        raise ValueError("AI returned empty JSON")

    reading_material = parsed["reading_material"]
    questions = parsed["questions"]

    if len(questions) != int(parsed["total_questions"]):
        raise ValueError("Literary question count mismatch")

    enriched_questions = []
    for i, q in enumerate(questions, start=1):
        enriched_questions.append({
            "question_id": f"LIT_Q{i}",
            "question_text": q["question_text"],
            "answer_options": q["answer_options"],
            "correct_answer": q["correct_answer"]
        })

    bundle = {
        "question_type": "literary",
        "passage_style": "literary",
        "topic": parsed["topic"],
        "reading_material": reading_material,
        "questions": enriched_questions
    }

    obj = QuestionReading(
        class_name=parsed["class_name"].lower(),
        subject=parsed["subject"],
        difficulty=parsed["difficulty"].lower(),
        topic=parsed["topic"],
        total_questions=len(enriched_questions),
        exam_bundle=bundle
    )

    db.add(obj)
    db.commit()
    db.refresh(obj)

    saved_ids.append(obj.id)

    print(f"💾 Literary exam saved | ID={obj.id}")
    return saved_ids

def parse_main_idea_block(block_text: str, db: Session) -> list[int]:
    import json

    saved_ids: list[int] = []

    print("🧠 [main_idea] START parsing block")

    SYSTEM_PROMPT = """
You are an exam integrity validation engine.

RULES:
- Exactly ONE MainREADING exam
- Questions MUST match Total_Questions
- Answer options MUST be A–G
- DO NOT infer content
- RETURN {"status":"ok"} if valid
- Otherwise RETURN {}
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": f"BEGIN_DOCUMENT\n{block_text}\nEND_DOCUMENT"
            }
        ]
    )

    result = json.loads(response.choices[0].message.content.strip())
    if result != {"status": "ok"}:
        raise ValueError("Main idea validation failed")

    parsed = parse_exam_block(block_text)

    # -----------------------------
    # Normalize metadata keys
    # -----------------------------
    raw_meta = parsed.get("metadata", {})
    meta = {k.lower().strip(): v for k, v in raw_meta.items()}

    # Alias handling
    if "class_name" in meta and "class" not in meta:
        meta["class"] = meta["class_name"]

    required = ["class", "subject", "difficulty", "topic"]
    for r in required:
        if r not in meta:
            raise ValueError(f"Missing METADATA field: {r}")

    rm = parsed["reading_material"]
    questions = parsed["questions"]
    opts = parsed["answer_options"]

    enriched_questions = []
    for i, q in enumerate(questions, start=1):
        enriched_questions.append({
            "question_id": f"MI_Q{i}",
            "paragraph": q["paragraph"],
            "question_text": f"Paragraph {q['paragraph']}",
            "correct_answer": q["correct_answer"]
        })

    bundle = {
        "question_type": "main_idea",
        "reading_material": rm,
        "answer_options": opts,
        "questions": enriched_questions
    }

    obj = QuestionReading(
        class_name=meta["class"].lower(),
        subject=meta["subject"],
        difficulty=meta["difficulty"].lower(),
        topic=meta["topic"],
        total_questions=len(enriched_questions),
        exam_bundle=bundle
    )

    db.add(obj)
    db.commit()
    db.refresh(obj)

    saved_ids.append(obj.id)
    print(f"💾 Main idea exam saved | ID={obj.id}")

    return saved_ids


import uuid

upload_id = str(uuid.uuid4())[:8]
print(f"🆔 Upload ID: {upload_id}")

def parse_comparative_block(block_text: str, db: Session) -> list[int]:
    import json
    import re

    # --------------------------------------------------
    # 0️⃣ NORMALIZE WORD ARTIFACTS (SAFE, NON-DESTRUCTIVE)
    # --------------------------------------------------
    block_text = (
        block_text
        .replace("\xa0", " ")
        .replace("\u200b", "")
        .replace("\t", " ")
    )

    saved_ids: list[int] = []

    # --------------------------------------------------
    # 1️⃣ METADATA (SEMANTIC, LAYOUT-AGNOSTIC)
    # --------------------------------------------------
    def extract_meta(key: str) -> str | None:
        match = re.search(
            rf"{key}\s*:\s*\"?(.*?)\"?(?=\s+[A-Z_]+\s*:|$)",
            block_text,
            re.IGNORECASE
        )
        return match.group(1).strip() if match else None

    class_name = extract_meta("CLASS_NAME") or extract_meta("CLASS")
    subject = extract_meta("SUBJECT")
    topic = extract_meta("TOPIC")
    difficulty = extract_meta("DIFFICULTY")

    if not all([class_name, subject, topic, difficulty]):
        raise ValueError("Missing required METADATA fields")

    # --------------------------------------------------
    # 2️⃣ TOTAL QUESTIONS (SEMANTIC)
    # --------------------------------------------------
    tq_match = re.search(
        r"TOTAL[_ ]?QUESTIONS\s*:\s*(\d+)",
        block_text,
        re.IGNORECASE
    )

    if not tq_match:
        raise ValueError("Total_Questions missing")

    expected_q_count = int(tq_match.group(1))

    # --------------------------------------------------
    # 3️⃣ EXTRACT LABELS (SEMANTIC, ROBUST)
    # --------------------------------------------------
    extract_header_pattern = re.compile(
        r"^\s*Extract\s+([A-Z])\b",
        re.IGNORECASE | re.MULTILINE
    )
    
    extract_keys = extract_header_pattern.findall(block_text)
    
    # Preserve document order + dedupe
    extract_keys = list(dict.fromkeys(extract_keys))
    
    if len(extract_keys) < 2:
        raise ValueError("Comparative exam requires at least 2 extracts")
    
    print("   → Extracts:", extract_keys)


    # --------------------------------------------------
    # 4️⃣ SUBTYPE DETECTION
    # --------------------------------------------------
    is_mcq = "ANSWER_OPTIONS:" in block_text
    print("   → Subtype:", "MCQ" if is_mcq else "Extract-selection")

    # --------------------------------------------------
    # 5️⃣ AI EXTRACTION (STRICT CONTRACT)
    # --------------------------------------------------
    system_prompt = """
You are an exam content extraction engine.

You MUST extract ONE COMPLETE COMPARATIVE ANALYSIS reading exam.

The document contains:
- class_name
- subject
- topic
- difficulty

The JSON output MUST contain EXACTLY:

{
  "class_name": string,
  "subject": string,
  "topic": string,
  "difficulty": string,
  "reading_material": {
    "extracts": { "A": string, "B": string }
  },
  "questions": [...]
}

RULES:
- Extract ONLY what exists
- Preserve wording EXACTLY
- DO NOT infer or generate content
- If ANY required field is missing, RETURN {}

QUESTION RULES:
- questions count MUST match Total_Questions
- EVERY question MUST have question_text and correct_answer

IF ANSWER_OPTIONS exist:
- answer_options required
- correct_answer must be one of A,B,C,D

IF NOT:
- correct_answer must reference an extract letter

OUTPUT:
- VALID JSON ONLY
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0,
        messages=[
            {"role": "system", "content": system_prompt},
            {
                "role": "user",
                "content": f"BEGIN_DOCUMENT\n{block_text}\nEND_DOCUMENT"
            }
        ]
    )

    raw_output = response.choices[0].message.content.strip()

    if not raw_output.startswith("{"):
        raise ValueError("AI returned non-JSON output")

    parsed = json.loads(raw_output)

    if not parsed:
        raise ValueError("AI returned empty JSON")

    # --------------------------------------------------
    # 6️⃣ HARD VALIDATION AGAINST DOCUMENT
    # --------------------------------------------------
    rm = parsed.get("reading_material", {})
    extracts = rm.get("extracts", {})

    if sorted(extracts.keys()) != extract_keys:
        raise ValueError(
            f"Extract mismatch. Document={extract_keys}, GPT={list(extracts.keys())}"
        )

    questions = parsed.get("questions", [])

    if len(questions) != expected_q_count:
        raise ValueError(
            f"Question count mismatch: expected {expected_q_count}, got {len(questions)}"
        )

    # --------------------------------------------------
    # 7️⃣ QUESTION VALIDATION
    # --------------------------------------------------
    for i, q in enumerate(questions, start=1):
        if "question_text" not in q:
            raise ValueError(f"Question {i} missing question_text")

        if "correct_answer" not in q:
            raise ValueError(f"Question {i} missing correct_answer")

        if is_mcq:
            opts = q.get("answer_options")
            if not opts or set(opts.keys()) != {"A", "B", "C", "D"}:
                raise ValueError(f"Question {i} invalid answer_options")

            if q["correct_answer"] not in opts:
                raise ValueError(f"Question {i} correct_answer not in options")
        else:
            if q["correct_answer"] not in extract_keys:
                raise ValueError(
                    f"Question {i} invalid extract reference {q['correct_answer']}"
                )

    # --------------------------------------------------
    # 8️⃣ ENRICH FOR FRONTEND
    # --------------------------------------------------
    enriched_questions = []

    for i, q in enumerate(questions, start=1):
        if not is_mcq:
            q["answer_options"] = {
                k: f"Extract {k}" for k in extract_keys
            }

        enriched_questions.append({
            "question_id": f"CA_Q{i}",
            "question_text": q["question_text"],
            "answer_options": q["answer_options"],
            "correct_answer": q["correct_answer"]
        })

    bundle = {
        "question_type": "comparative_analysis",
        "topic": topic,
        "reading_material": rm,
        "questions": enriched_questions
    }

    if is_mcq:
        bundle["answer_options"] = enriched_questions[0]["answer_options"]

    # --------------------------------------------------
    # 9️⃣ SAVE
    # --------------------------------------------------
    obj = QuestionReading(
        class_name=class_name.lower(),
        subject=subject,
        difficulty=difficulty.lower(),
        topic=topic,
        total_questions=len(enriched_questions),
        exam_bundle=bundle
    )

    db.add(obj)
    db.commit()
    db.refresh(obj)

    saved_ids.append(obj.id)

    print(f"✅ Comparative exam saved | ID={obj.id}")
    return saved_ids

@app.post("/upload-word-reading-unified")
async def upload_word_reading_unified(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    import uuid
    import traceback

    upload_id = str(uuid.uuid4())[:8]

    print("\n" + "=" * 80)
    print(f"📄 START: UNIFIED READING WORD UPLOAD | ID={upload_id}")
    print("=" * 80)

    # --------------------------------------------------
    # Tracking containers (DEFINED ONCE)
    # --------------------------------------------------
    saved_ids = []
    failed_blocks = []
    progress = []

    # --------------------------------------------------
    # 1️⃣ File validation
    # --------------------------------------------------
    print(f"📥 [{upload_id}] STEP 1: File validation")
    print(f"   → Filename: {file.filename}")
    print(f"   → Content-Type: {file.content_type}")

    if file.content_type != (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        print(f"❌ [{upload_id}] Invalid file type")
        raise HTTPException(status_code=400, detail="File must be .docx")

    raw = await file.read()
    print(f"✅ [{upload_id}] File read | bytes={len(raw)}")

    # --------------------------------------------------
    # 2️⃣ Text extraction
    # --------------------------------------------------
    print(f"\n📄 [{upload_id}] STEP 2: Extracting text")

    try:
        full_text = extract_comparative_text_from_docx(raw)
        full_text = normalize_doc_text(full_text)
     
    except Exception:
        print(f"❌ [{upload_id}] DOCX extraction failed")
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail="Failed to extract document text"
        )

    if not full_text or len(full_text.strip()) < 300:
        print(f"❌ [{upload_id}] Document too short or empty")
        raise HTTPException(
            status_code=400,
            detail="Invalid or empty document"
        )

    print(f"✅ [{upload_id}] Text extracted | chars={len(full_text)}")

    # --------------------------------------------------
    # 3️⃣ Split EXAM blocks
    # --------------------------------------------------
    print(f"\n🧩 [{upload_id}] STEP 3: Splitting EXAM blocks")

    blocks = split_exam_blocks(full_text)
    print(f"   → Blocks detected: {len(blocks)}")

    if not blocks:
        raise HTTPException(
            status_code=400,
            detail="No valid EXAM blocks found"
        )

    # --------------------------------------------------
    # 4️⃣ Process each block
    # --------------------------------------------------
    for idx, block in enumerate(blocks, start=1):
        print("\n" + "-" * 80)
        print(f"📦 [{upload_id}] Processing block {idx}/{len(blocks)}")
        print(f"   → Block length: {len(block)}")
        print(f"   → Preview: {block[:200].replace(chr(10), ' ')}...")

        qtype = detect_question_type(block)
        print(f"   → Detected question_type: {qtype}")

        progress.append({
            "block": idx,
            "status": "processing",
            "question_type": qtype
        })

        if not qtype:
            failed_blocks.append({
                "block": idx,
                "reason": "Missing question_type"
            })
            progress.append({
                "block": idx,
                "status": "failed",
                "reason": "Missing question_type"
            })
            continue

        try:
            before_count = len(saved_ids)

            # -------- ROUTING (CORRECT & SAFE) --------
            if qtype == "comparative_analysis":
                ids = parse_comparative_block(block, db)

            elif qtype == "gapped_text":
                ids = parse_gapped_block(block, db)

            elif qtype == "main_idea":
                ids = parse_main_idea_block(block, db)

            elif qtype == "literary":
                ids = parse_literary_block(block, db)

            else:
                raise ValueError(f"Unknown question_type: {qtype}")

            # -------- SAVE RESULTS --------
            saved_ids.extend(ids)
            added = len(saved_ids) - before_count

            progress.append({
                "block": idx,
                "status": "saved",
                "exam_ids": ids
            })

            print(
                f"✅ [{upload_id}] Block {idx} saved | "
                f"New exams={added}"
            )

        except Exception as e:
            print(f"❌ [{upload_id}] Block {idx} failed")
            print(f"   → {type(e).__name__}: {str(e)}")
            traceback.print_exc()

            failed_blocks.append({
                "block": idx,
                "question_type": qtype,
                "reason": str(e)
            })

            progress.append({
                "block": idx,
                "status": "failed",
                "reason": str(e)
            })

    # --------------------------------------------------
    # 5️⃣ Final response
    # --------------------------------------------------
    print("\n" + "=" * 80)
    print(
        f"🏁 [{upload_id}] COMPLETE | "
        f"Blocks={len(blocks)} | "
        f"Saved={len(saved_ids)} | "
        f"Failed={len(failed_blocks)}"
    )
    print("=" * 80)

    if not saved_ids:
        status = "failed"
    elif failed_blocks:
        status = "partial_success"
    else:
        status = "success"

    return {
        "status": status,
        "upload_id": upload_id,
        "total_blocks": len(blocks),
        "saved_exam_ids": saved_ids,
        "failed_blocks": failed_blocks,
        "progress": progress
    }


@app.post("/upload-word-reading-literary-ai")
async def upload_word_reading_literary_ai(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    print("\n" + "=" * 70)
    print("📄 START: LITERARY READING WORD UPLOAD (AI - DEBUG MODE)")
    print("=" * 70)

    # --------------------------------------------------
    # 1️⃣ File validation
    # --------------------------------------------------
    print("📥 STEP 1: File validation")
    print("   → Filename:", file.filename)
    print("   → Content-Type:", file.content_type)

    if file.content_type != (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        print("❌ INVALID FILE TYPE")
        raise HTTPException(status_code=400, detail="File must be .docx")

    raw = await file.read()
    print("✅ File read successfully | bytes:", len(raw))

    # --------------------------------------------------
    # 2️⃣ Text extraction
    # --------------------------------------------------
    print("\n📄 STEP 2: Extracting text from DOCX")

    try:
        full_text = extract_text_from_docx(raw)
        full_text = normalize_doc_text(full_text)
    except Exception as e:
        print("❌ DOCX TEXT EXTRACTION FAILED:", str(e))
        raise HTTPException(status_code=500, detail="Failed to extract document text")

    if not full_text or len(full_text.strip()) < 300:
        print("❌ Invalid or empty document")
        raise HTTPException(status_code=400, detail="Invalid document")

    print("✅ Text extracted | characters:", len(full_text))

    # --------------------------------------------------
    # 3️⃣ Exam block detection
    # --------------------------------------------------
    print("\n🧩 STEP 3: Detecting exam blocks")

    start_token = "=== EXAM START ==="
    end_token = "=== EXAM END ==="

    parts = full_text.split(start_token)
    print("   → START tokens found:", len(parts) - 1)

    blocks = []
    for idx, part in enumerate(parts[1:], start=1):
        print(f"   → Inspecting block {idx}")

        if end_token not in part:
            print("     ⚠️ END token missing, skipped")
            continue

        block = part.split(end_token)[0].strip()
        print("     → Block length:", len(block))

        if len(block) < 500:
            print("     ⚠️ Block too short, skipped")
            continue

        blocks.append(block)
        print("     ✅ Block accepted")

    if not blocks:
        print("❌ NO VALID EXAM BLOCKS FOUND")
        raise HTTPException(
            status_code=400,
            detail="No exam blocks found. Missing EXAM START / EXAM END markers."
        )

    print("✅ Total valid exam blocks:", len(blocks))
    saved_ids = []

    # --------------------------------------------------
    # 4️⃣ AI extraction prompt (STRICT CONTRACT)
    # --------------------------------------------------
    print("\n🤖 STEP 4: Preparing AI extraction prompt")

    system_prompt = """
You are an exam content extraction engine.

You MUST extract ONE COMPLETE LITERARY READING exam.

RETURN JSON WITH EXACT KEYS:
- class_name (string)
- subject (string)
- topic (string)
- difficulty (string)
- total_questions (number)
- reading_material (string)
- questions (array)

Each question MUST contain:
- question_text (string)
- answer_options (object with keys A, B, C, D)
- correct_answer (one of A, B, C, D)

CRITICAL RULES (FAIL HARD):
- DO NOT generate, infer, summarize, or rewrite
- Extract ONLY what exists in the document
- total_questions MUST equal number of questions
- If ANY required field is missing or empty, RETURN {}

OUTPUT:
- VALID JSON ONLY
- No markdown
- No explanations
"""

    # --------------------------------------------------
    # 5️⃣ Process each exam block
    # --------------------------------------------------
    for block_idx, block_text in enumerate(blocks, start=1):
        print("\n" + "-" * 70)
        print(f"🔍 STEP 5: Processing block {block_idx}/{len(blocks)}")

        # Normalize known metadata variations to stabilize extraction
        normalized_block = block_text
        normalized_block = normalized_block.replace("Total_Questions", "total_questions")
        normalized_block = normalized_block.replace("TOPIC", "topic")
        normalized_block = normalized_block.replace("DIFFICULTY", "difficulty")
        normalized_block = normalized_block.replace("Reading_Material:", "reading_material:")

        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                temperature=0,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {
                        "role": "user",
                        "content": f"BEGIN_DOCUMENT\n{normalized_block}\nEND_DOCUMENT"
                    }
                ]
            )
        except Exception as e:
            print("❌ OpenAI API CALL FAILED:", str(e))
            continue

        raw_output = response.choices[0].message.content.strip()
        print("🧠 AI response length:", len(raw_output))

        if not raw_output.startswith("{"):
            print("❌ NON-JSON AI OUTPUT:", raw_output[:300])
            continue

        try:
            parsed = json.loads(raw_output)
            print("✅ JSON parsed successfully")
        except Exception as e:
            print("❌ JSON PARSE FAILED:", str(e))
            continue

        if not parsed:
            print("⚠️ EMPTY JSON returned by AI")
            continue

        # --------------------------------------------------
        # 6️⃣ Hard validation
        # --------------------------------------------------
        print("\n🧪 STEP 6: Validating extracted content")

        required_root_keys = [
            "class_name",
            "subject",
            "topic",
            "difficulty",
            "total_questions",
            "reading_material",
            "questions"
        ]

        missing_root = [k for k in required_root_keys if k not in parsed]
        if missing_root:
            print("❌ Missing root fields:", missing_root)
            continue

        reading_material = parsed["reading_material"]
        questions = parsed["questions"]

        try:
            expected_total = int(parsed["total_questions"])
        except Exception:
            print("❌ total_questions is not numeric")
            continue

        if isinstance(reading_material, dict):
            print("⚠️ reading_material returned as object, normalizing")
            reading_material = "\n".join(reading_material.values())

        if not isinstance(reading_material, str) or len(reading_material.strip()) < 200:
            print("❌ reading_material invalid or too short")
            continue

        if not isinstance(questions, list) or not questions:
            print("❌ Invalid or empty questions list")
            continue

        if len(questions) != expected_total:
            print(
                f"❌ Question count mismatch: expected {expected_total}, got {len(questions)}"
            )
            continue

        invalid_q = False
        for i, q in enumerate(questions, start=1):
            print(f"   → Validating question {i}")

            if (
                "question_text" not in q
                or "answer_options" not in q
                or "correct_answer" not in q
            ):
                print("     ❌ Missing required question fields")
                invalid_q = True
                break

            opts = q["answer_options"]
            if not isinstance(opts, dict):
                print("     ❌ answer_options not an object")
                invalid_q = True
                break

            for opt in ["A", "B", "C", "D"]:
                if opt not in opts or not isinstance(opts[opt], str) or not opts[opt].strip():
                    print(f"     ❌ Invalid or missing option {opt}")
                    invalid_q = True
                    break

            if q["correct_answer"] not in ["A", "B", "C", "D"]:
                print("     ❌ Invalid correct_answer:", q["correct_answer"])
                invalid_q = True
                break

        if invalid_q:
            continue

        print("✅ Validation passed")

        # --------------------------------------------------
        # 7️⃣ Enrich questions
        # --------------------------------------------------
        print("\n🧩 STEP 7: Enriching questions")

        enriched_questions = []
        for i, q in enumerate(questions, start=1):
            enriched_questions.append({
                "question_id": f"LIT_Q{i}",
                "question_text": q["question_text"].strip(),
                "answer_options": q["answer_options"],
                "correct_answer": q["correct_answer"]
            })

        # --------------------------------------------------
        # 8️⃣ Bundle (RENDER-SAFE)
        # --------------------------------------------------
        bundle = {
            "question_type": "main_idea",
            "passage_style": "literary",
            "topic": parsed["topic"],
            "reading_material": reading_material,
            "questions": enriched_questions
        }


        # --------------------------------------------------
        # 9️⃣ Save to DB
        # --------------------------------------------------
        print("\n💾 STEP 8: Saving exam to database")

        try:
            obj = QuestionReading(
                class_name=parsed["class_name"].lower(),
                subject=parsed["subject"],
                difficulty=parsed["difficulty"].lower(),
                topic=parsed["topic"],
                total_questions=len(enriched_questions),
                exam_bundle=bundle
            )

            db.add(obj)
            db.commit()
            db.refresh(obj)

            saved_ids.append(obj.id)
            print(f"✅ Block {block_idx} saved | ID={obj.id}")

        except Exception as e:
            db.rollback()
            print("❌ DATABASE SAVE FAILED:", str(e))
            continue

    # --------------------------------------------------
    # 🔟 Final response
    # --------------------------------------------------
    print("\n" + "=" * 70)
    print("🎉 UPLOAD COMPLETE")
    print("   → Total saved exams:", len(saved_ids))
    print("   → IDs:", saved_ids)
    print("=" * 70)

    return {
        "message": "Literary Reading documents processed",
        "saved_count": len(saved_ids),
        "bundle_ids": saved_ids
    }

# ================================
# Question Type Registry
# ================================

class QuestionHandler:
    question_type: int

    def parse(self, ctx):
        raise NotImplementedError

    def validate(self, parsed, *, context=None):
        raise NotImplementedError

    def build_exam_bundle(self, parsed):
        raise NotImplementedError


QUESTION_HANDLERS = {}
def register(handler_cls):
    QUESTION_HANDLERS[handler_cls.question_type] = handler_cls()
    return handler_cls

def parse_block(ctx, key, stop_keys, *, required=True):
    def normalize(s):
        return (
            s.strip()
             .lstrip("-•– ")
             .upper()
             .replace(" ", "_")
        )

    if not ctx.peek() or normalize(ctx.peek()) != normalize(f"{key}:"):
        if required:
            raise ValueError(f"MISSING_{key}")
        return None

    ctx.next()  # consume header

    lines = []
    stop_norm = {normalize(k + ":") for k in stop_keys}

    while ctx.peek():
        if normalize(ctx.peek()) in stop_norm:
            break
        lines.append(ctx.next())

    content = "\n".join(lines).strip()
    return content or None


def parse_options(ctx):
    def normalize(s):
        return (
            s.strip()
             .lstrip("-•– ")   # 🔥 THIS IS THE MISSING PIECE
             .upper()
             .replace(" ", "_")
        )

    # Seek ANSWER_OPTIONS header
    while ctx.peek() and normalize(ctx.peek()) != "ANSWER_OPTIONS:":
        ctx.next()

    if not ctx.peek():
        raise ValueError("MISSING_OPTIONS")

    ctx.next()  # consume ANSWER_OPTIONS:

    options = {}

    while ctx.peek():
        line = ctx.peek().strip()

        # Stop at correct answer
        if normalize(line).startswith("CORRECT_ANSWER"):
            break

        clean = line.lstrip("-•– ").strip()

        if ":" in clean:
            k, v = clean.split(":", 1)
            options[k.strip()] = v.strip()
            ctx.next()
            continue

        break

    if len(options) < 2:
        raise ValueError("INVALID_OPTIONS")

    return options



def is_question_boundary(line: str) -> bool:
    line = line.strip().upper()
    return (
        line == "--- QUESTION END ---"
        or line == "--- QUESTION START ---"
        or line.startswith("QUESTION_TYPE:")
    )

def parse_correct_answers(ctx):
    answers = []

    # Seek CORRECT_ANSWER header
    while ctx.peek() and not ctx.peek().strip().upper().startswith("CORRECT_ANSWER"):
        ctx.next()

    if not ctx.peek():
        raise ValueError("MISSING_CORRECT_ANSWER")

    ctx.next()  # consume CORRECT_ANSWER:

    while ctx.peek():
        raw_line = ctx.peek()
        line = raw_line.strip()

        # 🔒 HARD STOP at question boundary
        if is_question_boundary(line):
            break

        # Bullet answers
        if line.startswith(("-", "•")):
            value = ctx.next().lstrip("-• ").strip()

            # ✅ NEW: split comma-separated answers
            parts = [v.strip() for v in value.split(",") if v.strip()]
            answers.extend(parts)
            continue

        # Single-token answer (rare but supported)
        if line and " " not in line:
            value = ctx.next().strip()
            parts = [v.strip() for v in value.split(",") if v.strip()]
            answers.extend(parts)
            continue

        break

    if not answers:
        raise ValueError("EMPTY_CORRECT_ANSWER")

    return answers

class ParseContext:
    def __init__(self, lines):
        self.lines = lines
        self.ptr = 0

    def peek(self):
        return self.lines[self.ptr] if self.ptr < len(self.lines) else None

    def next(self):
        if self.ptr >= len(self.lines):
            return None
        val = self.lines[self.ptr]
        self.ptr += 1
        return val


# --------------------------------------------------
# Base mixin for instruction handling
# --------------------------------------------------

class InstructionAwareHandler(QuestionHandler):
    def _validate_instruction(self, parsed, *, context=None):
        if context != "reading" and not parsed.get("instruction"):
            raise ValueError("MISSING_QUESTION_INSTRUCTION")

    def _instruction_block(self, parsed):
        if parsed.get("instruction"):
            return [{"type": "text", "content": parsed["instruction"]}]
        return []

# --------------------------------------------------
# 1️⃣ Single MCQ
# --------------------------------------------------

@register
class SingleMCQHandler(InstructionAwareHandler):
    question_type = 1

    def parse(self, ctx):
        instruction = parse_block(
            ctx,
            "QUESTION_INSTRUCTION",
            ["QUESTION_TEXT"],
            required=False
        )

        question = parse_block(
            ctx,
            "QUESTION_TEXT",
            ["ANSWER_OPTIONS"],
            required=True
        )

        options = parse_options(ctx)
        correct = parse_correct_answers(ctx)

        return {
            "instruction": instruction,
            "question": question,
            "options": options,
            "correct": correct
        }

    def validate(self, parsed, *, context=None):
        self._validate_instruction(parsed, context=context)

        if len(parsed["correct"]) != 1:
            raise ValueError("INVALID_CORRECT_ANSWER_COUNT")

    def build_exam_bundle(self, parsed):
        return {
            "question_type": 1,
            "question_blocks": (
                self._instruction_block(parsed)
                + [{"type": "text", "content": parsed["question"]}]
            ),
            "options": parsed["options"],
            "correct_answer": parsed["correct"]
        }


# --------------------------------------------------
# 2️⃣ Multi-select MCQ
# --------------------------------------------------

@register
class MultiSelectMCQHandler(InstructionAwareHandler):
    question_type = 2

    def parse(self, ctx):
        instruction = parse_block(
            ctx,
            "QUESTION_INSTRUCTION",
            ["QUESTION_TEXT"],
            required=False
        )

        question = parse_block(
            ctx,
            "QUESTION_TEXT",
            ["ANSWER_OPTIONS"],
            required=True
        )

        options = parse_options(ctx)
        correct = parse_correct_answers(ctx)

        return {
            "instruction": instruction,
            "question": question,
            "options": options,
            "correct": correct
        }

    def validate(self, parsed, *, context=None):
        self._validate_instruction(parsed, context=context)

        # 🔒 Must select at least TWO answers
        if len(parsed["correct"]) < 2:
            raise ValueError("MULTI_SELECT_REQUIRES_MULTIPLE_ANSWERS")

        # 🔒 Every correct answer must exist in options
        for ans in parsed["correct"]:
            if ans not in parsed["options"]:
                raise ValueError("INVALID_CORRECT_OPTION")

    def build_exam_bundle(self, parsed):
        return {
            "question_type": 2,
            "question_blocks": (
                self._instruction_block(parsed)
                + [{"type": "multi_select", "content": parsed["question"]}]
            ),
            "options": parsed["options"],
            "correct_answer": parsed["correct"]
        }


# --------------------------------------------------
# 3️⃣ Image MCQ
# --------------------------------------------------

@register
class ImageMCQHandler(InstructionAwareHandler):
    question_type = 4

    def parse(self, ctx):
        instruction = parse_block(
            ctx,
            "QUESTION_INSTRUCTION",
            ["QUESTION_TEXT"],
            required=False
        )

        question_text = parse_block(
            ctx,
            "QUESTION_TEXT",
            ["IMAGE_OPTIONS"],
            required=True
        )

        image_options = parse_image_options(ctx)
        correct = parse_correct_answers(ctx)

        return {
            "instruction": instruction,
            "question_text": question_text,
            "image_options": image_options,
            "correct": correct
        }

    def validate(self, parsed, *, context=None):
        self._validate_instruction(parsed, context=context)

        if len(parsed["correct"]) != 1:
            raise ValueError("INVALID_CORRECT_ANSWER_COUNT")

        if parsed["correct"][0] not in parsed["image_options"]:
            raise ValueError("CORRECT_ANSWER_NOT_IN_IMAGE_OPTIONS")

    def build_exam_bundle(self, parsed):
        return {
            "question_type": 4,
            "question_blocks": (
                self._instruction_block(parsed)
                + [{"type": "text", "content": parsed["question_text"]}]
            ),
            "image_options": parsed["image_options"],
            "correct_answer": parsed["correct"]
        }


# --------------------------------------------------
# 4️⃣ True / False
# --------------------------------------------------

@register
class TrueFalseHandler(InstructionAwareHandler):
    question_type = 5

    def parse(self, ctx):
        instruction = parse_block(
            ctx, "QUESTION_INSTRUCTION", ["STATEMENTS"], required=False
        )
        statements = parse_statements(ctx)
        correct = parse_true_false_answers(ctx)

        return {
            "instruction": instruction,
            "statements": statements,
            "correct": correct
        }

    def validate(self, parsed, *, context=None):
        self._validate_instruction(parsed, context=context)

        if len(parsed["correct"]) != len(parsed["statements"]):
            raise ValueError("INVALID_CORRECT_ANSWER_COUNT")
    def build_exam_bundle(self, parsed):
        return {
            "question_type": 5,
            "question_blocks": (
                self._instruction_block(parsed)
                + [{
                    "type": "true_false",
                    "statements": parsed["statements"]
                }]
            ),
            "correct_answer": parsed["correct"]
        }


# --------------------------------------------------
# 5️⃣ Single Gap
# --------------------------------------------------

@register
class SingleGapHandler(InstructionAwareHandler):
    question_type = 6

    def parse(self, ctx):
        instruction = parse_block(
            ctx,
            "QUESTION_INSTRUCTION",
            ["QUESTION_TEXT"],
            required=False
        )

        question_text = parse_block(
            ctx,
            "QUESTION_TEXT",
            ["WORD_OPTIONS"],
            required=True
        )

        options = parse_word_options(ctx)
        correct = parse_correct_answers(ctx)

        return {
            "instruction": instruction,
            "question_text": question_text,
            "options": options,
            "correct": correct
        }

    def validate(self, parsed, *, context=None):
        self._validate_instruction(parsed, context=context)

        if parsed["question_text"].count("[BLANK]") != 1:
            raise ValueError("INVALID_BLANK_COUNT")

        if len(parsed["correct"]) != 1:
            raise ValueError("INVALID_CORRECT_ANSWER_COUNT")

        if parsed["correct"][0] not in parsed["options"]:
            raise ValueError("CORRECT_ANSWER_NOT_IN_OPTIONS")

    def build_exam_bundle(self, parsed):
        return {
            "question_type": 6,
            "question_blocks": (
                self._instruction_block(parsed)
                + [{
                    "type": "single_gap",
                    "content": parsed["question_text"],
                    "options": parsed["options"]
                }]
            ),
            "correct_answer": parsed["correct"]
        }


# --------------------------------------------------
# 6️⃣ Gap Fill
# --------------------------------------------------

@register
class GapFillHandler(InstructionAwareHandler):
    question_type = 3

    def parse(self, ctx):
        instruction = parse_block(
            ctx,
            "QUESTION_INSTRUCTION",
            ["QUESTION_TEXT"],
            required=False
        )

        question_text = parse_block(
            ctx,
            "QUESTION_TEXT",
            ["WORD_BANK"]
        )

        word_bank = parse_word_bank(ctx)
        correct = parse_blank_answers(ctx)

        return {
            "instruction": instruction,
            "question_text": question_text,
            "word_bank": word_bank,
            "correct": correct
        }

    def validate(self, parsed, *, context=None):
        self._validate_instruction(parsed, context=context)

        # 🔹 Find blanks exactly as used in NAPLAN Reading
        blanks = re.findall(r"\[BLANK\]", parsed["question_text"])
        if not blanks:
            raise ValueError("NO_BLANKS_FOUND")

        # 🔹 Must have one correct answer per blank
        if len(parsed["correct"]) != len(blanks):
            raise ValueError("INVALID_CORRECT_ANSWER_COUNT")

        # 🔹 Every correct answer must exist in word bank
        for ans in parsed["correct"]:
            if ans not in parsed["word_bank"]:
                raise ValueError("ANSWER_NOT_IN_WORD_BANK")

    def build_exam_bundle(self, parsed):
        return {
            "question_type": 3,
            "question_blocks": (
                self._instruction_block(parsed)
                + [{
                    "type": "gap_fill",
                    "content": parsed["question_text"],
                    "word_bank": parsed["word_bank"]
                }]
            ),
            "correct_answer": parsed["correct"]
        }

def parse_common_sections_naplan_reading(ctx):
    """
    Parses shared NAPLAN Reading sections:

    - METADATA
    - READING_MATERIAL with support for MULTIPLE extracts
      - TEXT_TITLE starts a new extract
      - IMAGES belong to the current extract

    Leaves ctx.ptr positioned at the first QUESTION marker.

    Returns:
        meta: dict
        extracts: list[dict]
    """

    # ==================================================
    # METADATA
    # ==================================================
    if not ctx.peek() or ctx.peek().strip().upper() != "METADATA:":
        raise ValueError("MISSING_METADATA")

    ctx.next()  # consume METADATA:

    meta = {}

    while ctx.peek():
        peek = ctx.peek().strip().upper()

        if peek == "READING_MATERIAL:":
            break

        line = ctx.next()
        if ":" in line:
            k, v = line.split(":", 1)
            meta[k.strip().upper()] = v.strip().strip('"')

    required = {"CLASS", "SUBJECT", "TOPIC", "DIFFICULTY"}
    if not required.issubset(meta.keys()):
        raise ValueError("INCOMPLETE_METADATA")

    # ==================================================
    # READING MATERIAL (MULTI-EXTRACT)
    # ==================================================
    line = ctx.next()
    if not line or line.strip().upper() != "READING_MATERIAL:":
        raise ValueError("MISSING_READING_MATERIAL")

    extracts = []
    current = None
    extract_index = 0

    def start_new_extract(title=None):
        nonlocal extract_index, current
        extract_index += 1
        current = {
            "extract_id": chr(64 + extract_index),  # A, B, C...
            "title": title,
            "text_lines": [],
            "images": []
        }
        extracts.append(current)

    while ctx.peek():
        raw = ctx.peek()
        line = raw.strip()
        upper = line.upper()

        # -------------------------------
        # STOP at first question
        # -------------------------------
        if upper.startswith("QUESTION_") or upper.startswith("--- QUESTION"):
            break

        # -------------------------------
        # TEXT_TITLE → new extract
        # -------------------------------
        if upper.startswith("TEXT_TITLE:"):
            ctx.next()
            title = line.split(":", 1)[1].strip().strip('"')
            start_new_extract(title=title)
            continue

        # -------------------------------
        # IMAGES section
        # -------------------------------
        if upper.startswith("IMAGES"):
            ctx.next()

            if current is None:
                start_new_extract(title=None)

            # Inline: IMAGES: a.png, b.png
            if ":" in line:
                _, rest = line.split(":", 1)
                current["images"].extend(
                    [i.strip().lower() for i in rest.split(",") if i.strip()]
                )

            # Multiline images
            while ctx.peek():
                peek = ctx.peek().strip()
                up = peek.upper()

                if (
                    not peek
                    or up.startswith("QUESTION_")
                    or up.startswith("--- QUESTION")
                    or up.startswith("TEXT_TITLE:")
                ):
                    break

                current["images"].append(
                    ctx.next().lstrip("-• ").strip().lower()
                )

            continue

        # -------------------------------
        # Regular reading text
        # -------------------------------
        if current is None:
            # Backward compatibility:
            # reading text before any TEXT_TITLE
            start_new_extract(title=None)

        current["text_lines"].append(ctx.next())

    # ==================================================
    # VALIDATION + NORMALIZATION
    # ==================================================
    if not extracts:
        raise ValueError("READING_TOO_SHORT")

    normalized = []

    for e in extracts:
        content = "\n".join(e["text_lines"]).strip()
        if content:
            normalized.append({
                "extract_id": e["extract_id"],
                "title": e["title"],
                "content": content,
                "images": e["images"]
            })

    if not normalized:
        raise ValueError("READING_TOO_SHORT")

    return meta, normalized

def build_blocks_naplan_reading_from_extracts(extracts, db):
    """
    Builds renderer-ready reading blocks from multiple extracts.

    Each extract may contain:
    - title (optional)
    - text content
    - images (validated via UploadedImage)

    Returns a single shared reading block suitable
    for prepending to all questions.
    """

    reading_block = {
        "type": "reading",
        "extracts": []
    }

    for e in extracts:
        extract_block = {
            "extract_id": e["extract_id"],
            "title": e.get("title"),
            "content": e["content"],
            "images": []
        }

        for img in e.get("images", []):
            normalized = img.lstrip("-• ").strip().lower()

            record = (
                db.query(UploadedImage)
                .filter(
                    func.lower(func.trim(UploadedImage.original_name)) == normalized
                )
                .first()
            )

            if not record:
                raise ValueError(f"IMAGE_NOT_UPLOADED: {normalized}")

            extract_block["images"].append(record.gcs_url)

        reading_block["extracts"].append(extract_block)

    return [reading_block]

def parse_year(meta):
    try:
        return int(meta["CLASS"].replace("Year", "").strip())
    except Exception:
        raise ValueError("INVALID_CLASS_YEAR")
def parse_word_bank(ctx):
    words = []

    def normalize(s):
        return (
            s.strip()
             .lstrip("-•– ")
             .upper()
             .replace(" ", "_")
        )

    # Seek WORD_BANK header
    while ctx.peek() and normalize(ctx.peek()) != "WORD_BANK:":
        ctx.next()

    if not ctx.peek():
        raise ValueError("EMPTY_WORD_BANK")

    ctx.next()  # consume WORD BANK:

    while ctx.peek():
        line = ctx.peek().strip()

        # Stop on next section
        if normalize(line).startswith((
            "CORRECT_ANSWER",
            "QUESTION_"
        )):
            break

        if line.startswith(("-", "•")):
            words.append(ctx.next().lstrip("-• ").strip())
            continue

        break

    if not words:
        raise ValueError("EMPTY_WORD_BANK")

    return words

def parse_blank_answers(ctx):
    line = ctx.next()
    if not line.strip().upper().startswith("CORRECT_ANSWER"):
        raise ValueError("MISSING_CORRECT_ANSWER")

    answers = []

    # Inline answer support (rare but allowed)
    if ":" in line:
        _, v = line.split(":", 1)
        v = v.strip()
        if v:
            answers.append(v)

    while ctx.peek():
        raw = ctx.peek()
        line = raw.strip()

        # 🚨 HARD STOP: question boundary
        if is_question_boundary(line):
            break

        # valid bullet answer
        if line.startswith("-"):
            answers.append(
                ctx.next().lstrip("- ").strip()
            )
            continue

        # anything else ends the answer block
        break

    if not answers:
        raise ValueError("EMPTY_CORRECT_ANSWER")

    return answers
def parse_image_options(ctx):
    if ctx.next() != "IMAGE_OPTIONS:":
        raise ValueError("MISSING_IMAGE_OPTIONS")

    options = {}

    while ctx.peek() and not ctx.peek().startswith("CORRECT_ANSWER"):
        line = ctx.next()
        if ":" in line:
            k, v = line.split(":", 1)
            options[k.strip()] = v.strip().lower()

    if len(options) != 4:
        raise ValueError("INVALID_IMAGE_OPTIONS")

    return options

def resolve_image_options(image_options, db):
    resolved = {}
    missing = None

    for key, img in image_options.items():
        record = (
            db.query(UploadedImage)
            .filter(func.lower(func.trim(UploadedImage.original_name)) == img)
            .first()
        )

        if not record:
            missing = img
            break

        resolved[key] = record.gcs_url

    if missing:
        raise ValueError(f"IMAGE_OPTION_NOT_UPLOADED: {missing}")

    return resolved
def parse_statements(ctx):
    statements = []

    # Seek STATEMENTS header
    while ctx.peek() and ctx.peek().strip().upper() != "STATEMENTS:":
        ctx.next()

    if not ctx.peek():
        raise ValueError("EMPTY_STATEMENTS")

    ctx.next()  # consume STATEMENTS:

    while ctx.peek():
        raw = ctx.peek()
        line = raw.strip()

        # 🚨 HARD STOP: question boundary ONLY
        if is_question_boundary(line):
            break

        # valid statement
        if line.startswith(("-", "•")):
            statements.append(ctx.next().lstrip("-• ").strip())
            continue

        # anything else ends statements
        break

    if not statements:
        raise ValueError("EMPTY_STATEMENTS")

    return statements
 
def parse_true_false_answers(ctx):
    line = ctx.next()
    if not line.strip().upper().startswith("CORRECT_ANSWER"):
        raise ValueError("MISSING_CORRECT_ANSWER")

    answers = []

    while ctx.peek():
        raw = ctx.peek()
        line = raw.strip()

        # 🚨 HARD STOP: question boundary
        if is_question_boundary(line):
            break

        # valid bullet answer
        if line.startswith("-"):
            val = ctx.next().lstrip("- ").strip().lower()

            if val not in ("true", "false"):
                raise ValueError("INVALID_TRUE_FALSE_VALUE")

            answers.append(val.capitalize())
            continue

        break

    if not answers:
        raise ValueError("EMPTY_CORRECT_ANSWER")

    return answers

def parse_word_options(ctx):
    if ctx.next() != "WORD_OPTIONS:":
        raise ValueError("MISSING_WORD_OPTIONS")

    options = {}

    while ctx.peek() and not ctx.peek().startswith("CORRECT_ANSWER"):
        line = ctx.next()
        if ":" in line:
            k, v = line.split(":", 1)
            options[k.strip()] = v.strip()

    if len(options) < 2:
        raise ValueError("INVALID_WORD_OPTIONS")

    return options
def is_cloze_question(blocks: list) -> bool:
    for b in blocks:
        if b.get("type") == "text":
            content = b.get("content", "").lower()

            # Explicit declaration
            if "question_type: 5" in content:
                return True

            # Structural detection
            if "cloze:" in content and "{{dropdown}}" in content:
                return True

    return False

def is_cloze_question(blocks: list) -> bool:
    for b in blocks:
        if b.get("type") == "text":
            content = b.get("content", "").lower()

            # Explicit declaration
            if "question_type: 5" in content or "question_type:5" in content:
                return True

            # Structural detection
            if "cloze:" in content and "{{dropdown}}" in content:
                return True

    return False
def extract_cloze_options_from_blocks(block_elements):
    """
    Extract CLOZE OPTIONS from a Word-derived block list.

    Supports:
    - OPTIONS appearing in a single collapsed text block
    - OPTIONS split across multiple blocks
    - Option formats: A: text | A. text | A) text

    Fails safely by returning {} if no valid options found.
    """

    options = {}
    in_options = False

    print("🔎 [CLOZE DEBUG] Starting options extraction")

    for idx, el in enumerate(block_elements):
        raw = el.get("content")

        print(
            f"🔎 [CLOZE DEBUG] Block {idx} | "
            f"type={el.get('type')} | "
            f"content={repr(raw)}"
        )

        if not isinstance(raw, str):
            continue

        # Word often collapses multiple lines into one block
        lines = [
            line.strip()
            for line in raw.splitlines()
            if line.strip()
        ]

        for line in lines:
            upper = line.upper()

            print(
                f"🔎 [CLOZE DEBUG] Line={repr(line)} | "
                f"in_options={in_options}"
            )

            # --------------------------------------------------
            # Detect start of OPTIONS section
            # --------------------------------------------------
            if upper == "OPTIONS:":
                in_options = True
                print("✅ [CLOZE DEBUG] OPTIONS section detected")
                continue

            # --------------------------------------------------
            # Parse option lines once inside OPTIONS
            # --------------------------------------------------
            if in_options:
                match = re.match(
                    r"^([A-Z])[\:\.\)]\s*(.+)$",
                    line
                )

                if match:
                    key, value = match.groups()
                    options[key] = value.strip()

                    print(
                        f"✅ [CLOZE DEBUG] Parsed option "
                        f"{key}: {value.strip()}"
                    )
                    continue

                # --------------------------------------------------
                # Stop parsing when OPTIONS section ends
                # --------------------------------------------------
                if options:
                    print(
                        "🛑 [CLOZE DEBUG] End of OPTIONS section"
                    )
                    return options

    if not options:
        print(
            "❌ [CLOZE DEBUG] No CLOZE options found"
        )

    return options

def handle_cloze_question(
    q,
    question_block,
    meta,
    db,
    request_id,
    summary,
    block_idx
):
    print(
        f"[{request_id}] 🧩 [CLOZE] Persisting CLOZE | block={block_idx}"
    )

    persist_question(
        q=q,
        question_type=5,
        question_block=question_block,
        meta=meta,
        db=db,
        request_id=request_id,
        summary=summary,
        block_idx=block_idx,
    )


ClozeQuestionSchema = {
    "type": "object",
    "properties": {
        "class_name": {"type": "string"},
        "year": {"type": "integer"},
        "subject": {"type": "string"},
        "topic": {"type": "string"},
        "difficulty": {"type": "string"},
        "cloze_text": {"type": "string"},
        "correct_answer": {"type": "string"},
        "answer_type": {
            "type": "string",
            "enum": ["CLOZE_DROPDOWN"],
        },
    },
    "required": [
        "cloze_text",
        "correct_answer",
        "answer_type",
    ],
    "additionalProperties": False,
}
import re

import re

def extract_cloze_from_exam_block(block_elements: list[dict]) -> dict:
    """
    Deterministic extractor for CLOZE (question_type = 5).

    Design guarantees:
    - One CLOZE per exam block
    - One {{dropdown}}
    - Options are A:, B:, C: (Word whitespace tolerant)
    - CORRECT_ANSWER follows CORRECT_ANSWER:
    - Order independent
    """

    print("🔍 [CLOZE] START extract_cloze_from_exam_block")

    try:
        # --------------------------------------------------
        # 1. Flatten exam block into normalized text lines
        # --------------------------------------------------
        lines: list[str] = []

        for el in block_elements:
            text = el.get("content")
            if not isinstance(text, str):
                continue

            for raw_line in text.splitlines():
                # Normalize Word weirdness
                line = raw_line.replace("\u00a0", " ")  # non-breaking space
                line = re.sub(r"\s+", " ", line)        # collapse whitespace
                line = line.strip()

                if line:
                    lines.append(line)

        print(f"🔍 [CLOZE] Flattened lines ({len(lines)}):")
        for l in lines:
            print(f"   • {repr(l)}")

        # --------------------------------------------------
        # 2. Extract CLOZE text (must be exactly one)
        # --------------------------------------------------
        cloze_lines = [l for l in lines if "{{dropdown}}" in l]

        if len(cloze_lines) != 1:
            raise ValueError(
                f"CLOZE must contain exactly one {{dropdown}} "
                f"(found {len(cloze_lines)})"
            )

        cloze_text = cloze_lines[0]
        print(f"🧩 [CLOZE] cloze_text = {cloze_text}")

        # --------------------------------------------------
        # 3. Extract OPTIONS (Word-proof)
        # --------------------------------------------------
        option_pattern = re.compile(r"^([A-Z])\s*:\s*(.+)$")
        options: dict[str, str] = {}

        for line in lines:
            match = option_pattern.match(line)
            if match:
                key = match.group(1)
                value = match.group(2).strip()
                options[key] = value
                print(f"🅾️ [CLOZE] option {key} = {value}")

        if not options:
            raise ValueError("CLOZE missing options")

        # --------------------------------------------------
        # 4. Extract CORRECT_ANSWER
        # --------------------------------------------------
        correct_answer = None

        for i, line in enumerate(lines):
            if line.upper() == "CORRECT_ANSWER:":
                if i + 1 >= len(lines):
                    raise ValueError("CORRECT_ANSWER value missing")
                correct_answer = lines[i + 1].strip()
                break

        if not correct_answer:
            raise ValueError("CLOZE missing correct_answer")

        if correct_answer not in options:
            raise ValueError(
                f"Correct answer '{correct_answer}' not in options {list(options.keys())}"
            )

        print(f"🎯 [CLOZE] correct_answer = {correct_answer}")

        # --------------------------------------------------
        # 5. Extract METADATA (order-agnostic)
        # --------------------------------------------------
        meta = {
            "class_name": None,
            "year": None,
            "subject": None,
            "topic": None,
            "difficulty": None,
        }

        for line in lines:
            if ":" not in line:
                continue

            key, value = line.split(":", 1)
            key = key.strip().lower()
            value = value.strip().strip('"')

            if key in {"class", "class_name"}:
                meta["class_name"] = value
            elif key == "year":
                meta["year"] = int(value)
            elif key == "subject":
                meta["subject"] = value
            elif key == "topic":
                meta["topic"] = value
            elif key == "difficulty":
                meta["difficulty"] = value

        print("📦 [CLOZE] metadata =", meta)

        # --------------------------------------------------
        # 6. Final payload
        # --------------------------------------------------
        result = {
            "class_name": meta["class_name"],
            "year": meta["year"],
            "subject": meta["subject"],
            "topic": meta["topic"],
            "difficulty": meta["difficulty"],
            "cloze_text": cloze_text,
            "options": options,
            "correct_answer": correct_answer,
            "answer_type": "CLOZE_DROPDOWN",
            "options_source": "document",
        }

        print("✅ [CLOZE] FINAL EXTRACTED QUESTION:")
        for k, v in result.items():
            print(f"   {k}: {v}")

        print("🔍 [CLOZE] END extract_cloze_from_exam_block")
        return result

    except Exception as e:
        print("❌ [CLOZE] EXTRACTOR FAILED")
        print(f"❌ [CLOZE] Error: {e}")
        print("❌ [CLOZE] Raw block contents:")
        for el in block_elements:
            print(repr(el.get("content")))
        raise

def parse_cloze_from_document(block_elements: list[dict]) -> dict:
    """
    Deterministic CLOZE (type 5) parser.

    Rules:
    - Exactly ONE line must contain {{dropdown}} → cloze_text
    - OPTIONS and CORRECT_ANSWER are mandatory
    - No GPT, no layout assumptions
    """

    print("🔍 [CLOZE PARSER] START")
    print(f"🔍 [CLOZE PARSER] Block elements = {len(block_elements)}")

    data = {
        "class_name": None,
        "year": None,
        "subject": None,
        "topic": None,
        "difficulty": None,
        "cloze_text": None,
        "options": {},
        "correct_answer": None,
        "answer_type": "CLOZE_DROPDOWN",
        "options_source": "document",
    }

    current_section = None

    for idx, el in enumerate(block_elements, start=1):
        text = el.get("content")
        if not isinstance(text, str):
            continue

        print(f"\n🔎 [CLOZE PARSER] Block {idx} raw content:\n{text}")

        lines = [l.strip() for l in text.splitlines() if l.strip()]

        for line in lines:
            upper = line.upper()

            print(
                f"➡️ [CLOZE PARSER] Line='{line}' "
                f"(section={current_section})"
            )

            # ==================================================
            # SECTION HEADERS (MUST RUN FIRST)
            # ==================================================
            if upper == "METADATA:":
                current_section = "metadata"
                print("📌 [CLOZE PARSER] Entered METADATA")
                continue

            if upper == "OPTIONS:":
                current_section = "options"
                print("📌 [CLOZE PARSER] Entered OPTIONS")
                continue

            if upper == "CORRECT_ANSWER:":
                current_section = "answer"
                print("📌 [CLOZE PARSER] Entered CORRECT_ANSWER")
                continue

            # ==================================================
            # CLOZE TEXT (AUTHORITATIVE, SECTION-AGNOSTIC)
            # ==================================================
            if "{{dropdown}}" in line:
                if data["cloze_text"] is not None:
                    raise ValueError(
                        "Multiple CLOZE lines detected "
                        "(multiple {{dropdown}} found)"
                    )

                data["cloze_text"] = line
                print(
                    f"🧩 [CLOZE PARSER] Detected cloze_text = "
                    f"{data['cloze_text']}"
                )
                continue

            # ==================================================
            # METADATA
            # ==================================================
            if current_section == "metadata":
                if upper.startswith("CLASS:"):
                    data["class_name"] = line.split(":", 1)[1].strip().strip('"')
                    print(f"✅ class_name = {data['class_name']}")

                elif upper.startswith("YEAR:"):
                    data["year"] = int(line.split(":", 1)[1].strip())
                    print(f"✅ year = {data['year']}")

                elif upper.startswith("SUBJECT:"):
                    data["subject"] = line.split(":", 1)[1].strip().strip('"')
                    print(f"✅ subject = {data['subject']}")

                elif upper.startswith("TOPIC:"):
                    data["topic"] = line.split(":", 1)[1].strip().strip('"')
                    print(f"✅ topic = {data['topic']}")

                elif upper.startswith("DIFFICULTY:"):
                    data["difficulty"] = line.split(":", 1)[1].strip().strip('"')
                    print(f"✅ difficulty = {data['difficulty']}")

            # ==================================================
            # OPTIONS
            # ==================================================
            elif current_section == "options":
                if ":" in line:
                    key, value = line.split(":", 1)
                    key = key.strip()
                    value = value.strip()

                    if len(key) == 1 and key.isupper():
                        data["options"][key] = value
                        print(f"🅾️ option {key} = {value}")
                    else:
                        print(
                            f"⚠️ Ignored invalid option line: {line}"
                        )

            # ==================================================
            # CORRECT ANSWER
            # ==================================================
            elif current_section == "answer":
                data["correct_answer"] = line
                print(f"🎯 correct_answer = {data['correct_answer']}")

    print("\n📦 [CLOZE PARSER] FINAL PARSED DATA")
    for k, v in data.items():
        print(f"  {k}: {v}")

    print("🔍 [CLOZE PARSER] END")
    return data


@app.post("/upload-word-naplan-reading")
async def upload_word_naplan_reading(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    request_id = str(uuid.uuid4())[:8]

    print("\n" + "=" * 80)
    print(f"📘 NAPLAN READING UPLOAD START [{request_id}]")
    print("=" * 80)

    # --------------------------------------------------
    # 1️⃣ Read + normalize document
    # --------------------------------------------------
    raw = await file.read()
    text = normalize_doc_text(extract_text_from_docx(raw))

    print(f"[{request_id}] 📄 Extracted characters: {len(text)}")

    exam_blocks = split_exam_blocks(text)
    print(f"[{request_id}] 🧩 EXAM blocks detected: {len(exam_blocks)}")

    saved_ids = []
    skipped = []

    # --------------------------------------------------
    # 2️⃣ Process each EXAM
    # --------------------------------------------------
    for idx, block in enumerate(exam_blocks, start=1):
        print("\n" + "-" * 70)
        print(f"[{request_id}] 🔍 Processing EXAM {idx}")
        print("-" * 70)

        try:
            lines = [l.rstrip() for l in block.splitlines() if l.strip()]
            print(f"[{request_id}] 📄 Lines after cleanup: {len(lines)}")

            ctx = ParseContext(lines)
            passage_id = str(uuid.uuid4())
            print(f"[{request_id}] 🔗 passage_id = {passage_id}")


           

            # -------------------------------
            # Common sections
            # -------------------------------
            print(f"[{request_id}] 📘 Parsing METADATA / Reading / Extracts")
            meta, extracts = parse_common_sections_naplan_reading(ctx)
            
            print(f"[{request_id}] 📘 METADATA: {meta}")
            print(f"[{request_id}] 📖 Extracts detected: {len(extracts)}")
            shared_blocks = build_blocks_naplan_reading_from_extracts(extracts, db)


            question_index = 0

            while ctx.peek():
                peek = ctx.peek().strip().upper()
            
                # Skip until QUESTION START
                if not peek.startswith("--- QUESTION START"):
                    ctx.next()
                    continue
            
                ctx.next()  # consume --- QUESTION START ---
                question_index += 1
            
                print(f"[{request_id}] ➤ Parsing QUESTION {question_index}")
            
                try:
                    # -------------------------------
                    # QUESTION_TYPE
                    # -------------------------------
                    qt_line = ctx.next()
                    if not qt_line or not qt_line.strip().upper().startswith("QUESTION_TYPE"):
                        raise ValueError("MISSING_QUESTION_TYPE")
            
                    qt = int(qt_line.split(":", 1)[1].strip())
                    handler = QUESTION_HANDLERS.get(qt)
            
                    if not handler:
                        raise ValueError(f"UNSUPPORTED_QUESTION_TYPE: {qt}")
            
                    # -------------------------------
                    # Parse + validate question
                    # -------------------------------
                    parsed = handler.parse(ctx)
                    handler.validate(parsed, context="reading")

            
                    exam_bundle = handler.build_exam_bundle(parsed)
            
                    # Resolve image options if needed
                    if "image_options" in exam_bundle:
                        exam_bundle["image_options"] = resolve_image_options(
                            exam_bundle["image_options"], db
                        )
            
                    # Prepend shared reading blocks
                    exam_bundle["question_blocks"] = (
                        shared_blocks + exam_bundle["question_blocks"]
                    )
            
                    # -------------------------------
                    # Save QUESTION
                    # -------------------------------
                    obj = QuestionNaplanReading(
                        passage_id=passage_id,
                        class_name=meta["CLASS"].lower(),
                        subject=meta["SUBJECT"].lower(),
                        year=parse_year(meta),
                        difficulty=meta["DIFFICULTY"].lower(),
                        topic=meta["TOPIC"],
                        question_type=qt,
                        exam_bundle=exam_bundle
                    )
            
                    db.add(obj)
                    db.commit()
                    db.refresh(obj)
            
                    saved_ids.append(obj.id)
            
                    print(
                        f"[{request_id}] 💾 SAVED QUESTION {question_index} | ID={obj.id}"
                    )
            
                except Exception as q_err:
                    db.rollback()
                    skipped.append(
                       (f"exam {idx} question {question_index}", str(q_err))
                    )

                    print(
                        f"[{request_id}] ❌ SKIPPED QUESTION {question_index} | {q_err}"
                    )
            
                # -------------------------------
                # Consume QUESTION END safely
                # -------------------------------
                while ctx.peek() and not ctx.peek().strip().upper().startswith("--- QUESTION END"):
                    ctx.next()
            
                if ctx.peek():
                    ctx.next()  # consume --- QUESTION END ---


            
            
            #saved_ids.append(obj.id)
            #print(f"[{request_id}] 💾 SAVED EXAM {idx} | ID={obj.id}")

        except Exception as e:
            db.rollback()
            error_msg = str(e)

            print(f"[{request_id}] ❌ FAILED EXAM {idx}")
            print(f"[{request_id}] ❌ Reason: {error_msg}")

            # Pointer visibility (very useful)
            if "ctx" in locals():
                print(
                    f"[{request_id}] 🧭 Parser state | "
                    f"ptr={ctx.ptr}, "
                    f"current_line={ctx.peek()}"
                )

            skipped.append((idx, error_msg))

    # --------------------------------------------------
    # 3️⃣ Final decision
    # --------------------------------------------------
    print("\n" + "=" * 80)
    print(f"[{request_id}] 📊 Upload summary")
    print(f"[{request_id}] ✅ Saved: {len(saved_ids)}")
    print(f"[{request_id}] ❌ Skipped: {len(skipped)}")
    print(f"[{request_id}] Details: {skipped}")
    print("=" * 80)

    if not saved_ids:
        raise HTTPException(
            status_code=422,
            detail={
                "error": "NO_EXAMS_SAVED",
                "skipped": skipped
            }
        )

    return {
        "status": "success" if not skipped else "partial_success",
        "saved": len(saved_ids),
        "skipped": skipped,
        "ids": saved_ids
    }


import re

def parse_exam_block(block_text: str):
    """
    Deterministic parser for Exam Format v2.
    The document is the single source of truth.
    """

    # --------------------------------------------------
    # Helper: extract section by label
    # --------------------------------------------------
    def section(label: str) -> str:
        """
        Extract a labeled section safely.
        METADATA ends ONLY at READING_MATERIAL.
        Other sections end at the next ALL-CAPS label.
        """
        if label == "METADATA":
            pattern = rf"(?ms)^METADATA\s*:\s*(.*?)(?=^READING_MATERIAL\s*:|\Z)"
        else:
            pattern = rf"(?ms)^{label}\s*:\s*(.*?)(?=^[A-Z_]+\s*:|\Z)"
    
        match = re.search(pattern, block_text)
        if not match:
            raise ValueError(f"Missing required section: {label}")
    
        return match.group(1).strip()

    # --------------------------------------------------
    # 1️⃣ METADATA
    # --------------------------------------------------
    metadata_raw = section("METADATA")

    metadata = {}
    for line in metadata_raw.splitlines():
        if ":" not in line:
            continue
        k, v = line.split(":", 1)
        key = k.strip().lower()
        metadata[key] = v.strip().strip('"')

    required_meta = ["class", "subject", "topic", "difficulty", "total_questions"]
    for k in required_meta:
        if k not in metadata:
            raise ValueError(f"Missing METADATA field: {k}")

    metadata_parsed = {
        "class_name": metadata["class"],
        "subject": metadata["subject"],
        "topic": metadata["topic"],
        "difficulty": metadata["difficulty"],
        "total_questions": int(metadata["total_questions"]),
    }

    # --------------------------------------------------
    # 2️⃣ READING MATERIAL
    # --------------------------------------------------
    instructions = section("INSTRUCTIONS")
    title = section("TITLE")
    paragraphs_raw = section("PARAGRAPHS")

    paragraphs = {}
    for line in paragraphs_raw.splitlines():
        if ":" not in line:
            continue
        k, v = line.split(":", 1)
        if k.strip().isdigit():
            paragraphs[k.strip()] = v.strip()

    if not paragraphs:
        raise ValueError("No paragraphs parsed")

    # --------------------------------------------------
    # 3️⃣ ANSWER OPTIONS
    # --------------------------------------------------
    options_raw = section("ANSWER_OPTIONS")

    answer_options = {}
    for line in options_raw.splitlines():
        if "." not in line:
            continue
        k, v = line.split(".", 1)
        k = k.strip()
        if k in ["A", "B", "C", "D", "E", "F", "G"]:
            answer_options[k] = v.strip()

    if len(answer_options) != 7:
        raise ValueError("Answer options must contain A–G")

    
    # --------------------------------------------------
    # 4️⃣ QUESTIONS (PIPE FORMAT — DETERMINISTIC)
    # --------------------------------------------------
    questions_raw = section("QUESTIONS")
    
    print("\n🧪 RAW QUESTIONS SECTION (repr):")
    print(repr(questions_raw))
    print()
    
    questions = []
    
    for idx, raw_line in enumerate(questions_raw.splitlines(), start=1):
        print(f"🧪 Q-LINE {idx} (raw): {repr(raw_line)}")
    
        # Normalize DOCX / Unicode junk
        line = (
            raw_line
            .strip()
            .replace("│", "|")   # Unicode pipe
            .replace("＝", "=")   # Unicode equals
            .replace("\r", "")
        )
    
        print(f"🧪 Q-LINE {idx} (normalized): {repr(line)}")
    
        if not line:
            print("   ↳ skipped (empty)")
            continue
    
        if "|" not in line:
            print("   ↳ skipped (no pipe)")
            continue
    
        m = re.match(
            r"""
            ^\s*\d+\s*                # question index
            \|\s*paragraph\s*=\s*(\d+) # paragraph number
            \s*\|\s*correct_answer\s*=\s*([A-Ga-g])\s*$
            """,
            line,
            re.VERBOSE
        )
    
        if not m:
            raise ValueError(f"Invalid QUESTION line format: {line}")
    
        print("   ✅ parsed")
    
        questions.append({
            "paragraph": int(m.group(1)),
            "correct_answer": m.group(2).upper()
        })
    
    if not questions:
        raise ValueError("No questions parsed")


    # --------------------------------------------------
    # 5️⃣ FINAL SHAPE
    # --------------------------------------------------
    return {
        "metadata": metadata_parsed,
        "reading_material": {
            "title": title,
            "instructions": [
                l.strip() for l in instructions.splitlines() if l.strip()
            ],
            "paragraphs": paragraphs,
        },
        "answer_options": answer_options,
        "questions": questions,
    }


@app.post("/upload-word-reading-main-idea-ai")
async def upload_word_reading_main_idea_ai(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    print("\n" + "=" * 70)
    print("📄 START: MAIN IDEA & SUMMARY WORD UPLOAD")
    print("=" * 70)

    # --------------------------------------------------
    # 1️⃣ File validation
    # --------------------------------------------------
    print("📥 STEP 1: File validation")

    if file.content_type != (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        raise HTTPException(status_code=400, detail="File must be .docx")

    raw = await file.read()
    print("✅ File read | bytes:", len(raw))

    # --------------------------------------------------
    # 2️⃣ Text extraction
    # --------------------------------------------------
    print("\n📄 STEP 2: Extracting text")

    try:
        full_text = extract_text_from_docx(raw)
        full_text = normalize_doc_text(full_text)
    except Exception as e:
        raise HTTPException(status_code=500, detail="Failed to extract document text")

    if not full_text or len(full_text.strip()) < 300:
        raise HTTPException(status_code=400, detail="Invalid or empty document")

    # --------------------------------------------------
    # 3️⃣ Detect exam blocks
    # --------------------------------------------------
    print("\n🧩 STEP 3: Detecting exam blocks")

    start_token = "=== EXAM START ==="
    end_token = "=== EXAM END ==="

    blocks = []
    for part in full_text.split(start_token)[1:]:
        if end_token not in part:
            continue
        block = part.split(end_token)[0].strip()
        if len(block) >= 500:
            blocks.append(block)

    if not blocks:
        raise HTTPException(
            status_code=400,
            detail="No valid exam blocks found"
        )

    print("✅ Exam blocks found:", len(blocks))
    saved_ids = []

    # --------------------------------------------------
    # 4️⃣ AI validation prompt
    # --------------------------------------------------
    system_prompt = """
You are an exam integrity validation engine.

The document is already structured and must be treated as the single source of truth.

YOUR ROLE:
- VERIFY the exam content only
- DO NOT extract, rewrite, infer, summarise, or restructure anything
- DO NOT create missing fields
- DO NOT correct mistakes

VALIDATION RULES (FAIL HARD):
- Exactly ONE Main Idea & Summary reading exam must be present
- Total number of questions MUST match Total_Questions
- Each question MUST reference an existing paragraph
- All answer options MUST be shared and include A–G only
- Each correct_answer MUST be one of A–G
- No required section may be missing or empty

OUTPUT RULES:
- Return ONLY valid JSON
- If ALL checks pass, return: {"status": "ok"}
- If ANY check fails, return: {}
- No markdown
- No explanations
"""

    # --------------------------------------------------
    # 5️⃣ Process each exam block
    # --------------------------------------------------
    for block_idx, block_text in enumerate(blocks, start=1):
        print("\n" + "-" * 70)
        print(f"🔍 Processing block {block_idx}/{len(blocks)}")

        # ---------------- AI integrity validation ----------------
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                temperature=0,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {
                        "role": "user",
                        "content": f"BEGIN_DOCUMENT\n{block_text}\nEND_DOCUMENT"
                    }
                ]
            )
        except Exception as e:
            print("❌ AI call failed:", str(e))
            continue

        try:
            ai_result = json.loads(response.choices[0].message.content.strip())
        except Exception:
            print("❌ AI returned invalid JSON")
            continue

        if ai_result != {"status": "ok"}:
            print("❌ AI integrity validation failed")
            continue

        print("✅ AI integrity validation passed")

        # ---------------- Parse document ----------------
        try:
            parsed = parse_exam_block(block_text)
        except Exception as e:
            print("❌ Document parsing failed:", str(e))
            continue

        meta = parsed["metadata"]
        rm = parsed["reading_material"]
        paragraphs = rm["paragraphs"]
        questions = parsed["questions"]
        opts = parsed["answer_options"]

        # ---------------- Backend hard validation ----------------
        print("🧪 Backend validation")

        if len(paragraphs) != meta["total_questions"]:
            print("❌ Paragraph count mismatch")
            continue

        if len(questions) != meta["total_questions"]:
            print("❌ Question count mismatch")
            continue

        required_opts = ["A", "B", "C", "D", "E", "F", "G"]
        if any(k not in opts for k in required_opts):
            print("❌ Missing answer options")
            continue

        valid = True
        for q in questions:
            if str(q["paragraph"]) not in paragraphs:
                valid = False
                break
            if q["correct_answer"] not in required_opts:
                valid = False
                break

        if not valid:
            print("❌ Question validation failed")
            continue

        print("✅ Backend validation passed")

        # --------------------------------------------------
        # 6️⃣ Enrich questions
        # --------------------------------------------------
        enriched_questions = []
        for i, q in enumerate(questions, start=1):
            enriched_questions.append({
                "question_id": f"MI_Q{i}",
                "paragraph": q["paragraph"],
                "question_text": f"Paragraph {q['paragraph']}",
                "correct_answer": q["correct_answer"]
            })

        # --------------------------------------------------
        # 7️⃣ Bundle
        # --------------------------------------------------
        bundle = {
            "question_type": "main_idea",
            "reading_material": rm,
            "answer_options": opts,
            "questions": enriched_questions
        }

        # --------------------------------------------------
        # 8️⃣ Save to DB
        # --------------------------------------------------
        try:
            obj = QuestionReading(
                class_name=meta["class_name"].lower(),
                subject=meta["subject"],
                difficulty=meta["difficulty"].lower(),
                topic=meta["topic"],
                total_questions=meta["total_questions"],
                exam_bundle=bundle
            )

            db.add(obj)
            db.commit()
            db.refresh(obj)

            saved_ids.append(obj.id)
            print(f"✅ Saved exam | ID={obj.id}")

        except Exception as e:
            db.rollback()
            print("❌ Database save failed:", str(e))

    # --------------------------------------------------
    # 9️⃣ Final response
    # --------------------------------------------------
    print("\n" + "=" * 70)
    print("🎉 UPLOAD COMPLETE")
    print("   → Saved exams:", len(saved_ids))
    print("   → IDs:", saved_ids)
    print("=" * 70)

    return {
        "message": "Main Idea & Summary documents processed",
        "saved_count": len(saved_ids),
        "bundle_ids": saved_ids
    }
    


#here line 8026
@app.post("/upload-word-reading-comparative-ai")
async def upload_word_reading_comparative_ai(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    print("\n" + "=" * 70)
    print("📄 START: COMPARATIVE ANALYSIS WORD UPLOAD (AI - DEBUG MODE)")
    print("=" * 70)

    # --------------------------------------------------
    # 1️⃣ File validation
    # --------------------------------------------------
    print("📥 STEP 1: File validation")
    print("   → Filename:", file.filename)
    print("   → Content-Type:", file.content_type)

    if file.content_type != (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        print("❌ INVALID FILE TYPE")
        raise HTTPException(status_code=400, detail="File must be .docx")

    raw = await file.read()
    print("✅ File read successfully | bytes:", len(raw))

    # --------------------------------------------------
    # 2️⃣ Text extraction
    # --------------------------------------------------
    print("\n📄 STEP 2: Extracting text from DOCX")

    try:
        full_text = extract_text_from_docx(raw)
        full_text = normalize_doc_text(full_text)
    except Exception as e:
        print("❌ DOCX TEXT EXTRACTION FAILED")
        print("   → Exception:", str(e))
        raise HTTPException(status_code=500, detail="Failed to extract document text")

    if not full_text:
        print("❌ Extracted text is EMPTY")
        raise HTTPException(status_code=400, detail="Empty document")

    print("✅ Text extracted")
    print("   → Character count:", len(full_text))

    if len(full_text.strip()) < 300:
        print("❌ Document too short")
        raise HTTPException(status_code=400, detail="Invalid document")

    # --------------------------------------------------
    # 3️⃣ Exam block detection
    # --------------------------------------------------
    print("\n🧩 STEP 3: Detecting exam blocks")

    start_token = "=== EXAM START ==="
    end_token = "=== EXAM END ==="

    parts = full_text.split(start_token)
    print("   → START tokens found:", len(parts) - 1)

    blocks = []

    for idx, part in enumerate(parts[1:], start=1):
        print(f"   → Inspecting block candidate {idx}")

        if end_token not in part:
            print("     ⚠️ END token missing — skipped")
            continue

        block = part.split(end_token)[0].strip()
        print("     → Block length:", len(block))

        if len(block) < 200:
            print("     ⚠️ Block too short — skipped")
            continue

        blocks.append(block)
        print("     ✅ Block accepted")

    if not blocks:
        print("❌ NO VALID EXAM BLOCKS FOUND")
        raise HTTPException(
            status_code=400,
            detail="No exam blocks found. Missing EXAM START / EXAM END markers."
        )

    print(f"✅ Total valid exam blocks: {len(blocks)}")

    saved_ids = []

    # --------------------------------------------------
    # 4️⃣ AI extraction prompt
    # --------------------------------------------------
    print("\n🤖 STEP 4: Preparing AI extraction prompt")

    system_prompt = """
  You are an exam content extraction engine.
  
  You MUST extract ONE COMPLETE COMPARATIVE ANALYSIS reading exam.
  
  The document contains a METADATA section with the following fields:
  - class_name
  - subject
  - topic
  - difficulty
  
  You MUST extract these fields EXACTLY as written and include them
  as top-level keys in the JSON output.
  
  ────────────────────────────────────────
  CRITICAL SCHEMA REQUIREMENTS (FAIL HARD)
  ────────────────────────────────────────
  
  The JSON output MUST ALWAYS contain:
  
  {
    "class_name": string,
    "subject": string,
    "topic": string,
    "difficulty": string,
    "reading_material": {
      "extracts": {
        "A": string,
        "B": string,
        "...": string
      }
    },
    "questions": [...]
  }
  
  ❗ Missing reading_material or extracts is INVALID.
  
  ────────────────────────────────────────
  EXTRACT RULES (STRICT)
  ────────────────────────────────────────
  - There MUST be AT LEAST 2 extracts
  - Extract labels MUST be consecutive capital letters starting from A
    (A,B) or (A,B,C) or (A,B,C,D)
  - Extract labels MUST EXACTLY match those in the document
  - DO NOT invent, merge, split, or rename extracts
  - Extract text MUST be preserved exactly as written
  - Empty extract text is INVALID
  
  ────────────────────────────────────────
  QUESTION RULES (STRICT)
  ────────────────────────────────────────
  
  - questions MUST match Total_Questions EXACTLY
  
  IF the document contains the token "ANSWER_OPTIONS:":
  - This is an MCQ-based comparative exam
  - EVERY question MUST contain:
    - question_text
    - answer_options (object with EXACT keys A, B, C, D)
    - correct_answer (A|B|C|D)
  - correct_answer MUST be one of the answer_options keys
  - answer_options text MUST match the document EXACTLY
  
  IF the document does NOT contain "ANSWER_OPTIONS:":
  - This is an extract-selection comparative exam
  - EVERY question MUST contain:
    - question_text
    - correct_answer
  - questions MUST NOT contain answer_options
  - correct_answer MUST match one of the extract labels present
  
  ────────────────────────────────────────
  FAIL CONDITIONS
  ────────────────────────────────────────
  - DO NOT infer, generate, repair, or guess missing fields
  - If ANY rule above is violated, RETURN {}
  - If ANY required field is missing or empty, RETURN {}
  
  ────────────────────────────────────────
  CRITICAL RULES
  ────────────────────────────────────────
  - DO NOT generate, infer, rewrite, summarize, or explain content
  - Extract ONLY what exists in the document
  - Preserve wording, punctuation, capitalization, and ordering EXACTLY
  
  ────────────────────────────────────────
  OUTPUT RULES
  ────────────────────────────────────────
  - VALID JSON ONLY
  - No markdown
  - No explanations
  """





    # --------------------------------------------------
    # 5️⃣ Process each exam block
    # --------------------------------------------------
    for block_idx, block_text in enumerate(blocks, start=1):
         
        # --------------------------------------------------
        # 🔢 Extract Total_Questions from document (STRICT)
        # --------------------------------------------------
        match = re.search(
            r"^\s*Total_Questions\s*:\s*(\d+)\s*$",
            block_text,
            re.MULTILINE
        )
        
        if not match:
            print("❌ Total_Questions missing in document")
            continue
        
        expected_q_count = int(match.group(1))
        
        # 🔒 Allow ONLY 8 or 10
        ALLOWED_QUESTION_COUNTS = {8, 10}

        if expected_q_count not in ALLOWED_QUESTION_COUNTS:
            print(
                f"❌ Invalid Total_Questions value: {expected_q_count}. "
                f"Allowed values: {sorted(ALLOWED_QUESTION_COUNTS)}"
            )
            continue
         
        print(f"✅ Total_Questions detected: {expected_q_count}")        

        print("\n" + "-" * 70)
        print(f"🔍 STEP 5: Processing block {block_idx}/{len(blocks)}")
        print("   → Block length:", len(block_text))
        is_mcq_comparative = "ANSWER_OPTIONS:" in block_text


        # --------------------------------------------------
        # 🔠 Detect extract labels from document (AUTHORITATIVE)
        # --------------------------------------------------
        doc_extract_matches = re.findall(
            r"^\s*Extract\s+([A-Z])\s*$",
            block_text,
            re.MULTILINE
        )
        
        doc_extract_keys = sorted(dict.fromkeys(doc_extract_matches))
        
        if len(doc_extract_keys) < 2:
            print("❌ Less than 2 extracts declared in document")
            continue
        
        print(f"✅ Extracts declared in document: {doc_extract_keys}")


        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                temperature=0,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {
                        "role": "user",
                        "content": f"BEGIN_DOCUMENT\n{block_text}\nEND_DOCUMENT"
                    }
                ]
            )
        except Exception as e:
            print("❌ OpenAI API CALL FAILED")
            print("   → Exception:", str(e))
            continue

        raw_output = response.choices[0].message.content.strip()
        print("🧠 AI response length:", len(raw_output))

        if not raw_output.startswith("{"):
            print("❌ NON-JSON AI OUTPUT")
            print("   → Preview:", raw_output[:200])
            continue

        try:
            parsed = json.loads(raw_output)
            print("✅ JSON parsed successfully")
        except Exception as e:
            print("❌ JSON PARSE FAILED")
            print("   → Exception:", str(e))
            continue

        if not parsed:
            print("⚠️ EMPTY JSON returned by AI")
            continue

        # --------------------------------------------------
        # 6️⃣ Hard validation
        # --------------------------------------------------
        print("\n🧪 STEP 6: Validating extracted content")

        rm = parsed.get("reading_material", {})
        extracts = rm.get("extracts", {})
        gpt_extract_keys = sorted(extracts.keys())
        # 🔒 GPT must extract ALL document-declared extracts
        if gpt_extract_keys != doc_extract_keys:
            print(
                f"❌ Extract mismatch. "
                f"Document: {doc_extract_keys}, "
                f"GPT returned: {gpt_extract_keys}"
            )
            continue
        extract_keys = doc_extract_keys  # authoritative source
        # 🔹 Ensure extract text is not empty
        empty_extracts = [
            k for k in extract_keys
            if not extracts.get(k) or not extracts[k].strip()
        ]
        
        if empty_extracts:
            print("❌ Empty extracts found:", empty_extracts)
            continue    
        
        
        
        # --------------------------------------------------
        # 🔹 Question validation (UPDATED)
        # --------------------------------------------------
        questions = parsed.get("questions", [])

        has_explicit_options = any("answer_options" in q for q in questions)

        # 🔒 Enforce question count matches metadata
        if len(questions) != expected_q_count:
            print(
                f"❌ Question count mismatch: expected {expected_q_count}, got {len(questions)}"
            )
            continue

        
        print("   → Questions found:", len(questions))
        
        if not questions:
            print("❌ No questions extracted")
            continue
        
        # NEW — supports extract-based AND MCQ-based comparative
        invalid_q = False

        for i, q in enumerate(questions, start=1):
            missing_keys = {"question_text", "correct_answer"} - q.keys()
            if missing_keys:
                print(f"❌ Question {i} missing keys: {missing_keys}")
                invalid_q = True
                break
        
            if is_mcq_comparative:
               # MCQ-based comparative MUST have answer_options
               opts = q.get("answer_options")
           
               if not opts:
                   print(
                       f"❌ Question {i} missing answer_options in MCQ comparative"
                   )
                   invalid_q = True
                   break
           
               if not isinstance(opts, dict) or set(opts.keys()) != {"A", "B", "C", "D"}:
                   print(
                       f"❌ Question {i} has invalid answer_options keys: {opts}"
                   )
                   invalid_q = True
                   break
           
               if q["correct_answer"] not in opts:
                   print(
                       f"❌ Question {i} correct_answer not in answer_options"
                   )
                   invalid_q = True
                   break

            else:
                # Extract-selection comparative
                if q["correct_answer"] not in extract_keys:
                    print(
                        f"❌ Question {i} has invalid extract correct_answer: "
                        f"{q['correct_answer']} (allowed: {extract_keys})"
                    )
                    invalid_q = True
                    break
        
        if invalid_q:
            continue


        # --------------------------------------------------
        # 7️⃣ Enrich bundle (RENDER-SAFE)
        # --------------------------------------------------
        print("\n🧩 STEP 7: Enriching bundle")

        if not is_mcq_comparative:
            # Extract-selection comparative ONLY
            for q in questions:
                if "answer_options" not in q:
                    q["answer_options"] = {
                        key: f"Extract {key}" for key in extract_keys
                    }
             
        for i, q in enumerate(questions, start=1):
            q["question_id"] = f"CA_Q{i}"

        bundle = {
            "question_type": "comparative_analysis",
            "topic": parsed["topic"],
            "reading_material": rm,
            "questions": questions
        }
        
        # Promote answer_options for frontend if MCQ
        if is_mcq_comparative:
            bundle["answer_options"] = questions[0]["answer_options"]



        # --------------------------------------------------
        # 8️⃣ Save to DB
        # --------------------------------------------------
        print("\n💾 STEP 8: Saving exam to database")
        subject = "reading_comprehension"



        try:
            obj = QuestionReading(
                class_name=parsed["class_name"].lower(),
                subject=subject,
                difficulty=parsed["difficulty"].lower(),
                topic=parsed["topic"],
                total_questions=len(questions),
                exam_bundle=bundle
            )

            db.add(obj)
            db.commit()
            db.refresh(obj)

            saved_ids.append(obj.id)
            print(f"✅ Block {block_idx} saved | ID={obj.id}")

        except Exception as e:
            db.rollback()
            print("❌ DATABASE SAVE FAILED")
            print("   → Exception:", str(e))
            continue

    # --------------------------------------------------
    # 9️⃣ Final response
    # --------------------------------------------------
    print("\n" + "=" * 70)
    print("🎉 UPLOAD COMPLETE")
    print("   → Total saved exams:", len(saved_ids))
    print("   → IDs:", saved_ids)
    print("=" * 70)
    total_detected = len(blocks)
    total_saved = len(saved_ids)
    total_failed = total_detected - total_saved
    
    status = "success"
    if total_saved == 0:
        status = "failed"
    elif total_failed > 0:
        status = "partial_success"

    return {
       "status": status,
       "message": "Upload complete",
       "summary": {
           "total_exam_blocks_detected": total_detected,
           "total_exams_saved": total_saved,
           "total_exams_failed": total_failed
       },
       "saved_exam_ids": saved_ids
   }

#here line 8452
@app.post("/api/admin/create-reading-config")
def create_reading_config(payload: ReadingExamConfigCreate, db: Session = Depends(get_db)):

    # Validate topics
    if len(payload.topics) == 0:
        raise HTTPException(status_code=400, detail="At least one topic must be provided.")

    num_topics = len(payload.topics)

    # Build JSON structure WITHOUT topic_id
    topics_json = [
        {
            "name": t.name,
            "num_questions": t.num_questions
        }
        for t in payload.topics
    ]
    print("\n--- Deleting dependent Reading generated exams ---")

    db.query(GeneratedExamReading).delete(synchronize_session=False)
    print("\n--- Deleting previous Reading exam configs ---")

    db.query(ReadingExamConfig).delete(synchronize_session=False)
    db.commit()

    print("🗑️ Previous reading configs deleted")

    print("\n--- Creating new Reading exam config ---")

    # Create DB record
    new_config = ReadingExamConfig(
        class_name=payload.class_name,
        subject=payload.subject,
        difficulty=payload.difficulty,
        num_topics=num_topics,
        topics=topics_json
    )

    db.add(new_config)
    db.commit()
    db.refresh(new_config)

    return {
        "status": "success",
        "config_id": new_config.id,
        "message": "Reading exam configuration saved successfully.",
        "created_at": new_config.created_at
    }


@app.post("/api/student/start-exam")
def start_exam(
    req: StartExamRequest = Body(...),
    db: Session = Depends(get_db)
):
    print("\n🚀 START-EXAM REQUEST")
    print("➡ payload:", req.dict())

    # --------------------------------------------------
    # 1️⃣ Resolve student (external → internal)
    # --------------------------------------------------
    student = (
        db.query(Student)
        .filter(
            func.lower(Student.student_id) ==
            func.lower(req.student_id.strip())
        )
        .first()
    )
    
    if not student:
        print(f"❌ Student not found (raw={repr(req.student_id)})")
        raise HTTPException(status_code=404, detail="Student not found")


    print(f"✅ Student resolved: id={student.id}")

    
    # --------------------------------------------------
    # 🟢 CASE A — COMPLETED attempt exists → SHOW REPORT
    # --------------------------------------------------
    

    # --------------------------------------------------
    # 3️⃣ Load quiz (class-based)
    # --------------------------------------------------
    quiz = (
        db.query(Quiz)
        .filter(func.lower(Quiz.class_name) == func.lower(student.class_name))
        .order_by(Quiz.id.desc())
        .first()
    )

    if not quiz:
        print("❌ Quiz not found for class")
        raise HTTPException(status_code=404, detail="Quiz not found")

    # --------------------------------------------------
    # 4️⃣ Load exam
    # --------------------------------------------------
    exam = (
        db.query(Exam)
        .filter(
            func.lower(Exam.class_name) == func.lower(student.class_name),
            Exam.subject == "mathematical_reasoning"
        )
        .order_by(Exam.created_at.desc())
        .first()
    )


    if not exam:
        print("❌ Exam not generated")
        raise HTTPException(status_code=404, detail="Exam not generated")
    def normalize_questions(raw_questions):
        normalized = []
    
        for q in raw_questions or []:
            fixed = dict(q)
    
            # -----------------------------
            # ✅ NORMALIZE QUESTION CONTENT
            # -----------------------------
            question_text = (
                fixed.get("question")
                or fixed.get("question_text")
                or fixed.get("text")
                or ""
            )
    
            fixed["blocks"] = [
                {
                    "type": "text",
                    "content": question_text
                }
            ]
    
            # -----------------------------
            # ✅ NORMALIZE OPTIONS
            # -----------------------------
            opts = fixed.get("options")
    
            if isinstance(opts, dict):
                fixed["options"] = {
                    k: {"content": v} for k, v in opts.items()
                }
            elif isinstance(opts, list):
                fixed["options"] = {
                    chr(65 + i): {"content": v}
                    for i, v in enumerate(opts)
                }
            else:
                fixed["options"] = {}
    
            normalized.append(fixed)
    
        return normalized
    # --------------------------------------------------
    # 🧮 MATHEMATICAL REASONING ATTEMPT HANDLING
    # --------------------------------------------------
    print("🔍 Checking Mathematical Reasoning attempts...")

    math_attempt = (
        db.query(StudentExamMathematicalReasoning)
        .filter(
            StudentExamMathematicalReasoning.student_id == student.id
        )
        .order_by(StudentExamMathematicalReasoning.started_at.desc())
        .first()
    )
    
    if math_attempt:
        print("📘 Latest math attempt found")
        print("   ├─ attempt.id:", math_attempt.id)
        print("   ├─ started_at:", math_attempt.started_at)
        print("   └─ completed_at:", math_attempt.completed_at)
    else:
        print("📘 No previous math attempts found")
    # --------------------------------------------------
    # 🟢 CASE A — COMPLETED math attempt exists → SHOW REPORT
    # --------------------------------------------------
    if math_attempt and math_attempt.completed_at is not None:
        print("✅ Completed math attempt exists → returning completed=true")
        return {"completed": True}
 
    # --------------------------------------------------
    # 🟡 CASE B — ACTIVE math attempt → RESUME
    # --------------------------------------------------
    if math_attempt and math_attempt.completed_at is None:
        print("⏳ Active math attempt → resuming")
    
        # 🔒 DEFENSIVE: ensure response rows exist
        existing_count = (
            db.query(StudentExamResponseMathematicalReasoning)
            .filter(
                StudentExamResponseMathematicalReasoning.exam_attempt_id == math_attempt.id
            )
            .count()
        )
    
        if existing_count == 0:
            print("⚠️ Active attempt has no response rows — rehydrating")
            for q in exam.questions or []:
                db.add(
                    StudentExamResponseMathematicalReasoning(
                        student_id=student.id,
                        exam_id=exam.id,
                        exam_attempt_id=math_attempt.id,
                        q_id=q["q_id"],
                        topic=q.get("topic"),
                        selected_option=None,
                        correct_option=q.get("correct"),
                        is_correct=None
                    )
                )
            db.commit()
    
        started_at = math_attempt.started_at
        if started_at.tzinfo is None:
            started_at = started_at.replace(tzinfo=timezone.utc)
    
        now = datetime.now(timezone.utc)
        elapsed = int((now - started_at).total_seconds())
        remaining = max(0, 40 * 60 - elapsed)
    
        print(f"⏱ elapsed={elapsed}s, remaining={remaining}s")
    
        if remaining == 0:
            print("⛔ Time expired → marking completed")
            math_attempt.completed_at = now
            db.commit()
            return {"completed": True}
    
        return {
            "completed": False,
            "questions": normalize_questions(exam.questions),
            "remaining_time": remaining
        }
    
    # --------------------------------------------------
    # 🔵 CASE C — NO active attempt → START NEW
    # --------------------------------------------------
    print("🆕 No active math attempt → starting new exam")
    
    new_attempt = StudentExamMathematicalReasoning(
        student_id=student.id,
        exam_id=exam.id,
        started_at=datetime.now(timezone.utc)
    )
    
    db.add(new_attempt)
    db.commit()
    
    print("✅ New math attempt created")
    print("   ├─ attempt.id:", new_attempt.id)
    print("   └─ started_at:", new_attempt.started_at)
    
    # --------------------------------------------------
    # 🧱 PRE-CREATE RESPONSE ROWS
    # --------------------------------------------------
    print("🧱 Pre-creating response rows")
    
    for q in exam.questions or []:
        db.add(
            StudentExamResponseMathematicalReasoning(
                student_id=student.id,
                exam_id=exam.id,
                exam_attempt_id=new_attempt.id,
                q_id=q["q_id"],              # ✅ CORRECT
                topic=q.get("topic"),
                selected_option=None,
                correct_option=q.get("correct"),
                is_correct=None
            )
        )

    
    db.commit()
    print(
       "🧱 Response rows created:",
       db.query(StudentExamResponseMathematicalReasoning)
         .filter(
             StudentExamResponseMathematicalReasoning.exam_attempt_id == new_attempt.id
         )
         .count()
   )
    
    return {
        "completed": False,
        "questions": normalize_questions(exam.questions),
        "remaining_time": 40 * 60
    }
 
def normalize_naplan_numeracy_questions_live(raw_questions):
    normalized = []

    for q in raw_questions or []:

        blocks = q.get("question_blocks") or []
        correct_answer = str(q.get("correct_answer")).strip()
        display_blocks = []

        for block in blocks:
            content = (block.get("content") or "").strip()
            lowered = content.lower()

            if lowered.startswith("question_type"):
                continue
            if lowered.startswith("year:"):
                continue
            if lowered.startswith("answer_type"):
                continue
            if lowered.startswith("correct_answer"):
                continue

            if content == correct_answer:
                continue

            display_blocks.append(block)

        normalized.append({
            "id": q["id"],  # ← FIXED
            "question_type": q.get("question_type"),  # ← FIXED
            "topic": q.get("topic"),
            "difficulty": q.get("difficulty"),
            "question_blocks": display_blocks,  # ← FIXED
            "options": q.get("options")
        })

    return normalized

   
def normalize_correct_option_for_db(correct):
    if correct is None:
        return None

    # { value: "B" }
    if isinstance(correct, dict) and "value" in correct:
        return str(correct["value"])

    # ['A', 'C']
    if isinstance(correct, list):
        return ",".join(map(str, correct))

    # "['A','C']"
    if isinstance(correct, str):
        return correct

    return str(correct)
def serialize_type1_question_for_exam(q):
    """
    Build exam-safe payload for Type 1 (MCQ single).
    Strips authoring junk and NEVER sends correct_answer.
    """
    return {
        "id": q["id"],
        "question_type": q["question_type"],
        "topic": q.get("topic"),
        "difficulty": q.get("difficulty"),

        # 🔑 Use ALREADY normalized blocks (do NOT rebuild)
        "question_blocks": normalize_question_blocks_backend(
            q.get("question_blocks")
        ),

        # MCQ options are needed
        "options": q.get("options"),
    }


def normalize_images_in_question(question: dict, image_map: dict):
    # Normalize images in options
    options = question.get("options")
    if options:
        for key, value in options.items():
            if value in image_map:
                options[key] = image_map[value]

    # Normalize images inside question_blocks
    blocks = question.get("question_blocks") or []
    for block in blocks:
        images = block.get("images")
        if not images:
            continue

        normalized_images = []
        for img in images:
            normalized_images.append(image_map.get(img, img))

        block["images"] = normalized_images
def normalize_type2_image_multiselect(question: dict):
    if question.get("question_type") != 2:
        return

    blocks = question.get("question_blocks") or []
    options = question.get("options") or {}

    image_block = None
    text_blocks = []

    for block in blocks:
        if block.get("type") == "image-selection":
            image_block = block
        else:
            text_blocks.append(block)

    if not image_block:
        return

    images = image_block.get("images") or []
    option_items = list(options.items())

    if len(images) != len(option_items):
        raise ValueError("TYPE2_IMAGE_OPTION_COUNT_MISMATCH")

    visual_options = []
    for (opt_id, label), image_url in zip(option_items, images):
        visual_options.append({
            "id": opt_id,
            "label": label,
            "image": image_url
        })

    # ✅ PRESERVE text + replace image block only
    question["question_blocks"] = [
        *text_blocks,
        {
            "type": "image-multi-select",
            "maxSelections": image_block.get("maxSelections", 1),
            "options": visual_options
        }
    ]

    # Remove legacy options
    question.pop("options", None)

import ast

def normalize_type2_correct_answer(question: dict):
    if question.get("question_type") != 2:
        return

    raw = question.get("correct_answer")

    if isinstance(raw, dict):
        return

    if isinstance(raw, str):
        question["correct_answer"] = {
            "value": ast.literal_eval(raw)
        }

@app.post("/api/student/start-exam/naplan-numeracy")
def start_naplan_numeracy_exam(
    req: StartExamRequest = Body(...),
    db: Session = Depends(get_db)
):
    print("\n================ START NAPLAN NUMERACY EXAM =================")
    print("📥 Incoming payload:", req.dict())

    student = (
        db.query(Student)
        .filter(
            func.lower(Student.student_id) ==
            func.lower(req.student_id.strip())
        )
        .first()
    )

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    attempt = (
        db.query(StudentExamNaplanNumeracy)
        .filter(StudentExamNaplanNumeracy.student_id == student.id)
        .order_by(StudentExamNaplanNumeracy.started_at.desc())
        .first()
    )

    MAX_DURATION = timedelta(minutes=40)
    now = datetime.now(timezone.utc)

    uploaded_images = db.query(UploadedImage).all()
    image_map = {img.original_name: img.gcs_url for img in uploaded_images}

    if attempt:
        started_at = attempt.started_at
        if started_at.tzinfo is None:
            started_at = started_at.replace(tzinfo=timezone.utc)

        expires_at = started_at + MAX_DURATION
        elapsed = int((now - started_at).total_seconds())

        if attempt.completed_at is None and now > expires_at:
            attempt.completed_at = expires_at
            db.commit()
            return {"completed": True}

        if attempt.completed_at is not None:
            return {"completed": True}

        remaining = max(0, attempt.duration_minutes * 60 - elapsed)

        exam = (
            db.query(ExamNaplanNumeracy)
            .filter(
                func.lower(ExamNaplanNumeracy.class_name) ==
                func.lower(student.class_name),
                func.lower(ExamNaplanNumeracy.subject) == "numeracy"
            )
            .order_by(ExamNaplanNumeracy.created_at.desc())
            .first()
        )

        raw_questions = normalize_naplan_numeracy_questions_live(exam.questions or [])

        normalized_questions = []
        for q in raw_questions:
            if q.get("question_type") == 1:
                q = serialize_type1_question_for_exam(q)

            normalize_images_in_question(q, image_map)
            normalize_type2_image_multiselect(q)
            normalize_type2_correct_answer(q)
            normalized_questions.append(q)

        return {
            "completed": False,
            "questions": normalized_questions,
            "remaining_time": remaining
        }

    # --------------------------------------------------
    # 🆕 FIRST ATTEMPT
    # --------------------------------------------------
    exam = (
        db.query(ExamNaplanNumeracy)
        .filter(
            func.lower(ExamNaplanNumeracy.class_name) ==
            func.lower(student.class_name),
            func.lower(ExamNaplanNumeracy.subject) == "numeracy"
        )
        .order_by(ExamNaplanNumeracy.created_at.desc())
        .first()
    )

    raw_questions = normalize_naplan_numeracy_questions_live(exam.questions or [])

    normalized_questions = []
    for q in raw_questions:
        if q.get("question_type") == 1:
            q = serialize_type1_question_for_exam(q)

        normalize_images_in_question(q, image_map)
        normalize_type2_image_multiselect(q)
        normalize_type2_correct_answer(q)

        normalized_questions.append(q)

    new_attempt = StudentExamNaplanNumeracy(
        student_id=student.id,
        exam_id=exam.id,
        started_at=now,
        duration_minutes=40
    )

    db.add(new_attempt)
    db.commit()
    db.refresh(new_attempt)

    for original_q in exam.questions or []:
        normalized_correct_option = normalize_correct_option_for_db(
            original_q.get("correct_answer")
        )
        db.add(
            StudentExamResponseNaplanNumeracy(
                student_id=student.id,
                exam_id=exam.id,
                exam_attempt_id=new_attempt.id,
                q_id=original_q["id"],
                topic=original_q.get("topic"),
                selected_option=None,
                correct_option=normalized_correct_option,
                is_correct=None
            )
        )

    db.commit()

    return {
        "completed": False,
        "questions": normalized_questions,
        "remaining_time": new_attempt.duration_minutes * 60
    }
  
# @app.post("/api/student/start-exam-thinkingskills")
# def start_exam(
#     req: StartExamRequest = Body(...),
#     db: Session = Depends(get_db)
# ):
#     print("\n🚀 START-EXAM REQUEST")
#     print("➡ payload:", req.dict())
#
#     # --------------------------------------------------
#     # 1️⃣ Resolve student
#     # --------------------------------------------------
#     student = (
#         db.query(Student)
#         .filter(
#             func.lower(Student.student_id) ==
#             func.lower(req.student_id.strip())
#         )
#         .first()
#     )
#
#     if not student:
#         raise HTTPException(status_code=404, detail="Student not found")
#
#     print(f"✅ Student resolved: id={student.id}")
#
#     # --------------------------------------------------
#     # 2️⃣ Get latest THINKING SKILLS attempt
#     # --------------------------------------------------
#     attempt = (
#         db.query(StudentExamThinkingSkills)
#         .filter(StudentExamThinkingSkills.student_id == student.id)
#         .order_by(StudentExamThinkingSkills.started_at.desc())
#         .first()
#     )
#
#     if attempt and attempt.completed_at:
#         print("✅ Exam already completed")
#         return {"completed": True}
#
#     # --------------------------------------------------
#     # 3️⃣ Load latest THINKING SKILLS exam
#     # --------------------------------------------------
#     exam = (
#         db.query(Exam)
#         .filter(
#             func.lower(Exam.class_name) ==
#             func.lower(student.class_name),
#             Exam.subject == "thinking_skills"
#         )
#         .order_by(Exam.created_at.desc())
#         .first()
#     )
#
#     if not exam:
#         raise HTTPException(status_code=404, detail="Thinking Skills exam not found")
#
#     # --------------------------------------------------
#     # 🔧 Normalize questions for frontend
#     # --------------------------------------------------
#     def normalize_questions(raw_questions):
#         normalized = []
#
#         for q in raw_questions or []:
#             fixed = dict(q)
#
#             qid = fixed.get("q_id")
#
#             # ==================================================
#             # ✅ Normalize OPTIONS (ALL supported formats)
#             # ==================================================
#             opts = fixed.get("options")
#             normalized_opts = {}
#
#             if isinstance(opts, dict):
#                 normalized_opts = opts
#
#             elif isinstance(opts, list):
#                 if opts and isinstance(opts[0], str) and ")" in opts[0]:
#                     for raw in opts:
#                         key, payload = raw.split(")", 1)
#                         key = key.strip()
#
#                         try:
#                             option_obj = ast.literal_eval(payload.strip())
#                             if isinstance(option_obj, dict):
#                                 normalized_opts[key] = option_obj
#                         except Exception:
#                             normalized_opts[key] = {
#                                 "type": "text",
#                                 "content": payload.strip()
#                             }
#                 else:
#                     for idx, text in enumerate(opts):
#                         key = chr(ord("A") + idx)
#                         normalized_opts[key] = {
#                             "type": "text",
#                             "content": text
#                         }
#
#             fixed["options"] = normalized_opts
#
#             if not normalized_opts:
#                 print(f"[WARN] No options parsed for question {qid}")
#
#             # ==================================================
#             # ✅ Normalize CORRECT ANSWER
#             # ==================================================
#             fixed["correct_answer"] = (
#                 fixed.get("correct_answer")
#                 or fixed.get("correct")
#             )
#
#             if not fixed["correct_answer"]:
#                 print(f"[WARN] No correct_answer for question {qid}")
#
#             # ==================================================
#             # 🔥 Normalize BLOCKS (text + images)
#             # ==================================================
#             blocks = fixed.get("blocks") or fixed.get("question_blocks") or []
#
#             for block in blocks:
#                 if block.get("type") != "image":
#                     continue
#
#                 src = block.get("src")
#                 if not src:
#                     print(f"[WARN] Image block missing src in question {qid}")
#                     continue
#
#                 if src.startswith("http"):
#                     continue
#
#                 img_norm = src.strip().lower()
#
#                 record = (
#                     db.query(UploadedImage)
#                     .filter(
#                         func.lower(UploadedImage.original_name) == img_norm
#                     )
#                     .first()
#                 )
#
#                 if not record:
#                     print(f"[WARN] Unresolved image '{src}' in question {qid}")
#                     continue
#
#                 block["src"] = record.gcs_url
#
#             fixed["blocks"] = blocks
#             normalized.append(fixed)
#
#         return normalized
#
#     raw_questions = exam.questions or []
#     normalized_questions = normalize_questions(raw_questions)
#
#     # --------------------------------------------------
#     # 🟡 Resume active attempt
#     # --------------------------------------------------
#     if attempt and attempt.completed_at is None:
#         print("⏳ Resuming active attempt")
#
#         started_at = attempt.started_at
#         if started_at.tzinfo is None:
#             started_at = started_at.replace(tzinfo=timezone.utc)
#
#         now = datetime.now(timezone.utc)
#         elapsed = int((now - started_at).total_seconds())
#         remaining = max(0, attempt.duration_minutes * 60 - elapsed)
#
#         if remaining == 0:
#             attempt.completed_at = now
#             db.commit()
#             return {"completed": True}
#
#         return {
#             "completed": False,
#             "questions": normalized_questions,
#             "remaining_time": remaining
#         }
#
#     # --------------------------------------------------
#     # 🔵 Start new attempt
#     # --------------------------------------------------
#     print("🆕 Starting new Thinking Skills attempt")
#
#     new_attempt = StudentExamThinkingSkills(
#         student_id=student.id,
#         exam_id=exam.id,
#         started_at=datetime.now(timezone.utc),
#         duration_minutes=40
#     )
#
#     db.add(new_attempt)
#     db.commit()
#     db.refresh(new_attempt)
#
#     for q in normalized_questions:
#         db.add(
#             StudentExamResponseThinkingSkills(
#                 student_id=student.id,
#                 exam_id=exam.id,
#                 exam_attempt_id=new_attempt.id,
#                 q_id=q["q_id"],
#                 topic=q.get("topic"),
#                 selected_option=None,
#                 correct_option=q["correct_answer"],
#                 is_correct=None
#             )
#         )
#
#     db.commit()
#
#     return {
#         "completed": False,
#         "questions": normalized_questions,
#         "remaining_time": new_attempt.duration_minutes * 60
#     }
    

@app.post("/api/student/start-exam-thinkingskills")
def start_exam(
    req: StartExamRequest = Body(...),
    db: Session = Depends(get_db)
):
    print("\n================ START THINKING SKILLS EXAM =================")
    print("📥 Incoming payload:", req.dict())

    # --------------------------------------------------
    # 1️⃣ Resolve student
    # --------------------------------------------------
    student = (
        db.query(Student)
        .filter(
            func.lower(Student.student_id) ==
            func.lower(req.student_id.strip())
        )
        .first()
    )

    if not student:
        print("❌ Student not found:", repr(req.student_id))
        raise HTTPException(status_code=404, detail="Student not found")

    print(f"✅ Student resolved | student_id={student.student_id} | internal_id={student.id}")

    # --------------------------------------------------
    # 2️⃣ Fetch the ONE (and only) Thinking Skills attempt
    # --------------------------------------------------
    attempt = (
        db.query(StudentExamThinkingSkills)
        .filter(StudentExamThinkingSkills.student_id == student.id)
        .order_by(StudentExamThinkingSkills.started_at.desc())
        .first()
    )

    MAX_DURATION = timedelta(minutes=40)
    now = datetime.now(timezone.utc)

    if attempt:
        print(
            "🧠 Existing attempt found | "
            f"attempt_id={attempt.id} | "
            f"started_at={attempt.started_at} | "
            f"completed_at={attempt.completed_at} | "
            f"duration_minutes={attempt.duration_minutes}"
        )

        started_at = attempt.started_at
        if started_at.tzinfo is None:
            started_at = started_at.replace(tzinfo=timezone.utc)

        expires_at = started_at + MAX_DURATION
        elapsed = int((now - started_at).total_seconds())

        print(
            "⏱️ Timer evaluation | "
            f"now={now} | "
            f"started_at={started_at} | "
            f"expires_at={expires_at} | "
            f"elapsed_seconds={elapsed}"
        )

        # --------------------------------------------------
        # ⛔ Timeout → FINAL completion
        # --------------------------------------------------
        if attempt.completed_at is None and now > expires_at:
            print(
                "⛔ Attempt expired | "
                f"attempt_id={attempt.id} | "
                f"auto_completed_at={expires_at}"
            )

            attempt.completed_at = expires_at
            db.commit()

            print("➡️ Returning: completed=true (timeout)")
            return {"completed": True}

        # --------------------------------------------------
        # ✅ Already completed → show report
        # --------------------------------------------------
        if attempt.completed_at is not None:
            print(
                "✅ Attempt already completed | "
                f"attempt_id={attempt.id} | "
                f"completed_at={attempt.completed_at}"
            )

            print("➡️ Returning: completed=true (already completed)")
            return {"completed": True}

        # --------------------------------------------------
        # ▶ Resume active attempt
        # --------------------------------------------------
        remaining = max(0, attempt.duration_minutes * 60 - elapsed)

        print(
            "▶ Resuming active attempt | "
            f"attempt_id={attempt.id} | "
            f"remaining_seconds={remaining}"
        )

        # Load exam only when resuming
        exam = (
            db.query(Exam)
            .filter(
                func.lower(Exam.class_name) ==
                func.lower(student.class_name),
                Exam.subject == "thinking_skills"
            )
            .order_by(Exam.created_at.desc())
            .first()
        )

        if not exam:
            print("❌ Exam not found while resuming")
            raise HTTPException(
                status_code=404,
                detail="Thinking Skills exam not found"
            )

        normalized_questions = normalize_thinking_skills_questions(
            exam.questions or [],
            db
        )

        print(
            "➡️ Returning: resume exam | "
            f"questions={len(normalized_questions)} | "
            f"remaining_seconds={remaining}"
        )

        return {
            "completed": False,
            "questions": normalized_questions,
            "remaining_time": remaining
        }

    # --------------------------------------------------
    # 🆕 FIRST AND ONLY ATTEMPT (no attempt exists)
    # --------------------------------------------------
    print("🆕 No existing attempt found → creating FIRST and ONLY attempt")

    exam = (
        db.query(Exam)
        .filter(
            func.lower(Exam.class_name) ==
            func.lower(student.class_name),
            Exam.subject == "thinking_skills"
        )
        .order_by(Exam.created_at.desc())
        .first()
    )

    if not exam:
        print("❌ Exam not found for first attempt")
        raise HTTPException(
            status_code=404,
            detail="Thinking Skills exam not found"
        )

    normalized_questions = normalize_thinking_skills_questions(
        exam.questions or [],
        db
    )

    new_attempt = StudentExamThinkingSkills(
        student_id=student.id,
        exam_id=exam.id,
        started_at=now,
        duration_minutes=40
    )

    db.add(new_attempt)
    db.commit()
    db.refresh(new_attempt)

    print(
        "🆕 New attempt created | "
        f"attempt_id={new_attempt.id} | "
        f"started_at={new_attempt.started_at}"
    )

    # --------------------------------------------------
    # Pre-create response rows (ONLY here)
    # --------------------------------------------------
    for q in normalized_questions:
        db.add(
            StudentExamResponseThinkingSkills(
                student_id=student.id,
                exam_id=exam.id,
                exam_attempt_id=new_attempt.id,
                q_id=q["q_id"],
                topic=q.get("topic"),
                selected_option=None,
                correct_option=q["correct_answer"],
                is_correct=None
            )
        )

    db.commit()

    print(
        "➡️ Returning: new exam started | "
        f"attempt_id={new_attempt.id} | "
        f"questions={len(normalized_questions)} | "
        f"remaining_seconds={new_attempt.duration_minutes * 60}"
    )

    print("================ END START THINKING SKILLS EXAM ================\n")

    return {
        "completed": False,
        "questions": normalized_questions,
        "remaining_time": new_attempt.duration_minutes * 60
    }





@app.get("/api/student/get-exam")
def get_exam(session_id: int, db: Session = Depends(get_db)):

    session = (
        db.query(StudentExam)
        .filter(StudentExam.id == session_id)
        .first()
    )

    # ❌ Invalid or deleted session → return safe payload
    if not session:
        return {
            "questions": [],
            "total_questions": 0,
            "remaining_time": 0,
            "completed": True
        }

    exam = (
        db.query(Exam)
        .filter(Exam.id == session.exam_id)
        .first()
    )

    if not exam:
        return {
            "questions": [],
            "total_questions": 0,
            "remaining_time": 0,
            "completed": True
        }

    # -------------------------------------
    # TIME CALCULATION (timezone-safe)
    # -------------------------------------
    now = datetime.now(timezone.utc)

    started_at = session.started_at
    if started_at.tzinfo is None:
        started_at = started_at.replace(tzinfo=timezone.utc)

    elapsed = (now - started_at).total_seconds()
    remaining = session.duration_minutes * 60 - elapsed

    if remaining <= 0:
        session.completed_at = now
        db.commit()
        remaining = 0

    # -------------------------------------
    # NORMALIZE QUESTIONS
    # -------------------------------------
    normalized_questions = []

    for q in exam.questions:
        fixed = dict(q)
        opts = fixed.get("options")

        if opts is None:
            fixed["options"] = []
        elif isinstance(opts, dict):
            fixed["options"] = [f"{k}) {v}" for k, v in opts.items()]
        elif not isinstance(opts, list):
            fixed["options"] = []

        normalized_questions.append(fixed)

    return {
        "questions": normalized_questions,
        "total_questions": len(normalized_questions),
        "remaining_time": int(max(remaining, 0)),
        "completed": session.completed_at is not None
    }
def admin_report_exists(db, exam_attempt_id, exam_type):
    report = (
        db.query(AdminExamReport)
        .filter(
            AdminExamReport.exam_attempt_id == exam_attempt_id,
            AdminExamReport.exam_type == exam_type
        )
        .first()
    )

    if report:
        print("🟡 admin_report_exists = TRUE", {
            "exam_attempt_id": exam_attempt_id,
            "exam_type": exam_type,
            "admin_report_id": report.id,
            "created_at": report.created_at
        })
        return True

    print("🟢 admin_report_exists = FALSE", {
        "exam_attempt_id": exam_attempt_id,
        "exam_type": exam_type
    })
    return False
 
def generate_admin_exam_report(
    db: Session,
    student: Student,
    exam_attempt: StudentExamThinkingSkills,
    subject: str,
    accuracy: float,
    correct: int,
    wrong: int,
    total_questions: int
):
    """
    Ensures a single immutable admin report snapshot per exam attempt.
    Safe to call multiple times.
    MUST NOT commit. Caller controls the transaction.
    """

    print("\n🧠 ================== GENERATE ADMIN EXAM REPORT START ==================")
    print("🧾 Input payload:", {
        "student_external_id": student.student_id,
        "student_internal_id": student.id,
        "exam_attempt_id": exam_attempt.id,
        "exam_subject": subject,
        "accuracy": accuracy,
        "correct": correct,
        "wrong": wrong,
        "total_questions": total_questions
    })

    # --------------------------------------------------
    # 0️⃣ Idempotency guard (CRITICAL)
    # --------------------------------------------------
    print("🔎 Checking for existing AdminExamReport (idempotency guard)")

    existing_report = (
        db.query(AdminExamReport)
        .filter(
            AdminExamReport.exam_attempt_id == exam_attempt.id,
            AdminExamReport.exam_type == subject
        )
        .first()
    )

    if existing_report:
        print("⚠️ IDPOTENCY HIT: Admin report already exists", {
            "admin_report_id": existing_report.id,
            "exam_attempt_id": exam_attempt.id,
            "exam_type": existing_report.exam_type,
            "created_at": existing_report.created_at
        })
        print("🧠 ================== GENERATE ADMIN EXAM REPORT END ==================\n")
        return existing_report

    print("✅ No existing admin report found — proceeding with creation")

    # --------------------------------------------------
    # 1️⃣ Determine performance & readiness bands
    # --------------------------------------------------
    print("📊 Determining performance and readiness bands")

    if accuracy >= 80:
        performance_band = "A"
        readiness_band = "Strong Selective Potential"
        school_guidance = "Mid-tier Selective"
    elif accuracy >= 60:
        performance_band = "B"
        readiness_band = "Borderline Selective"
        school_guidance = "Selective Preparation Required"
    else:
        performance_band = "C"
        readiness_band = "Not Yet Selective Ready"
        school_guidance = "General Stream Recommended"

    print("🎯 Bands resolved:", {
        "performance_band": performance_band,
        "readiness_band": readiness_band,
        "school_guidance": school_guidance
    })

    # --------------------------------------------------
    # 2️⃣ Create main admin report
    # --------------------------------------------------
    print("📝 Creating AdminExamReport row")

    admin_report = AdminExamReport(
        student_id=student.student_id,
        exam_attempt_id=exam_attempt.id,
        exam_type=subject,
        overall_score=accuracy,
        readiness_band=readiness_band,
        school_guidance_level=school_guidance,
        summary_notes=f"Thinking Skills accuracy: {accuracy}%"
    )

    db.add(admin_report)

    print("💾 AdminExamReport added to session — flushing to obtain ID")
    db.flush()  # critical point

    print("✅ AdminExamReport flushed successfully", {
        "admin_report_id": admin_report.id,
        "exam_attempt_id": exam_attempt.id,
        "exam_type": admin_report.exam_type
    })

    # --------------------------------------------------
    # 3️⃣ Create section result
    # --------------------------------------------------
    print("📐 Creating AdminExamSectionResult")

    section_result = AdminExamSectionResult(
        admin_report_id=admin_report.id,
        section_name="thinking",
        raw_score=accuracy,
        performance_band=performance_band,
        strengths_summary="Logical reasoning applied correctly",
        improvement_summary="Improve speed and multi-step reasoning"
    )

    db.add(section_result)

    print("✅ Section result staged", {
        "admin_report_id": admin_report.id,
        "section_name": "thinking",
        "raw_score": accuracy,
        "performance_band": performance_band
    })

    # --------------------------------------------------
    # 4️⃣ Store applied readiness rule
    # --------------------------------------------------
    print("📜 Recording applied readiness rule")

    rule_result = "passed" if accuracy >= 60 else "failed"

    rule_applied = AdminReadinessRuleApplied(
        admin_report_id=admin_report.id,
        rule_code="THINKING_SKILLS_ACCURACY",
        rule_description="Thinking Skills accuracy used for readiness classification",
        rule_result=rule_result
    )

    db.add(rule_applied)

    print("✅ Readiness rule staged", {
        "admin_report_id": admin_report.id,
        "rule_code": "THINKING_SKILLS_ACCURACY",
        "rule_result": rule_result
    })

    print("🧠 ================== GENERATE ADMIN EXAM REPORT END ==================\n")

    return admin_report
def resolve_option_value(raw_opt, db):
    """
    Normalize a single option value into:
    - {type: image, src: full_url}
    - {type: text, content: ...}
    """

    # Already normalized
    if isinstance(raw_opt, dict):
        return raw_opt

    # Image filename or URL
    if isinstance(raw_opt, str) and raw_opt.lower().endswith(
        (".png", ".jpg", ".jpeg", ".webp")
    ):
        # full URL already
        if raw_opt.startswith("http"):
            return {
                "type": "image",
                "src": raw_opt,
                "name": raw_opt.split("/")[-1]
            }

        # filename only → resolve via DB
        record = (
            db.query(UploadedImage)
            .filter(
                func.lower(UploadedImage.original_name)
                == raw_opt.strip().lower()
            )
            .first()
        )

        if record:
            return {
                "type": "image",
                "src": record.gcs_url,
                "name": raw_opt
            }

        # fallback (still image-shaped)
        return {
            "type": "image",
            "src": raw_opt
        }

    # Plain text
    return {
        "type": "text",
        "content": raw_opt
    }

 
def normalize_thinking_skills_questions(raw_questions, db):
    normalized = []

    for q in raw_questions or []:
        fixed = dict(q)
        qid = fixed.get("q_id")

        # -----------------------------
        # Normalize OPTIONS
        # -----------------------------
        opts = fixed.get("options")
        normalized_opts = {}

        if isinstance(opts, dict):
           for key, raw_opt in opts.items():
               normalized_opts[key] = resolve_option_value(raw_opt, db)

        elif isinstance(opts, list):
            if opts and isinstance(opts[0], str) and ")" in opts[0]:
                for raw in opts:
                    key, payload = raw.split(")", 1)
                    key = key.strip()

                    try:
                        option_obj = ast.literal_eval(payload.strip())
                        if isinstance(option_obj, dict):
                            normalized_opts[key] = resolve_option_value(option_obj, db)

                    except Exception:
                        normalized_opts[key] = {
                            "type": "text",
                            "content": payload.strip()
                        }
            else:
                for idx, text in enumerate(opts):
                    key = chr(ord("A") + idx)
                    normalized_opts[key] = {
                        "type": "text",
                        "content": text
                    }

        fixed["options"] = normalized_opts

        # -----------------------------
        # Normalize CORRECT ANSWER
        # -----------------------------
        fixed["correct_answer"] = (
            fixed.get("correct_answer") or fixed.get("correct")
        )

        # -----------------------------
        # Normalize BLOCKS (images)
        # -----------------------------
        blocks = fixed.get("blocks") or fixed.get("question_blocks") or []

        for block in blocks:
            if block.get("type") != "image":
                continue

            src = block.get("src")
            if not src or src.startswith("http"):
                continue

            record = (
                db.query(UploadedImage)
                .filter(func.lower(UploadedImage.original_name) == src.strip().lower())
                .first()
            )

            if record:
                block["src"] = record.gcs_url

        fixed["blocks"] = blocks
        normalized.append(fixed)

    return normalized

@app.post("/api/student/finish-exam/thinking-skills")
def finish_thinking_skills_exam(
    req: FinishExamRequest,
    db: Session = Depends(get_db)
):
    print("\n================ FINISH THINKING SKILLS EXAM START ================")
    print("📥 Incoming payload:", req.dict())

    # --------------------------------------------------
    # 1️⃣ Validate student
    # --------------------------------------------------
    student = (
        db.query(Student)
        .filter(Student.student_id == req.student_id)
        .first()
    )

    if not student:
        print("❌ ABORT: Student not found for student_id =", req.student_id)
        raise HTTPException(status_code=404, detail="Student not found")

    print("✅ Student resolved:", {
        "external_id": student.student_id,
        "internal_id": student.id
    })

    # --------------------------------------------------
    # 2️⃣ Fetch MOST RECENT Thinking Skills attempt
    #     (CRITICAL: do NOT filter by completed_at)
    # --------------------------------------------------
    attempt = (
        db.query(StudentExamThinkingSkills)
        .filter(StudentExamThinkingSkills.student_id == student.id)
        .order_by(StudentExamThinkingSkills.started_at.desc())
        .first()
    )

    if not attempt:
        print("❌ ABORT: No Thinking Skills attempt exists for student")
        raise HTTPException(
            status_code=404,
            detail="No Thinking Skills exam attempt found"
        )

    print("✅ Latest attempt found:", {
        "attempt_id": attempt.id,
        "exam_id": attempt.exam_id,
        "started_at": attempt.started_at,
        "completed_at": attempt.completed_at
    })

    # --------------------------------------------------
    # 3️⃣ Load exam definition
    # --------------------------------------------------
    exam = (
        db.query(Exam)
        .filter(
            Exam.id == attempt.exam_id,
            Exam.subject == "thinking_skills"
        )
        .first()
    )

    if not exam:
        print("❌ ABORT: Exam not found or subject mismatch", {
            "exam_id": attempt.exam_id,
            "expected_subject": "thinking_skills"
        })
        raise HTTPException(status_code=404, detail="Exam not found")

    questions = exam.questions or []
    total_questions = len(questions)
    question_map = {q["q_id"]: q for q in questions}

    print("📊 Exam loaded:", {
        "total_questions": total_questions
    })

    # --------------------------------------------------
    # 4️⃣ Update responses (idempotent)
    # --------------------------------------------------
    correct = 0
    saved_responses = 0

    for q_id_str, selected in req.answers.items():
        try:
            q_id = int(q_id_str)
        except ValueError:
            print("⚠️ Skipping invalid q_id:", q_id_str)
            continue

        q = question_map.get(q_id)
        if not q:
            print("⚠️ Question not found in exam JSON:", q_id)
            continue

        is_correct = selected == q.get("correct")
        if is_correct:
            correct += 1

        response = (
            db.query(StudentExamResponseThinkingSkills)
            .filter(
                StudentExamResponseThinkingSkills.exam_attempt_id == attempt.id,
                StudentExamResponseThinkingSkills.q_id == q_id
            )
            .first()
        )

        if not response:
            print("❌ Response row missing (DATA INTEGRITY ISSUE):", {
                "attempt_id": attempt.id,
                "q_id": q_id
            })
            continue

        response.selected_option = selected
        response.correct_option = q.get("correct")
        response.is_correct = is_correct
        saved_responses += 1

    wrong = saved_responses - correct
    accuracy = round((correct / saved_responses) * 100, 2) if saved_responses else 0

    print("📈 Scoring summary:", {
        "attempted": saved_responses,
        "correct": correct,
        "wrong": wrong,
        "accuracy": accuracy
    })

    # --------------------------------------------------
    # 5️⃣ Save result row (idempotent)
    # --------------------------------------------------
    existing_result = (
        db.query(StudentExamResultsThinkingSkills)
        .filter(
            StudentExamResultsThinkingSkills.exam_attempt_id == attempt.id
        )
        .first()
    )

    if existing_result:
        print("ℹ️ Result row already exists:", {
            "result_id": existing_result.id
        })
    else:
        print("💾 Creating result row for attempt:", attempt.id)
        db.add(
            StudentExamResultsThinkingSkills(
                student_id=student.id,
                exam_attempt_id=attempt.id,
                total_questions=total_questions,
                correct_answers=correct,
                wrong_answers=wrong,
                accuracy_percent=accuracy
            )
        )

    # --------------------------------------------------
    # 6️⃣ ADMIN REPORT GENERATION (CRITICAL DIAGNOSTICS)
    # --------------------------------------------------
    print("🔎 Checking admin_exam_reports existence:", {
        "exam_attempt_id": attempt.id,
        "exam_type": "thinking_skills"
    })

    report_exists = admin_report_exists(
        db=db,
        exam_attempt_id=attempt.id,
        exam_type="thinking_skills"
    )

    if report_exists:
        print("⚠️ Admin report already exists → generation skipped")
    else:
        print("🚨 Admin report DOES NOT EXIST → generating now")
        print("🧠 Admin report input payload:", {
            "student_internal_id": student.id,
            "student_external_id": student.student_id,
            "attempt_id": attempt.id,
            "exam_type": "thinking_skills",
            "accuracy": accuracy,
            "correct": correct,
            "wrong": wrong,
            "total": total_questions
        })

        generate_admin_exam_report(
            db=db,
            student=student,
            exam_attempt=attempt,
            subject="thinking_skills",
            accuracy=accuracy,
            correct=correct,
            wrong=wrong,
            total_questions=total_questions
        )

        # 🔍 POST-CONDITION CHECK (THIS IS THE SMOKING GUN)
        post_check = admin_report_exists(
            db=db,
            exam_attempt_id=attempt.id,
            exam_type="thinking_skills"
        )

        if post_check:
            print("✅ CONFIRMED: Admin exam report successfully created")
        else:
            print("🔥 CRITICAL FAILURE: Admin report generation FAILED")
            print("🔥 INVESTIGATE generate_admin_exam_report IMMEDIATELY")

    # --------------------------------------------------
    # 7️⃣ Mark attempt completed (LAST STEP)
    # --------------------------------------------------
    if attempt.completed_at is None:
        attempt.completed_at = datetime.now(timezone.utc)
        print("✅ Attempt marked completed")

    db.commit()

    print("================ FINISH THINKING SKILLS EXAM END =================\n")

    return {
        "status": "completed",
        "total_questions": total_questions,
        "attempted": saved_responses,
        "correct": correct,
        "wrong": wrong,
        "accuracy": accuracy
    }


def aggregate_math_sections(db: Session, exam_attempt_id: int):
    """
    Aggregates Mathematical Reasoning responses into topic-based sections.

    Returns:
    {
        "Arithmetic": { "accuracy": 75.0 },
        "Word Problems": { "accuracy": 60.0 }
    }
    """

    responses = (
        db.query(StudentExamResponse)
        .filter(
            StudentExamResponse.exam_attempt_id == exam_attempt_id
        )
        .all()
    )

    section_totals = {}

    for r in responses:
        section = r.topic  # ✅ USE topic as section

        if not section:
            section = "General"

        if section not in section_totals:
            section_totals[section] = {"correct": 0, "total": 0}

        section_totals[section]["total"] += 1
        if r.is_correct:
            section_totals[section]["correct"] += 1

    section_results = {}

    for section, stats in section_totals.items():
        accuracy = (
            round((stats["correct"] / stats["total"]) * 100, 1)
            if stats["total"] else 0
        )

        section_results[section] = {
            "accuracy": accuracy
        }

    return section_results


def generate_admin_exam_report_math(
    db: Session,
    student: Student,
    exam_attempt,
    accuracy: float,
    correct: int,
    wrong: int,
    total_questions: int
):
    """
    Generates admin report + section results for Mathematical Reasoning.
    MUST NOT commit.
    """

    # -------------------------------
    # 1️⃣ Readiness & guidance rules
    # -------------------------------
    if accuracy >= 80:
        readiness_band = "Strong Selective Potential"
        school_guidance = "Mid-tier Selective"
    elif accuracy >= 60:
        readiness_band = "Borderline Selective"
        school_guidance = "Selective Preparation Required"
    else:
        readiness_band = "Not Yet Selective Ready"
        school_guidance = "General Stream Recommended"

    # -------------------------------
    # 2️⃣ Create admin report snapshot
    # -------------------------------
    admin_report = AdminExamReport(
        student_id=student.student_id,  # external ID (IMPORTANT)
        exam_attempt_id=exam_attempt.id,
        exam_type="mathematical_reasoning",
        overall_score=accuracy,
        readiness_band=readiness_band,
        school_guidance_level=school_guidance,
        summary_notes=f"Mathematical Reasoning accuracy: {accuracy}%"
    )

    db.add(admin_report)
    db.flush()  # get admin_report.id

    # -------------------------------
    # 3️⃣ Aggregate maths sections
    # -------------------------------
    section_stats = aggregate_math_sections(
        db=db,
        exam_attempt_id=exam_attempt.id
    )

    # -------------------------------
    # 4️⃣ Insert section rows
    # -------------------------------
    for section_name, stats in section_stats.items():
        raw_score = stats["accuracy"]

        if raw_score >= 80:
            band = "A"
        elif raw_score >= 60:
            band = "B"
        else:
            band = "C"

        db.add(
            AdminExamSectionResult(
                admin_report_id=admin_report.id,
                section_name=section_name,
                raw_score=raw_score,
                performance_band=band,
                strengths_summary=None,
                improvement_summary=None
            )
        )

    # -------------------------------
    # 5️⃣ Store applied rule (audit)
    # -------------------------------
    db.add(
        AdminReadinessRuleApplied(
            admin_report_id=admin_report.id,
            rule_code="MATHS_ACCURACY",
            rule_description="Mathematical Reasoning accuracy used for readiness classification",
            rule_result="passed" if accuracy >= 60 else "failed"
        )
    )


@app.post("/api/student/finish-exam")
def finish_exam(
    req: FinishExamRequest,
    db: Session = Depends(get_db)
):
    print("\n================ FINISH MATHEMATICAL REASONING EXAM START ================")
    print("📥 Incoming payload:", req.dict())

    # --------------------------------------------------
    # 1️⃣ Resolve student (external → internal)
    # --------------------------------------------------
    student = (
        db.query(Student)
        .filter(
            func.lower(Student.student_id) ==
            func.lower(req.student_id.strip())
        )

        .first()
    )

    if not student:
        print("❌ Student NOT FOUND:", req.student_id)
        raise HTTPException(status_code=404, detail="Student not found")

    print("✅ Student resolved → id:", student.id)

    # --------------------------------------------------
    # 2️⃣ Get active Mathematical Reasoning attempt
    # --------------------------------------------------
    attempt = (
        db.query(StudentExamMathematicalReasoning)
        .filter(
            StudentExamMathematicalReasoning.student_id == student.id,
            StudentExamMathematicalReasoning.completed_at.is_(None)
        )
        .order_by(StudentExamMathematicalReasoning.started_at.desc())
        .first()
    )
    if not attempt:
        print("⚠️ No active mathematical reasoning attempt found")
        return {"status": "completed"}
    print("✅ Active mathematical reasoning attempt found → id:", attempt.id)

    existing_result = (
        db.query(StudentExamResultsMathematicalReasoning)
        .filter(StudentExamResultsMathematicalReasoning.exam_attempt_id == attempt.id)
        .first()
    )
    
    if existing_result:
        print("⚠️ Result already exists → returning idempotent response")
        return {"status": "completed"}




    
    
    

    # --------------------------------------------------
    # 3️⃣ Load exam (STRICT subject guard)
    # --------------------------------------------------
    exam = (
        db.query(Exam)
        .filter(
            Exam.id == attempt.exam_id,
            Exam.subject == "mathematical_reasoning"
        )
        .first()
    )

    if not exam:
        raise HTTPException(
            status_code=400,
            detail="Finish endpoint called for non-mathematical reasoning exam"
        )

    questions = exam.questions or []
    total_questions = len(questions)

    print("📘 Exam loaded → questions:", total_questions)

    question_map = {q["q_id"]: q for q in questions}

    # --------------------------------------------------
    # Clear previous reports for this student + subject
    # --------------------------------------------------
    
    
    # --------------------------------------------------
    # 5️⃣ Update student responses (NO inserts)
    # --------------------------------------------------
    correct = 0
    saved_responses = 0

    for q_id_str, selected in req.answers.items():
        try:
            q_id = int(q_id_str)
        except ValueError:
            continue
    
        q = question_map.get(q_id)
        if not q:
            continue
    
        correct_answer = q.get("correct")
    
        # 🔍 DEBUG LOGS (THIS IS THE KEY)
        print("🧪 EVALUATING QUESTION")
        print("Question ID:", q_id)
        print("Student Answer:", selected)
        print("Correct Answer (DB):", correct_answer)
    
        is_correct = selected == correct_answer
    
        print("➡️ MATCH RESULT:", is_correct)
        print("------------------------------------")
    
        if is_correct:
            correct += 1
        response = (
            db.query(StudentExamResponseMathematicalReasoning)
            .filter(
                StudentExamResponseMathematicalReasoning.exam_attempt_id == attempt.id,
                StudentExamResponseMathematicalReasoning.q_id == q_id
            )
            .first()
        )


        if not response:
            print(f"⚠️ Missing response row for q_id={q_id}, attempt_id={attempt.id}")
            continue


        response.selected_option = selected
        response.correct_option = q.get("correct")
        response.is_correct = is_correct

        saved_responses += 1

    wrong = saved_responses - correct
    accuracy = round((correct / total_questions) * 100, 2) if saved_responses else 0

    print("📊 Result computed → correct:", correct, "wrong:", wrong)
    existing_raw = (
        db.query(AdminExamRawScore)
        .filter(
            AdminExamRawScore.exam_attempt_id == attempt.id,
            AdminExamRawScore.subject == "mathematical_reasoning"
        )
        .first()
    )
    
    if not existing_raw:
        raw_score = AdminExamRawScore(
            student_id=student.id,
            exam_attempt_id=attempt.id,
            subject="mathematical_reasoning",
            total_questions=total_questions,
            correct_answers=correct,
            wrong_answers=wrong,
            accuracy_percent=accuracy
        )
        db.add(raw_score)


    # --------------------------------------------------
    # 6️⃣ Save summary (NEW TABLE)
    # --------------------------------------------------
    result_row = StudentExamResultsMathematicalReasoning(
        student_id=student.id,
        exam_attempt_id=attempt.id,
        total_questions=total_questions,
        correct_answers=correct,
        wrong_answers=wrong,
        accuracy_percent=accuracy
    )

    db.add(result_row)

    # --------------------------------------------------
    # 7️⃣ Mark attempt completed
    # --------------------------------------------------
    attempt.completed_at = datetime.now(timezone.utc)
    # --------------------------------------------------
    # 8️⃣ Generate Admin Report Snapshot (ADMIN)
    # --------------------------------------------------
    if not admin_report_exists(
        db=db,
        exam_attempt_id=attempt.id,
        exam_type="mathematical_reasoning"

    ):
        generate_admin_exam_report_math(
            db=db,
            student=student,
            exam_attempt=attempt,
            accuracy=accuracy,
            correct=correct,
            wrong=wrong,
            total_questions=total_questions
        )
    
    # --------------------------------------------------
    # 9️⃣ Final commit (single transaction)
    # --------------------------------------------------
    db.commit()

    print("================ FINISH MATHEMATICAL REASONING EXAM END =================\n")
    

    return {
        "status": "completed",
        "total_questions": total_questions,
        "attempted": saved_responses,
        "correct": correct,
        "wrong": wrong,
        "accuracy": accuracy
    }



@app.get("/api/student/get-quiz")
def get_quiz(student_id: str, subject: str, difficulty: str, db: Session = Depends(get_db)):
    """
    Returns full quiz questions for the student based on class, subject, difficulty.
    Includes detailed debug logs for diagnosing issues.
    """

    print("\n===========================")
    print("📥 API CALL: /api/student/get-quiz")
    print("===========================")
    print("🔸 Incoming parameters:")
    print(f"   student_id: {student_id}")
    print(f"   subject: {subject}")
    print(f"   difficulty: {difficulty}")
    print("---------------------------")

    # -------------------------------
    # VALIDATE STUDENT
    # -------------------------------
    student = db.query(Student).filter(Student.student_id == student_id).first()

    if not student:
        print("❌ ERROR: Student not found in database.")
        print(f"   student_id searched: {student_id}")
        raise HTTPException(status_code=404, detail="Student not found")

    print("✅ Student found:")
    print(f"   internal_id: {student.id}")
    print(f"   class_name: {student.class_name}")
    print("---------------------------")

    # -------------------------------
    # FIND MATCHING QUIZ
    # -------------------------------
    print("🔎 Searching for quiz definition matching:")
    print(f"   class_name: {student.class_name.lower()}")
    print(f"   subject: {subject}")
    print(f"   difficulty: {difficulty}")

    quiz = (
        db.query(Quiz)
        .filter(
            func.lower(Quiz.class_name) == student.class_name.lower(),
            Quiz.subject == subject,
            Quiz.difficulty == difficulty
        )
        .order_by(Quiz.id.desc())
        .first()
    )

    if not quiz:
        print("❌ ERROR: No quiz definition found matching given params.")
        raise HTTPException(status_code=404, detail="No quiz found")

    print("✅ Quiz definition found:")
    print(f"   quiz_id: {quiz.id}")
    print(f"   topics: {quiz.topics}")
    print("---------------------------")

    # -------------------------------
    # GET EXAM INSTANCE
    # -------------------------------
    print("🔎 Looking for generated exam linked to quiz...")

    exam = (
        db.query(Exam)
        .filter(Exam.quiz_id == quiz.id)
        .order_by(Exam.id.desc())
        .first()
    )

    if not exam:
        print("❌ ERROR: Exam instance not found — quiz not generated yet.")
        raise HTTPException(status_code=404, detail="Quiz not generated")

    print("✅ Exam found:")
    print(f"   exam_id: {exam.id}")
    print(f"   number_of_questions: {len(exam.questions) if exam.questions else 0}")
    print("---------------------------")

    # -------------------------------
    # RETURN QUESTIONS
    # -------------------------------
    print("📤 Preparing JSON response...")
    response = {
        "exam_id": exam.id,
        "quiz_id": quiz.id,
        "questions": exam.questions
    }

    print("📦 Returning response:")
    print(f"   exam_id: {response['exam_id']}")
    print(f"   quiz_id: {response['quiz_id']}")
    print(f"   total_questions: {len(response['questions'])}")
    print("===========================\n")

    return response

@app.get("/list-images")
def list_images_from_bucket():
    print("\n====================== LIST IMAGES START ======================")

    try:
        blobs = gcs_bucket.list_blobs()

        images = []

        for blob in blobs:
            if blob.name.endswith("/"):
                continue

            images.append({
                "image_name": blob.name,
                "url": blob.public_url
            })

        images.sort(key=lambda x: x["image_name"].lower())

        print(f"[INFO] Total images found: {len(images)}")
        print("====================== LIST IMAGES END ======================\n")

        return {
            "status": "success",
            "total": len(images),
            "images": images
        }

    except Exception:
        print("[ERROR] Failed to fetch images from bucket")
        print(traceback.format_exc())
        raise HTTPException(
            status_code=500,
            detail="Failed to fetch images"
        )



@app.post("/upload-image-folder")
async def upload_image_folder(
    images: List[UploadFile] = File(...),
    db: Session = Depends(get_db)
):
    print("\n====================== FOLDER UPLOAD START ======================")
    print(f"[INFO] Number of received files: {len(images)}")

    if not images:
        raise HTTPException(status_code=400, detail="No images received.")

    uploaded_urls = []

    for file in images:
        try:
            print("\n-----------------------------------------------------------")
            print(f"[FILE] Incoming filename: {file.filename}")
            print(f"[FILE] Content type: {file.content_type}")

            # ✅ Normalize + sanitize filename (CRITICAL)
            original_name = sanitize_filename(file.filename)
            print(f"[FILE] Sanitized original_name: {original_name}")

            # 🔍 Check DB FIRST (idempotent)
            existing = (
                db.query(UploadedImage)
                .filter(UploadedImage.original_name == original_name)
                .first()
            )

            if existing:
                print("[SKIP] Image already exists in DB")
                print(f"   • original_name : {existing.original_name}")
                print(f"   • gcs_url       : {existing.gcs_url}")

                uploaded_urls.append({
                    "original_name": existing.original_name,
                    "url": existing.gcs_url,
                    "status": "already_exists"
                })
                continue

            # ⬆️ Only upload if NOT already present
            file_bytes = await file.read()
            gcs_url = upload_to_gcs(file_bytes, original_name)

            print("[UPLOAD SUCCESS]")
            print(f"   • original_name : {original_name}")
            print(f"   • gcs_url       : {gcs_url}")

            # 💾 Persist mapping
            record = UploadedImage(
                original_name=original_name,
                gcs_url=gcs_url
            )
            db.add(record)
            db.commit()

            uploaded_urls.append({
                "original_name": original_name,
                "url": gcs_url,
                "status": "uploaded"
            })

        except Exception as e:
            print(f"[ERROR] Failed to process file '{file.filename}'")
            print(traceback.format_exc())
            raise HTTPException(status_code=500, detail=str(e))

    print("\n====================== FOLDER UPLOAD END ========================")

    return {
        "status": "success",
        "message": f"{len(uploaded_urls)} images processed.",
        "files": uploaded_urls
    }
 

#old upload-word code
# @app.post("/upload-word")
# async def upload_word(
#     file: UploadFile = File(...),
#     db: Session = Depends(get_db)
# ):
#     request_id = str(uuid.uuid4())[:8]
#
#     print(f"\n========== GPT-UPLOAD-WORD START [{request_id}] ==========")
#     print(f"[{request_id}] 📄 Incoming file: {file.filename}")
#     print(f"[{request_id}] 📄 Content-Type: {file.content_type}")
#
#     allowed = {
#         "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         "text/plain"
#     }
#
#     if file.content_type not in allowed:
#         print(f"[{request_id}] ❌ Invalid file type")
#         raise HTTPException(status_code=400, detail="Invalid file type.")
#
#     try:
#         content = await file.read()
#         doc = docx.Document(BytesIO(content))
#         paragraphs = []
#
#         for p in doc.paragraphs:
#             text = p.text.rstrip()
#             if text:
#                 paragraphs.append(text)
#
#         print(f"[{request_id}] 📄 Extracted paragraphs: {len(paragraphs)}")
#
#     except Exception as e:
#         print(f"[{request_id}] ❌ DOCX read error: {e}")
#         raise HTTPException(status_code=400, detail="Failed to read file.")
#
#     blocks = chunk_by_question(paragraphs)
#     print(f"[{request_id}] 📦 Total blocks: {len(blocks)}")
#
#     all_questions = []
#
#     for block_idx, block in enumerate(blocks, start=1):
#         print(
#             f"\n[{request_id}] ▶️ GPT BLOCK {block_idx}/{len(blocks)} "
#             f"| chars={len(block)}"
#         )
#
#         result = await parse_with_gpt(block)
#         questions = result.get("questions", [])
#
#         print(f"[{request_id}] 🧩 GPT returned {len(questions)} questions")
#
#         for qi, q in enumerate(questions, start=1):
#             print(
#                 f"[{request_id}] 🔎 B{block_idx}-Q{qi} | "
#                 f"partial={q.get('partial')} | "
#                 f"class={q.get('class_name')} | "
#                 f"topic={q.get('topic')} | "
#                 f"images={q.get('images')}"
#             )
#
#         all_questions.extend(questions)
#
#     print(f"\n[{request_id}] 📊 Total questions detected: {len(all_questions)}")
#
#     saved_ids = []
#     skipped_partial = 0
#
#     for idx, q in enumerate(all_questions, start=1):
#         qid = f"Q{idx}"
#
#         required_fields = ["question_text", "correct_answer"]
#
#         if q.get("options") is not None:
#             required_fields.append("options")
#
#         missing = [f for f in required_fields if not q.get(f)]
#
#         if missing:
#             skipped_partial += 1
#             print(
#                 f"[{request_id}] ⚠ SKIP {qid} | "
#                 f"missing_fields={missing} | "
#                 f"partial={q.get('partial')}"
#             )
#             continue
#
#         print(
#             f"[{request_id}] ✅ READY {qid} | "
#             f"options={list(q['options'].keys())} | "
#             f"correct={q['correct_answer']}"
#         )
#
#         resolved_images = []
#
#         for img in q.get("images") or []:
#             img_raw = img
#             img_norm = img.strip().lower()
#
#             record = (
#                 db.query(UploadedImage)
#                 .filter(func.lower(func.trim(UploadedImage.original_name)) == img_norm)
#                 .first()
#             )
#
#             if not record:
#                 raise HTTPException(
#                     status_code=400,
#                     detail=f"Image '{img_raw}' not uploaded yet"
#                 )
#
#             resolved_images.append(record.gcs_url)
#
#         question_blocks = q.get("question_blocks")
#         if not question_blocks and q.get("question_text"):
#             question_blocks = [
#                 {"type": "text", "content": q["question_text"]}
#             ]
#
#         new_q = Question(
#             class_name=q.get("class_name"),
#             subject=q.get("subject"),
#             topic=q.get("topic"),
#             difficulty=q.get("difficulty"),
#             question_type="multi_image_diagram_mcq",
#
#             question_text=q.get("question_text"),
#             question_blocks=question_blocks,
#
#             images=resolved_images,
#             options=q.get("options"),
#             correct_answer=q.get("correct_answer")
#         )
#
#         db.add(new_q)
#         db.commit()
#         db.refresh(new_q)
#
#         saved_ids.append(new_q.id)
#         print(f"[{request_id}] 💾 SAVED {qid} | db_id={new_q.id}")
#
#     print(
#         f"\n[{request_id}] 🏁 UPLOAD COMPLETE | "
#         f"saved={len(saved_ids)} | "
#         f"partial_skipped={skipped_partial} | "
#         f"total_detected={len(all_questions)}"
#     )
#     print(f"========== GPT-UPLOAD-WORD END [{request_id}] ==========\n")
#
#     return {
#         "status": "success",
#         "count": len(saved_ids),
#         "skipped_partial": skipped_partial,
#         "message": (
#             f"{len(saved_ids)} questions added successfully. "
#             f"{skipped_partial} questions were skipped due to missing data."
#         )
#     }

def chunk_by_question(blocks):
    """
    Input:
      blocks = [
        {"type": "text", "content": "..."},
        {"type": "image", "name": "..."},
        ...
      ]

    Output:
      [
        [ block, block, block ],   # Question 1
        [ block, block, block ],   # Question 2
      ]
    """

    questions = []
    current = []

    for block in blocks:
        # Detect question boundary
        if (
            block["type"] == "text"
            and block["content"].strip() == "METADATA:"
        ):
            if current:
                questions.append(current)
                current = []

        current.append(block)

    if current:
        questions.append(current)

    return questions
 
def extract_images_from_text(block: str) -> list[str]:
    images = []
    for line in block.splitlines():
        if line.strip().upper().startswith("IMAGES:"):
            img_part = line.split(":", 1)[1]
            images = [i.strip() for i in img_part.split(",") if i.strip()]
    return images

def parse_docx_to_ordered_blocks(doc):
    """
    Returns a linear list of blocks:
    [{type: 'text', content: ...}, {type: 'image', name: ...}]
    preserving DOCX order exactly.
    Supports BOTH embedded images and IMAGES: text declarations.
    """
    blocks = []

    for element in doc.element.body:
        if element.tag.endswith("}p"):
            texts = []
            image_names = []

            for run in element.iter():
                # Text
                if run.tag.endswith("}t") and run.text:
                    texts.append(run.text)

                # Embedded image
                if run.tag.endswith("}blip"):
                    r_id = run.attrib.get(qn("r:embed"))
                    if r_id:
                        image_part = doc.part.related_parts[r_id]
                        image_name = image_part.partname.split("/")[-1]
                        image_names.append(image_name)

            if texts:
                text = "".join(texts).strip()

                # 👇 NEW: Convert IMAGES: lines into image blocks
                if text.upper().startswith("IMAGES:"):
                    img_part = text.split(":", 1)[1]
                    for img in img_part.split(","):
                        blocks.append({
                            "type": "image",
                            "name": img.strip()
                        })
                else:
                    blocks.append({
                        "type": "text",
                        "content": text
                    })

            # Embedded images (rare in your docs, but supported)
            for img in image_names:
                blocks.append({
                    "type": "image",
                    "name": img
                })

    return blocks

def filter_display_blocks(blocks: list[dict]) -> list[dict]:
    """
    Removes authoring-only text from question blocks.
    Keeps ONLY student-visible question stem + images.
    """
    filtered = []

    DROP_PREFIXES = (
        "METADATA:",
        "CLASS:",
        "SUBJECT:",
        "TOPIC:",
        "DIFFICULTY:",
        "QUESTION_TEXT:",
        "OPTIONS:",
        "CORRECT_ANSWER:",
        "CORRECT ANSWER:",
    )

    for block in blocks:
        # Always keep images
        if block["type"] == "image":
            filtered.append(block)
            continue

        text = block.get("content", "").strip()

        # Drop empty text
        if not text:
            continue

        upper = text.upper()

        # Drop metadata / labels
        if any(upper.startswith(p) for p in DROP_PREFIXES):
            continue

        # Drop "Question X:" lines
        if re.match(r"^QUESTION\s+\d+\s*:", upper):
            continue

        # Drop option lines accidentally leaking (A. / B) etc)
        if re.match(r"^[A-D][\.\)]\s+", text):
            continue

        filtered.append({
            "type": "text",
            "content": text
        })

    return filtered




def resolve_option_value(value: str, db: Session) -> dict:
    value = value.strip()

    # Detect image option
    if value.lower().endswith((".png", ".jpg", ".jpeg")):
        img_name = value.lower()

        record = (
            db.query(UploadedImage)
            .filter(
                func.lower(func.trim(
                    UploadedImage.original_name
                )) == img_name
            )
            .first()
        )

        if not record:
            raise HTTPException(
                status_code=400,
                detail=f"Option image '{img_name}' not uploaded yet"
            )

        return {
            "type": "image",
            "src": record.gcs_url
        }

    # Fallback → text option
    return {
        "type": "text",
        "content": value
    }


def looks_like_question(blocks):
    markers = {
        "question_text",
        "correct_answer",
        "question_type",
        "answer_type",
        "cloze",
        "options",
    }


    for b in blocks:
        if b.get("type") == "text":
            content = b.get("content", "").lower()
            if any(m in content for m in markers):
                return True
    return False


def extract_exam_block(ordered_blocks):
    exam_blocks = []
    inside_exam = False

    for block in ordered_blocks:
        if block["type"] == "text":
            text = block["content"].strip()

            if "=== EXAM START ===" in text:
                inside_exam = True
                continue

            if "=== EXAM END ===" in text:
                inside_exam = False
                break

        if inside_exam:
            exam_blocks.append(block)

    if not exam_blocks:
        raise ValueError("No EXAM block found")

    return exam_blocks
EXAM_START_PATTERN = re.compile(r"^={3,}\s*exam\s+start\s*={3,}$", re.IGNORECASE)
EXAM_END_PATTERN = re.compile(r"^={3,}\s*exam\s+end\s*={3,}$", re.IGNORECASE)


def is_exam_start(block):
    if block.get("type") != "text":
        return False

    text = block.get("content", "").strip()
    return bool(EXAM_START_PATTERN.match(text))


def is_exam_end(block):
    if block.get("type") != "text":
        return False

    text = block.get("content", "").strip()
    return bool(EXAM_END_PATTERN.match(text))

def chunk_by_exam(blocks):
    exams = []
    current = []
    inside_exam = False

    for block in blocks:
        if is_exam_start(block):
            current = []
            inside_exam = True
            continue

        if is_exam_end(block):
            if inside_exam and current:
                exams.append(current)
            current = []
            inside_exam = False
            continue

        if inside_exam:
            current.append(block)

    # Handle missing EXAM END at EOF
    if inside_exam and current:
        exams.append(current)

    return exams
def count_question_markers(blocks):
    count = 0
    for block in blocks:
        if block["type"] == "text":
            text = block["content"].lower()
            if text.strip().startswith("question "):
                count += 1
    return count

#new upload-word code
#new upload-word code
@app.post("/upload-word")
async def upload_word(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    request_id = str(uuid.uuid4())[:8]
    report = []

    print(f"\n========== GPT-UPLOAD-WORD START [{request_id}] ==========")
    print(f"[{request_id}] 📄 Incoming file: {file.filename}")
    print(f"[{request_id}] 📄 Content-Type: {file.content_type}")

    # -------------------------------------------------
    # Validate file type
    # -------------------------------------------------
    allowed_types = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/plain",
    }
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Invalid file type")

    # -------------------------------------------------
    # Load DOCX → ordered blocks
    # -------------------------------------------------
    try:
        content = await file.read()
        doc = docx.Document(BytesIO(content))
        ordered_blocks = parse_docx_to_ordered_blocks(doc)
    except Exception as e:
        raise HTTPException(status_code=400, detail="Failed to read file")

    # -------------------------------------------------
    # EXAM = SOURCE OF TRUTH
    # -------------------------------------------------
    exam_blocks = chunk_by_exam(ordered_blocks)

    saved = 0
    skipped = 0

    # =================================================
    # PROCESS EACH EXAM (1 exam → 1 question max)
    # =================================================
    for exam_idx, exam_block in enumerate(exam_blocks, start=1):
        question_count = count_question_markers(exam_block)

        if question_count > 1:
            raise HTTPException(
                status_code=422,
                detail={
                    "error": "INVALID_EXAM_FORMAT",
                    "message": (
                        f"Exam {exam_idx} contains {question_count} questions. "
                        "Expected exactly ONE question per exam. "
                        "Please split questions into separate EXAM blocks."
                    )
                }
            )

        # -----------------------------
        # GPT parse (expect ONE question)
        # -----------------------------
        try:
            result = await parse_with_gpt({"blocks": exam_block})
            question = result.get("question")
        except Exception as e:
            skipped += 1
            report.append({
                "exam": exam_idx,
                "status": "failed",
                "error_code": "GPT_ERROR",
                "details": str(e)
            })
            continue

        if not question:
            skipped += 1
            report.append({
                "exam": exam_idx,
                "status": "skipped",
                "error_code": "NO_QUESTION_PARSED"
            })
            continue

        if isinstance(question, list):
            skipped += 1
            report.append({
                "exam": exam_idx,
                "status": "failed",
                "error_code": "MULTIPLE_QUESTIONS_RETURNED"
            })
            continue

        # -----------------------------
        # Validate required fields
        # -----------------------------
        if not question.get("options") or not question.get("correct_answer"):
            skipped += 1
            report.append({
                "exam": exam_idx,
                "status": "skipped",
                "error_code": "INCOMPLETE_QUESTION"
            })
            continue

        # -----------------------------
        # Resolve images (exam-level)
        # -----------------------------
        resolved_blocks = []
        missing_image = None

        for block in exam_block:
            if block["type"] != "image":
                resolved_blocks.append(block)
                continue

            raw_name = (block.get("name") or block.get("src") or "").strip().lower()

            image_record = (
                db.query(UploadedImage)
                .filter(func.lower(func.trim(UploadedImage.original_name)) == raw_name)
                .first()
            )

            if not image_record:
                missing_image = raw_name
                break

            resolved_blocks.append({
                **block,
                "src": image_record.gcs_url
            })

        if missing_image:
            skipped += 1
            report.append({
                "exam": exam_idx,
                "status": "skipped",
                "error_code": "IMAGE_NOT_UPLOADED",
                "details": missing_image
            })
            continue

        # -----------------------------
        # Guard: must contain text
        # -----------------------------
        if not any(b["type"] == "text" for b in resolved_blocks):
            skipped += 1
            report.append({
                "exam": exam_idx,
                "status": "skipped",
                "error_code": "IMAGE_ONLY_EXAM"
            })
            continue
        # -----------------------------
        # Guard: question length sanity
        # -----------------------------
        MIN_QUESTION_CHARS = 20
        MAX_QUESTION_CHARS = 2000
        
        question_text_length = sum(
            len(b["content"])
            for b in resolved_blocks
            if b["type"] == "text"
        )
        
        if question_text_length < MIN_QUESTION_CHARS:
            skipped += 1
            report.append({
                "exam": exam_idx,
                "status": "skipped",
                "error_code": "QUESTION_TOO_SHORT",
                "details": f"{question_text_length} chars"
            })
            continue
        
        if question_text_length > MAX_QUESTION_CHARS:
            skipped += 1
            report.append({
                "exam": exam_idx,
                "status": "skipped",
                "error_code": "QUESTION_TOO_LONG",
                "details": f"{question_text_length} chars"
            })
            continue

        FORBIDDEN_TOKENS = [
            "=== EXAM START ===",
            "=== EXAM END ===",
        ]
        exam_marker_found = False

        for b in resolved_blocks:
            if b["type"] == "text":
                for token in FORBIDDEN_TOKENS:
                    if token in b["content"]:
                        exam_marker_found = True
                        break
            if exam_marker_found:
                break
        
        if exam_marker_found:
            skipped += 1
            report.append({
                "exam": exam_idx,
                "status": "skipped",
                "error_code": "EXAM_MARKER_IN_QUESTION"
            })
            continue

        # -----------------------------
        # Save question
        # -----------------------------
        try:
            question_text = "\n\n".join(
                b["content"] for b in resolved_blocks if b["type"] == "text"
            )

            new_question = Question(
                class_name=question.get("class_name"),
                subject=question.get("subject"),
                topic=question.get("topic"),
                difficulty=question.get("difficulty"),
                question_type=question.get("question_type") or "multi_image_diagram_mcq",
                question_text=question_text,
                question_blocks=filter_display_blocks(resolved_blocks),
                options=question["options"],
                correct_answer=question["correct_answer"]
            )

            db.add(new_question)
            db.commit()
            db.refresh(new_question)

            saved += 1
            report.append({
                "exam": exam_idx,
                "question": f"Q{saved}",
                "status": "success"
            })

        except Exception as e:
            skipped += 1
            report.append({
                "exam": exam_idx,
                "status": "failed",
                "error_code": "DB_ERROR",
                "details": str(e)
            })

    # -------------------------------------------------
    # Final response
    # -------------------------------------------------
    overall_status = "success" if skipped == 0 else "partial_success"

    return {
        "status": overall_status,
        "summary": {
            "total_exams": len(exam_blocks),
            "saved": saved,
            "skipped": skipped
        },
        "exams": report
    }





def chunk_by_exam_markers(blocks: list[dict]) -> list[list[dict]]:
    exams = []
    current_exam = []
    in_exam = False

    for block in blocks:
        if block["type"] != "text":
            if in_exam:
                current_exam.append(block)
            continue

        raw = block.get("content", "")
        lines = [line.strip() for line in raw.splitlines() if line.strip()]

        for text in lines:
            if text == "=== EXAM START ===":
                if in_exam:
                    raise ValueError("Nested EXAM START detected")
                in_exam = True
                current_exam = []
                continue

            if text == "=== EXAM END ===":
                if not in_exam:
                    raise ValueError("EXAM END without EXAM START")
                exams.append(current_exam)
                current_exam = []
                in_exam = False
                continue

            if in_exam:
                current_exam.append({"type": "text", "content": text})

    if in_exam:
        raise ValueError("Missing EXAM END")

    return exams


#async def read_and_validate_file(file, request_id) -> bytes:
 ##      "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
   #     "text/plain",
    #}

    #if file.content_type not in allowed:
     #   raise HTTPException(400, "Invalid file type")

    #content = await file.read()
    #print(f"[{request_id}] 📦 File read ({len(content)} bytes)")
    #return content


async def read_and_validate_file(file, request_id) -> bytes:
    allowed = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    if file.content_type not in allowed:
        raise HTTPException(
            status_code=400,
            detail="Invalid file type. Only .docx Word files are allowed."
        )

    content = await file.read()

    if not content:
        raise HTTPException(
            status_code=400,
            detail="Uploaded file is empty."
        )

    print(f"[{request_id}] 📦 File read ({len(content)} bytes)")
    return content

def is_cloze_exam_document(doc):
    for p in doc.paragraphs:
        if "question_type: 5" in p.text.lower():
            return True
    return False

def parse_docx_to_ordered_blocks_numeracy(doc):
    """
    Returns a linear list of ordered blocks preserving DOCX order exactly.

    Block types emitted:
    - text:     {"type": "text", "content": "..."}
    - image:    {"type": "image", "name": "..."}
    - cloze:    {"type": "cloze", "content": "..."}
    - options:  {"type": "options", "options": [{"id": "A", "text": "..."}]}

    Guarantees:
    - EXAM markers (=== EXAM START / END ===) are preserved verbatim
    - CLOZE and OPTIONS are emitted as structured blocks
    - IMAGES declarations are converted to image blocks
    """

    blocks = []

    current_mode = None    # None | "cloze" | "options"
    buffer = []

    def flush_buffer():
        nonlocal buffer, current_mode

        if not buffer or not current_mode:
            buffer = []
            current_mode = None
            return

        if current_mode == "cloze":
            blocks.append({
                "type": "cloze",
                "content": " ".join(buffer).strip()
            })

        elif current_mode == "options":
            options = []
            for line in buffer:
                if ":" in line:
                    label, text = line.split(":", 1)
                    options.append({
                        "id": label.strip(),
                        "text": text.strip()
                    })

            blocks.append({
                "type": "options",
                "options": options
            })

        buffer = []
        current_mode = None

    for element in doc.element.body:
        if not element.tag.endswith("}p"):
            continue

        texts = []
        image_names = []

        for run in element.iter():
            # Text runs
            if run.tag.endswith("}t") and run.text:
                texts.append(run.text)

            # Embedded images
            if run.tag.endswith("}blip"):
                r_id = run.attrib.get(qn("r:embed"))
                if r_id:
                    image_part = doc.part.related_parts[r_id]
                    image_name = image_part.partname.split("/")[-1]
                    image_names.append(image_name)

        if texts:
            text = "".join(texts).strip()
            upper = text.upper()

            # --------------------------------------------------
            # EXAM markers (must remain untouched)
            # --------------------------------------------------
            if upper.startswith("==="):
                flush_buffer()
                blocks.append({
                    "type": "text",
                    "content": text
                })
                continue

            # --------------------------------------------------
            # Section headers
            # --------------------------------------------------
            if upper == "CLOZE:":
                flush_buffer()
                current_mode = "cloze"
                continue

            if upper == "OPTIONS:":
                flush_buffer()
                current_mode = "options"
                continue

            # --------------------------------------------------
            # Inside CLOZE / OPTIONS sections
            # --------------------------------------------------
            if current_mode in {"cloze", "options"}:
                buffer.append(text)
                continue

            # --------------------------------------------------
            # IMAGES declaration
            # --------------------------------------------------
            if upper.startswith("IMAGES:"):
                flush_buffer()
                img_part = text.split(":", 1)[1]
                for img in img_part.split(","):
                    blocks.append({
                        "type": "image",
                        "name": img.strip()
                    })
                continue

            # --------------------------------------------------
            # Default text
            # --------------------------------------------------
            blocks.append({
                "type": "text",
                "content": text
            })

        # --------------------------------------------------
        # Embedded images (still supported)
        # --------------------------------------------------
        for img in image_names:
            blocks.append({
                "type": "image",
                "name": img
            })

    # Flush any remaining buffered content
    flush_buffer()

    return blocks
def parse_docx_to_flat_text_blocks(doc):
    blocks = []

    for idx, paragraph in enumerate(doc.paragraphs):
        raw = paragraph.text
        if not raw:
            continue

        content = raw.strip()
        if not content:
            continue

        blocks.append({
            "type": "text",
            "content": content,   # 🔑 THIS IS THE KEY
            "source": "docx",
            "paragraph_index": idx
        })

    return blocks

def parse_docx_blocks(content: bytes, request_id):
    doc = docx.Document(BytesIO(content))

    # 👇 Detect CLOZE exam early
    if is_cloze_exam_document(doc):
        blocks = parse_docx_to_flat_text_blocks(doc)
    else:
        blocks = parse_docx_to_ordered_blocks_numeracy(doc)

    return blocks

async def parse_questions_with_gpt_naplan_numeracy_lc(
    question_block: list,
    request_id: str
):
    """
    Thin adapter around `parse_with_gpt_numeracy_lc`.

    Responsibilities:
    - Logging
    - Payload shaping
    - Basic input/output validation

    Must NOT:
    - Modify GPT behavior
    - Infer missing data
    - Transform question semantics
    """

    print(f"[{request_id}] 🤖 GPT-NAPLAN-NUMERACY-LC parse START")

    # --------------------------------------------------
    # Defensive input checks
    # --------------------------------------------------
    if not isinstance(question_block, list):
        raise ValueError(
            "question_block must be a list of parsed document blocks"
        )

    print(
        f"[{request_id}] 🤖 Incoming blocks = {len(question_block)} | "
        f"preview = {str(question_block[:1])[:120]}"
    )

    # --------------------------------------------------
    # Call existing GPT parser (DO NOT MODIFY)
    # --------------------------------------------------
    result = await parse_with_gpt_numeracy_lc({
        "blocks": question_block
    })

    if not isinstance(result, dict):
        raise ValueError("GPT parser returned non-dict result")

    questions = result.get("questions", [])

    # --------------------------------------------------
    # Defensive output checks + logging
    # --------------------------------------------------
    print(
        f"[{request_id}] 🤖 GPT-NAPLAN-NUMERACY-LC parse END | "
        f"questions_returned = {len(questions)}"
    )

    if questions:
        print(
            f"[{request_id}] 🧠 First question preview | "
            f"answer_type={questions[0].get('answer_type')} | "
            f"keys={list(questions[0].keys())}"
        )


    return questions

ANSWER_TYPE_TO_QUESTION_TYPE = {
    "MCQ_SINGLE": 1,
    "MCQ_MULTI": 2,
    "NUMERIC_INPUT": 3,
    "TEXT_INPUT": 4,
    "CLOZE_DROPDOWN": 5,   # 👈 NEW
}
def validate_cloze_deterministic(q: dict):
    """
    Strict validation for deterministic CLOZE questions.
    Raises ValueError with clear messages on failure.
    """

    # ------------------------------
    # CLOZE text
    # ------------------------------
    cloze_text = q.get("cloze_text")
    if not cloze_text:
        raise ValueError("CLOZE missing cloze_text")

    token_count = cloze_text.count("{{dropdown}}")
    if token_count != 1:
        raise ValueError(
            f"CLOZE must contain exactly one {{dropdown}} "
            f"(found {token_count})"
        )

    # ------------------------------
    # Options
    # ------------------------------
    options = q.get("options")
    if not isinstance(options, dict) or not options:
        raise ValueError("CLOZE missing options")

    for key, value in options.items():
        if not (len(key) == 1 and key.isupper()):
            raise ValueError(f"Invalid option key: {key}")
        if not value or not value.strip():
            raise ValueError(f"Empty option value for {key}")

    # ------------------------------
    # Correct answer
    # ------------------------------
    correct = q.get("correct_answer")
    if not correct:
        raise ValueError("CLOZE missing correct_answer")

    if correct not in options:
        raise ValueError(
            f"Correct answer '{correct}' not in options {list(options.keys())}"
        )

def is_visual_counting_exam(block):
    """
    Detects Visual Counting (Type 6) exams.

    We intentionally use loose matching here because DOCX text
    may contain extra spaces, line breaks, or formatting noise.
    """
    for item in block:
        if not isinstance(item, dict):
            continue

        if item.get("type") != "text":
            continue

        text = (item.get("content") or item.get("text") or "")
        text = text.strip().lower()

        # Robust detection
        if "question_type" in text and "6" in text:
            return True

    return False
def vc_extract_question_text(block):
    collecting = False
    lines = []

    for item in block:
        if item.get("type") != "text":
            continue

        raw = item.get("text") or item.get("content") or ""
        text = raw.strip()

        if text == "QUESTION_TEXT:":
            collecting = True
            continue

        if collecting:
            if text == "OPTIONS:":
                break
            if text:
                lines.append(text)

    if not lines:
        raise ValueError("VC: QUESTION_TEXT section not found or empty")

    return " ".join(lines)

import re
import unicodedata

def normalize_text(text: str) -> str:
    # Normalize unicode (fix full-width chars, etc.)
    text = unicodedata.normalize("NFKC", text)

    # Remove zero-width / non-breaking spaces
    text = text.replace("\u00a0", " ")
    text = text.replace("\u200b", "")

    # Collapse whitespace
    text = re.sub(r"\s+", " ", text)

    return text.strip()


def vc_extract_image_options(block):
    options = {}

    print("\n================ VC DEBUG START ================")
    print(f"Total block items: {len(block)}")

    for idx, item in enumerate(block):
        raw = item.get("text") or item.get("content") or ""
        text = normalize_text(raw)

        print(f"[VC RAW {idx}] repr(raw) = {repr(raw)}")
        print(f"[VC NORM {idx}] text = {repr(text)}")

        if not text:
            continue

        match = re.match(
            r"^([A-D])\s*[:\.]\s*(image_[\w\-.]+)",
            text,
            flags=re.IGNORECASE
        )

        if match:
            label = match.group(1).upper()
            image_ref = match.group(2)

            print(
                f"✅ VC MATCH FOUND → "
                f"label={label}, image={image_ref}"
            )

            if label in options:
                print(
                    f"⚠️ DUPLICATE LABEL IGNORED → {label}"
                )
            else:
                options[label] = {
                    "label": label,
                    "image_ref": image_ref,
                }
                print(
                    f"➕ OPTION STORED → {label}"
                )

        else:
            if text.startswith(("A:", "B:", "C:", "D:")):
                print(
                    f"❌ LOOKS LIKE OPTION BUT REGEX DID NOT MATCH → {repr(text)}"
                )

    print("\n📊 VC OPTIONS COLLECTED:")
    for k, v in options.items():
        print(f"  {k}: {v['image_ref']}")

    print(
        f"📊 VC OPTION COUNT = {len(options)} "
        f"(expected 4)"
    )
    print("================ VC DEBUG END ==================\n")

    if len(options) != 4:
        raise ValueError(
            f"VC: Expected 4 options, found {len(options)}"
        )

    return [options[k] for k in ["A", "B", "C", "D"]]

def normalize_text(text: str) -> str:
    if not text:
        return ""

    # Normalize unicode (fix smart chars, mixed forms)
    text = unicodedata.normalize("NFKC", text)

    # Replace non-breaking spaces
    text = text.replace("\u00A0", " ")

    # Collapse whitespace
    text = re.sub(r"\s+", " ", text)

def vc_extract_correct_answer(block):
    print("🧪 VC CORRECT_ANSWER DEBUG START")

    for idx, item in enumerate(block):
        if item.get("type") != "text":
            print(f"  [{idx}] NON-TEXT ITEM → skipped")
            continue

        raw = item.get("text") or item.get("content") or ""
        print(f"  [{idx}] RAW TEXT: {repr(raw)}")

        text = normalize_text(raw)
        print(f"       NORMALIZED: {repr(text)}")

        if text.startswith("CORRECT_ANSWER"):
            print("       ↑ FOUND CORRECT_ANSWER LABEL")

            parts = text.split(":", 1)
            if len(parts) == 2 and parts[1].strip():
                value = parts[1].strip()
                print(f"       INLINE VALUE: {repr(value)}")
                return value

            print("       INLINE VALUE EMPTY → expecting next line")
            expecting_value = True
            continue

        if 'expecting_value' in locals() and expecting_value and text:
            print(f"       NEXT LINE VALUE: {repr(text)}")
            return text

    print("🧪 VC CORRECT_ANSWER DEBUG END")
    raise ValueError("VC: CORRECT_ANSWER not found")

def vc_extract_metadata(block):
    metadata = {}
    collecting = False

    for item in block:
        if item.get("type") != "text":
            continue

        raw = item.get("text") or item.get("content") or ""
        text = raw.strip()

        if text == "METADATA:":
            collecting = True
            continue

        if collecting:
            if text == "QUESTION_TEXT:":
                break

            if ":" in text:
                key, value = text.split(":", 1)
                metadata[key.strip()] = value.strip().strip('"')

    if not metadata:
        raise ValueError("VC: METADATA section not found")

    return metadata
def vc_extract_correct_answer_from_docx(file_bytes: bytes) -> str:
    doc = docx.Document(BytesIO(file_bytes))

    for para in doc.paragraphs:
        text = normalize_text(para.text)
        if not text:
            continue

        if text.startswith("CORRECT_ANSWER"):
            parts = text.split(":", 1)
            if len(parts) == 2 and parts[1].strip():
                return parts[1].strip()

    raise ValueError("VC: CORRECT_ANSWER not found in DOCX")

def vc_validate_block(parsed):
    """
    Validation gatekeeper for VISUAL_COUNTING (Type 6).

    This function must be called AFTER parse_visual_counting_block()
    and BEFORE persist_visual_counting_question().

    It enforces structural correctness only.
    No image understanding. No counting. No AI.
    """

    # --------------------------------------------------
    # Required top-level keys
    # --------------------------------------------------
    required_keys = {
        "QUESTION_TEXT",
        "OPTIONS",
        "CORRECT_ANSWER",
        "METADATA",
    }

    missing = required_keys - parsed.keys()
    if missing:
        raise ValueError(
            f"VC: Missing required keys: {sorted(missing)}"
        )

    # --------------------------------------------------
    # QUESTION_TEXT
    # --------------------------------------------------
    question_text = parsed["QUESTION_TEXT"]
    if not isinstance(question_text, str) or not question_text.strip():
        raise ValueError("VC: QUESTION_TEXT must be a non-empty string")

    # --------------------------------------------------
    # OPTIONS
    # --------------------------------------------------
    options = parsed["OPTIONS"]

    if not isinstance(options, list):
        raise ValueError("VC: OPTIONS must be a list")

    if len(options) != 4:
        raise ValueError(
            f"VC: VISUAL_COUNTING requires exactly 4 options "
            f"(found {len(options)})"
        )

    labels_seen = set()

    for opt in options:
        if not isinstance(opt, dict):
            raise ValueError("VC: Each option must be a dict")

        label = opt.get("label")
        image_url = opt.get("image_url")

        if label not in {"A", "B", "C", "D"}:
            raise ValueError(
                f"VC: Invalid option label '{label}'"
            )

        if label in labels_seen:
            raise ValueError(
                f"VC: Duplicate option label '{label}'"
            )

        labels_seen.add(label)

        if not isinstance(image_url, str) or not image_url.strip():
            raise ValueError(
                f"VC: Option {label} must have a valid image_url"
            )

    # --------------------------------------------------
    # CORRECT_ANSWER
    # --------------------------------------------------
    correct = parsed["CORRECT_ANSWER"]

    if correct not in labels_seen:
        raise ValueError(
            f"VC: CORRECT_ANSWER '{correct}' "
            f"does not match any option label"
        )

    # --------------------------------------------------
    # METADATA
    # --------------------------------------------------
    metadata = parsed["METADATA"]

    if not isinstance(metadata, dict):
        raise ValueError("VC: METADATA must be a dict")

    required_metadata_fields = {
        "CLASS",
        "Year",
        "SUBJECT",
        "TOPIC",
        "DIFFICULTY",
    }

    missing_meta = required_metadata_fields - metadata.keys()
    if missing_meta:
        raise ValueError(
            f"VC: Missing METADATA fields: {sorted(missing_meta)}"
        )

    # --------------------------------------------------
    # Passed
    # --------------------------------------------------
    return True


def parse_visual_counting_block(block):
    return {
        "QUESTION_TEXT": vc_extract_question_text(block),        
        "METADATA": vc_extract_metadata(block),
    }
def persist_visual_counting_question(
    block,
    db,
    request_id,
    summary,
):
    """
    Persist a VISUAL_COUNTING (Type 6) question.

    Assumes:
    - block has already been validated by vc_validate_block
    - OPTIONS contain resolved image_url values
    """

    # --------------------------------------------------
    # Build options map: {"A": url, "B": url, ...}
    # --------------------------------------------------
    options_map = {
        opt["label"]: opt["image_url"]
        for opt in block["OPTIONS"]
    }

    # --------------------------------------------------
    # Create Question row
    # --------------------------------------------------
    question = QuestionNumeracyLC(
        question_type=6,
        class_name=block["METADATA"]["CLASS"],
        year=int(block["METADATA"]["Year"]),
        subject=block["METADATA"]["SUBJECT"],
        topic=block["METADATA"].get("TOPIC"),
        difficulty=block["METADATA"]["DIFFICULTY"],

        question_text=block["QUESTION_TEXT"],

        # Optional but useful for rendering/debugging
        question_blocks=block.get("QUESTION_BLOCKS"),

        options=options_map,

        # Stored as JSON for consistency across question types
        correct_answer={
            "value": block["CORRECT_ANSWER"]
        },
    )

    db.add(question)
    db.commit()
    db.refresh(question)

    summary.saved += 1

    print(
        f"[{request_id}] 💾 VC: Question saved successfully | "
        f"id={question.id}"
    )

from docx import Document
from io import BytesIO
import re
import unicodedata


def normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\u00a0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def vc_extract_options_from_docx(content: bytes) -> list[dict]:
    from docx import Document
    from io import BytesIO
    print("🔥🔥🔥 HIT vc_extract_options_from_docx FROM helpers/vc_options.py 🔥🔥🔥")
    doc = Document(BytesIO(content))

    print("\n🧾 VC DOCX FULL DUMP START")

    print("➡️ PARAGRAPHS:")
    for i, p in enumerate(doc.paragraphs):
        print(f"  [P{i}] {repr(p.text)}")

    print("\n➡️ TABLES:")
    for ti, table in enumerate(doc.tables):
        print(f"  TABLE {ti}:")
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    print(
                        f"    [T{ti} R{ri} C{ci} P{pi}] {repr(p.text)}"
                    )

    print("🧾 VC DOCX FULL DUMP END\n")

    
    options = []

    def try_parse_line(text: str):
        match = re.match(r"^([A-D])\s*:\s*(.+)?$", text)
        if match:
            options.append({
                "label": match.group(1),
                "image_ref": match.group(2),
            })

    # ---------------------------------------
    # 1. Scan normal paragraphs
    # ---------------------------------------
    for para in doc.paragraphs:
        text = normalize_text(para.text or "")
        if text:
            try_parse_line(text)

    # ---------------------------------------
    # 2. Scan tables (THIS WAS MISSING)
    # ---------------------------------------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = normalize_text(para.text or "")
                    if text:
                        try_parse_line(text)

    if not options:
        raise ValueError("VC: No options detected")

    if len(options) != 4:
        raise ValueError(
            f"VC: Expected 4 options, found {len(options)}"
        )

    return options
def vc_resolve_option_images(
    options: list[dict],
    db,
    request_id: str,
) -> list[dict]:
    resolved = []

    for opt in options:
        image_ref = opt["image_ref"]

        record = (
            db.query(UploadedImage)
            .filter(UploadedImage.original_name == image_ref)
            .first()
        )

        if not record:
            raise ValueError(
                f"VC: Image '{image_ref}' not found in uploaded_images"
            )

        resolved.append({
            "label": opt["label"],
            "image_url": record.gcs_url,
        })

        print(
            f"[{request_id}] 🖼️ VC: resolved {image_ref} → {record.gcs_url}"
        )

    return resolved
def vc_extract_options_from_block(question_block):
    options = []

    for item in question_block:
        if not isinstance(item, dict):
            continue

        text = (item.get("content") or "").strip()

        match = re.match(r"^([A-D])\s*[:.]\s*(.+)$", text)
        if match:
            options.append({
                "label": match.group(1),
                "image_ref": match.group(2),
            })

    if len(options) != 4:
        raise ValueError(
            f"VC: Expected 4 options, found {len(options)}"
        )

    return options
def vc_extract_correct_answer_from_block(question_block):
    """
    Extracts correct answer for Visual Counting (Type 6)
    from the already-isolated exam block.
    """

    for item in question_block:
        if not isinstance(item, dict):
            continue

        text = (item.get("content") or "").strip()

        if text.upper().startswith("CORRECT_ANSWER"):
            # Move to next non-empty line
            continue

        # Match single-letter answer (A–D)
        if text in {"A", "B", "C", "D"}:
            return text

    raise ValueError("VC: CORRECT_ANSWER not found in block")

def process_visual_counting_exam(
    block_idx,
    question_block,
    db,
    request_id,
    summary,
):
    print(
        f"[{request_id}] 🖼️ TYPE 6 detected | "
        f"processing visual counting question"
    )
    print(f"[{request_id}] 🔥 VC STEP 1: entered process_visual_counting_exam")

    # --------------------------------------------------
    # 1. Parse NON-option, NON-answer block content only
    # --------------------------------------------------
    print(f"[{request_id}] 🔥 VC STEP 2: parsing visual counting block")
    parsed = parse_visual_counting_block(question_block)

    print(
        f"[{request_id}] 🔥 VC STEP 2a: parsed base fields | "
        f"keys={list(parsed.keys())}"
    )

    # 🚨 DO NOT validate here

    # --------------------------------------------------
    # 2. Force DOCX-level extraction (OPTIONS)
    # --------------------------------------------------
    print(f"[{request_id}] 🔥 VC STEP 3: extracting OPTIONS from DOCX")
    raw_options = vc_extract_options_from_block(question_block)

    print(
        f"[{request_id}] 🔥 VC STEP 3a: raw options extracted | "
        f"count={len(raw_options)} | raw={raw_options}"
    )

    parsed["OPTIONS"] = vc_resolve_option_images(
        raw_options,
        db=db,
        request_id=request_id,
    )

    print(
        f"[{request_id}] 🔥 VC STEP 3b: resolved option images | "
        f"count={len(parsed['OPTIONS'])} | options={parsed['OPTIONS']}"
    )

    # --------------------------------------------------
    # 3. Extract CORRECT_ANSWER
    # --------------------------------------------------
    print(f"[{request_id}] 🔥 VC STEP 4: extracting CORRECT_ANSWER from DOCX")
    parsed["CORRECT_ANSWER"] = vc_extract_correct_answer_from_block(question_block)

    print(
        f"[{request_id}] 🔥 VC STEP 4a: correct answer extracted | "
        f"value={parsed['CORRECT_ANSWER']}"
    )

    # --------------------------------------------------
    # 4. Validate (single gatekeeper)
    # --------------------------------------------------
    print(f"[{request_id}] 🔥 VC STEP 5: validating visual counting block")
    vc_validate_block(parsed)

    print(f"[{request_id}] 🔥 VC STEP 5a: validation PASSED")

    # --------------------------------------------------
    # 5. Persist
    # --------------------------------------------------
    print(f"[{request_id}] 🔥 VC STEP 6: persisting visual counting question")
    question_id = persist_visual_counting_question(
        block=parsed,
        db=db,
        request_id=request_id,
        summary=summary,
    )

    print(
        f"[{request_id}] 🔥 VC STEP 6a: persistence complete | "
        f"question_id={question_id}"
    )

    summary.block_success(block_idx, [question_id])
    print(
        f"[{request_id}] 🟢 BLOCK {block_idx} SUCCESS (VISUAL COUNTING)"
    )


def debug_detect_question_types(question_block):
    types = set()

    for item in question_block:
        if not isinstance(item, dict):
            continue

        text = (item.get("content") or "").lower()
        if "question_type" in text:
            parts = text.replace(":", " ").split()
            for p in parts:
                if p.isdigit():
                    types.add(p)

    return types

WORD_RE = re.compile(r"^[A-Za-z]+$")

def normalize_ws_word(raw: str) -> str:
    # Strip whitespace (including Word junk)
    word = raw.strip()

    # Remove leading/trailing punctuation
    word = re.sub(r"^[^\w]+|[^\w]+$", "", word)

    return word

HEADER_RE = re.compile(r"^[A-Z_]+\s*:")

def is_header_line(line: str) -> bool:
    if not line:
        return False
    return bool(HEADER_RE.match(line.strip()))

def ws_extract_selectable_words(ctx):
    """
    Extracts and normalizes SELECTABLE_WORDS using a ParsingCursor.
    Returns: List[str]
    """

    # Seek SELECTABLE_WORDS header
    while ctx.peek() and not ctx.peek().strip().upper().startswith("SELECTABLE_WORDS"):
        ctx.next()

    if not ctx.peek():
        raise ValueError("MISSING_SELECTABLE_WORDS")

    ctx.next()  # consume SELECTABLE_WORDS:

    # Skip blank lines after header
    while ctx.peek() and not ctx.peek().strip():
        ctx.next()

    raw_words = []

    # Collect words until next header
    while ctx.peek():
        current = ctx.peek().strip()

        if is_header_line(current):
            break

        raw_words.append(current)
        ctx.next()

    if not raw_words:
        raise ValueError("SELECTABLE_WORDS cannot be empty")

    # Normalize words
    selectable_words = []
    for raw in raw_words:
        word = normalize_ws_word(raw)
        if word:
            selectable_words.append(word)

    # Validate single-word tokens
    for w in selectable_words:
        if not WORD_RE.match(w):
            raise ValueError(
                "WORD_SELECTION validation failed: SELECTABLE_WORDS must be single words only"
            )

    return selectable_words



QUESTION_SECTION_HEADERS = {
    "QUESTION_TEXT",
    "SENTENCE",
    "SELECTABLE_WORDS",
    "CORRECT_ANSWER",
    "ANSWER_TYPE",
}

def ws_extract_metadata_from_block(block_text: str) -> dict:
    metadata = {}
    in_meta = False

    print("🧪 ENTER ws_extract_metadata_from_block")

    for idx, raw_line in enumerate(block_text.splitlines(), start=1):
        line = raw_line.strip()
        print(f"🧪 META line {idx}: {repr(line)} | in_meta={in_meta}")

        if not line:
            continue

        # Start metadata
        if line.upper() == "METADATA:":
            in_meta = True
            print("🧪 META START detected")
            continue

        if not in_meta:
            continue

        # Stop ONLY at real question sections
        header_name = line.split(":", 1)[0].strip().upper()
        if header_name in QUESTION_SECTION_HEADERS:
            print(f"🧪 META END detected at section header: {repr(line)}")
            break

        # Parse metadata key-value
        if ":" in line:
            k, v = line.split(":", 1)
            key = k.strip().lower()
            value = v.strip().strip('"')
            metadata[key] = value
            print(f"🧪 META PARSED: {key} = {value}")

    print("🧪 META FINAL DICT:", metadata)
    return metadata

async def process_exam_block(
    block_idx: int,
    question_block: list,
    db: Session,
    request_id: str,
    summary
):
    print("\n" + "-" * 60)
    print(f"[{request_id}] ▶️ BLOCK {block_idx} START")
    print(f"[{request_id}] 📦 Block elements = {len(question_block)}")
    detected_types = debug_detect_question_types(question_block)
    print(
        f"[{request_id}] 🧪 Block {block_idx} detected question_types = {detected_types}"
    )
    validate_single_question_type(question_block)

    # 🔒 Determine question_type ONCE
    # ==================================================
    # 🧩 CLOZE BRANCH (QUESTION TYPE 5 ONLY)
    # ==================================================
    if is_cloze_question(question_block):
        print(
            f"[{request_id}] 🧩 CLOZE detected | block={block_idx}"
        )
    
        try:
            q = extract_cloze_from_exam_block(question_block)
            validate_cloze_deterministic(q)
    
            handle_cloze_question(
                q=q,
                question_block=question_block,
                meta=q,
                db=db,
                request_id=request_id,
                summary=summary,
                block_idx=block_idx,
            )

    
            summary.block_success(block_idx, [q])
            print(
                f"[{request_id}] 🟢 BLOCK {block_idx} SUCCESS (CLOZE)"
            )
            return
    
        except Exception as e:
            error_msg = str(e)
            print(
                f"[{request_id}] ❌ BLOCK {block_idx} FAILED (CLOZE) | "
                f"error={error_msg}"
            )
            summary.block_failure(block_idx, error_msg)
    
        return


    # ==================================================
    # 🖼️ TYPE 6 — VISUAL COUNTING (SEALED)
    # ==================================================
    if is_visual_counting_exam(question_block):
        print(
            f"[{request_id}] 🖼️ TYPE 6 detected | "
            f"processing visual counting question"
        )
        process_visual_counting_exam(
            block_idx=block_idx,
            question_block=question_block,
            db=db,
            request_id=request_id,
            summary=summary,
        )
        return  # 🚨 DO NOT FALL THROUGH

    # ==================================================
    # 🔤 TYPE 7 — WORD_SELECTION (DETERMINISTIC)
    # ==================================================
    block_text = "\n".join(
        item["content"]
        for item in question_block
        if isinstance(item, dict) and "content" in item
    )
    
    if is_word_selection_exam(block_text):
        metadata = ws_extract_metadata_from_block(block_text)
        ctx = ParsingCursor(block_text.splitlines())

        print("🧪 DEBUG cursor fast-forward start")
        
        steps = 0
        while ctx.peek():
            current = ctx.peek()
            print(f"   ↪ peek[{steps}] = {repr(current)}")
        
            if current.strip().upper().startswith("QUESTION_TEXT"):
                print("   🎯 FOUND QUESTION_TEXT header")
                break
        
            ctx.next()
            steps += 1
        
        print("🧪 DEBUG cursor fast-forward end")
        print(f"🧪 DEBUG cursor final peek = {repr(ctx.peek())}")

        
        question_text = ws_extract_question_text(ctx)
        sentence = ws_extract_sentence(ctx)
        selectable_words = ws_extract_selectable_words(ctx)
        # ⬇️ NEW: skip ANSWER_TYPE
        while ctx.peek() and not ctx.peek().strip().upper().startswith("CORRECT_ANSWER"):
            ctx.next()
        correct_answer = ws_extract_correct_answer(ctx)
    
        parsed = {
            "question_text": question_text,
            "sentence": sentence,
            "selectable_words": selectable_words,
            "correct_answer": correct_answer,
        }
    
        ws_validate_block(parsed)
        print("🧪 DEBUG metadata dict:", metadata)
        print("🧪 DEBUG metadata keys:", list(metadata.keys()))

    
        persist_word_selection_question(
            db=db,
            metadata=metadata,
            question_text=question_text,
            sentence=sentence,
            selectable_words=selectable_words,
            correct_answer=correct_answer,
        )
    
        summary.block_success(block_idx, [1])
        return  # 🚨 DO NOT FALL THROUGH
    
     
    
        # --------------------------------------------------
        # Quick structural sanity check
        # --------------------------------------------------
        if not looks_like_question(question_block):
            print(
                f"[{request_id}] ⚠️ Block {block_idx} skipped "
                f"(not a question)"
            )
            return

    
    
        
    # ==================================================
    # 🧠 LEGACY BRANCH (QUESTION TYPES 1–4)
    # ==================================================
    try:
        print(
            f"[{request_id}] 🤖 Calling legacy GPT parser "
            f"for block {block_idx}"
        )

        questions = await parse_questions_with_gpt_naplan_numeracy_lc(
            question_block=question_block,
            request_id=request_id
        )

        print(
            f"[{request_id}] 🤖 GPT returned "
            f"{len(questions)} question(s)"
        )

        if not questions:
            raise ValueError("GPT returned zero valid questions")

        meta = validate_common_metadata(
            questions=questions,
            block_idx=block_idx,
            request_id=request_id
        )

        print(
            f"[{request_id}] 🏷️ Metadata validated | "
            f"subject={meta.get('subject')}"
        )

        for i, q in enumerate(questions, start=1):
            question_type = q.get("question_type")

            if not question_type:
                raise ValueError(
                    "Legacy question missing question_type"
                )

            print(
                f"[{request_id}] ➕ Persisting question "
                f"{i}/{len(questions)} (type={question_type})"
            )

            persist_question(
                q=q,
                question_type=question_type,
                question_block=question_block,
                meta=meta,
                db=db,
                request_id=request_id,
                summary=summary,
                block_idx=block_idx
            )

        summary.block_success(block_idx, questions)
        print(
            f"[{request_id}] 🟢 BLOCK {block_idx} SUCCESS"
        )

    except Exception as e:
        error_msg = str(e)
        print(
            f"[{request_id}] ❌ BLOCK {block_idx} FAILED | "
            f"error={error_msg}"
        )
        summary.block_failure(block_idx, error_msg)

def validate_cloze_dropdown(q: dict):
    if q.get("answer_type") != "CLOZE_DROPDOWN":
        raise ValueError("Invalid answer_type for CLOZE")

    cloze_text = q.get("cloze_text")
    if not cloze_text or "{{dropdown}}" not in cloze_text:
        raise ValueError("CLOZE text must contain {{dropdown}}")

    options = q.get("options")
    if not options or not isinstance(options, dict):
        raise ValueError("CLOZE requires options dictionary")

    correct = q.get("correct_answer")
    if not isinstance(correct, str):
        raise ValueError("CLOZE correct_answer must be a string")

    if correct not in options:
        raise ValueError("CLOZE correct_answer must be a valid option key")

def validate_common_metadata(questions, block_idx, request_id):
    required = [
        "class_name",
        "year",
        "subject",
        "topic",
        "difficulty",
    ]



    missing = [f for f in required if not questions[0].get(f)]
    if missing:
        raise ValueError(f"Missing metadata: {missing}")

    subject = questions[0]["subject"].strip().title()
    if subject not in {"Numeracy", "Language Conventions"}:
        raise ValueError(f"Invalid subject: {subject}")

    return {"subject": subject}
 
def validate_text_input(q):
    # Text input questions are permissive for now
    return

def persist_question(
    q,
    question_type,   # 👈 ADD THIS
    question_block,
    meta,
    db,
    request_id,
    summary,
    block_idx
):
    
    validate_question_by_type(question_type, q)

    resolve_images(q, db, request_id)

    if question_type == 5:
        question_text = q.get("cloze_text")
    else:
        question_text = "\n\n".join(
            b["content"] for b in question_block if b["type"] == "text"
        )


    obj = QuestionNumeracyLC(
        question_type=question_type,
        class_name=q["class_name"],
        year=q["year"],
        subject=meta["subject"],
        topic=q["topic"],
        difficulty=q["difficulty"],
        question_text=question_text,
        question_blocks=filter_display_blocks(question_block),
        options=q.get("options"),
        correct_answer=str(q["correct_answer"]).strip(),
    )

    db.add(obj)
    db.commit()
    db.refresh(obj)

    summary.saved += 1
    print(f"[{request_id}] ✅ SAVED question (id={obj.id})")

def log_start(request_id, file):
    print("\n" + "=" * 70)
    print(f"🚀 GPT-UPLOAD-NAPLAN START | request_id={request_id}")
    print("=" * 70)
    print(f"[{request_id}] 📄 Filename     : {file.filename}")
    print(f"[{request_id}] 📄 Content-Type : {file.content_type}")


def log_exam_count(request_id, blocks):
    print(f"[{request_id}] 🔪 Chunked into {len(blocks)} exam blocks")


def log_end(request_id, summary):
    print("\n" + "=" * 70)
    print(f"🏁 GPT-UPLOAD-NAPLAN END | request_id={request_id}")
    print(f"📊 Saved questions : {summary.saved}")
    print(f"📊 Skipped partial : {summary.skipped_partial}")
    print("=" * 70)

def validate_question_by_type(qt: int, q: dict):
    """
    Validates a question payload based on question_type.

    NOTE:
    - Types 1 & 2 (MCQs) are lenient for legacy data
    - Type 5 (CLOZE_DROPDOWN) is strict
    """

    # --------------------------------------------------
    # Type 1: Single-choice MCQ (lenient for legacy)
    # --------------------------------------------------
    if qt == 1:
        options = q.get("options") or {}
        if not options:
            print(
                "⚠️ Legacy MCQ_SINGLE without options allowed"
            )

        if not q.get("correct_answer"):
            raise ValueError("question_type=1 requires correct_answer")

        if not isinstance(q["correct_answer"], str):
            raise ValueError(
                "question_type=1 correct_answer must be a string"
            )

    # --------------------------------------------------
    # Type 2: Multi-select MCQ (lenient for legacy)
    # --------------------------------------------------
    elif qt == 2:
        options = q.get("options") or {}
        if not options:
            print(
                "⚠️ Legacy MCQ_MULTI without options allowed"
            )

        if not q.get("correct_answer"):
            raise ValueError("question_type=2 requires correct_answer")

        if not isinstance(q["correct_answer"], list):
            raise ValueError(
                "question_type=2 correct_answer must be a list"
            )

        if not all(isinstance(a, str) for a in q["correct_answer"]):
            raise ValueError(
                "question_type=2 correct_answer list must contain strings only"
            )

    # --------------------------------------------------
    # Type 3: Numeric input (strict)
    # --------------------------------------------------
    elif qt == 3:
        if q.get("correct_answer") is None:
            raise ValueError("question_type=3 requires correct_answer")

    # --------------------------------------------------
    # Type 4: Text input (strict)
    # --------------------------------------------------
    elif qt == 4:
        if not q.get("correct_answer"):
            raise ValueError("question_type=4 requires correct_answer")

        if not isinstance(q["correct_answer"], str):
            raise ValueError(
                "question_type=4 correct_answer must be a string"
            )

    # --------------------------------------------------
    # Type 5: CLOZE_DROPDOWN (STRICT)
    # --------------------------------------------------
    elif qt == 5:
        if not q.get("cloze_text"):
            raise ValueError("question_type=5 requires cloze_text")

        if "{{dropdown}}" not in q["cloze_text"]:
            raise ValueError(
                "question_type=5 cloze_text must contain {{dropdown}}"
            )

        options = q.get("options") or []
        if not options:
            raise ValueError("question_type=5 requires options")

        if not q.get("correct_answer"):
            raise ValueError("question_type=5 requires correct_answer")

    # --------------------------------------------------
    # Unsupported
    # --------------------------------------------------
    else:
        raise ValueError(f"Unsupported question_type: {qt}")

def resolve_images(q: dict, db: Session, request_id: str):
    """
    Resolves image blocks in q["question_blocks"] by mapping
    original image names to stored GCS URLs.

    Mutates q["question_blocks"] in place.
    """

    blocks = q.get("question_blocks", [])

    for block in blocks:
        if block.get("type") != "image":
            continue

        raw_name = (
            block.get("name")
            or block.get("src")
            or ""
        ).strip().lower()

        print(f"[{request_id}] 🖼️ Resolving image: '{raw_name}'")

        record = (
            db.query(UploadedImage)
            .filter(
                func.lower(func.trim(UploadedImage.original_name))
                == raw_name
            )
            .first()
        )

        if not record:
            raise ValueError(
                f"Image '{raw_name}' not uploaded yet"
            )

        # Replace image reference
        block.pop("name", None)
        block["src"] = record.gcs_url

        print(
            f"[{request_id}] ✅ Image resolved → {record.gcs_url}"
        )
def is_cloze_exam(exam_block: list[dict]) -> bool:
    for el in exam_block:
        text = el.get("content", "")
        if isinstance(text, str) and "question_type: 5" in text.lower():
            return True
    return False

import re


import re


def ws_validate_block(parsed: dict) -> None:
    """
    Validates a parsed Type 7 — WORD_SELECTION block.

    Raises ValueError on any validation failure.
    Returns None on success.
    """

    if not isinstance(parsed, dict):
        raise ValueError("Parsed block must be a dict")

    question_text = parsed.get("question_text")
    sentence = parsed.get("sentence")
    selectable_words = parsed.get("selectable_words")
    correct_answer = parsed.get("correct_answer")

    # Rule 1: QUESTION_TEXT exists
    if not question_text or not question_text.strip():
        raise ValueError(
            "WORD_SELECTION validation failed: QUESTION_TEXT is missing or empty"
        )

    # Rule 2: SENTENCE exists
    if not sentence or not sentence.strip():
        raise ValueError(
            "WORD_SELECTION validation failed: SENTENCE is missing or empty"
        )

    # Rule 3: SELECTABLE_WORDS exists and is a list
    if not isinstance(selectable_words, list) or not selectable_words:
        raise ValueError(
            "WORD_SELECTION validation failed: SELECTABLE_WORDS must be a non-empty list"
        )

    # Normalize selectable words
    normalized_words = []
    for w in selectable_words:
        if not isinstance(w, str) or not w.strip():
            raise ValueError(
                "WORD_SELECTION validation failed: SELECTABLE_WORDS contains invalid entries"
            )
        if " " in w.strip():
            raise ValueError(
                "WORD_SELECTION validation failed: SELECTABLE_WORDS must be single words only"
            )
        normalized_words.append(w.strip())

    # Rule 4: CORRECT_ANSWER exists
    if not correct_answer or not correct_answer.strip():
        raise ValueError(
            "WORD_SELECTION validation failed: CORRECT_ANSWER is missing or empty"
        )

    answer = correct_answer.strip()

    # Rule 5: CORRECT_ANSWER must be a single word
    if " " in answer:
        raise ValueError(
            "WORD_SELECTION validation failed: CORRECT_ANSWER must be a single word"
        )

    # Rule 6: CORRECT_ANSWER must be one of SELECTABLE_WORDS
    if answer not in normalized_words:
        raise ValueError(
            "WORD_SELECTION validation failed: "
            "CORRECT_ANSWER must be one of SELECTABLE_WORDS"
        )

    # Rule 7: Every selectable word must appear in sentence (word-boundary safe)
    for word in normalized_words:
        pattern = rf"\b{re.escape(word)}\b"
        if not re.search(pattern, sentence, flags=re.IGNORECASE):
            raise ValueError(
                f"WORD_SELECTION validation failed: "
                f"Selectable word '{word}' does not appear as a standalone word in SENTENCE"
            )



def persist_word_selection_question(
    *,
    db,
    metadata: dict,
    question_text: str,
    sentence: str,
    selectable_words: list[str],
    correct_answer: str,
):
    """
    Persists a validated Type 7 — WORD_SELECTION question.

    Assumes ws_validate_block has already been called.
    """
    class_name = (
        metadata.get("class_name")   # legacy / internal
        or metadata.get("class")     # document-authored
    )
    print("🧪 FINAL question_blocks payload:", {
        "sentence": sentence,
        "selectable_words": selectable_words,
    })

    record = QuestionNumeracyLC(
        question_type=7,
        class_name=class_name,
        year=int(metadata.get("year")) if metadata.get("year") else None,
        subject=metadata.get("subject"),
        topic=metadata.get("topic"),
        difficulty=metadata.get("difficulty"),
        question_text=question_text,
        question_blocks={
            "sentence": sentence,
            "selectable_words": selectable_words,
        },
        options=None,
        correct_answer={"value": correct_answer},
    )

    db.add(record)
    db.commit()
    db.refresh(record)

    return record
def is_word_selection_exam(block: str) -> bool:
    """
    Detects whether a full exam block represents
    Question Type 7 — WORD_SELECTION.

    This function is:
    - Deterministic
    - Text-only
    - Side-effect free
    - Safe to call before any parsing
    """

    if not block or not isinstance(block, str):
        return False

    text = block.lower()

    required_markers = [
        "question_type: 7",
        'answer_type: "word_selection"',
        "sentence:",
        "correct_answer:",
    ]

    return all(marker in text for marker in required_markers)

def ws_extract_metadata(block: str) -> dict:
    """
    Extracts METADATA fields for Type 7 — WORD_SELECTION.

    Returns a dict with keys:
    - class_name
    - year
    - subject
    - topic
    - difficulty

    Missing fields are returned as None.
    """

    metadata = {
        "class_name": None,
        "year": None,
        "subject": None,
        "topic": None,
        "difficulty": None,
    }

    if not block or not isinstance(block, str):
        return metadata

    # Isolate METADATA section
    try:
        meta_section = block.split("METADATA:", 1)[1]
    except IndexError:
        return metadata

    # Stop at the next known section boundary
    stop_markers = [
        "\nQUESTION_TEXT:",
        "\nSENTENCE:",
        "\nANSWER_TYPE:",
        "\nCORRECT_ANSWER:",
    ]

    for marker in stop_markers:
        if marker in meta_section:
            meta_section = meta_section.split(marker, 1)[0]

    # Line-by-line parse
    for line in meta_section.splitlines():
        line = line.strip()
        if ":" not in line:
            continue

        key, value = line.split(":", 1)
        key = key.strip().lower()
        value = value.strip().strip('"')

        if key == "class":
            metadata["class_name"] = value
        elif key == "year":
            metadata["year"] = value
        elif key == "subject":
            metadata["subject"] = value
        elif key == "topic":
            metadata["topic"] = value
        elif key == "difficulty":
            metadata["difficulty"] = value

    return metadata
def ws_extract_question_text(ctx):
    print("🧪 ENTER ws_extract_question_text")
    print(f"🧪 peek at entry = {repr(ctx.peek())}")

    # Expect QUESTION_TEXT header
    if not ctx.peek() or not ctx.peek().strip().upper().startswith("QUESTION_TEXT"):
        return None

    ctx.next()  # consume QUESTION_TEXT:

    # ⬇️ IMPORTANT: skip blank lines
    while ctx.peek() and not ctx.peek().strip():
        ctx.next()

    lines = []

    while ctx.peek():
        current = ctx.peek().strip()

        # Stop at next section header
        if is_header_line(current):
            break

        lines.append(current)
        ctx.next()

    question_text = " ".join(lines).strip()
    print(f"🧪 extracted QUESTION_TEXT = {repr(question_text)}")

    return question_text if question_text else None

#here123
def ws_extract_sentence(ctx):
    """
    Extracts SENTENCE using a ParsingCursor.
    Cursor is expected to be positioned at SENTENCE:
    """

    # Expect SENTENCE header
    if not ctx.peek() or not ctx.peek().strip().upper().startswith("SENTENCE"):
        return None

    ctx.next()  # consume SENTENCE:

    # Skip blank lines after header
    while ctx.peek() and not ctx.peek().strip():
        ctx.next()

    lines = []

    while ctx.peek():
        current = ctx.peek().strip()

        # Stop at next section header
        if is_header_line(current):
            break

        lines.append(current)
        ctx.next()

    sentence = " ".join(lines).strip()
    return sentence if sentence else None
def ws_extract_correct_answer(ctx):
    """
    Extracts CORRECT_ANSWER using a ParsingCursor.
    Cursor is expected to be positioned at CORRECT_ANSWER:
    """

    # Expect CORRECT_ANSWER header
    if not ctx.peek() or not ctx.peek().strip().upper().startswith("CORRECT_ANSWER"):
        return None

    ctx.next()  # consume CORRECT_ANSWER:

    # Skip blank lines after header
    while ctx.peek() and not ctx.peek().strip():
        ctx.next()

    lines = []

    while ctx.peek():
        current = ctx.peek().strip()

        # Stop at next section header or exam boundary
        if is_header_line(current) or current.startswith("==="):
            break

        lines.append(current)
        ctx.next()

    answer = " ".join(lines).strip()
    return answer if answer else None


def validate_single_question_type(exam_block):
    found = set()

    for item in exam_block:
        if item.get("type") != "text":
            continue

        text = (item.get("content") or "").lower()

        if "question_type:" in text:
            # extract the number crudely but safely
            found.add(text.strip())

    if len(found) > 1:
        raise ValueError(
            f"Invalid exam block: multiple question_type declarations found: {found}"
        )



@app.post("/upload-word-naplan")
async def upload_word_naplan(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    request_id = str(uuid.uuid4())[:8]

    print("\n" + "=" * 80)
    print(f"🚀 upload-word-naplan START | request_id={request_id}")
    print("=" * 80)

    # --------------------------------------------------
    # STEP 1: Log incoming file
    # --------------------------------------------------
    print(f"[{request_id}] 📥 Incoming file received")
    print(f"[{request_id}] 📄 filename      = {file.filename}")
    print(f"[{request_id}] 📄 content_type  = {file.content_type}")

    # --------------------------------------------------
    # STEP 2: Read + validate file
    # --------------------------------------------------
    print(f"[{request_id}] ▶️ STEP 2: read_and_validate_file START")

    content = await read_and_validate_file(file, request_id)

    print(
        f"[{request_id}] ✅ STEP 2 COMPLETE | "
        f"bytes_read={len(content)}"
    )

    # --------------------------------------------------
    # STEP 3: Parse DOCX into ordered blocks
    # --------------------------------------------------
    print(f"[{request_id}] ▶️ STEP 3: parse_docx_blocks START")

    ordered_blocks = parse_docx_blocks(content, request_id)

    print(
        f"[{request_id}] ✅ STEP 3 COMPLETE | "
        f"ordered_blocks_count={len(ordered_blocks)}"
    )

    # Defensive visibility
    if ordered_blocks:
        print(
            f"[{request_id}] 🧩 First block preview: "
            f"type={ordered_blocks[0].get('type')} "
            f"content_preview={str(ordered_blocks[0])[:120]}"
        )

    # --------------------------------------------------
    # STEP 4: Chunk blocks into EXAM sections
    # --------------------------------------------------
    print(f"[{request_id}] ▶️ STEP 4: chunk_by_exam_markers START")

    exam_blocks = chunk_by_exam_markers(ordered_blocks)

    print(
        f"[{request_id}] ✅ STEP 4 COMPLETE | "
        f"exam_blocks_detected={len(exam_blocks)}"
    )

    for i, blk in enumerate(exam_blocks, start=1):
        print(
            f"[{request_id}] 🧪 Exam block {i} | "
            f"elements={len(blk)}"
        )

    # --------------------------------------------------
    # STEP 5: Initialise summary
    # --------------------------------------------------
    print(f"[{request_id}] ▶️ STEP 5: Initialise UploadSummary")

    summary = UploadSummary()
    summary.file_bytes = content
    
    print(f"[{request_id}] ✅ UploadSummary initialised")

    # --------------------------------------------------
    # STEP 6: Process each EXAM block
    # --------------------------------------------------
    print(f"[{request_id}] ▶️ STEP 6: Begin EXAM processing loop")

    for idx, block in enumerate(exam_blocks, start=1):
        print("\n" + "-" * 70)
        print(f"[{request_id}] 🔄 Processing EXAM BLOCK {idx}")
        print(f"[{request_id}] 📦 Block element count = {len(block)}")
    
        if not isinstance(block, list):
            print(
                f"[{request_id}] ❌ Block {idx} is not a list | "
                f"type={type(block)}"
            )
            continue
    
        try:
            # --------------------------------------------------
            # 🧩 CLOZE exam → full exam block
            # --------------------------------------------------
            validate_single_question_type(block)
            if is_cloze_exam(block):
                print(
                    f"[{request_id}] 🧩 CLOZE exam detected | "
                    f"passing full exam block"
                )
    
                await process_exam_block(
                    block_idx=idx,
                    question_block=block,
                    db=db,
                    request_id=request_id,
                    summary=summary
                )
    
                continue  # 🚨 do NOT fall through
    
            # --------------------------------------------------
            # 🧠 Legacy exams → unchanged behavior
            # --------------------------------------------------
            await process_exam_block(
                block_idx=idx,
                question_block=block,
                db=db,
                request_id=request_id,
                summary=summary
            )
    
        except Exception as e:
            print(
                f"[{request_id}] ❌ EXCEPTION in block {idx} | "
                f"error={e}"
            )
            summary.block_failure(idx, str(e))


    # --------------------------------------------------
    # STEP 7: Final logging
    # --------------------------------------------------
    print("\n" + "=" * 80)
    print(f"🏁 upload-word-naplan END | request_id={request_id}")
    print(f"[{request_id}] 📊 SUMMARY")
    print(f"[{request_id}]   saved           = {summary.saved}")
    print(f"[{request_id}]   skipped_partial = {summary.skipped_partial}")
    print(f"[{request_id}]   blocks_reported = {len(summary.blocks)}")
    print("=" * 80)

    response = summary.response()

    print(
        f"[{request_id}] 📤 Response payload preview: "
        f"{str(response)[:300]}"
    )

    return response




@app.get("/api/student/exam-status")
def exam_status(student_id: str, subject: str, db: Session = Depends(get_db)):
    """
    Returns:
    {
        "started": true/false,
        "completed": true/false,
        "attempts_used": 1,
        "attempts_allowed": 3,
        "total_questions": 40,
        "exam_id": 12
    }
    """

    # 1. Validate student exists
    student = db.query(Student).filter(Student.student_id == student_id).first()
    if not student:
        raise HTTPException(status_code=404, detail="Student not found.")

    # 2. Find a quiz requirement for this student's class and subject
    quiz = (
        db.query(Quiz)
        .filter(
            func.lower(Quiz.class_name) == student.class_name.lower(),
            Quiz.subject == subject
        )
        .order_by(Quiz.id.desc())
        .first()
    )


    if not quiz:
        return {
            "started": False,
            "completed": False,
            "attempts_used": 0,
            "attempts_allowed": 3,
            "total_questions": 0,
            "exam_id": None
        }

    # Total number of questions in this quiz
    total_questions = sum(t["total"] for t in quiz.topics)

    # 3. Check if student has any exam attempts for this quiz
    attempts = (
        db.query(StudentExam)
        .filter(StudentExam.student_id == student.id, StudentExam.exam_id.isnot(None))
        .all()
    )

    attempts_used = len(attempts)
    attempts_allowed = 3  # your rule (frontend shows 1/3)

    # 4. Find the latest attempt for status
    latest_attempt = (
        db.query(StudentExam)
        .filter(StudentExam.student_id == student.id)
        .order_by(StudentExam.id.desc())
        .first()
    )

    # If no attempts at all
    if not latest_attempt:
        return {
            "started": False,
            "completed": False,
            "attempts_used": attempts_used,
            "attempts_allowed": attempts_allowed,
            "total_questions": total_questions,
            "quiz_id": quiz.id,
            "exam_id": None
        }

    # 5. Load the exam linked to latest attempt
    exam = db.query(Exam).filter(Exam.id == latest_attempt.exam_id).first()

    # Exam may not exist yet if not generated
    if not exam:
        return {
            "started": False,
            "completed": False,
            "attempts_used": attempts_used,
            "attempts_allowed": attempts_allowed,
            "total_questions": total_questions,
            "quiz_id": quiz.id,
            "exam_id": None
        }

    # Determine started & completed status
    started = latest_attempt.started_at is not None
    completed = latest_attempt.completed_at is not None

    return {
        "started": started,
        "completed": completed,
        "attempts_used": attempts_used,
        "attempts_allowed": attempts_allowed,
        "total_questions": len(exam.questions),
        "quiz_id": quiz.id,
        "exam_id": exam.id
    }


@app.post("/api/student-exams/start")
def start_exam(student_id: int, exam_id: int, db: Session = Depends(get_db)):
    """
    A student starts an exam.
    We create a StudentExam row and return the exam questions to frontend.
    """
    exam = db.query(Exam).filter(Exam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")

    # Create student_exam entry
    student_exam = StudentExam(
        student_id=student_id,
        exam_id=exam.id,
        status="started",
        started_at=datetime.now(timezone.utc)
    )

    db.add(student_exam)
    db.commit()
    db.refresh(student_exam)

    return {
        "message": "Exam started",
        "student_exam_id": student_exam.id,
        "exam_id": exam.id,
        "questions": exam.questions  # send questions to frontend
    }

@app.post("/api/student-exams/answer")
def submit_answer(
    student_exam_id: int,
    question_id: int,
    student_answer: str,
    db: Session = Depends(get_db)
):
    """
    Save an answer for a question and evaluate correctness.
    """
    # 1. Get the exam session
    student_exam = db.query(StudentExam).filter(
        StudentExam.id == student_exam_id
    ).first()

    if not student_exam:
        raise HTTPException(status_code=404, detail="Student exam session not found")

    # 2. Find the exam the student is taking
    exam = db.query(Exam).filter(Exam.id == student_exam.exam_id).first()

    # 3. Look for the question inside exam.questions JSON
    correct_answer = None
    for q in exam.questions:
        if q["q_id"] == question_id:
            correct_answer = q["correct"]
            break

    if correct_answer is None:
        raise HTTPException(status_code=400, detail="Invalid question ID")

    # 4. Save answer in DB
    answer_row = StudentExamAnswer(
        student_exam_id=student_exam_id,
        question_id=question_id,
        student_answer=student_answer,
        correct_answer=correct_answer,
        is_correct=(student_answer == correct_answer)
    )

    db.add(answer_row)
    db.commit()

    return {
        "message": "Answer saved",
        "correct": student_answer == correct_answer
    }




@app.post("/retrieve-week-number")
def retrieve_week_number(request: WeekRequest, db: Session = Depends(get_db)):    
    try:
        # Convert input date string to date object
        input_date = datetime.strptime(request.date, "%Y-%m-%d").date()
        # ------------------ Save or replace admin date ------------------
        existing_entry = db.query(AdminDate).first()
        if existing_entry:
            existing_entry.date = request.date
        else:
            existing_entry = AdminDate(date=request.date)
            db.add(existing_entry)

        db.commit()

        # ------------------ Calculate week number ------------------

        # Get the Monday of the week of the input date
        input_monday = input_date - timedelta(days=input_date.weekday())  # weekday(): Mon=0 ... Sun=6

        # Get the current date
        today = datetime.now(timezone.utc).date()
        current_monday = today - timedelta(days=today.weekday())

        # Calculate number of weeks between input week and current week
        delta_days = (current_monday - input_monday).days
        if delta_days < 0:
            raise HTTPException(status_code=400, detail="Selected date is in the future")

        week_number = (delta_days // 7) + 1  # first week is 1

        return {"week_number": week_number}
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD")

@app.post("/login-exam-module")
def login_exam_module(login_data: StudentLogin, db: Session = Depends(get_db)):
    """
    Login a student using student_id and password
    """
    # Look up student by student_id
    student = db.query(Student).filter(Student.student_id == login_data.student_id).first()
    
    if not student:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Simple password check (plaintext; replace with hashed check in production)
    if login_data.password != student.password:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Return student info (without password)
    return {
        "student_id": student.student_id,
        "name": student.name,
        "parent_email": student.parent_email,
        "class_name": student.class_name,
        "class_day": student.class_day
    }


@app.get("/student-name")
def get_student_name(student_id: int = Query(...), db: Session = Depends(get_db)):
    """
    Return student name for a given student_id.
    """
    student = db.query(User).filter(User.id == student_id).first()

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    return {"student_id": student.id, "student_name": student.name}
    
@app.get("/run-scheduler-now")
def run_scheduler_now():
    generate_quizzes()
    return {"message": "Scheduler task executed"}

@app.get("/get-pending-quiz/{user_id}")
def get_pending_quiz(user_id: int, db: Session = Depends(get_db)):
    quiz = db.query(StudentQuiz).filter_by(student_id=user_id, status="pending").first()
    if not quiz:
        return {"message": "No pending quiz found"}
    return {
        "quiz_id": quiz.quiz_id,
        "activity_id": quiz.activity_id,
        "quiz_json": quiz.quiz_json,
        "status": quiz.status,
        "created_at": quiz.created_at
    }

#endpoint to save a certain exam for Exam module
@app.post("/api/quizzes")
def create_quiz(quiz: QuizCreate, db: Session = Depends(get_db)):
    """
    Create a new quiz with extensive debugging.
    """

    print("\n========== QUIZ CREATION START ==========")
    
    # Print the raw QuizCreate object
    print("🔍 Incoming QuizCreate object:", quiz)

    # Convert quiz to dict to inspect content
    try:
        quiz_dict = quiz.dict()
        print("📦 Parsed quiz payload:", quiz_dict)
    except Exception as e:
        print("❌ Failed to convert quiz to dict:", e)
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=f"Invalid quiz model: {str(e)}")

    # Debug each field
    print("➡️ class_name:", quiz.class_name, type(quiz.class_name))
    print("➡️ subject:", quiz.subject, type(quiz.subject))
    print("➡️ difficulty:", quiz.difficulty, type(quiz.difficulty))
    print("➡️ num_topics:", quiz.num_topics, type(quiz.num_topics))
    print("➡️ topics:", quiz.topics, type(quiz.topics))

    # Validate topics field
    if not isinstance(quiz.topics, list):
        print("❌ ERROR: topics is not a list")
        raise HTTPException(status_code=400, detail="topics must be a list")

    print("📝 Topics count:", len(quiz.topics))
    for i, t in enumerate(quiz.topics):
        print(f"   └─ Topic {i}: {t}")

    try:
        print("\n--- Deleting previous data (Exam → Quiz) ---")

        # 1️⃣ Delete dependent exams (StudentExam deleted automatically via cascade)
        db.query(Exam).filter(Exam.subject == "thinking_skills").delete(
            synchronize_session=False
        )

        # 2️⃣ Delete quizzes
        db.query(Quiz).delete(synchronize_session=False)

        db.commit()
        print("🗑️ All previous exams and quizzes deleted")

        print("\n--- Creating SQLAlchemy Quiz object ---")

        print("\n--- Creating SQLAlchemy Quiz object ---")
        
        new_quiz = Quiz(
            class_name = quiz.class_name,
            subject = quiz.subject,
            difficulty = quiz.difficulty,
            num_topics = quiz.num_topics,
            topics = [t.dict() for t in quiz.topics]   # Convert to JSON-safe dict
        )
        print("✅ SQLAlchemy object created:", new_quiz)

        print("\n--- Adding to DB session ---")
        db.add(new_quiz)

        print("\n--- Attempting commit ---")
        db.commit()
        print("✅ Commit successful!")

        db.refresh(new_quiz)
        print("🔄 Refreshed object:", new_quiz)

        print("========== QUIZ CREATION COMPLETE ==========\n")

        return {
            "message": "Quiz created successfully",
            "quiz_id": new_quiz.id
        }

    except Exception as e:
        print("\n❌ EXCEPTION DURING DB OPERATION ❌")
        print("Error message:", str(e))
        traceback.print_exc()

        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error creating quiz: {str(e)}")


@app.get("/api/quizzes")
def get_quizzes(db: Session = Depends(get_db)):
    """
    Return a minimal list of quiz requirement entries for frontend selection.
    """
    try:
        quizzes = db.query(Quiz).order_by(Quiz.id.desc()).all()
        return [
            {
                "id": q.id,
                "class_name": q.class_name,
                "subject": q.subject,
                "difficulty": q.difficulty
            }
            for q in quizzes
        ]

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching quizzes: {str(e)}")


     
def generate_otp():
    return random.randint(100000, 999999)

@app.post("/send-otp")
def send_otp_endpoint(request: OTPRequest, db: Session = Depends(get_db)):
    email = request.email.strip().lower()
    print(f"[DEBUG] Received OTP request for email: {email}")

    if not email:
        raise HTTPException(status_code=400, detail="Email is required")

    user = db.query(User).filter(User.email == email).first()
    if not user:
        raise HTTPException(status_code=404, detail="Email not registered")

    otp = generate_otp()
    otp_store[email] = {"otp": otp, "expiry": time.time() + 300}

    # Call your email sending function
    try:
        send_otp_email(email, otp)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error sending email: {e}")

    return {"message": "OTP sent successfully"} 

@app.post("/retrieve-term-start-date", response_model=TermStartDateResponse)
def retrieve_term_start_date(db: Session = Depends(get_db)):
    try:
        stmt = select(AdminDate).limit(1)
        result = db.execute(stmt).scalars().first()

        if not result or not result.date:
            raise HTTPException(status_code=404, detail="Term start date not found")

        term_date = result.date
        # Convert to string if needed
        if isinstance(term_date, (date, datetime)):
            term_date_str = term_date.isoformat()
        else:
            term_date_str = str(term_date)  # already string

        return TermStartDateResponse(date=term_date_str)

    except Exception as e:
        print("Error retrieving term start date:", e)
        raise HTTPException(status_code=500, detail="Internal server error")

@app.get("/get_next_user_id_exam_module")
def get_next_user_id_exam_module(db: Session = Depends(get_db)):
    """
    Returns the next numeric student id (as string),
    safely handling string-based IDs.
    """

    try:
        last_id_int = (
            db.query(func.max(cast(Student.id, Integer)))
            .filter(Student.id.op("~")("^[0-9]+$"))  # only numeric strings
            .scalar()
        ) or 0

        next_id = last_id_int + 1
        return str(next_id)

    except Exception as e:
        print("[ERROR] Failed to get next student id:", e)
        raise HTTPException(
            status_code=500,
            detail="Could not generate next user ID",
        )
     
@app.post("/set-term-start-date")
def set_term_start_date(term_data: AdminDateSchema, db: Session = Depends(get_db)):
    try:
        # Clear any previous dates (since primary key only allows unique row)
        db.query(AdminDate).delete()
        
        # Add new date
        new_date = AdminDate(date=term_data.date)
        db.add(new_date)
        db.commit()
        db.refresh(new_date)

        return {"message": "Term start date updated successfully", "date": new_date.date}

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/set-term-data")
def set_term_data(term_data: TermDataSchema, db: Session = Depends(get_db)):
    try:
        # Convert strings to date objects
        term_start = datetime.fromisoformat(term_data.termStartDate).date()
        term_end = datetime.fromisoformat(term_data.termEndDate).date()

        # Optionally, delete old data or just insert new
        # db.query(AdminTerm).delete()  # if you want only 1 row

        new_term = AdminTerm(
            year=term_data.year,
            term_start_date=term_start,
            term_end_date=term_end
        )
        db.add(new_term)
        db.commit()
        db.refresh(new_term)

        return {
            "message": "Term data saved successfully",
            "term": {
                "year": new_term.year,
                "termStartDate": str(new_term.term_start_date),
                "termEndDate": str(new_term.term_end_date)
            }
        }

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/get-term-dates")
def get_term_dates(db: Session = Depends(get_db)):
    try:
        terms = db.query(AdminTerm).all()
        return [
            {
                "id": t.id,
                "year": t.year,
                "termStartDate": t.term_start_date.isoformat(),
                "termEndDate": t.term_end_date.isoformat()
            }
            for t in terms
        ]
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


        
@app.post("/verify-otp")
def verify_otp(request: OTPVerify, db: Session = Depends(get_db)):
    email = request.email.strip().lower()
    otp = str(request.otp).strip()

    if not email:
        raise HTTPException(status_code=400, detail="Email is required")

    # --- Check if OTP exists for this email ---
    record = otp_store.get(email)
    if not record:
        raise HTTPException(status_code=400, detail="OTP not sent")

    # --- Check expiry ---
    if time.time() > record["expiry"]:
        otp_store.pop(email, None)
        raise HTTPException(status_code=400, detail="OTP expired")

    # --- Check OTP match ---
    if otp != str(record["otp"]):
        raise HTTPException(status_code=401, detail="Invalid OTP")

    # --- Fetch user by email ---
    user = db.query(User).filter(User.email == email).first()
    if not user:
        raise HTTPException(status_code=404, detail="Email not registered")

    # --- Remove OTP after successful verification ---
    otp_store.pop(email, None)

    # --- Respond with user info ---
    return {
        "message": "OTP verified",
        "user": {
            "student_id": user.id,
            "email": user.email,
            "name": user.name,
            "class_name": user.class_name,
            "class_day": user.class_day
        }
    }
    
    
@app.post("/login")
def login(request: LoginRequest, response: Response, db: Session = Depends(get_db)):
    phone = request.phone_number.strip()
    print(f"[DEBUG] Received login request for phone: {phone}")

    if not phone:
        raise HTTPException(status_code=400, detail="Phone number is required")

    # --- Fetch user by phone ---
    user = db.query(User).filter(User.phone_number == phone).first()
    if not user:
        print(f"[WARNING] Phone number {phone} not registered")
        raise HTTPException(status_code=401, detail="Phone number not registered")

    # --- Retrieve OTP record ---
    record = otp_store.get(phone)
    if not record:
        print(f"[WARNING] No OTP record found for phone {phone}")
        raise HTTPException(status_code=400, detail="OTP not sent")

    if time.time() > record["expiry"]:
        print(f"[WARNING] OTP for phone {phone} has expired")
        otp_store.pop(phone, None)
        raise HTTPException(status_code=400, detail="OTP expired")

    if str(request.otp) != str(record["otp"]):
        print(f"[WARNING] Entered OTP ({request.otp}) does not match stored OTP ({record['otp']})")
        raise HTTPException(status_code=400, detail="Invalid OTP")

    print(f"[INFO] OTP for phone {phone} is valid")
    otp_store.pop(phone, None)  # clear OTP after verification

    # --- Clear previous user state if applicable ---
    if user.name in user_contexts:
        user_contexts[user.name] = []
        print(f"[DEBUG] Cleared previous context for user {user.name}")

    user_vectorstores_initialized[user.name] = False
    print(f"[DEBUG] vectorstores_initialized for user {user.name} set to False")

    # --- Remove existing session if using session management ---
    existing_session = db.query(SessionModel).filter(SessionModel.user_id == user.id).first()
    if existing_session:
        db.delete(existing_session)
        db.commit()
        print(f"[DEBUG] Cleared existing session for user {user.id}")

    # --- Set session cookie ---
    response.set_cookie(
        key="session_token",
        value=f"session_{user.id}",
        httponly=True,
        max_age=3600
    )

    # --- Prepare response ---
    user_info = {
        "id": user.id,
        "name": user.name,
        "phone_number": user.phone_number,
        "class_name": user.class_name,
        "class_day": user.class_day  # optional
    }

    print(f"[INFO] Login successful for user {user.name}")
    return {"message": "Login successful", "user": user_info}




@app.get("/get-activity")
def get_activity(student_id: int, db: Session = Depends(get_db)):
    activity = db.query(Activity).first()
    if not activity:
        raise HTTPException(status_code=404, detail="No activity found")
    return {
        "activity_id": activity.activity_id,
        "instructions": activity.instructions,
        "questions": activity.questions,
        "score_logic": activity.score_logic
    }


@app.get("/api/leaderboard/accumulative/{class_name}")
def accumulative_leaderboard(
    class_name: str = Path(..., description="Class name"),
    day: str = Query(..., description="Class day, e.g., Monday"),
    db: Session = Depends(get_db)
):
    """
    Returns accumulative leaderboard for a given class and class day
    within 10 weeks from admin start date.
    """
    # --- Get start date from AdminDate ---
    admin_date_entry = db.query(AdminDate).first()
    if not admin_date_entry or not admin_date_entry.date:
        raise HTTPException(status_code=400, detail="Admin start date not set")
    
    # Convert to datetime if stored as string
    if isinstance(admin_date_entry.date, str):
        try:
            start_date = datetime.strptime(admin_date_entry.date, "%Y-%m-%d")
        except ValueError:
            raise HTTPException(status_code=500, detail="Invalid date format in AdminDate")
    else:
        start_date = admin_date_entry.date

    # Calculate end date (10 weeks from start)
    end_date = start_date + timedelta(weeks=10)

    try:
        # Aggregate scores within date range and class day
        results = (
            db.query(
                QuizResult.student_id,
                QuizResult.student_name,
                func.sum(QuizResult.total_score).label("total_score")
            )
            .filter(QuizResult.class_name == class_name)
            .filter(QuizResult.class_day == day)
            .filter(QuizResult.submitted_at >= start_date)
            .filter(QuizResult.submitted_at <= end_date)
            .group_by(QuizResult.student_id, QuizResult.student_name)
            .order_by(func.sum(QuizResult.total_score).desc())
            .all()
        )

        leaderboard = [
            {
                "student_id": r.student_id,
                "student_name": r.student_name,
                "total_score": r.total_score
            }
            for r in results
        ]

        return {
            "class_name": class_name,
            "class_day": day,
            "start_date": start_date.strftime("%Y-%m-%d"),
            "end_date": end_date.strftime("%Y-%m-%d"),
            "leaderboard": leaderboard
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
        
@app.get("/quiz-results")
def get_quiz_results(
    class_name: str = Query(..., description="Class name to filter results"),
    db: Session = Depends(get_db)
):
    """
    Return all quiz results for a specific class, sorted by score (descending).
    """
    cleaned_class_name = class_name.strip()

    results = (
        db.query(QuizResult)
        .filter(QuizResult.class_name == cleaned_class_name)
        .order_by(QuizResult.total_score.desc())
        .all()
    )

    if not results:
        raise HTTPException(
            status_code=404,
            detail=f"No results found for class '{class_name}'"
        )

    response_data = [
        {
            "student_id": r.student_id,
            "student_name": r.student_name,
            "class_name": r.class_name,
            "total_score": r.total_score,
            "total_questions": r.total_questions,
            "submitted_at": r.submitted_at.isoformat()
        }
        for r in results
    ]

    return JSONResponse(content=response_data)


@app.post("/add-activities-from-csv")
async def add_activities_from_csv(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """
    Upload activities via CSV and create Activity entries identical to the /add-activity endpoint.
    """

    # --- Validate file type ---
    if not file.filename.endswith(".csv"):
        raise HTTPException(status_code=400, detail="File must be a CSV")

    # --- Read CSV ---
    try:
        df = pd.read_csv(file.file)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read CSV: {e}")

    required_columns = ["instructions", "score_logic", "class_name", "class_day", "week_number"]
    for col in required_columns:
        if col not in df.columns:
            raise HTTPException(status_code=400, detail=f"Missing required column: {col}")

    added_activities = []

    # --- Process each row ---
    for index, row in df.iterrows():

        instructions = str(row.get("instructions", "")).strip()
        score_logic = str(row.get("score_logic", "")).strip()
        class_name = str(row.get("class_name", "")).strip()
        class_day = str(row.get("class_day", "")).strip()
        week_number = row.get("week_number")

        # --- Validation (same as single-endpoint) ---
        if not instructions or not isinstance(instructions, str):
            continue
        if not score_logic or not isinstance(score_logic, str):
            continue
        if not class_name or not class_day:
            continue

        try:
            week_number = int(week_number)
        except:
            week_number = None

        # --- activity_type comes from score_logic ---
        activity_type = score_logic

        # --- SAME admin_prompt as /add-activity ---
        admin_prompt = (
            f"Create a gamified activity for students in {class_name}.\n"
            f"The activity format should be {activity_type}.\n"
            f"The topic taught is {instructions}."
        )

        # --- SAME default questions as /add-activity ---
        questions_json = [
            {
                "category": "General",
                "prompt": instructions,
                "options": ["A", "B", "C", "D"],
                "answer": "A"
            }
        ]

        # --- Create Activity entry ---
        activity = Activity(
            instructions=instructions,
            admin_prompt=admin_prompt,
            questions=questions_json,
            score_logic=score_logic,
            class_name=class_name,
            class_day=class_day,
            week_number=week_number
        )

        # --- Save ---
        try:
            db.add(activity)
            db.commit()
            db.refresh(activity)
            added_activities.append({
                "activity_id": activity.activity_id,
                "instructions": instructions,
                "week_number": week_number
            })
        except Exception as e:
            db.rollback()
            print(f"[ERROR] Could not save activity at row {index}: {e}")
            continue

    if not added_activities:
        raise HTTPException(
            status_code=400,
            detail="CSV processed but no valid activities were added."
        )

    return {
        "message": f"{len(added_activities)} activities added successfully",
        "activities": added_activities
    }

 
@app.post("/add-activity")
def add_activity(payload: ActivityPayload, db: Session = Depends(get_db)):
    """
    Add a new activity. Front end provides instructions and score_logic (used as activity type).
    Backend constructs admin_prompt for GPT use.
    """

    # --- Validate payload ---
    if not payload.instructions or not isinstance(payload.instructions, str):
        raise HTTPException(status_code=400, detail="Instructions must be a non-empty string")
    if not payload.score_logic or not isinstance(payload.score_logic, str):
        raise HTTPException(status_code=400, detail="score_logic (used as activity type) must be a non-empty string")
    if not payload.class_name or not payload.class_day:
        raise HTTPException(status_code=400, detail="class_name and class_day are required")
    if payload.week_number is not None and not isinstance(payload.week_number, int):
        raise HTTPException(status_code=400, detail="week_number must be an integer if provided")

    # --- Use score_logic as activity_type in admin_prompt ---
    activity_type = payload.score_logic

    admin_prompt = (
        f"Create a gamified activity for students in {payload.class_name}.\n"
        f"The activity format should be {activity_type}.\n"  # dynamic injection from score_logic
        f"The topic taught is {payload.instructions}."
    )

    # --- Default or provided questions ---
    questions_json = payload.questions or [
        {
            "category": "General",
            "prompt": payload.instructions,
            "options": ["A", "B", "C", "D"],  # placeholder options
            "answer": "A"  # placeholder answer
        }
    ]

    # --- Create Activity object ---
    activity = Activity(
        instructions=payload.instructions,
        admin_prompt=admin_prompt,   # save admin prompt for AI
        questions=questions_json,
        score_logic=payload.score_logic,
        class_name=payload.class_name,
        class_day=payload.class_day,
        week_number=payload.week_number
    )

    # --- Save to DB ---
    try:
        db.add(activity)
        db.commit()
        db.refresh(activity)
        return {
            "message": "Activity added successfully",
            "activity_id": activity.activity_id,
            "admin_prompt": activity.admin_prompt  # Optional: return for debugging
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to add activity: {e}")

@app.get("/get-quiz")
def get_quiz(
    class_name: str = Query(..., description="Class name of the quiz"),
    student_id: int = Query(..., description="ID of the student requesting the quiz"),
    db: Session = Depends(get_db)
):
    """
    Fetch the latest quiz for a given class_name.
    Trims whitespace and matches case-insensitively.
    Returns the quiz JSON.
    """
    # ------------------ Calculate week_number ------------------
    admin_date_entry = db.query(AdminDate).first()
    if not admin_date_entry or not admin_date_entry.date:
        raise HTTPException(status_code=400, detail="Admin date not set")
    
    date_to_use = admin_date_entry.date
    week_number = calculate_week_number(date_to_use)

    # ------------------ Check if student already attempted ------------------
    existing_result = db.query(QuizResult).filter(
        QuizResult.student_id == student_id,
        QuizResult.class_name == class_name.strip(),
        QuizResult.week_number == week_number
    ).first()

    if existing_result:
        return JSONResponse(
            status_code=200,
            content={"message": "You have already attempted this week's quiz."}
        )

    cleaned_class_name = class_name.strip()

    quiz = (
        db.query(StudentQuiz)
        .filter(func.lower(func.trim(StudentQuiz.class_name)) == cleaned_class_name.lower())
        .order_by(StudentQuiz.created_at.desc())
        .first()
    )

    if not quiz:
        raise HTTPException(status_code=404, detail=f"No quiz found for class '{class_name}'")

    # Ensure quiz_json is always returned as a dict
    quiz_data = quiz.quiz_json or {}

    return JSONResponse(content=quiz_data)
    

"""
@app.post("/submit-activity")
def submit_activity(submit: ActivitySubmit, db: Session = Depends(get_db)):
    activity = db.query(Activity).filter(Activity.activity_id == submit.activity_id).first()
    if not activity:
        raise HTTPException(status_code=404, detail="Activity not found")
    score = 0
    for i, answer in enumerate(submit.answers):
        if i < len(activity.questions) and answer == activity.questions[i]["answer"]:
            score += 1
    attempt = ActivityAttempt(
        student_id=submit.student_id,
        quiz_id=None,
        score=score,
        time_taken=random.randint(30, 180)
    )
    db.add(attempt)
    db.commit()
    return {"message": "Activity submitted", "score": score}
"""

@app.post("/submit-answer")
def submit_answer(
    student_id: int,
    quiz_id: int,
    question_index: int,
    selected_option: str,
    db: Session = Depends(get_db)
):
    """
    Record a student's answer for a specific quiz.
    Does nothing if a result already exists for this student and quiz.
    """
    # Check if the student already has a QuizResult for this quiz
    existing_result = db.query(QuizResult).filter(
        QuizResult.student_id == student_id,
        QuizResult.quiz_id == quiz_id
    ).first()
    
    if existing_result:
        return {
            "message": "Result already exists for this student and quiz",
            "current_score": existing_result.total_score
        }

    # Fetch the quiz by quiz_id
    quiz = db.query(StudentQuiz).filter(StudentQuiz.quiz_id == quiz_id).first()
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")

    # Initialize student_answers dictionary if missing
    if "student_answers" not in quiz.quiz_json:
        quiz.quiz_json["student_answers"] = {}

    # Initialize student-specific answers
    if str(student_id) not in quiz.quiz_json["student_answers"]:
        quiz.quiz_json["student_answers"][str(student_id)] = {}

    # Record the answer
    quiz.quiz_json["student_answers"][str(student_id)][f"q{question_index+1}"] = selected_option

    # Recalculate student's score
    student_answers = quiz.quiz_json["student_answers"][str(student_id)]
    score = 0
    for idx, q in enumerate(quiz.quiz_json.get("questions", [])):
        if student_answers.get(f"q{idx+1}") == q.get("answer"):
            score += 1

    # Save the score in student_answers
    quiz.quiz_json["student_answers"][str(student_id)]["score"] = score

    db.commit()
    return {
        "message": "Answer recorded",
        "current_score": score
    }

    

def send_otp_email(to_email: str, otp: str):
    message = Mail(
        from_email='noreply@gemkidsacademy.com.au',
        to_emails=to_email,
        subject='Your OTP Code',
        html_content=f'<p>Your OTP code is <strong>{otp}</strong>. It will expire in 5 minutes.</p>'
    )
    try:
        sg = SendGridAPIClient(SENDGRID_API_KEY)
        response = sg.send(message)
        print(f"[INFO] Email sent to {to_email}, status code {response.status_code}")
    except Exception as e:
        print(f"[ERROR] Failed to send email to {to_email}: {e}")
        raise


def calculate_week_number(date_str: str) -> int:
    try:
        input_date = datetime.strptime(date_str, "%Y-%m-%d").date()
        input_monday = input_date - timedelta(days=input_date.weekday())

        today = datetime.now(timezone.utc).date()
        current_monday = today - timedelta(days=today.weekday())

        delta_days = (current_monday - input_monday).days
        if delta_days < 0:
            raise HTTPException(status_code=400, detail="Selected date is in the future")

        week_number = (delta_days // 7) + 1
        return week_number
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD")


@app.post("/submit-quiz-answer")
def submit_quiz_answer(payload: AnswerPayload, db: Session = Depends(get_db)):
    """
    Submit a single answer for the current quiz question.
    Records answers in StudentQuiz, but creates only one QuizResult per student per quiz.
    """
    print("\n--- SUBMIT QUIZ ANSWER ---")
    print("Received payload:", payload)

    # Fetch the latest quiz for this class
    quiz = (
        db.query(StudentQuiz)
        .filter(StudentQuiz.class_name == payload.class_name)
        .order_by(StudentQuiz.created_at.desc())
        .first()
    )
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")

    # Start quiz if not started
    if not quiz.started_at:
        quiz.started_at = datetime.now(timezone.utc)

    # Load or initialize student_answers
    student_answers = quiz.quiz_json.get("student_answers") or {}

    # Record current answer
    q_key = f"q{payload.question_index + 1}"
    student_answers[q_key] = payload.selected_option
    quiz.quiz_json["student_answers"] = student_answers  # MutableDict tracks changes

    # Calculate current score
    correct_count = 0
    for idx, question in enumerate(quiz.quiz_json.get("questions", [])):
        q_key_check = f"q{idx+1}"
        student_answer = student_answers.get(q_key_check)
        if student_answer == question.get("answer"):
            correct_count += 1

    quiz.quiz_json["score"] = correct_count

    # Check if quiz is completed
    total_questions = len(quiz.quiz_json.get("questions", []))
    answered_questions = len(student_answers)
    quiz_completed = answered_questions >= total_questions

    if quiz_completed:
        quiz.status = "completed"
        quiz.completed_at = datetime.now(timezone.utc)

        # Check if a result already exists for this student for this quiz
        existing_result = db.query(QuizResult).filter(
            QuizResult.student_id == payload.student_id,
            QuizResult.quiz_id == quiz.quiz_id
        ).first()

        if existing_result:
            print("Result already exists, not creating a new row")
        else:
            admin_date_entry = db.query(AdminDate).first()
            date_to_use = admin_date_entry.date
            week_number = calculate_week_number(date_to_use)

            result = QuizResult(
                quiz_id=quiz.quiz_id,
                student_id=payload.student_id,
                student_name=payload.student_name,
                class_name=payload.class_name,
                class_day=getattr(payload, "class_day", None),
                week_number=week_number,   # <-- new
                total_score=correct_count,
                total_questions=total_questions,
                submitted_at=datetime.now(timezone.utc)
            )
            db.add(result)
            print("Result saved successfully")

    else:
        quiz.status = "in_progress"

    # Commit quiz updates (answers & score)
    db.commit()

    return {
        "message": "Answer recorded",
        "current_score": correct_count,
        "quiz_status": quiz.status
    }





